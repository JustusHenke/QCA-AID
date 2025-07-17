"""
QCA-AID: Qualitative Content Analysis with AI Support - Deductive Coding
========================================================================

A Python implementation of Mayring's Qualitative Content Analysis methodology,
enhanced with AI capabilities through the OpenAI API.

Version:
--------
0.9.18.2 (2025-07-15)

0.9.18.2 
- save complete console output to log-file in output
- fixes issues with intermediate export after aborting coding process

0.9.18.1 Hotfixes
- fixes broken manual review mode due to missing method after earlier refactoring 
- fixes a bug that prevents skipping PDF annotation
- fixes a but where ESC-button would not properly close and save intermediate coding

New in 0.9.18   
KATEGORIE-KONSISTENZ: Deduktiver Modus mit Hauptkategorie-Vorauswahl (1-3 wahrscheinlichste), 40-60% weniger Token, keine inkompatiblen Subkategorie-Zuordnungen
SUBKATEGORIE-VALIDIERUNG: Strikte Konsistenzprüfung mit automatischer Entfernung fremder Subkategorien, zweistufige Validierung, detailliertes Tracking
PERFORMANCE-OPTIMIERUNG: Fokussierte AI-Kodierung nur mit relevanten Kategorien, verbesserte Qualität durch kategorie-spezifischen Fokus, kompatibel mit allen Features
PYMUPDF-FIX: fitz.open() durch fitz.Document() ersetzt, robuste Fehlerbehandlung für PDF-Laden/-Speichern
CONFIDENCE-SCALES: Zentrale Klasse mit 5 spezialisierten Skalen (0.6+ definitiv, 0.8+ eindeutig), einheitliche textbelegte Konfidenz-Bewertungen in allen Prompts
EXPORT-FIX: Begründungen bei Nichtkodierung werden nun korrekt exportiert
CONFIG-Sheet: die Konfiguration des Tools im Codebook wird nun auch exportiert



Description:
-----------
This script provides a framework for conducting qualitative content analysis
following Mayring's approach, combining both deductive and inductive category
development. It supports automated coding through AI while maintaining
methodological rigor and transparency.

ATTENTION: 
Be aware that this script is still under development and not all features are available. Also keep
in mind that AI-Assistance is not perfect and the results depend on the quality of the input data.
Use the script at your own risk!
Feedback is welcome!

Key Features:
------------
- Automated text preprocessing and chunking
- Deductive category application with parallel processing
- Inductive category development (restructured and optimized)
- Abductive mode for subcategory enhancement
- Multi-coder support (AI and human) with parallel execution
- Graceful interruption handling with checkpoint saving
- Intercoder reliability calculation
- Comprehensive analysis export with performance metrics
- Real-time progress monitoring
- Detailed documentation of the coding process

Requirements:
------------
- Python 3.8+
- OpenAI API key
- Required packages: see requirements.txt

Usage:
------
1. Place interview transcripts in the 'input' directory
2. Configure .env file with OpenAI API key
3. Adjust CONFIG settings if needed
4. Run the script (ESC to interrupt and save progress)
5. Results will be exported to the 'output' directory

File Naming Convention:
---------------------
Interview files should follow the pattern: attribute1_attribute2_whatever_you_want.extension
Example: university-type_position_2024-01-01.txt

Performance Notes:
-----------------
- Large datasets now process 4x faster due to parallel batch processing
- Memory usage optimized for datasets with 1000+ text segments
- Progress can be monitored in real-time
- Analysis can be safely interrupted and resumed

Author:
-------
Justus Henke 
Institut für Hochschulforschung Halle-Wittenberg (HoF)
Contact: justus.henke@hof.uni-halle.de

License:
--------
MIT License
Copyright (c) 2025 Justus Henke

Repository:
----------
https://github.com/JustusHenke/QCA-AID

Citation:
--------
If you use this software in your research, please cite:
Henke, J. (2025). QCA-AID: Qualitative Content Analysis with AI Support [Computer software].
https://github.com/JustusHenke/QCA-AID 

"""

# ============================
# 1. Import notwendiger Module
# ============================
import os        # Dateisystemzugriff
import re        # Reguläre Ausdrücke für deduktives Codieren
import openai    # OpenAI API-Integration
from openai import AsyncOpenAI
import httpx
from mistralai import Mistral
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Union, Any
import json      # Export/Import von Daten (z.B. CSV/JSON)
import pandas as pd  # Zum Aggregieren und Visualisieren der Ergebnisse
import logging   # Protokollierung
import markdown  # Für Markdown-Konvertierung
from datetime import datetime  # Für Datum und Zeit
from dotenv import load_dotenv  # Für das Laden von Umgebungsvariablen
import asyncio  # Für asynchrone Programmierung
import spacy
from typing import Dict, List, Optional, Union, Any, TYPE_CHECKING, Tuple
from typing import List, Tuple, Set
from dataclasses import dataclass
from openai import OpenAI
import sys
import docx
import PyPDF2
import itertools 
from typing import List, Dict, Tuple
import numpy as np
from collections import defaultdict
from collections import Counter
from difflib import SequenceMatcher
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl import load_workbook
from openpyxl import Workbook
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import threading
from typing import Dict, List, Optional
import platform
import time
import statistics
import traceback 


# ============================
# 1. Lokale Imports
# ============================


# lokale Imports 
from QCA_Utils import (
    TokenTracker, get_input_with_timeout, _calculate_multiple_coding_stats,
    _patch_tkinter_for_threaded_exit, ConfigLoader,
    LLMProvider, LLMProviderFactory, LLMResponse, MistralProvider, OpenAIProvider,
    _sanitize_text_for_excel, _generate_pastel_colors, _format_confidence,
    EscapeHandler, add_escape_handler_to_manager,
    MultiSelectListbox, ManualMultipleCodingDialog, create_multiple_coding_results, show_multiple_coding_info, 
    setup_manual_coding_window_enhanced, validate_multiple_selection,
    DocumentReader, ManualReviewGUI, ManualReviewComponent, validate_category_specific_segments, analyze_multiple_coding_impact, export_multiple_coding_report,
    ConsoleLogger, TeeWriter
)
try:
    from QCA_Utils import PDFAnnotator, DocumentToPDFConverter
    pdf_annotation_available = True
except ImportError:
    pdf_annotation_available = False
from QCA_Prompts import QCAPrompts, ConfidenceScales # Prompt Bibliothek

# Instanziierung des globalen Token-Counters
token_counter = TokenTracker()



# ============================
# 2. Globale Variablen
# ============================

# Definition der Forschungsfrage
FORSCHUNGSFRAGE = "Wie gestaltet sich [Phänomen] im Kontext von [Setting] und welche [Aspekt] lassen sich dabei identifizieren?"

# Allgemeine Kodierregeln
KODIERREGELN = {
    "general": [
        "Kodiere nur manifeste, nicht latente Inhalte",
        "Berücksichtige den Kontext der Aussage",
        "Bei Unsicherheit dokumentiere die Gründe",
        "Kodiere vollständige Sinneinheiten",
        "Prüfe Überschneidungen zwischen Kategorien"
    ],
    "format": [
        "Markiere relevante Textstellen",
        "Dokumentiere Begründung der Zuordnung",
        "Gib Konfidenzwert (1-5) an",
        "Notiere eventuelle Querverbindungen zu anderen Kategorien"
    ]
}

# Deduktive Kategorien mit integrierten spezifischen Kodierregeln
DEDUKTIVE_KATEGORIEN = {
    "Akteure": {
        "definition": "Erfasst alle handelnden Personen, Gruppen oder Institutionen sowie deren Rollen, Beziehungen und Interaktionen",
        "rules": "Codiere Aussagen zu: Individuen, Gruppen, Organisationen, Netzwerken",
        "subcategories": {
            "Individuelle_Akteure": "Einzelpersonen und deren Eigenschaften",
            "Kollektive_Akteure": "Gruppen, Organisationen, Institutionen",
            "Beziehungen": "Interaktionen, Hierarchien, Netzwerke",
            "Rollen": "Formelle und informelle Positionen"
        },
        "examples": {
            "Die Projektleiterin hat die Entscheidung eigenständig getroffen",
            "Die Arbeitsgruppe trifft sich wöchentlich zur Abstimmung"    ,
            "Als Vermittler zwischen den Parteien konnte er den Konflikt lösen",
            "Die beteiligten Organisationen haben eine Kooperationsvereinbarung unterzeichnet"
        }
    },
    "Kontextfaktoren": {
        "definition": "Umfasst die strukturellen, zeitlichen und räumlichen Rahmenbedingungen des untersuchten Phänomens",
        "subcategories": {
            "Strukturell": "Organisatorische und institutionelle Bedingungen",
            "Zeitlich": "Historische Entwicklung, Zeitpunkte, Perioden",
            "Räumlich": "Geografische und sozialräumliche Aspekte",
            "Kulturell": "Normen, Werte, Traditionen"
        }
    },
    "Kontextfaktoren": {
        "definition": "Umfasst die strukturellen, zeitlichen und räumlichen Rahmenbedingungen des untersuchten Phänomens",
        "subcategories": {
            "Strukturell": "Organisatorische und institutionelle Bedingungen",
            "Zeitlich": "Historische Entwicklung, Zeitpunkte, Perioden",
            "Räumlich": "Geografische und sozialräumliche Aspekte",
            "Kulturell": "Normen, Werte, Traditionen"
        }
    },
    "Prozesse": {
        "definition": "Erfasst Abläufe, Entwicklungen und Veränderungen über Zeit",
        "subcategories": {
            "Entscheidungsprozesse": "Formelle und informelle Entscheidungsfindung",
            "Entwicklungsprozesse": "Veränderungen und Transformationen",
            "Interaktionsprozesse": "Kommunikation und Austausch",
            "Konfliktprozesse": "Aushandlungen und Konflikte"
        }
    },
    "Ressourcen": {
        "definition": "Materielle und immaterielle Mittel und Kapazitäten",
        "subcategories": {
            "Materiell": "Finanzielle und physische Ressourcen",
            "Immateriell": "Wissen, Kompetenzen, soziales Kapital",
            "Zugang": "Verfügbarkeit und Verteilung",
            "Nutzung": "Einsatz und Verwertung"
        }
    },
    "Strategien": {
        "definition": "Handlungsmuster und -konzepte zur Zielerreichung",
        "subcategories": {
            "Formell": "Offizielle Strategien und Pläne",
            "Informell": "Ungeschriebene Praktiken",
            "Adaptiv": "Anpassungsstrategien",
            "Innovativ": "Neue Lösungsansätze"
        }
    },
    "Outcomes": {
        "definition": "Ergebnisse, Wirkungen und Folgen von Handlungen und Prozessen",
        "subcategories": {
            "Intendiert": "Beabsichtigte Wirkungen",
            "Nicht_intendiert": "Unbeabsichtigte Folgen",
            "Kurzfristig": "Unmittelbare Effekte",
            "Langfristig": "Nachhaltige Wirkungen"
        }
    },
    "Herausforderungen": {
        "definition": "Probleme, Hindernisse und Spannungsfelder",
        "subcategories": {
            "Strukturell": "Systemische Barrieren",
            "Prozessual": "Ablaufbezogene Schwierigkeiten",
            "Individuell": "Persönliche Herausforderungen",
            "Kontextuell": "Umfeldbezogene Probleme"
        }
    },
    "Legitimation": {
        "definition": "Begründungen, Rechtfertigungen und Deutungsmuster",
        "subcategories": {
            "Normativ": "Wertbasierte Begründungen",
            "Pragmatisch": "Praktische Rechtfertigungen",
            "Kognitiv": "Wissensbasierte Erklärungen",
            "Emotional": "Gefühlsbezogene Deutungen"
        }
    }
}


VALIDATION_THRESHOLDS = {
    'MIN_DEFINITION_WORDS': 15,
    'MIN_EXAMPLES': 2,
    'SIMILARITY_THRESHOLD': 0.7,
    'MIN_SUBCATEGORIES': 2,
    'MAX_NAME_LENGTH': 50,
    'MIN_NAME_LENGTH': 3
}


# ------------------------
# Konfigurationskonstanten
# ------------------------

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG = {
    'MODEL_PROVIDER': 'OpenAI',
    'MODEL_NAME': 'gpt-4o-mini',
    'DATA_DIR': os.path.join(SCRIPT_DIR, 'input'),
    'OUTPUT_DIR': os.path.join(SCRIPT_DIR, 'output'),
    'CHUNK_SIZE': 1200,
    'CHUNK_OVERLAP': 50,
    'BATCH_SIZE': 8,
    'CODE_WITH_CONTEXT': False,
    'MULTIPLE_CODINGS': True, 
    'MULTIPLE_CODING_THRESHOLD': 0.85,  # Schwellenwert für zusätzliche Relevanz
    'ANALYSIS_MODE': 'deductive',
    'REVIEW_MODE': 'consensus',
    'ATTRIBUTE_LABELS': {
        'attribut1': 'Attribut1',
        'attribut2': 'Attribut2',
        'attribut3': 'Attribut3'  
    },
    'EXPORT_ANNOTATED_PDFS': True,  # Aktiviert/deaktiviert PDF-Annotation
    'PDF_ANNOTATION_FUZZY_THRESHOLD': 0.85,  # Schwellenwert für Fuzzy-Text-Matching
    'PDF_SIDEBAR_BAR_WIDTH': 8,  # Breite der Sidebar-Marker in Pixeln
    'PDF_SIDEBAR_SPACING': 2,  # Abstand zwischen Sidebar-Markern
    'CODER_SETTINGS': [
        {
            'temperature': 0.3,
            'coder_id': 'auto_1'
        },
        {
            'temperature': 0.5,
            'coder_id': 'auto_2'
        }
    ]
}

# Stelle sicher, dass die Verzeichnisse existieren
os.makedirs(CONFIG['DATA_DIR'], exist_ok=True)
os.makedirs(CONFIG['OUTPUT_DIR'], exist_ok=True)

# Lade Umgebungsvariablen
env_path = os.path.join(os.path.expanduser("~"), '.environ.env')
load_dotenv(env_path)



# ============================
# 3. Klassen und Funktionen
# ============================

@dataclass
class CategoryDefinition:
    """Datenklasse für eine Kategorie im Kodiersystem"""
    name: str
    definition: str
    examples: List[str]
    rules: List[str]
    subcategories: Dict[str, str]
    added_date: str
    modified_date: str

    def replace(self, **changes) -> 'CategoryDefinition':
        """
        Erstellt eine neue Instanz mit aktualisierten Werten.
        Ähnlich wie _replace bei namedtuples.
        """
        new_values = {
            'name': self.name,
            'definition': self.definition,
            'examples': self.examples.copy(),
            'rules': self.rules.copy(),
            'subcategories': self.subcategories.copy(),
            'added_date': self.added_date,
            'modified_date': datetime.now().strftime("%Y-%m-%d")
        }
        new_values.update(changes)
        return CategoryDefinition(**new_values)

    def update_examples(self, new_examples: List[str]) -> None:
        """Fügt neue Beispiele hinzu ohne Duplikate."""
        self.examples = list(set(self.examples + new_examples))
        self.modified_date = datetime.now().strftime("%Y-%m-%d")

    def update_rules(self, new_rules: List[str]) -> None:
        """Fügt neue Kodierregeln hinzu ohne Duplikate."""
        self.rules = list(set(self.rules + new_rules))
        self.modified_date = datetime.now().strftime("%Y-%m-%d")

    def add_subcategories(self, new_subcats: Dict[str, str]) -> None:
        """Fügt neue Subkategorien hinzu."""
        self.subcategories.update(new_subcats)
        self.modified_date = datetime.now().strftime("%Y-%m-%d")

    def to_dict(self) -> Dict:
        """Konvertiert die Kategorie in ein Dictionary."""
        return {
            'name': self.name,
            'definition': self.definition,
            'examples': self.examples,
            'rules': self.rules,
            'subcategories': self.subcategories,
            'added_date': self.added_date,
            'modified_date': self.modified_date
        }

@dataclass(frozen=True)  # Macht die Klasse immutable und hashable
class CodingResult:
    """Datenklasse für ein Kodierungsergebnis"""
    category: str
    subcategories: Tuple[str, ...]  # Änderung von List zu Tuple für Hashability
    justification: str
    confidence: Dict[str, Union[float, Tuple[str, ...]]]  # Ändere List zu Tuple
    text_references: Tuple[str, ...]  # Änderung von List zu Tuple
    uncertainties: Optional[Tuple[str, ...]] = None  # Änderung von List zu Tuple
    paraphrase: str = ""
    keywords: str = "" 

    def __post_init__(self):
        # Konvertiere Listen zu Tupeln, falls nötig
        object.__setattr__(self, 'subcategories', tuple(self.subcategories) if isinstance(self.subcategories, list) else self.subcategories)
        object.__setattr__(self, 'text_references', tuple(self.text_references) if isinstance(self.text_references, list) else self.text_references)
        if self.uncertainties is not None:
            object.__setattr__(self, 'uncertainties', tuple(self.uncertainties) if isinstance(self.uncertainties, list) else self.uncertainties)
        
        # Konvertiere confidence Listen zu Tupeln
        if isinstance(self.confidence, dict):
            new_confidence = {}
            for k, v in self.confidence.items():
                if isinstance(v, list):
                    new_confidence[k] = tuple(v)
                else:
                    new_confidence[k] = v
            object.__setattr__(self, 'confidence', new_confidence)

    def to_dict(self) -> Dict:
        """Konvertiert das CodingResult in ein Dictionary"""
        return {
            'category': self.category,
            'subcategories': list(self.subcategories),  # Zurück zu Liste für JSON-Serialisierung
            'justification': self.justification,
            'confidence': self.confidence,
            'text_references': list(self.text_references),  # Zurück zu Liste
            'uncertainties': list(self.uncertainties) if self.uncertainties else None,
            'paraphrase': self.paraphrase ,
            'keywords': self.keywords         }

@dataclass
class CategoryChange:
    """Dokumentiert eine Änderung an einer Kategorie"""
    category_name: str
    change_type: str  # 'add', 'modify', 'delete', 'merge', 'split'
    description: str
    timestamp: str
    old_value: Optional[dict] = None
    new_value: Optional[dict] = None
    affected_codings: List[str] = None
    justification: str = ""

class CategoryValidator:
    """
    FIX: Verbesserte Utility-Klasse für alle Kategorie- und Subkategorie-Validierungen
    Mit robuster Kategorie-Struktur-Erkennung
    """
    
    @staticmethod
    def validate_subcategories_for_category(subcategories: List[str], 
                                           main_category: str, 
                                           categories_dict: Dict[str, Any],
                                           warn_only: bool = True) -> List[str]:
        """
        FIX: Robuste Subkategorien-Validierung mit verbesserter Fehlerbehandlung
        
        Args:
            subcategories: Liste der zu validierenden Subkategorien
            main_category: Name der Hauptkategorie
            categories_dict: Dictionary mit Kategorie-Definitionen
            warn_only: Wenn True, nur warnen statt entfernen
            
        Returns:
            List[str]: Validierte Subkategorien-Liste
        """
        if not subcategories:
            return []
            
        if not categories_dict:
            print(f"⚠️ KRITISCH: Kein categories_dict verfügbar für Validierung!")
            return [] if not warn_only else subcategories
            
        if main_category not in categories_dict:
            if warn_only:
                print(f"⚠️ WARNUNG: Hauptkategorie '{main_category}' nicht im Kategoriensystem gefunden")
                print(f"   Verfügbare Kategorien: {list(categories_dict.keys())[:5]}...")
                return subcategories
            else:
                print(f"🔧 FIX: Hauptkategorie '{main_category}' ungültig - alle Subkategorien entfernt")
                return []
        
        # FIX: Verbesserte Extraktion der gültigen Subkategorien
        main_cat_def = categories_dict[main_category]
        valid_subcats = CategoryValidator._extract_valid_subcategories(main_cat_def)
        
        if not valid_subcats:
            if warn_only:
                print(f"⚠️ WARNUNG: Keine Subkategorien definiert für '{main_category}'")
                return subcategories
            else:
                print(f"🔧 FIX: Keine gültigen Subkategorien für '{main_category}' - alle entfernt")
                return []
        
        # FIX: Validiere jede Subkategorie einzeln
        validated = []
        invalid = []
        
        for subcat in subcategories:
            if subcat in valid_subcats:
                validated.append(subcat)
            else:
                invalid.append(subcat)
        
        # FIX: Detailliertes Logging
        if invalid:
            if warn_only:
                print(f"⚠️ WARNUNG: {len(invalid)} ungültige Subkategorien für '{main_category}': {invalid}")
                print(f"   Gültige Subkategorien: {sorted(list(valid_subcats))}")
                return subcategories  # Behalte alle
            else:
                print(f"🔧 FIX: Entfernt {len(invalid)} ungültige Subkategorien für '{main_category}': {invalid}")
                print(f"   Behalten: {validated}")
                print(f"   Gültige Optionen: {sorted(list(valid_subcats))}")
                return validated
        
        return subcategories
    
    @staticmethod
    def _extract_valid_subcategories(category_def: Any) -> Set[str]:
        """
        FIX: Robuste Extraktion von Subkategorien aus verschiedenen Datenstrukturen
        
        Args:
            category_def: Kategorie-Definition (kann verschiedene Typen sein)
            
        Returns:
            Set[str]: Set der gültigen Subkategorien-Namen
        """
        valid_subcats = set()
        
        try:
            # FIX: Versuche verschiedene mögliche Strukturen
            
            # 1. Object mit subcategories-Attribut (häufigster Fall)
            if hasattr(category_def, 'subcategories'):
                subcats = category_def.subcategories
                
                if isinstance(subcats, dict):
                    # Dict mit Subkategorie-Name -> Definition
                    valid_subcats.update(subcats.keys())
                elif isinstance(subcats, (list, set, tuple)):
                    # Liste/Set von Subkategorie-Namen
                    valid_subcats.update(str(sub) for sub in subcats)
                elif isinstance(subcats, str):
                    # Einzelner String
                    valid_subcats.add(subcats)
                    
            # 2. Dictionary-Struktur direkt
            elif isinstance(category_def, dict):
                if 'subcategories' in category_def:
                    sub_def = category_def['subcategories']
                    if isinstance(sub_def, dict):
                        valid_subcats.update(sub_def.keys())
                    elif isinstance(sub_def, (list, set)):
                        valid_subcats.update(str(sub) for sub in sub_def)
                        
            # 3. Liste von Subkategorien direkt
            elif isinstance(category_def, (list, set, tuple)):
                valid_subcats.update(str(sub) for sub in category_def)
                
        except Exception as e:
            print(f"⚠️ FEHLER bei Subkategorien-Extraktion: {str(e)}")
            print(f"   Kategorie-Definition Typ: {type(category_def)}")
            if hasattr(category_def, '__dict__'):
                print(f"   Verfügbare Attribute: {list(category_def.__dict__.keys())}")
        
        return valid_subcats
    
    @staticmethod
    def validate_coding_consistency(coding_result: Dict[str, Any], 
                                  categories_dict: Dict[str, Any],
                                  fix_inconsistencies: bool = True) -> Dict[str, Any]:
        """
        FIX: Verbesserte Konsistenz-Validierung für einzelne Kodierungen
        """
        if not coding_result:
            return coding_result
        
        main_category = coding_result.get('category', '')
        subcategories = coding_result.get('subcategories', [])
        
        # FIX: Hauptkategorie validieren
        if main_category and main_category not in categories_dict:
            print(f"⚠️ INKONSISTENZ: Hauptkategorie '{main_category}' existiert nicht im Kategoriensystem")
            if fix_inconsistencies:
                coding_result['category'] = 'Nicht kodiert'
                coding_result['subcategories'] = []
                coding_result['justification'] = f"[KORRIGIERT] Ungültige Kategorie '{main_category}' → 'Nicht kodiert'"
                return coding_result
        
        # FIX: Subkategorien validieren - nur wenn Hauptkategorie gültig
        if main_category and main_category in categories_dict and subcategories:
            validated_subcats = CategoryValidator.validate_subcategories_for_category(
                subcategories, main_category, categories_dict, warn_only=not fix_inconsistencies
            )
            
            if fix_inconsistencies:
                removed_subcats = set(subcategories) - set(validated_subcats)
                if removed_subcats:
                    coding_result['subcategories'] = validated_subcats
                    original_justification = coding_result.get('justification', '')
                    coding_result['justification'] = f"{original_justification} [FIX: Entfernt ungültige Subkategorien: {list(removed_subcats)}]"
        
        return coding_result
    
    @staticmethod
    def validate_multiple_codings(codings: List[Dict[str, Any]], 
                                categories_dict: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Validiert eine Liste von Kodierungen (z.B. bei Mehrfachkodierung)
        
        Args:
            codings: Liste von Kodierungsergebnissen
            categories_dict: Verfügbares Kategoriensystem
            
        Returns:
            List[Dict[str, Any]]: Liste validierter Kodierungen
        """
        validated_codings = []
        
        for coding in codings:
            validated_coding = CategoryValidator.validate_coding_consistency(
                coding, categories_dict, fix_inconsistencies=True
            )
            validated_codings.append(validated_coding)
        
        return validated_codings
    
    @staticmethod
    def get_category_statistics(categories_dict: Dict[str, Any]) -> Dict[str, Any]:
        """
        Erstellt Statistiken über das Kategoriensystem
        
        Args:
            categories_dict: Kategoriensystem-Dictionary
            
        Returns:
            Dict[str, Any]: Statistiken über Kategorien und Subkategorien
        """
        stats = {
            'total_main_categories': len(categories_dict),
            'categories_with_subcategories': 0,
            'total_subcategories': 0,
            'subcategories_per_category': {},
            'average_subcategories': 0.0,
            'category_details': {}
        }
        
        total_subcats = 0
        
        for cat_name, cat_def in categories_dict.items():
            valid_subcats = CategoryValidator._extract_valid_subcategories(cat_def)
            subcat_count = len(valid_subcats)
            
            if subcat_count > 0:
                stats['categories_with_subcategories'] += 1
                total_subcats += subcat_count
            
            stats['subcategories_per_category'][cat_name] = subcat_count
            stats['category_details'][cat_name] = {
                'subcategory_count': subcat_count,
                'subcategories': sorted(list(valid_subcats))
            }
        
        stats['total_subcategories'] = total_subcats
        stats['average_subcategories'] = total_subcats / max(1, len(categories_dict))
        
        return stats
    
    @staticmethod
    def find_category_conflicts(codings: List[Dict[str, Any]]) -> Dict[str, List[str]]:
        """
        Findet Konflikte zwischen verschiedenen Kodierungen
        
        Args:
            codings: Liste von Kodierungsergebnissen für das gleiche Segment
            
        Returns:
            Dict[str, List[str]]: Dictionary mit Konflikt-Typen und Details
        """
        conflicts = {
            'main_category_conflicts': [],
            'subcategory_conflicts': [],
            'confidence_discrepancies': []
        }
        
        if len(codings) <= 1:
            return conflicts
        
        # Hauptkategorien-Konflikte
        main_categories = [coding.get('category', '') for coding in codings]
        if len(set(main_categories)) > 1:
            conflicts['main_category_conflicts'] = list(set(main_categories))
        
        # Subkategorien-Konflikte (nur bei gleicher Hauptkategorie)
        if len(set(main_categories)) == 1:
            subcategory_sets = []
            for coding in codings:
                subcats = set(coding.get('subcategories', []))
                subcategory_sets.append(subcats)
            
            if len(set(frozenset(s) for s in subcategory_sets)) > 1:
                conflicts['subcategory_conflicts'] = [list(s) for s in subcategory_sets]
        
        # Konfidenz-Diskrepanzen
        confidences = []
        for coding in codings:
            conf = coding.get('confidence', {})
            if isinstance(conf, dict):
                total_conf = conf.get('total', 0.0)
            else:
                total_conf = float(conf) if conf else 0.0
            confidences.append(total_conf)
        
        if confidences and (max(confidences) - min(confidences)) > 0.3:
            conflicts['confidence_discrepancies'] = confidences
        
        return conflicts

# --- Klasse: MaterialLoader ---
# Aufgabe: Laden und Vorbereiten des Analysematerials (Textdokumente, output)
# WICHTIG: Lange Texte werden mittels Chunking in überschaubare Segmente zerlegt.
class MaterialLoader:
    """Lädt und verarbeitet Interviewdokumente."""
    
    def __init__(self, data_dir: str = CONFIG['DATA_DIR'], 
                 chunk_size: int = CONFIG['CHUNK_SIZE'], 
                 chunk_overlap: int = CONFIG['CHUNK_OVERLAP']):
        """
        Initialisiert den MaterialLoader.
        
        Args:
            data_dir (str): Verzeichnis mit den Dokumenten
            chunk_size (int): Ungefähre Anzahl der Zeichen pro Chunk
            chunk_overlap (int): Überlappung zwischen Chunks in Zeichen
        """
        self.data_dir = data_dir
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Lade das deutsche Sprachmodell für spaCy
        try:
            import spacy
            self.nlp = spacy.load("de_core_news_sm")
        except Exception as e:
            print("Bitte installieren Sie das deutsche Sprachmodell:")
            print("python -m spacy download de_core_news_sm")
            raise e

    def chunk_text(self, text: str) -> List[str]:
        """
        Teilt Text in überlappende Chunks basierend auf Satzgrenzen.
        
        Args:
            text (str): Zu teilender Text
            
        Returns:
            List[str]: Liste der Text-Chunks
        """
        # Debug-Ausgabe
        # print(f"\nChunking-Parameter:")
        # print(f"- Chunk Size: {self.chunk_size}")
        # print(f"- Chunk Overlap: {self.chunk_overlap}")
        # print(f"- Gesamtlänge Text: {len(text)} Zeichen")
        
        # Verarbeite den Text mit spaCy
        doc = self.nlp(text)
        sentences = list(doc.sents)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_text = sentence.text.strip()
            sentence_length = len(sentence_text)
            
            # Wenn der aktuelle Chunk mit diesem Satz zu groß würde
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Speichere aktuellen Chunk
                chunk_text = ' '.join(current_chunk)
                chunks.append(chunk_text)
                # print(f"- Neuer Chunk erstellt: {len(chunk_text)} Zeichen")
                
                # Starte neuen Chunk mit Überlappung
                # Berechne wie viele Sätze wir für die Überlappung behalten
                overlap_length = 0
                overlap_sentences = []
                for sent in reversed(current_chunk):
                    overlap_length += len(sent)
                    if overlap_length > self.chunk_overlap:
                        break
                    overlap_sentences.insert(0, sent)
                
                current_chunk = overlap_sentences
                current_length = sum(len(s) for s in current_chunk)
            
            # Füge den Satz zum aktuellen Chunk hinzu
            current_chunk.append(sentence_text)
            current_length += sentence_length
        
        # Letzten Chunk hinzufügen, falls vorhanden
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(chunk_text)
            # print(f"- Letzter Chunk: {len(chunk_text)} Zeichen")
        
        print(f"\nChunking Ergebnis:")
        print(f"- Anzahl Chunks: {len(chunks)}")
        print(f"- Durchschnittliche Chunk-Länge: {sum(len(c) for c in chunks)/len(chunks):.0f} Zeichen")
        
        return chunks

    def clean_problematic_characters(self, text: str) -> str:
        """
        Bereinigt Text von problematischen Zeichen, die später beim Excel-Export
        zu Fehlern führen könnten.
        
        Args:
            text (str): Zu bereinigender Text
            
        Returns:
            str: Bereinigter Text
        """
        if not text:
            return ""
            
        # Entferne problematische Steuerzeichen
        import re
        cleaned_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\uFFFE\uFFFF]', '', text)
        
        # Ersetze bekannte problematische Sonderzeichen
        problematic_chars = ['☺', '☻', '♥', '♦', '♣', '♠']
        for char in problematic_chars:
            cleaned_text = cleaned_text.replace(char, '')
        
        return cleaned_text

    def preprocess_text(self, text: str) -> str:
        """
        Bereinigt und normalisiert Textdaten.
        
        Args:
            text (str): Zu bereinigender Text
            
        Returns:
            str: Bereinigter Text
        """
        if not text:
            return ""
            
        # Entferne überflüssige Whitespaces
        text = ' '.join(text.split())
        
        # Ersetze verschiedene Anführungszeichen durch einheitliche
        text = text.replace('"', '"').replace('"', '"')
        
        # Ersetze verschiedene Bindestriche durch einheitliche
        text = text.replace('–', '-').replace('—', '-')
        
        # Entferne spezielle Steuerzeichen und problematische Zeichen
        text = self.clean_problematic_characters(text)
        
        return text

    def load_documents(self) -> dict:
        """
        Sammelt und lädt alle unterstützten Dateien aus dem Verzeichnis.
        
        Returns:
            dict: Dictionary mit Dateinamen als Schlüssel und Dokumenteninhalt als Wert
        """
        documents = {}
        supported_extensions = {'.txt', '.docx', '.doc', '.pdf'}
        
        try:
            # Prüfe ob Verzeichnis existiert
            if not os.path.exists(self.data_dir):
                print(f"Warnung: Verzeichnis {self.data_dir} existiert nicht")
                os.makedirs(self.data_dir)
                return documents

            # Durchsuche Verzeichnis nach Dateien
            for file in os.listdir(self.data_dir):
                file_path = os.path.join(self.data_dir, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                if not os.path.isfile(file_path):
                    continue
                    
                print(f"Gefunden: {file} ({file_ext})")
                
                try:
                    if file_ext == '.txt':
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                    elif file_ext == '.docx':
                        from docx import Document
                        doc = Document(file_path)
                        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    elif file_ext == '.doc':
                        # Benötigt antiword oder ähnliches Tool
                        import subprocess
                        text = subprocess.check_output(['antiword', file_path]).decode('utf-8')
                    elif file_ext == '.pdf':
                        import PyPDF2
                        with open(file_path, 'rb') as f:
                            pdf = PyPDF2.PdfReader(f)
                            text = '\n'.join([page.extract_text() for page in pdf.pages])
                    else:
                        print(f"Überspringe nicht unterstützte Datei: {file}")
                        continue
                        
                    # Bereinige und speichere Text
                    text = self.preprocess_text(text)
                    if text.strip():  # Nur nicht-leere Texte speichern
                        documents[file] = text
                        print(f"Erfolgreich geladen: {file} ({len(text)} Zeichen)")
                    else:
                        print(f"Warnung: Leerer Text in {file}")
                        
                except Exception as e:
                    print(f"Fehler beim Laden von {file}: {str(e)}")
                    continue
                    
            if not documents:
                print(f"\nKeine Dokumente gefunden in {self.data_dir}")
                print(f"Unterstützte Formate: {', '.join(supported_extensions)}")
            
        except Exception as e:
            print(f"Fehler beim Durchsuchen des Verzeichnisses: {str(e)}")
            
        return documents

# --- Klasse: DevelopmentHistory ---
# Aufgabe: Dokumentation der Kategorienentwicklung
class DevelopmentHistory:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.history_file = os.path.join(output_dir, "development_history.json")
        
        # Various aspects of history
        self.category_changes = []     # Category changes
        self.coding_history = []       # Coding results
        self.batch_history = []        # Batch processing
        self.saturation_checks = []    # Saturation checks
        self.validation_history = []   # System validations
        self.analysis_events = []      # Analysis events
        self.errors = []              # Error events
        self.category_development = [] # Category development tracking
        self.analysis_start = None    # Analysis start information
        
        # Performance tracking
        self.performance_metrics = {
            'batch_times': [],
            'coding_times': [],
            'processing_rates': []
        }

    def log_category_development(self, phase: str, **metrics) -> None:
        """
        Logs the progress of category development.
        
        Args:
            phase: Current phase of development
            **metrics: Additional metrics to log
        """
        entry = {
            'timestamp': datetime.now().isoformat(),
            'phase': phase,
            **metrics
        }
        self.category_development.append(entry)
        self._save_history()

    def log_batch_processing(self, batch_size: int, processing_time: float, results: dict) -> None:
        """Logs information about batch processing."""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'batch_size': batch_size,
            'processing_time': processing_time,
            'results': results
        }
        self.batch_history.append(entry)
        self._save_history()

    def log_saturation_check(self, material_percentage: float, result: str, metrics: Optional[dict]) -> None:
        """Logs the results of a saturation check."""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'material_percentage': material_percentage,
            'result': result,
            'metrics': metrics
        }
        self.saturation_checks.append(entry)
        self._save_history()

    def log_saturation_reached(self, material_percentage: float, metrics: dict) -> None:
        """Logs when saturation is reached."""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'event': 'saturation_reached',
            'material_percentage': material_percentage,
            'metrics': metrics
        }
        self.analysis_events.append(entry)
        self._save_history()

    def log_analysis_completion(self, final_categories: dict, total_time: float, total_codings: int) -> None:
        """Logs the completion of the analysis."""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'event': 'analysis_complete',
            'total_time': total_time,
            'total_categories': len(final_categories),
            'total_codings': total_codings
        }
        self.analysis_events.append(entry)
        self._save_history()

    def log_analysis_start(self, total_segments: int, total_categories: int) -> None:
        """
        Logs the start of analysis with initial metrics.
        
        Args:
            total_segments: Total number of text segments to analyze
            total_categories: Initial number of categories
        """
        self.analysis_start = {
            'timestamp': datetime.now().isoformat(),
            'total_segments': total_segments,
            'total_categories': total_categories,
            'start_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        self.analysis_events.append({
            'event': 'analysis_start',
            **self.analysis_start
        })
        self._save_history()

    def log_error(self, error_type: str, error_message: str) -> None:
        """Logs an error event."""
        entry = {
            'timestamp': datetime.now().isoformat(),
            'event': 'error',
            'type': error_type,
            'message': error_message
        }
        self.errors.append(entry)
        self._save_history()

    def _save_history(self) -> None:
        """Saves the current history to the JSON file."""
        try:
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump({
                    'category_changes': self.category_changes,
                    'coding_history': self.coding_history,
                    'batch_history': self.batch_history,
                    'saturation_checks': self.saturation_checks,
                    'validation_history': self.validation_history,
                    'analysis_events': self.analysis_events,
                    'errors': self.errors,
                    'category_development': self.category_development,
                    'performance_metrics': self.performance_metrics,
                'analysis_start': self.analysis_start
                }, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving history: {str(e)}")

    def generate_development_report(self) -> str:
        """
        Generates a comprehensive report of the development process.
        
        Returns:
            str: Formatted report
        """
        report = ["=== Category Development Report ===\n"]
        
        # Phase Analysis
        phases = {}
        for entry in self.category_development:
            phase = entry['phase']
            if phase not in phases:
                phases[phase] = []
            phases[phase].append(entry)
        
        for phase, entries in phases.items():
            report.append(f"\nPhase: {phase}")
            report.append("-" * (len(phase) + 7))
            
            # Summarize metrics for this phase
            metrics_summary = {}
            for entry in entries:
                for key, value in entry.items():
                    if key not in ['timestamp', 'phase']:
                        if key not in metrics_summary:
                            metrics_summary[key] = []
                        metrics_summary[key].append(value)
            
            # Report metrics
            for key, values in metrics_summary.items():
                if all(isinstance(v, (int, float)) for v in values):
                    avg = sum(values) / len(values)
                    report.append(f"{key}: avg = {avg:.2f}")
                else:
                    report.append(f"{key}: {len(values)} entries")
        
        # Error Analysis
        if self.errors:
            report.append("\nErrors:")
            report.append("-------")
            for error in self.errors:
                report.append(f"- {error['type']}: {error['message']}")
        
        return "\n".join(report)



# --- Klasse: CategoryManager ---
class CategoryManager:
    """Verwaltet das Speichern und Laden von Kategorien"""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.categories_file = os.path.join(output_dir, 'category_system.json')
        
    def save_categories(self, categories: Dict[str, CategoryDefinition]) -> None:
        """Speichert das aktuelle Kategoriensystem als JSON"""
        try:
            # Konvertiere CategoryDefinition Objekte in serialisierbares Format
            categories_dict = {
                name: {
                    'definition': cat.definition,
                    'examples': cat.examples,
                    'rules': cat.rules,
                    'subcategories': cat.subcategories,
                    'timestamp': datetime.now().isoformat()
                } for name, cat in categories.items()
            }
            
            # Metadata hinzufügen
            output_data = {
                'timestamp': datetime.now().isoformat(),
                'version': '1.0',
                'categories': categories_dict
            }
            
            with open(self.categories_file, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, indent=2, ensure_ascii=False)
                
            print(f"\nKategoriensystem gespeichert in: {self.categories_file}")
            
        except Exception as e:
            print(f"Fehler beim Speichern des Kategoriensystems: {str(e)}")

    def load_categories(self) -> Optional[Dict[str, CategoryDefinition]]:
        """
        Lädt gespeichertes Kategoriensystem falls vorhanden.
        
        Returns:
            Optional[Dict[str, CategoryDefinition]]: Dictionary mit Kategorienamen und deren Definitionen,
            oder None wenn kein gespeichertes System verwendet werden soll
        """
        try:
            if os.path.exists(self.categories_file):
                with open(self.categories_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Zeige Informationen zum gefundenen Kategoriensystem
                timestamp = datetime.fromisoformat(data['timestamp'])
                print(f"\nGefundenes Kategoriensystem vom {timestamp.strftime('%d.%m.%Y %H:%M')}")
                print(f"Enthält {len(data['categories'])} Kategorien:")
                for cat_name in data['categories'].keys():
                    print(f"- {cat_name}")
                
                # Frage Benutzer
                while True:
                    answer = input("\nMöchten Sie dieses Kategoriensystem verwenden? (j/n): ").lower()
                    if answer in ['j', 'n']:
                        break
                    print("Bitte antworten Sie mit 'j' oder 'n'")
                
                if answer == 'j':
                    # Konvertiere zurück zu CategoryDefinition Objekten
                    categories = {}
                    for name, cat_data in data['categories'].items():
                        # Stelle sicher, dass die Zeitstempel existieren
                        if 'timestamp' in cat_data:
                            added_date = modified_date = cat_data['timestamp'].split('T')[0]
                        else:
                            # Verwende aktuelle Zeit für fehlende Zeitstempel
                            current_date = datetime.now().strftime("%Y-%m-%d")
                            added_date = modified_date = current_date

                        categories[name] = CategoryDefinition(
                            name=name,
                            definition=cat_data['definition'],
                            examples=cat_data.get('examples', []),
                            rules=cat_data.get('rules', []),
                            subcategories=cat_data.get('subcategories', {}),
                            added_date=cat_data.get('added_date', added_date),
                            modified_date=cat_data.get('modified_date', modified_date)
                        )
                    
                    print(f"\nKategoriensystem erfolgreich geladen.")
                    return categories
                
            return None
                
        except Exception as e:
            print(f"Fehler beim Laden des Kategoriensystems: {str(e)}")
            print("Folgende Details könnten hilfreich sein:")
            import traceback
            traceback.print_exc()
            return None
    
    def save_codebook(self, categories: Dict[str, CategoryDefinition], filename: str = "codebook_inductive.json") -> None:
        """Speichert das vollständige Codebook inkl. deduktiver, induktiver und grounded Kategorien"""
        try:
            codebook_data = {
                "metadata": {
                    "created": datetime.now().isoformat(),
                    "version": "2.0",
                    "total_categories": len(categories),
                    "research_question": FORSCHUNGSFRAGE,
                    "analysis_mode": CONFIG.get('ANALYSIS_MODE', 'deductive')  # Speichere den Analysemodus
                },
                "categories": {}
            }
            
            for name, category in categories.items():
                # Bestimme den Kategorietyp je nach Analysemodus
                if name in DEDUKTIVE_KATEGORIEN:
                    development_type = "deductive"
                elif CONFIG.get('ANALYSIS_MODE') == 'grounded':
                    development_type = "grounded"  # Neue Markierung für grounded Kategorien
                else:
                    development_type = "inductive"
                    
                codebook_data["categories"][name] = {
                    "definition": category.definition,
                    # Wandle examples in eine Liste um, falls es ein Set ist
                    "examples": list(category.examples) if isinstance(category.examples, set) else category.examples,
                    # Wandle rules in eine Liste um, falls es ein Set ist
                    "rules": list(category.rules) if isinstance(category.rules, set) else category.rules,
                    # Wandle subcategories in ein Dictionary um, falls nötig
                    "subcategories": dict(category.subcategories) if isinstance(category.subcategories, set) else category.subcategories,
                    "development_type": development_type,
                    "added_date": category.added_date,
                    "last_modified": category.modified_date
                }
            
            output_path = os.path.join(self.output_dir, filename)
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(codebook_data, f, indent=2, ensure_ascii=False)
                
            print(f"\nCodebook gespeichert unter: {output_path}")
            print(f"- Deduktive Kategorien: {sum(1 for c in codebook_data['categories'].values() if c['development_type'] == 'deductive')}")
            print(f"- Induktive Kategorien: {sum(1 for c in codebook_data['categories'].values() if c['development_type'] == 'inductive')}")
            print(f"- Grounded Kategorien: {sum(1 for c in codebook_data['categories'].values() if c['development_type'] == 'grounded')}")
            
        except Exception as e:
            print(f"Fehler beim Speichern des Codebooks: {str(e)}")
            # Zusätzliche Fehlerdiagnose
            import traceback
            traceback.print_exc()


# --- Klasse: RelevanceChecker ---
# Aufgabe: Zentrale Klasse für Relevanzprüfungen mit Caching und Batch-Verarbeitung
class RelevanceChecker:
    """
    Zentrale Klasse für Relevanzprüfungen mit Caching und Batch-Verarbeitung.
    Reduziert API-Calls durch Zusammenfassung mehrerer Segmente.
    """
    
    def __init__(self, model_name: str, batch_size: int = 5):
        self.model_name = model_name
        self.batch_size = batch_size

        # Hole Provider aus CONFIG
        provider_name = CONFIG.get('MODEL_PROVIDER', 'openai').lower()  # Fallback zu OpenAI
        try:
            # Wir lassen den LLMProvider sich selbst um die Client-Initialisierung kümmern
            self.llm_provider = LLMProviderFactory.create_provider(provider_name)
        except Exception as e:
            print(f"Fehler bei Provider-Initialisierung: {str(e)}")
            raise
        
        # Cache für Relevanzprüfungen
        self.relevance_cache = {}
        self.relevance_details = {}
        
        # Tracking-Metriken
        self.total_segments = 0
        self.relevant_segments = 0
        self.api_calls = 0

        # Hole Ausschlussregeln aus KODIERREGELN
        self.exclusion_rules = KODIERREGELN.get('exclusion', [])
        print("\nRelevanceChecker initialisiert:")
        print(f"- {len(self.exclusion_rules)} Ausschlussregeln geladen")

        # Hole Mehrfachkodierungsparameter aus CONFIG
        self.multiple_codings_enabled = CONFIG.get('MULTIPLE_CODINGS', True)
        self.multiple_threshold = CONFIG.get('MULTIPLE_CODING_THRESHOLD', 0.7)

        # Cache für Mehrfachkodierungen
        self.multiple_coding_cache = {}
        
        print("\nRelevanceChecker initialisiert:")
        print(f"- {len(self.exclusion_rules)} Ausschlussregeln geladen")
        print(f"- Mehrfachkodierung: {'Aktiviert' if self.multiple_codings_enabled else 'Deaktiviert'}")
        if self.multiple_codings_enabled:
            print(f"- Mehrfachkodierung-Schwellenwert: {self.multiple_threshold}")
        
        # Prompt-Handler
        self.prompt_handler = QCAPrompts(
            forschungsfrage=FORSCHUNGSFRAGE,
            kodierregeln=KODIERREGELN,
            deduktive_kategorien=DEDUKTIVE_KATEGORIEN
        )

    # FIX: Hilfsmethode für Segment-Formatierung hinzufügen
    def _format_segments_for_batch(self, segments: List[Tuple[str, str]]) -> str:
        """
        Formatiert Segmente für Batch-Verarbeitung in der Relevanzprüfung
        
        Args:
            segments: Liste von (segment_id, text) Tupeln
            
        Returns:
            str: Formatierte Segmente für den Prompt
        """
        formatted_segments = []
        for i, (segment_id, text) in enumerate(segments, 1):
            # Begrenze Textlänge für bessere Performance
            truncated_text = text[:800] + "..." if len(text) > 800 else text
            # Entferne problematische Zeichen
            clean_text = truncated_text.replace('\n', ' ').replace('\r', ' ').strip()
            formatted_segments.append(f"SEGMENT {i}:\n{clean_text}\n")
        
        return "\n".join(formatted_segments)
    
    async def check_multiple_category_relevance(self, segments: List[Tuple[str, str]], 
                                            categories: Dict[str, CategoryDefinition]) -> Dict[str, List[Dict]]:
        """
        PARALLELISIERTE VERSION: Prüft ob Segmente für mehrere Hauptkategorien relevant sind.
        """
        if not self.multiple_codings_enabled:
            return {}
            
        # Filtere bereits gecachte Segmente
        uncached_segments = [
            (sid, text) for sid, text in segments 
            if sid not in self.multiple_coding_cache
        ]
        
        if not uncached_segments:
            return {sid: self.multiple_coding_cache[sid] for sid, _ in segments}

        try:
            print(f"🚀 Parallele Mehrfachkodierungs-Prüfung: {len(uncached_segments)} Segmente")
            
            # Bereite Kategorien-Kontext vor
            category_descriptions = []
            for cat_name, cat_def in categories.items():
                if cat_name not in ["Nicht kodiert", "Kein Kodierkonsens"]:
                    category_descriptions.append({
                        'name': cat_name,
                        'definition': cat_def.definition[:200] + '...' if len(cat_def.definition) > 200 else cat_def.definition,
                        'examples': cat_def.examples[:2] if cat_def.examples else []
                    })

            # 🚀 PARALLEL: Hilfsfunktion für einzelnes Segment
            async def check_single_segment_multiple(segment_id: str, text: str) -> Tuple[str, List[Dict]]:
                """Prüft ein einzelnes Segment auf Mehrfachkodierung parallel."""
                try:
                    prompt = self.prompt_handler.get_multiple_category_relevance_prompt(
                        segments_text=f"SEGMENT:\n{text}",
                        category_descriptions=category_descriptions,
                        multiple_threshold=self.multiple_threshold
                    )
                    
                    token_counter.start_request()
                    
                    response = await self.llm_provider.create_completion(
                        model=self.model_name,
                        messages=[
                            {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                            {"role": "user", "content": prompt}
                        ],
                        temperature=0.3,
                        response_format={"type": "json_object"}
                    )
                    
                    llm_response = LLMResponse(response)
                    result = json.loads(llm_response.content)
                    
                    
                    token_counter.track_response(response, self.model_name)
                    
                    # Verarbeite Ergebnisse - erwarte single segment format
                    segment_result = result.get('segment_results', [{}])[0] if result.get('segment_results') else result
                    
                    # Filtere Kategorien nach Schwellenwert
                    relevant_categories = []
                    for cat in segment_result.get('relevant_categories', []):
                        if cat.get('relevance_score', 0) >= self.multiple_threshold:
                            relevant_categories.append(cat)
                    
                    return segment_id, relevant_categories
                    
                except Exception as e:
                    print(f"⚠️ Fehler bei Mehrfachkodierungs-Prüfung {segment_id}: {str(e)}")
                    return segment_id, []  # Leer = keine Mehrfachkodierung
            
            # 🚀 Erstelle Tasks für alle Segmente
            tasks = [
                check_single_segment_multiple(segment_id, text) 
                for segment_id, text in uncached_segments
            ]
            
            # 🚀 Führe alle parallel aus
            start_time = time.time()
            results = await asyncio.gather(*tasks, return_exceptions=True)
            processing_time = time.time() - start_time
            
            # Verarbeite Ergebnisse und aktualisiere Cache
            multiple_coding_results = {}
            successful_checks = 0
            segments_with_multiple = 0
            
            for result in results:
                if isinstance(result, Exception):
                    print(f"⚠️ Ausnahme bei Mehrfachkodierungs-Prüfung: {result}")
                    continue
                
                segment_id, relevant_categories = result
                
                # Cache aktualisieren
                self.multiple_coding_cache[segment_id] = relevant_categories
                multiple_coding_results[segment_id] = relevant_categories
                
                successful_checks += 1
                
                # Debug-Ausgabe für Mehrfachkodierung
                if len(relevant_categories) > 1:
                    segments_with_multiple += 1
                    print(f"  🔄 Mehrfachkodierung: {segment_id}")
                    for cat in relevant_categories:
                        print(f"    - {cat['category']}: {cat['relevance_score']:.2f}")
            
            print(f"⚡ {successful_checks} Mehrfachkodierungs-Prüfungen in {processing_time:.2f}s")
            print(f"🔄 {segments_with_multiple} Segmente mit Mehrfachkodierung identifiziert")
            
            # Kombiniere mit Cache für alle ursprünglichen Segmente
            final_results = {}
            for segment_id, text in segments:
                if segment_id in multiple_coding_results:
                    final_results[segment_id] = multiple_coding_results[segment_id]
                else:
                    final_results[segment_id] = self.multiple_coding_cache.get(segment_id, [])
            
            return final_results

        except Exception as e:
            print(f"Fehler bei paralleler Mehrfachkodierungs-Prüfung: {str(e)}")
            import traceback
            traceback.print_exc()
            return {}
    
    
    async def check_relevance_batch(self, segments: List[Tuple[str, str]]) -> Dict[str, bool]:
        """
        KORRIGIERTE VERSION: Prüft die Relevanz mehrerer Segmente parallel mit Batch-Verarbeitung.
        Behält die bewährte Batch-Logik bei und fügt Parallelisierung nur für große Batches hinzu.
        
        Args:
            segments: Liste von (segment_id, text) Tupeln
            
        Returns:
            Dict[str, bool]: Mapping von segment_id zu Relevanz
        """
        # Filtere bereits gecachte Segmente
        uncached_segments = [
            (sid, text) for sid, text in segments 
            if sid not in self.relevance_cache
        ]
        
        if not uncached_segments:
            return {sid: self.relevance_cache[sid] for sid, _ in segments}

        try:
            print(f"🔍 Relevanzprüfung: {len(uncached_segments)} neue Segmente")
            
            # STRATEGIE: Kleine Batches (≤5) → bewährte Batch-Methode
            #           Große Batches (>5) → Parallelisierung in Sub-Batches
            
            if len(uncached_segments) <= 5:
                # BEWÄHRTE BATCH-METHODE für kleine Gruppen
                print(f"   📦 Verwende normale Batch-Methode für {len(uncached_segments)} Segmente")
                
                # Erstelle formatierten Text für Batch-Verarbeitung
                segments_text = "\n\n=== SEGMENT BREAK ===\n\n".join(
                    f"SEGMENT {i + 1}:\n{text}" 
                    for i, (_, text) in enumerate(uncached_segments)
                )

                prompt = self.prompt_handler.get_relevance_check_prompt(
                    segments_text=segments_text,
                    exclusion_rules=self.exclusion_rules
                )
                
                token_counter.start_request()
                
                # Ein API-Call für alle Segmente
                response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    response_format={"type": "json_object"}
                )
                
                llm_response = LLMResponse(response)
                results = json.loads(llm_response.content)
                
                
                token_counter.track_response(response, self.model_name)
                
                # Verarbeite Ergebnisse
                relevance_results = {}
                segment_results = results.get('segment_results', [])
                
                if len(segment_results) != len(uncached_segments):
                    print(f"⚠️ Warnung: Anzahl Ergebnisse ({len(segment_results)}) != Anzahl Segmente ({len(uncached_segments)})")
                    # Fallback: Alle als relevant markieren
                    for segment_id, _ in uncached_segments:
                        relevance_results[segment_id] = True
                        self.relevance_cache[segment_id] = True
                else:
                    for i, (segment_id, _) in enumerate(uncached_segments):
                        segment_result = segment_results[i]
                        is_relevant = segment_result.get('is_relevant', True)  # Default: relevant
                        
                        # Cache-Aktualisierung
                        self.relevance_cache[segment_id] = is_relevant
                        # FIX: Erweiterte Details-Speicherung mit Begründung
                        self.relevance_details[segment_id] = {
                            'confidence': segment_result.get('confidence', 0.8),
                            'key_aspects': segment_result.get('key_aspects', []),
                            'justification': segment_result.get('justification', ''),
                            'reasoning': segment_result.get('reasoning', 'Keine Begründung verfügbar'),
                            'main_themes': segment_result.get('main_themes', []),
                            'exclusion_match': segment_result.get('exclusion_match', False)
                        }
                        # FIX: Ende
                        
                        relevance_results[segment_id] = is_relevant
                        
                        # Tracking-Aktualisierung
                        self.total_segments += 1
                        if is_relevant:
                            self.relevant_segments += 1
                
                self.api_calls += 1
                
            else:
                # PARALLELISIERUNG für große Batches
                print(f"   🚀 Verwende Parallelisierung in Sub-Batches für {len(uncached_segments)} Segmente")
                
                # Teile in Sub-Batches von je 3-4 Segmenten
                sub_batch_size = 3
                sub_batches = [
                    uncached_segments[i:i + sub_batch_size] 
                    for i in range(0, len(uncached_segments), sub_batch_size)
                ]
                
                async def process_sub_batch(sub_batch: List[Tuple[str, str]]) -> Dict[str, bool]:
                    """Verarbeitet einen Sub-Batch mit der bewährten Methode."""
                    try:
                        # Verwende dieselbe Batch-Logik wie oben
                        segments_text = "\n\n=== SEGMENT BREAK ===\n\n".join(
                            f"SEGMENT {i + 1}:\n{text}" 
                            for i, (_, text) in enumerate(sub_batch)
                        )

                        prompt = self.prompt_handler.get_relevance_check_prompt(
                            segments_text=segments_text,
                            exclusion_rules=self.exclusion_rules
                        )
                        
                        token_counter.start_request()
                        
                        response = await self.llm_provider.create_completion(
                            model=self.model_name,
                            messages=[
                                {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                                {"role": "user", "content": prompt}
                            ],
                            temperature=0.3,
                            response_format={"type": "json_object"}
                        )
                        
                        llm_response = LLMResponse(response)
                        results = json.loads(llm_response.content)
                        
                        
                        token_counter.track_response(response, self.model_name)
                        
                        # Verarbeite Sub-Batch Ergebnisse
                        sub_batch_results = {}
                        segment_results = results.get('segment_results', [])
                        
                        if len(segment_results) != len(sub_batch):
                            # Fallback bei Mismatch
                            for segment_id, _ in sub_batch:
                                sub_batch_results[segment_id] = True
                        else:
                            for i, (segment_id, _) in enumerate(sub_batch):
                                segment_result = segment_results[i]
                                is_relevant = segment_result.get('is_relevant', True)
                                sub_batch_results[segment_id] = is_relevant
                                
                                # FIX: Erweiterte Details-Speicherung mit Begründung
                                self.relevance_details[segment_id] = {
                                    'confidence': segment_result.get('confidence', 0.8),
                                    'key_aspects': segment_result.get('key_aspects', []),
                                    'justification': segment_result.get('justification', ''),
                                    'reasoning': segment_result.get('reasoning', 'Keine Begründung verfügbar'),
                                    'main_themes': segment_result.get('main_themes', []),
                                    'exclusion_match': segment_result.get('exclusion_match', False)
                                }
                                # FIX: Ende
                        
                        return sub_batch_results
                        
                    except Exception as e:
                        print(f"⚠️ Fehler in Sub-Batch: {str(e)}")
                        # Fallback: Alle als relevant markieren
                        return {segment_id: True for segment_id, _ in sub_batch}
                
                # Führe alle Sub-Batches parallel aus
                start_time = time.time()
                tasks = [process_sub_batch(sub_batch) for sub_batch in sub_batches]
                sub_batch_results = await asyncio.gather(*tasks, return_exceptions=True)
                processing_time = time.time() - start_time
                
                # Kombiniere Ergebnisse
                relevance_results = {}
                successful_batches = 0
                
                for result in sub_batch_results:
                    if isinstance(result, Exception):
                        print(f"⚠️ Sub-Batch Ausnahme: {result}")
                        continue
                    
                    if isinstance(result, dict):
                        relevance_results.update(result)
                        successful_batches += 1
                
                # Aktualisiere Cache und Tracking
                for segment_id, is_relevant in relevance_results.items():
                    self.relevance_cache[segment_id] = is_relevant
                    self.total_segments += 1
                    if is_relevant:
                        self.relevant_segments += 1
                
                # API-Calls = Anzahl der Sub-Batches
                self.api_calls += successful_batches
                
                print(f"   ⚡ {successful_batches} Sub-Batches in {processing_time:.2f}s verarbeitet")
            
            # Debug-Ausgabe der Ergebnisse
            relevant_count = sum(1 for is_relevant in relevance_results.values() if is_relevant)
            print(f"   📊 Relevanz-Ergebnisse: {relevant_count}/{len(relevance_results)} als relevant eingestuft")
            
            # Kombiniere mit bereits gecachten Ergebnissen für finale Antwort
            final_results = {}
            for segment_id, text in segments:
                if segment_id in relevance_results:
                    final_results[segment_id] = relevance_results[segment_id]
                else:
                    final_results[segment_id] = self.relevance_cache.get(segment_id, True)
            
            return final_results

        except Exception as e:
            print(f"Fehler bei paralleler Relevanzprüfung: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback: Alle als relevant markieren bei Fehler
            return {sid: True for sid, _ in segments}
    

    async def check_relevance_with_category_preselection(self, segments: List[Tuple[str, str]], 
                                                        categories: Dict[str, CategoryDefinition] = None,
                                                        mode: str = 'deductive') -> Dict[str, Dict]:
        """
        Erweiterte Relevanzprüfung mit Hauptkategorie-Vorauswahl (nur für deduktiven Modus)
        
        Returns:
            Dict mit segment_id als Key und Dict mit 'is_relevant', 'preferred_categories', 'relevance_scores'
        """
        if mode != 'deductive' or not categories:
            # Fallback auf Standard-Relevanzprüfung für andere Modi
            standard_results = await self.check_relevance_batch(segments)
            return {
                sid: {'is_relevant': is_relevant, 'preferred_categories': [], 'relevance_scores': {}}
                for sid, is_relevant in standard_results.items()
            }
        
        print(f"🎯 Erweiterte Relevanzprüfung mit Kategorie-Vorauswahl für {len(segments)} Segmente...")
        
        # Cache-Key für erweiterte Relevanzprüfung
        cache_key_base = "extended_relevance"
        
        results = {}
        segments_to_process = []
        
        # Cache-Prüfung
        for segment_id, text in segments:
            cache_key = f"{cache_key_base}_{hash(text)}"
            if cache_key in self.relevance_cache:
                results[segment_id] = self.relevance_cache[cache_key]
            else:
                segments_to_process.append((segment_id, text))
        
        if not segments_to_process:
            return results
        
        # Batch-Verarbeitung für nicht-gecachte Segmente
        batch_text = self._format_segments_for_batch(segments_to_process)
        
        prompt = self.prompt_handler.get_relevance_with_category_preselection_prompt(
            batch_text, categories
        )
        
        try:
            self.api_calls += 1
            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            
            llm_response = LLMResponse(response)
            batch_results = json.loads(llm_response.content)
            
            # Verarbeite Batch-Ergebnisse
            segment_results = batch_results.get('segment_results', [])
            
            for i, (segment_id, text) in enumerate(segments_to_process):
                if i < len(segment_results):
                    segment_result = segment_results[i]
                    
                    result = {
                        'is_relevant': segment_result.get('is_relevant', False),
                        'preferred_categories': segment_result.get('preferred_categories', []),
                        'relevance_scores': segment_result.get('relevance_scores', {}),
                        'reasoning': segment_result.get('reasoning', '')
                    }
                    
                    # Cache speichern
                    cache_key = f"{cache_key_base}_{hash(text)}"
                    self.relevance_cache[cache_key] = result
                    
                    results[segment_id] = result
                else:
                    # Fallback bei unvollständigen Ergebnissen
                    results[segment_id] = {
                        'is_relevant': False,
                        'preferred_categories': [],
                        'relevance_scores': {},
                        'reasoning': 'Unvollständiges Batch-Ergebnis'
                    }
            
            print(f"✅ Erweiterte Relevanzprüfung abgeschlossen: {len([r for r in results.values() if r['is_relevant']])} relevante Segmente")
            
        except Exception as e:
            print(f"Fehler bei erweiterter Relevanzprüfung: {str(e)}")
            # Fallback auf Standard-Relevanzprüfung
            standard_results = await self.check_relevance_batch(segments_to_process)
            for segment_id, is_relevant in standard_results.items():
                results[segment_id] = {
                    'is_relevant': is_relevant,
                    'preferred_categories': [],
                    'relevance_scores': {},
                    'reasoning': 'Fallback nach Fehler'
                }
        
        return results
    
    def get_relevance_details(self, segment_id: str) -> Optional[Dict]:
        """Gibt detaillierte Relevanzinformationen für ein Segment zurück."""
        return self.relevance_details.get(segment_id)

    def get_statistics(self) -> Dict:
        """Gibt Statistiken zur Relevanzprüfung zurück."""
        return {
            'total_segments': self.total_segments,
            'relevant_segments': self.relevant_segments,
            'relevance_rate': self.relevant_segments / max(1, self.total_segments),
            'api_calls': self.api_calls,
            'cache_size': len(self.relevance_cache)
        }

    
# --- Klasse: IntegratedAnalysisManager ---
# Aufgabe: Integriert die verschiedenen Analysephasen in einem zusammenhängenden Prozess

class IntegratedAnalysisManager:

    def __init__(self, config: Dict):
        # Bestehende Initialisierung
        self.config = config
        self.history = DevelopmentHistory(config['OUTPUT_DIR'])

        # Batch Size aus Config
        self.batch_size = config.get('BATCH_SIZE', 5) 

        # Prompt-Handler initialisieren
        self.prompt_handler = QCAPrompts(
            forschungsfrage=FORSCHUNGSFRAGE,
            kodierregeln=KODIERREGELN,
            deduktive_kategorien=DEDUKTIVE_KATEGORIEN

        )
        
        # Zentrale Relevanzprüfung
        self.relevance_checker = RelevanceChecker(
            model_name=config['MODEL_NAME'],
            batch_size=self.batch_size
        )
        
        # KORREKTUR: Initialisiere den verbesserten InductiveCoder
        self.inductive_coder = InductiveCoder(
            model_name=config['MODEL_NAME'],
            history=self.history,
            output_dir=config['OUTPUT_DIR'],
            config=config  # Übergebe config für verbesserte Initialisierung
        )

        self.deductive_coders = [
            DeductiveCoder(
                config['MODEL_NAME'], 
                coder_config['temperature'],
                coder_config['coder_id']
            )
            for coder_config in config['CODER_SETTINGS']
        ]
        
        # Tracking-Variablen (unverändert)
        self.processed_segments = set()
        self.coding_results = []
        self.analysis_log = [] 
        self.performance_metrics = {
            'batch_processing_times': [],
            'coding_times': [],
            'category_changes': []
        }

        # Konfigurationsparameter (unverändert)
        self.use_context = config.get('CODE_WITH_CONTEXT', True)
        print(f"\nKontextuelle Kodierung: {'Aktiviert' if self.use_context else 'Deaktiviert'}")

        # Dictionary für die Verwaltung der Document-Summaries (unverändert)
        self.document_summaries = {}

        # NEU: Grounded Mode Spezifische Variablen
        self.grounded_subcodes_collection = []  # Zentrale Sammlung aller Subcodes
        self.grounded_keywords_collection = []  # Zentrale Sammlung aller Keywords
        self.grounded_segment_analyses = []     # Zentrale Sammlung aller Segment-Analysen
        self.grounded_batch_history = []        # Historie der Batch-Ergebnisse
        self.grounded_saturation_counter = 0    # Zähler für Batches ohne neue Subcodes


        # NEU: Escape-Handler hinzufügen (unverändert)
        self.escape_handler = EscapeHandler(self)
        self._should_abort = False
        self._escape_abort_requested = False

        self.current_categories = {}  # FIX: Hinzufügen

        print(f"\n🔬 IntegratedAnalysisManager initialisiert:")
        print(f"   - Analysemodus: {config.get('ANALYSIS_MODE', 'inductive')}")
        if config.get('ANALYSIS_MODE') == 'grounded':
            print(f"   - Grounded Mode: Subcode-Sammlung aktiviert")
            print(f"   - Hauptkategorien werden erst am Ende generiert")

    async def _get_next_batch(self, 
                           segments: List[Tuple[str, str]], 
                           batch_size: float) -> List[Tuple[str, str]]:
        """
        Bestimmt den nächsten zu analysierenden Batch.
        
        Args:
            segments: Liste aller Segmente
            batch_size_percentage: Batch-Größe als Prozentsatz
            
        Returns:
            List[Tuple[str, str]]: Nächster Batch von Segmenten
        """
        remaining_segments = [
            seg for seg in segments 
            if seg[0] not in self.processed_segments
        ]
        
        if not remaining_segments:
            return []
            
        batch_size = max(1, batch_size)
        return remaining_segments[:batch_size]
    
    
    async def _process_batch_inductively(self, 
                                    batch: List[Tuple[str, str]], 
                                    current_categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        VEREINFACHT: Keine weitere Relevanzprüfung mehr nötig
        """
        # Die Segmente sind bereits in analyze_material gefiltert worden
        relevant_segments = [text for _, text in batch]  # Einfach die Texte extrahieren
        
        if not relevant_segments:
            print("   ℹ️ Keine Segmente in diesem Batch")
            return {}

        print(f"\n🔍 Entwickle Kategorien aus {len(relevant_segments)} Segmenten")
        
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'inductive')
        
        if analysis_mode == 'inductive':
            return await self._process_inductive_mode(relevant_segments, current_categories)
        elif analysis_mode == 'abductive':
            return await self._process_abductive_mode(relevant_segments, current_categories)
        elif analysis_mode == 'grounded':
            return await self._process_grounded_mode(relevant_segments, current_categories)
        else:
            return {}

    async def _process_inductive_mode(self, relevant_segments: List[str], 
                                    current_categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        INDUCTIVE MODE: Vollständige induktive Kategorienentwicklung (ehemals full mode)
        """
        print("🔄 INDUCTIVE MODE: Vollständige induktive Kategorienentwicklung")
        print("   - Entwickle eigenständiges induktives Kategoriensystem")
        print("   - Deduktive Kategorien werden ignoriert")
        
        # KORRIGIERT: Übergebe bestehende induktive Kategorien als Basis
        new_categories = await self.inductive_coder.develop_category_system(
            relevant_segments,
            current_categories  # ✅ Bestehende induktive als Basis!
        )
        
        print(f"✅ INDUCTIVE MODE: {len(new_categories)} Kategorien entwickelt")
        if current_categories:
            print(f"   (zusätzlich zu {len(current_categories)} bereits bestehenden)")
        return new_categories

    async def _process_abductive_mode(self, relevant_segments: List[str], 
                                    current_categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        ABDUCTIVE MODE: Nur Subkategorien zu bestehenden Hauptkategorien
        """
        print("🔄 ABDUCTIVE MODE: Erweitere bestehende Kategorien um Subkategorien")
        
        if not current_categories:
            print("⚠️ ABDUCTIVE MODE: Keine bestehenden Kategorien zum Erweitern")
            return {}
        
        # Spezielle abduktive Analyse
        extended_categories = await self._analyze_for_subcategories(
            relevant_segments, 
            current_categories
        )
        
        print(f"✅ ABDUCTIVE MODE: {len(extended_categories)} Kategorien erweitert")
        return extended_categories

    
    async def _process_grounded_mode(self, relevant_segments: List[str], 
                                current_categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        KORRIGIERT: Diese Methode wird in normalen Batches NICHT aufgerufen im Grounded Mode
        """
        print("⚠️ WARNUNG: _process_grounded_mode sollte nicht in separatem Grounded Mode aufgerufen werden!")
        return {}
    
    async def _assess_grounded_saturation(self, batch_count: int, total_batches: int) -> bool:
        """
        KORRIGIERTE Sättigungslogik für Grounded Mode.
        """
        try:
            # Berechne Material-Fortschritt
            if hasattr(self, 'chunks') and self.chunks:
                total_segments = sum(len(chunk_list) for chunk_list in self.chunks.values())
            elif hasattr(self, '_total_segments'):
                total_segments = self._total_segments
            else:
                total_segments = max(len(self.processed_segments) * 2, 20)
            
            material_percentage = (len(self.processed_segments) / total_segments) * 100 if total_segments > 0 else 0.0
            
            # KORRIGIERT: Verwende die richtige Sammlung
            if not hasattr(self, 'grounded_subcodes_collection'):
                self.grounded_subcodes_collection = []
            
            # Analyse der Subcode-Entwicklung
            subcode_diversity = len(set(sc['name'] for sc in self.grounded_subcodes_collection))
            
            if not hasattr(self, 'grounded_keywords_collection'):
                self.grounded_keywords_collection = []
            keyword_diversity = len(set(self.grounded_keywords_collection))
            
            # Berechne Sättigungsmetriken
            avg_subcodes_per_batch = len(self.grounded_subcodes_collection) / max(batch_count, 1)
            
            # Kriterien für Grounded Mode Sättigung
            criteria = {
                'min_batches': batch_count >= 3,  # Mindestens 3 Batches
                'material_coverage': material_percentage >= 70,  # 70% Material verarbeitet
                'subcodes_collected': len(self.grounded_subcodes_collection) >= 8,  # Min. 8 Subcodes
                'saturation_stability': self.grounded_saturation_counter >= 2,  # 2 Batches ohne neue
                'diversity_threshold': subcode_diversity >= 5,  # Mindestens 5 verschiedene Subcodes
                'keyword_richness': keyword_diversity >= 15,  # Mindestens 15 verschiedene Keywords
            }
            
            print(f"\n🔍 Grounded Mode Sättigungsprüfung (Batch {batch_count}/{total_batches}):")
            print(f"📊 Aktuelle Metriken:")
            print(f"   - Material-Fortschritt: {material_percentage:.1f}%")
            print(f"   - Gesammelte Subcodes: {len(self.grounded_subcodes_collection)}")
            print(f"   - Subcode-Diversität: {subcode_diversity}")
            print(f"   - Keyword-Diversität: {keyword_diversity}")
            print(f"   - Sättigungs-Counter: {self.grounded_saturation_counter}")
            print(f"   - Ø Subcodes/Batch: {avg_subcodes_per_batch:.1f}")
            
            print(f"\n🎯 Sättigungskriterien:")
            for criterion, met in criteria.items():
                status = "✅" if met else "❌"
                print(f"   {status} {criterion}: {met}")
            
            # Bestimme Sättigungsstatus
            critical_criteria = ['min_batches', 'subcodes_collected', 'saturation_stability']
            critical_met = all(criteria[crit] for crit in critical_criteria)
            
            # Vollständige Sättigung: Alle Kriterien oder kritische + Material fast vollständig
            full_saturation = all(criteria.values())
            partial_saturation = critical_met and (material_percentage >= 85 or criteria['material_coverage'])
            forced_saturation = material_percentage >= 100  # 100% Material = Zwangssättigung
            
            is_saturated = full_saturation or partial_saturation or forced_saturation
            
            if is_saturated:
                saturation_type = "Vollständig" if full_saturation else ("Partiell" if partial_saturation else "Material-bedingt")
                print(f"\n🎯 GROUNDED MODE SÄTTIGUNG erreicht ({saturation_type}):")
                print(f"   - Material: {material_percentage:.1f}% verarbeitet")
                print(f"   - Subcodes: {len(self.grounded_subcodes_collection)} gesammelt")
                print(f"   - Sättigungs-Counter: {self.grounded_saturation_counter}")
            else:
                print(f"\n⏳ Sättigung noch nicht erreicht - setze Subcode-Sammlung fort")
                missing_criteria = [k for k, v in criteria.items() if not v]
                print(f"   - Fehlende Kriterien: {', '.join(missing_criteria)}")
            
            return is_saturated
            
        except Exception as e:
            print(f"❌ Fehler bei Grounded Mode Sättigungsprüfung: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback: Bei Fehler weiter sammeln, außer 100% Material erreicht
            return material_percentage >= 100
        
    async def _analyze_for_subcategories(self, relevant_segments: List[str], 
                                       current_categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        SPEZIELLE ABDUKTIVE ANALYSE: Nur Subkategorien entwickeln
        """
        segments_text = "\n\n=== SEGMENT BREAK ===\n\n".join(
            f"SEGMENT {i + 1}:\n{text}" 
            for i, text in enumerate(relevant_segments)
        )

        # Erstelle Kategorien-Kontext für abduktive Analyse
        categories_context = []
        for cat_name, cat_def in current_categories.items():
            categories_context.append({
                'name': cat_name,
                'definition': cat_def.definition[:200],
                'existing_subcategories': list(cat_def.subcategories.keys())
            })

        prompt = self.prompt_handler.get_analyze_for_subcategories_prompt(
                segments_text=segments_text,
                categories_context=categories_context
            )
        

        try:
            token_counter.start_request()
            
            response = await self.inductive_coder.llm_provider.create_completion(
                model=self.inductive_coder.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Antworte auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )
            
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)
            
            
            token_counter.track_response(response, self.model_name)
            
            # Verarbeite Ergebnisse - erweitere bestehende Kategorien
            extended_categories = current_categories.copy()
            total_new_subcats = 0
            
            for main_cat_name, updates in result.get('extended_categories', {}).items():
                if main_cat_name in extended_categories:
                    current_cat = extended_categories[main_cat_name]
                    new_subcats = {}
                    
                    for sub_data in updates.get('new_subcategories', []):
                        if sub_data.get('confidence', 0) >= 0.7:  # Schwelle für Subkategorien
                            new_subcats[sub_data['name']] = sub_data['definition']
                            total_new_subcats += 1
                            print(f"✅ Neue Subkategorie: {main_cat_name} → {sub_data['name']}")
                    
                    if new_subcats:
                        # Erweitere bestehende Kategorie
                        extended_categories[main_cat_name] = current_cat.replace(
                            subcategories={**current_cat.subcategories, **new_subcats},
                            modified_date=datetime.now().strftime("%Y-%m-%d")
                        )
            
            print(f"📊 Abduktive Entwicklung: {total_new_subcats} neue Subkategorien")
            return extended_categories
            
        except Exception as e:
            print(f"Fehler bei abduktiver Analyse: {str(e)}")
            return current_categories

    async def _code_batch_deductively(self, 
                                     batch: List[Tuple[str, str]], 
                                     categories: Dict[str, CategoryDefinition],
                                     category_preselections: Dict[str, Dict] = None) -> List[Dict]:
        """
        Kodiert einen Batch parallel ohne progressive Kontext-Funktionalität.
        FIX: Erweitert um Kategorie-Vorauswahl für deduktiven Modus
        BUGFIX: Verwendet separate, lockere Relevanzprüfung für Kodierung.
        """
        print(f"\n🚀 PARALLEL-KODIERUNG: {len(batch)} Segmente gleichzeitig")
        start_time = time.time()
        
        # FIX: Standardwert für category_preselections
        if category_preselections is None:
            category_preselections = {}
        
        # FIX: Zeige Kategorie-Vorauswahl-Informationen
        if category_preselections:
            preselected_count = len([s for s in batch if s[0] in category_preselections])
            print(f"🎯 {preselected_count} Segmente haben Kategorie-Präferenzen")
            
            # Statistik der Kategorie-Präferenzen
            all_preferred = []
            for prefs in category_preselections.values():
                all_preferred.extend(prefs.get('preferred_categories', []))
            
            if all_preferred:
                from collections import Counter
                pref_stats = Counter(all_preferred)
                print(f"🎯 Häufigste Präferenzen: {dict(pref_stats.most_common(3))}")
        
        print(f"\n🔍 Prüfe Kodierungs-Relevanz...")
        coding_relevance_results = await self.relevance_checker.check_relevance_batch(batch)
        
        # Debug-Ausgaben
        print(f"\n🔍 Kodierungs-Relevanzprüfung Ergebnisse:")
        relevant_count = sum(1 for is_relevant in coding_relevance_results.values() if is_relevant)
        print(f"   - Segmente geprüft: {len(coding_relevance_results)}")
        print(f"   - Als kodierungsrelevant eingestuft: {relevant_count}")
        print(f"   - Als nicht kodierungsrelevant eingestuft: {len(coding_relevance_results) - relevant_count}")
        
        # 2. PARALLEL: Mehrfachkodierungs-Prüfung (wenn aktiviert)
        multiple_coding_results = {}
        if CONFIG.get('MULTIPLE_CODINGS', True):
            coding_relevant_segments = [
                (segment_id, text) for segment_id, text in batch
                if coding_relevance_results.get(segment_id, True)
            ]
            
            if coding_relevant_segments:
                print(f"  🔄 Prüfe {len(coding_relevant_segments)} kodierungsrelevante Segmente auf Mehrfachkodierung...")
                multiple_coding_results = await self.relevance_checker.check_multiple_category_relevance(
                    coding_relevant_segments, categories
                )
        
        # 3. PARALLEL: Kodierung aller Segmente
        async def code_single_segment_all_coders(segment_id: str, text: str) -> List[Dict]:
            """FIX: Kodiert ein einzelnes Segment mit gefilterten Kategorien basierend auf Vorauswahl."""
            
            # FIX: Bestimme effektive Kategorien für dieses Segment
            preselection = category_preselections.get(segment_id, {})
            preferred_cats = preselection.get('preferred_categories', [])
            
            if preferred_cats:
                # FIX: Verwende vollständige CategoryDefinition-Objekte aus categories
                effective_categories = {
                    name: cat for name, cat in categories.items() 
                    if name in preferred_cats and isinstance(cat, CategoryDefinition)  # FIX: Validiere CategoryDefinition
                }
                
                if not effective_categories:
                    print(f"    ⚠️ Keine gültigen CategoryDefinition-Objekte in preferred_cats - verwende alle Kategorien")
                    effective_categories = categories
                else:
                    print(f"    🎯 Segment {segment_id}: Fokus auf {len(effective_categories)} Kategorien: {', '.join(preferred_cats)}")
                    
                    # FIX: Validiere, dass effective_categories vollständige Definitionen hat
                    for name, cat in effective_categories.items():
                        if not hasattr(cat, 'subcategories'):
                            print(f"    ⚠️ KRITISCH: effective_categories['{name}'] fehlen Subkategorien - hole aus categories")
                            if name in categories:
                                effective_categories[name] = categories[name]
            else:
                # FIX: Fallback auf alle Kategorien wenn keine Vorauswahl
                effective_categories = categories
                print(f"    📝 Segment {segment_id}: Standard-Kodierung mit allen {len(categories)} Kategorien")
            
            
            is_coding_relevant = coding_relevance_results.get(segment_id, True)  # Default: True

            # Zusätzliche einfache Heuristik für offensichtlich irrelevante Inhalte
            if len(text.strip()) < 20:
                is_coding_relevant = False
                print(f"   🚫 Segment {segment_id} zu kurz für Kodierung")
                
            text_lower = text.lower()
            exclusion_patterns = [
                'seite ', 'page ', 'copyright', '©', 'datum:', 'date:',
                'inhaltsverzeichnis', 'table of contents', 'literaturverzeichnis',
                'bibliography', 'anhang', 'appendix'
            ]
            
            is_metadata = any(pattern in text_lower for pattern in exclusion_patterns)
            if is_metadata and len(text) < 100:
                is_coding_relevant = False
                print(f"   🚫 Segment {segment_id} als Metadaten erkannt")
            
            if not is_coding_relevant:
                print(f"   🚫 Segment {segment_id} wird als 'Nicht kodiert' markiert")
                
                # FIX: Hole Begründung aus RelevanceChecker falls verfügbar
                relevance_details = self.relevance_checker.get_relevance_details(segment_id)
                justification = "Nicht relevant für Kodierung (zu kurz oder Metadaten)"
                if relevance_details:
                    if 'reasoning' in relevance_details and relevance_details['reasoning']:
                        justification = relevance_details['reasoning']
                    elif 'justification' in relevance_details and relevance_details['justification']:
                        justification = relevance_details['justification']
                
                # Spezifische Fallback-Begründungen
                if len(text.strip()) < 20:
                    justification = "Segment zu kurz für sinnvolle Kodierung"
                elif is_metadata:
                    justification = "Segment als Metadaten (z.B. Seitenzahl, Inhaltsverzeichnis) erkannt"
                
                not_coded_results = []
                for coder in self.deductive_coders:
                    result = {
                        'segment_id': segment_id,
                        'coder_id': coder.coder_id,
                        'category': "Nicht kodiert",
                        'subcategories': [],
                        'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
                        'justification': justification,
                        'text': text,
                        'multiple_coding_instance': 1,
                        'total_coding_instances': 1,
                        'target_category': '',
                        'category_focus_used': False,
                        # FIX: Zusätzliche Kategorie-Vorauswahl-Info
                        'category_preselection_used': bool(preferred_cats),
                        'preferred_categories': preferred_cats,
                        'preselection_reasoning': preselection.get('reasoning', '')
                    }
                    not_coded_results.append(result)
                return not_coded_results
            
            # Bestimme Kodierungsinstanzen (für Mehrfachkodierung)
            coding_instances = []
            multiple_categories = multiple_coding_results.get(segment_id, [])
            
            if len(multiple_categories) > 1:
                # Mehrfachkodierung
                for i, category_info in enumerate(multiple_categories, 1):
                    coding_instances.append({
                        'instance': i,
                        'total_instances': len(multiple_categories),
                        'target_category': category_info['category'],
                        'category_context': category_info
                    })
            else:
                # Standardkodierung
                coding_instances.append({
                    'instance': 1,
                    'total_instances': 1,
                    'target_category': '',
                    'category_context': None
                })
            
            # 🚀 PARALLEL: Alle Kodierer für alle Instanzen
            async def code_with_coder_and_instance(coder, instance_info):
                """FIX: Kodiert mit einem Kodierer unter Verwendung der vollständigen CategoryDefinition-Objekte."""
                try:

                    # FIX: Bei Fokuskodierung die target_category zu effective_categories hinzufügen
                    enhanced_categories = effective_categories.copy()
                    target_cat = instance_info['target_category']
                    if target_cat:
                        if target_cat and target_cat not in enhanced_categories:
                            if target_cat in categories:  
                                enhanced_categories[target_cat] = categories[target_cat]  
                                print(f"    🎯 Fokuskategorie '{target_cat}' zu verfügbaren Kategorien hinzugefügt")
                            else:
                                print(f"    ⚠️ Fokuskategorie '{target_cat}' nicht in Kategorien vorhanden")

                    if target_cat:
                        coding = await coder.code_chunk_with_focus(
                            text, enhanced_categories,
                            focus_category=target_cat,
                            focus_context=instance_info['category_context']
                        )
                    else:
                        coding = await coder.code_chunk(text, enhanced_categories)


                    if coding and isinstance(coding, CodingResult):
                        main_category = coding.category
                        original_subcats = list(coding.subcategories)
                        
                        # FIX: Verwende enhanced_categories für Validierung
                        validated_subcats = original_subcats  # Fallback
                        validation_source = "keine"
                        
                        # 1. Priorität: enhanced_categories (gefilterte + Fokuskategorien)
                        if main_category in enhanced_categories and hasattr(enhanced_categories[main_category], 'subcategories'):
                            try:
                                validated_subcats = CategoryValidator.validate_subcategories_for_category(
                                    original_subcats, main_category, enhanced_categories, warn_only=False
                                )
                                validation_source = "enhanced_categories"
                            except Exception as e:
                                print(f"    ❌ Validierung mit enhanced_categories fehlgeschlagen: {str(e)}")
                                # FIX: Fallback zu self.current_categories
                                if hasattr(self, 'current_categories') and main_category in self.current_categories:
                                    try:
                                        validated_subcats = CategoryValidator.validate_subcategories_for_category(
                                            original_subcats, main_category, self.current_categories, warn_only=False
                                        )
                                        validation_source = "self.current_categories_fallback"
                                    except Exception as e2:
                                        print(f"    ❌ Auch Fallback-Validierung fehlgeschlagen: {str(e2)}")
                        else:
                            # FIX: Informative Meldung bei nicht verfügbaren Kategorien
                            if main_category not in enhanced_categories:
                                print(f"    ℹ️ Kategorie '{main_category}' nicht in verfügbaren Kategorien")
                                print(f"    🎯 Verfügbare Kategorien: {list(enhanced_categories.keys())}")
                            elif not hasattr(enhanced_categories[main_category], 'subcategories'):
                                print(f"    ℹ️ Keine Subkategorie-Definitionen verfügbar für '{main_category}'")
                        
                        # FIX: Debug-Ausgabe nur bei wichtigen Ereignissen
                        if len(original_subcats) != len(validated_subcats):
                            removed = set(original_subcats) - set(validated_subcats)
                            print(f"    🔧 Subkategorien bereinigt: {len(original_subcats)} → {len(validated_subcats)}")
                            if removed:
                                print(f"    🔧 Entfernt: {list(removed)} (Quelle: {validation_source})")
                        elif validation_source != "keine" and original_subcats:
                            print(f"    ✅ Alle {len(original_subcats)} Subkategorien gültig (Quelle: {validation_source})")
                        elif validation_source == "keine" and original_subcats:
                            print(f"    ℹ️ Subkategorien-Validierung übersprungen für '{main_category}' (Quelle: {validation_source})")
                        
                        return {
                            'segment_id': segment_id,
                            'coder_id': coder.coder_id,
                            'category': coding.category,
                            'subcategories': validated_subcats,  # FIX: Immer validierte oder ursprüngliche Subkategorien
                            'confidence': coding.confidence,
                            'justification': coding.justification,
                            'text': text,
                            'paraphrase': coding.paraphrase,
                            'keywords': coding.keywords,
                            'multiple_coding_instance': instance_info['instance'],
                            'total_coding_instances': instance_info['total_instances'],
                            'target_category': instance_info['target_category'],
                            'category_focus_used': bool(instance_info['target_category']),
                            # FIX: Verbesserte Debug-Informationen für enhanced_categories
                            'category_preselection_used': bool(preferred_cats),
                            'preferred_categories': preferred_cats,
                            'effective_categories_count': len(effective_categories),
                            'enhanced_categories_count': len(enhanced_categories),  # FIX: Neue Info
                            'preselection_reasoning': preselection.get('reasoning', ''),
                            'subcategories_validated': len(original_subcats) != len(validated_subcats),
                            'validation_source': validation_source,
                            'validation_successful': validation_source != "keine",
                            'category_in_enhanced': main_category in enhanced_categories,  # FIX: Verwende enhanced_categories
                            'enhanced_has_subcategories': main_category in enhanced_categories and hasattr(enhanced_categories[main_category], 'subcategories'),  # FIX: Neue Debug-Info
                            'focus_category_added': instance_info['target_category'] and instance_info['target_category'] not in effective_categories  # FIX: Neue Info
                        }
                        
                    else:
                        return None
                        
                except Exception as e:
                    print(f"    ⚠️ Kodierungsfehler {coder.coder_id}: {str(e)}")
                    return None
                
            # Erstelle Tasks für alle Kodierer × alle Instanzen
            coding_tasks = []
            for instance_info in coding_instances:
                for coder in self.deductive_coders:
                    task = code_with_coder_and_instance(coder, instance_info)
                    coding_tasks.append(task)
            
            # Führe alle Kodierungen für dieses Segment parallel aus
            coding_results = await asyncio.gather(*coding_tasks, return_exceptions=True)
            
            # Sammle erfolgreiche Ergebnisse
            successful_codings = []
            for result in coding_results:
                if not isinstance(result, Exception) and result:
                    successful_codings.append(result)
            
            return successful_codings
        
        # 🚀 Erstelle Tasks für alle Segmente des Batches
        segment_tasks = [
            code_single_segment_all_coders(segment_id, text) 
            for segment_id, text in batch
        ]
        
        print(f"🚀 Starte parallele Kodierung von {len(segment_tasks)} Segmenten...")
        
        # 🚀 Führe alle Segment-Kodierungen parallel aus
        all_segment_results = await asyncio.gather(*segment_tasks, return_exceptions=True)
        
        # Sammle alle Ergebnisse
        batch_results = []
        successful_segments = 0
        error_count = 0
        preselection_used_count = 0
        validation_performed_count = 0
        
        for segment_result in all_segment_results:
            if isinstance(segment_result, Exception):
                print(f"⚠️ Segment-Fehler: {segment_result}")
                error_count += 1
                continue
                
            if segment_result:  # Liste von Kodierungen für dieses Segment
                batch_results.extend(segment_result)
                successful_segments += 1
                
                # FIX: Sammle Statistiken über Kategorie-Vorauswahl-Nutzung
                for coding in segment_result:
                    if coding.get('category_preselection_used', False):
                        preselection_used_count += 1
                    if coding.get('subcategories_validated', False):
                        validation_performed_count += 1
        
        # Markiere verarbeitete Segmente
        for segment_id, text in batch:
            self.processed_segments.add(segment_id)
        
        processing_time = time.time() - start_time
        
        print(f"✅ PARALLEL-BATCH ABGESCHLOSSEN:")
        print(f"   ⚡ Zeit: {processing_time:.2f}s")
        if processing_time > 0:
            print(f"   🚀 Geschwindigkeit: {len(batch)/processing_time:.1f} Segmente/Sekunde")
        else:
            print(f"   🚀 Geschwindigkeit: {len(batch)} Segmente in <0.01s (sehr schnell)")
        print(f"   ✓ Erfolgreiche Segmente: {successful_segments}/{len(batch)}")
        print(f"   📊 Gesamte Kodierungen: {len(batch_results)}")
        # FIX: Zusätzliche Statistiken für Kategorie-Vorauswahl
        if category_preselections:
            print(f"   🎯 Kategorie-Vorauswahl genutzt: {preselection_used_count} Kodierungen")
            print(f"   🔧 Subkategorie-Validierung durchgeführt: {validation_performed_count} Kodierungen")
        if error_count > 0:
            print(f"   ⚠️ Fehler: {error_count}")
        
        return batch_results
    
    async def _code_with_category_focus(self, coder, text, categories, instance_info):
        """Kodiert mit optionalem Fokus auf bestimmte Kategorie"""
        
        if instance_info['target_category']:
            # Mehrfachkodierung mit Kategorie-Fokus
            return await coder.code_chunk_with_focus(
                text, categories, 
                focus_category=instance_info['target_category'],
                focus_context=instance_info['category_context']
            )
        else:
            # Standard-Kodierung
            return await coder.code_chunk(text, categories)

    async def _code_batch_with_context(self, batch: List[Tuple[str, str]], 
                                     categories: Dict[str, CategoryDefinition],
                                     category_preselections: Dict[str, Dict] = None) -> List[Dict]:
        """
        Kodiert einen Batch sequentiell mit progressivem Dokumentkontext und Mehrfachkodierung.
        FIX: Erweitert um category_preselections Parameter für gefilterte Kategorien
        """
        # FIX: Standardwert für category_preselections
        if category_preselections is None:
            category_preselections = {}
        
        batch_results = []
        
        # Debug-Info für Kategorie-Präferenzen
        if category_preselections:
            preselected_count = len([s for s in batch if s[0] in category_preselections])
            print(f"🎯 Kontext-Kodierung: {preselected_count} Segmente haben Kategorie-Präferenzen")
        
            # FIX: Erweiterte Statistik zur Kategorie-Vorauswahl wie im ohne-Kontext-Modus
            all_preferred = []
            for prefs in category_preselections.values():
                all_preferred.extend(prefs.get('preferred_categories', []))
            
            if all_preferred:
                from collections import Counter
                pref_stats = Counter(all_preferred)
                print(f"🎯 Häufigste Präferenzen: {dict(pref_stats.most_common(3))}")
        else:
            print("🎯 Kontext-Kodierung: Keine Kategorie-Präferenzen übertragen")

        # Prüfe Mehrfachkodierungs-Möglichkeiten für den ganzen Batch
        multiple_coding_results = {}
        if CONFIG.get('MULTIPLE_CODINGS', True):
            # Relevanzprüfung für ganzen Batch
            relevance_results = await self.relevance_checker.check_relevance_batch(batch)
            relevant_segments = [
                (segment_id, text) for segment_id, text in batch
                if relevance_results.get(segment_id, False)
            ]
            
            if relevant_segments:
                print(f"  🔄 Prüfe {len(relevant_segments)} relevante Segmente auf Mehrfachkodierung...")
                multiple_coding_results = await self.relevance_checker.check_multiple_category_relevance(
                    relevant_segments, categories
                )
        
        # Sequentielle Verarbeitung um Kontext aufzubauen
        for segment_id, text in batch:
            # Extrahiere Dokumentnamen und Chunk-ID
            doc_name, chunk_id = self._extract_doc_and_chunk_id(segment_id)
            position = f"Chunk {chunk_id}"
            
            # Hole aktuelles Summary oder initialisiere es
            current_summary = self.document_summaries.get(doc_name, "")
            
            # Segment-Informationen
            segment_info = {
                'doc_name': doc_name,
                'chunk_id': chunk_id,
                'position': position
            }
            
            # Kontext-Daten vorbereiten
            context_data = {
                'current_summary': current_summary,
                'segment_info': segment_info
            }
            
            # FIX: Bestimme gefilterte Kategorien für dieses Segment
            preselection = category_preselections.get(segment_id, {})
            preferred_cats = preselection.get('preferred_categories', [])
            
            if preferred_cats:
                # FIX: Gefilterte Kategorien für Kodierung verwenden
                filtered_categories = {
                    name: cat for name, cat in categories.items() 
                    if name in preferred_cats
                }
                print(f"\n🔍 Verarbeite Segment {segment_id} mit Kontext (🎯 Fokus auf {len(filtered_categories)} Kategorien: {', '.join(preferred_cats)})")
                effective_categories = filtered_categories
            else:
                # FIX: Fallback auf alle Kategorien
                print(f"\n🔍 Verarbeite Segment {segment_id} mit Kontext")
                effective_categories = categories
            
            # Prüfe Relevanz
            relevance_result = await self.relevance_checker.check_relevance_batch([(segment_id, text)])
            is_relevant = relevance_result.get(segment_id, False)
            
            if not is_relevant:
                print(f"  ↪ Segment als nicht relevant markiert - wird übersprungen")
                
                # FIX: Hole spezifische Begründung aus RelevanceChecker
                relevance_details = self.relevance_checker.get_relevance_details(segment_id)
                justification = "Nicht relevant für Forschungsfrage"
                if relevance_details:
                    if 'reasoning' in relevance_details and relevance_details['reasoning']:
                        justification = relevance_details['reasoning']
                    elif 'justification' in relevance_details and relevance_details['justification']:
                        justification = relevance_details['justification']
                # FIX: Ende
                
                # Erstelle "Nicht kodiert" Ergebnis für alle Kodierer
                for coder in self.deductive_coders:
                    result = {
                        'segment_id': segment_id,
                        'coder_id': coder.coder_id,
                        'category': "Nicht kodiert",
                        'subcategories': [],
                        'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
                        'justification': justification,  # FIX: Verwende spezifische Begründung
                        'text': text,
                        'context_summary': current_summary,
                        'multiple_coding_instance': 1,
                        'total_coding_instances': 1,
                        'target_category': '',
                        'category_focus_used': False,
                        # FIX: Neue Felder für Kategorie-Präferenzen
                        'category_preselection_used': bool(preferred_cats),
                        'preselected_categories': preferred_cats,
                        'category_filtering_applied': bool(preferred_cats)
                    }
                    batch_results.append(result)
                continue
            
            # Bestimme Kodierungsinstanzen
            coding_instances = []
            multiple_categories = multiple_coding_results.get(segment_id, [])
            
            if len(multiple_categories) > 1:
                print(f"  🔄 Mehrfachkodierung mit Kontext: {len(multiple_categories)} Kategorien")
                for i, category_info in enumerate(multiple_categories, 1):
                    coding_instances.append({
                        'instance': i,
                        'total_instances': len(multiple_categories),
                        'target_category': category_info['category'],
                        'category_context': category_info
                    })
            else:
                coding_instances.append({
                    'instance': 1,
                    'total_instances': 1,
                    'target_category': '',
                    'category_context': None
                })
            
            # Verarbeite relevante Segmente mit Kontext für ALLE Kodierer und Instanzen
            updated_summary = current_summary
            
            for instance_info in coding_instances:
                if instance_info['total_instances'] > 1:
                    print(f"\n    📝 Kontext-Kodierung {instance_info['instance']}/{instance_info['total_instances']}")
                    print(f"        Fokus: {instance_info['target_category']}")
            
                for coder_index, coder in enumerate(self.deductive_coders):
                    try:
                        # Bestimme ob Summary aktualisiert werden soll (nur beim ersten Kodierer der ersten Instanz)
                        should_update_summary = (coder_index == 0 and instance_info['instance'] == 1)

                        # FIX: Enhanced Categories Logic - füge Fokuskategorie zu effective_categories hinzu
                        enhanced_categories = effective_categories.copy()
                        target_cat = instance_info['target_category']
                        if target_cat:
                            if target_cat and target_cat not in enhanced_categories:
                                if target_cat in categories:  
                                    enhanced_categories[target_cat] = categories[target_cat]  
                                    print(f"    🎯 Fokuskategorie '{target_cat}' zu verfügbaren Kategorien hinzugefügt")
                                else:
                                    print(f"    ⚠️ Fokuskategorie '{target_cat}' nicht in Kategorien vorhanden")
                        
                        if instance_info['target_category']:
                            # FIX: Mehrfachkodierung mit Fokus und Kontext (mit gefilterten Kategorien)
                            combined_result = await coder.code_chunk_with_focus_and_context(
                                text, enhanced_categories,  
                                focus_category=instance_info['target_category'],
                                focus_context=instance_info['category_context'],
                                current_summary=updated_summary if should_update_summary else current_summary,
                                segment_info=segment_info,
                                update_summary=should_update_summary
                            )
                        else:
                            # FIX: Standard Kontext-Kodierung (mit gefilterten Kategorien)
                            combined_result = await coder.code_chunk_with_progressive_context(
                                text, 
                                enhanced_categories,  
                                updated_summary if should_update_summary else current_summary,
                                segment_info
                            )
                        
                        if combined_result:
                            # Extrahiere Kodierungsergebnis und ggf. aktualisiertes Summary
                            coding_result = combined_result.get('coding_result', {})
                            
                            # Summary nur beim ersten Kodierer der ersten Instanz aktualisieren
                            if should_update_summary:
                                updated_summary = combined_result.get('updated_summary', current_summary)
                                self.document_summaries[doc_name] = updated_summary
                                print(f"🔄 Summary aktualisiert: {len(updated_summary.split())} Wörter")
                            
                            # FIX: Erstelle erweiterten Kodierungseintrag mit Kategorie-Präferenzen
                            coding_entry = {
                                'segment_id': segment_id,
                                'coder_id': coder.coder_id,
                                'category': coding_result.get('category', ''),
                                'subcategories': coding_result.get('subcategories', []),
                                'justification': coding_result.get('justification', ''),
                                'confidence': coding_result.get('confidence', {}),
                                'text': text,
                                'paraphrase': coding_result.get('paraphrase', ''),
                                'keywords': coding_result.get('keywords', ''),
                                'context_summary': updated_summary,
                                'context_influence': coding_result.get('context_influence', ''),
                                'multiple_coding_instance': instance_info['instance'],
                                'total_coding_instances': instance_info['total_instances'],
                                'target_category': instance_info['target_category'],
                                'category_focus_used': bool(instance_info['target_category']),
                                # FIX: Neue Felder für Kategorie-Präferenzen
                                'category_preselection_used': bool(preferred_cats),
                                'preselected_categories': preferred_cats,
                                'category_filtering_applied': bool(preferred_cats),
                                'relevance_scores': preselection.get('relevance_scores', {}),
                                'preselection_reasoning': preselection.get('reasoning', '')
                            }
                            
                            # FIX: Validiere Subkategorien gegen die gewählte Hauptkategorie
                            main_category = coding_entry['category']
                            original_subcats = coding_entry['subcategories']
                            if main_category and main_category != 'Nicht kodiert':
                                # FIX: Verwende enhanced_categories für Validierung (schon vorhanden)
                                categories_dict_for_validation = enhanced_categories
                                
                                if not categories_dict_for_validation:
                                    print(f"⚠️ KRITISCH: Kein categories_dict verfügbar für Validierung!")
                                    # Fallback: verwende original categories
                                    categories_dict_for_validation = categories
                                
                                try:
                                    validated_subcats = CategoryValidator.validate_subcategories_for_category(
                                        original_subcats, main_category, categories_dict_for_validation, warn_only=False
                                    )
                                    if len(validated_subcats) != len(original_subcats):
                                        print(f"    ⚠️ Subkategorien bereinigt: {len(original_subcats)} → {len(validated_subcats)}")
                                    coding_entry['subcategories'] = validated_subcats
                                except Exception as e:
                                    print(f"    ❌ Subkategorien-Validierung fehlgeschlagen: {str(e)}")
                                    print(f"    📊 Debug: main_category='{main_category}', enhanced_categories_keys={list(enhanced_categories.keys())[:5]}")
                                    # Behalte ursprüngliche Subkategorien bei Validierungsfehlern
                            
                            
                            batch_results.append(coding_entry)
                            
                            if instance_info['total_instances'] > 1:
                                category_display = coding_entry['category']
                                if preferred_cats and category_display in preferred_cats:
                                    category_display += " 🎯"
                                print(f"        ✓ {coder.coder_id}: {category_display}")
                            else:
                                category_display = coding_entry['category']
                                if preferred_cats and category_display in preferred_cats:
                                    category_display += " 🎯"
                                print(f"  ✓ Kodierer {coder.coder_id}: {category_display}")
                        else:
                            print(f"  ⚠ Keine Kodierung von {coder.coder_id} erhalten")
                            
                    except Exception as e:
                        print(f"  ⚠ Fehler bei {coder.coder_id}: {str(e)}")
                        continue
        
        return batch_results
    
    def _extract_doc_and_chunk_id(self, segment_id: str) -> Tuple[str, str]:
        """Extrahiert Dokumentname und Chunk-ID aus segment_id."""
        parts = segment_id.split('_chunk_')
        if len(parts) == 2:
            return parts[0], parts[1]
        else:
            return segment_id, "unknown"

    def _extract_segment_sort_key(self, segment_id: str) -> Tuple[str, int]:
        """Erstellt einen Sortierschlüssel für die richtige Chunk-Reihenfolge."""
        try:
            doc_name, chunk_id = self._extract_doc_and_chunk_id(segment_id)
            return (doc_name, int(chunk_id) if chunk_id.isdigit() else 0)
        except Exception:
            return (segment_id, 0)
    
    async def _finalize_by_mode(self, analysis_mode: str, current_categories: Dict, 
                            deductive_categories: Dict, initial_categories: Dict) -> Dict:
        """
        KORRIGIERTE Finalisierung - gibt immer ein Dictionary zurück
        """
        try:
            if analysis_mode == 'inductive':
                print(f"\n🔄 INDUCTIVE MODE Finalisierung:")
                print(f"   - Deduktive Kategorien: IGNORIERT")
                print(f"   - Induktive Kategorien: {len(current_categories)}")
                print(f"   → Finales System: NUR {len(current_categories)} induktive Kategorien")
                return current_categories
                
            elif analysis_mode == 'grounded':
                # Im separaten Grounded Mode wurde bereits alles erledigt
                print(f"\n✅ GROUNDED MODE bereits vollständig abgeschlossen")
                return current_categories
                
            elif analysis_mode == 'abductive':
                print(f"\n🔄 ABDUCTIVE MODE Finalisierung:")
                print(f"   - Erweiterte deduktive Kategorien: {len(current_categories)}")
                return current_categories
                
            else:  # deductive oder andere
                print(f"\n🔄 {analysis_mode.upper()} MODE Finalisierung:")
                print(f"   - Kategorien: {len(current_categories)}")
                return current_categories
                
        except Exception as e:
            print(f"Fehler in _finalize_by_mode: {str(e)}")
            # Fallback: Gebe wenigstens die aktuellen Kategorien zurück
            return current_categories or initial_categories or {}

    def _show_final_development_stats(self, final_categories: Dict, initial_categories: Dict, batch_count: int):
        """
        Zeigt finale Entwicklungsstatistiken
        """
        print(f"\n{'='*80}")
        print(f"📊 KATEGORIENENTWICKLUNG ABGESCHLOSSEN")
        print(f"{'='*80}")
        
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'inductive')
        
        if analysis_mode == 'inductive':
            print(f"🔬 INDUCTIVE MODE - Eigenständiges induktives System:")
            print(f"   - Deduktive Kategorien: IGNORIERT")
            print(f"   - Entwickelte induktive Kategorien: {len(final_categories)}")
            print(f"   - Verarbeitete Batches: {batch_count}")
            
            # Subkategorien-Statistik
            total_subcats = sum(len(cat.subcategories) for cat in final_categories.values())
            print(f"   - Subkategorien: {total_subcats}")
            
        else:
            # Bestehende Logik für andere Modi - KORRIGIERT
            initial_count = len(initial_categories) if initial_categories else 0  # ✅ BUGFIX: len() hinzugefügt
            final_count = len(final_categories)
            new_count = final_count - initial_count  # ✅ Jetzt korrekt: int - int
            
            print(f"📈 Entwicklungsbilanz:")
            print(f"   - Verarbeitete Batches: {batch_count}")
            print(f"   - Initial: {initial_count} Kategorien")
            print(f"   - Neu entwickelt: {new_count} Kategorien")
            print(f"   - Final: {final_count} Kategorien")
            
            # Subkategorien-Statistik
            total_subcats = sum(len(cat.subcategories) for cat in final_categories.values())
            print(f"   - Subkategorien: {total_subcats}")

        if (hasattr(self, 'inductive_coder') and 
            self.inductive_coder and 
            hasattr(self.inductive_coder, 'theoretical_saturation_history') and
            self.inductive_coder.theoretical_saturation_history):
            
            final_saturation = self.inductive_coder.theoretical_saturation_history[-1]
            print(f"\n🎯 Finale Sättigung:")
            print(f"   - Theoretische Sättigung: {final_saturation['theoretical_saturation']:.1%}")
            print(f"   - Kategorienqualität: {final_saturation['category_quality']:.1%}")
            print(f"   - Diversität: {final_saturation['category_diversity']:.1%}")
        
        if (hasattr(self, 'inductive_coder') and 
            self.inductive_coder and 
            hasattr(self.inductive_coder, 'category_development_phases') and
            self.inductive_coder.category_development_phases):
            
            print(f"\n📊 Entwicklungsphasen:")
            for phase in self.inductive_coder.category_development_phases:
                print(f"   Batch {phase['batch']}: +{phase['new_categories']} → {phase['total_categories']} total")
    
    async def _save_grounded_checkpoint(self):
        """Speichere Grounded Mode Checkpoint zwischen Batches"""
        try:
            checkpoint_path = os.path.join(self.config['OUTPUT_DIR'], 'grounded_checkpoint.json')
            checkpoint_data = {
                'subcodes': self.grounded_subcodes_collection,
                'keywords': list(set(self.grounded_keywords_collection)),
                'batch_history': self.grounded_batch_history,
                'saturation_counter': self.grounded_saturation_counter,
                'segment_analyses_count': len(self.grounded_segment_analyses),
                'timestamp': datetime.now().isoformat(),
                'version': '1.0'
            }
            
            with open(checkpoint_path, 'w', encoding='utf-8') as f:
                json.dump(checkpoint_data, f, indent=2, ensure_ascii=False)
                
            print(f"💾 Grounded Checkpoint gespeichert: {len(self.grounded_subcodes_collection)} Subcodes")
            
        except Exception as e:
            print(f"⚠️ Fehler beim Speichern des Grounded Checkpoints: {str(e)}")

    async def _load_grounded_checkpoint(self):
        """Lade Grounded Mode Checkpoint falls vorhanden"""
        try:
            checkpoint_path = os.path.join(self.config['OUTPUT_DIR'], 'grounded_checkpoint.json')
            if os.path.exists(checkpoint_path):
                with open(checkpoint_path, 'r', encoding='utf-8') as f:
                    checkpoint_data = json.load(f)
                
                self.grounded_subcodes_collection = checkpoint_data.get('subcodes', [])
                self.grounded_keywords_collection = checkpoint_data.get('keywords', [])
                self.grounded_batch_history = checkpoint_data.get('batch_history', [])
                self.grounded_saturation_counter = checkpoint_data.get('saturation_counter', 0)
                
                print(f"💾 Grounded Checkpoint geladen: {len(self.grounded_subcodes_collection)} Subcodes")
                print(f"   - Keywords: {len(self.grounded_keywords_collection)}")
                print(f"   - Batch-Historie: {len(self.grounded_batch_history)} Einträge")
                return True
        except Exception as e:
            print(f"⚠️ Fehler beim Laden des Grounded Checkpoints: {str(e)}")
        
        return False
    
    async def analyze_material(self, 
                            chunks: Dict[str, List[str]], 
                            initial_categories: Dict,
                            skip_inductive: bool = False,
                            batch_size: Optional[int] = None) -> Tuple[Dict, List]:
        """
        KORRIGIERTE Hauptanalyse mit besserer Fehlerbehandlung
        """
        try:
            token_counter.reset_session() # Token-Session zurücksetzen für neue Analyse

            self.escape_handler.start_monitoring()
            self.start_time = datetime.now()
            print(f"\nAnalyse gestartet um {self.start_time.strftime('%H:%M:%S')}")

            # Berechne _total_segments ZUERST
            all_segments = self._prepare_segments(chunks)
            self._total_segments = len(all_segments)
            
            # Speichere chunks als Instanzvariable
            self.chunks = chunks
            
            analysis_mode = CONFIG.get('ANALYSIS_MODE', 'inductive')
            
            # Reset Tracking-Variablen
            self.coding_results = []
            self.processed_segments = set()
            
            if batch_size is None:
                batch_size = CONFIG.get('BATCH_SIZE', 5)
            
            total_segments = len(all_segments)
            print(f"Verarbeite {total_segments} Segmente mit Batch-Größe {batch_size}...")
            self.history.log_analysis_start(total_segments, len(initial_categories))

            # GROUNDED MODE: Spezielle Behandlung
            if analysis_mode == 'grounded':
                result = await self._analyze_grounded_mode(
                    chunks, initial_categories, all_segments, batch_size
                )
            else:
                # Normale Modi (inductive, abductive, deductive)
                result = await self._analyze_normal_modes(
                    chunks, initial_categories, all_segments, skip_inductive, batch_size, analysis_mode
                )
            
            # KORRIGIERT: Prüfe ob result ein Tupel ist
            if result is None:
                print("⚠️ Warnung: Analyse-Methode gab None zurück")
                return initial_categories, []
            
            if not isinstance(result, tuple) or len(result) != 2:
                print("⚠️ Warnung: Analyse-Methode gab kein gültiges Tupel zurück")
                return initial_categories, []
            
            final_categories, coding_results = result
            
            # Stoppe Escape-Handler
            self.escape_handler.stop_monitoring()
            self.end_time = datetime.now()

            return final_categories, coding_results
                    
        except Exception as e:
            self.end_time = datetime.now()
            print(f"Fehler in der Analyse: {str(e)}")
            traceback.print_exc()
            if hasattr(self, 'escape_handler'):
                self.escape_handler.stop_monitoring()
            raise

    async def _analyze_normal_modes(self, 
                                chunks: Dict[str, List[str]], 
                                initial_categories: Dict,
                                all_segments: List,
                                skip_inductive: bool,
                                batch_size: int,
                                analysis_mode: str) -> Tuple[Dict, List]:
        """
        Analysiert normale Modi (inductive, abductive, deductive)
        KORRIGIERT: Gibt immer ein Tupel zurück
        """
        
        # Kategoriensystem-Behandlung
        if analysis_mode == 'inductive':
            print(f"\n🔄 INDUCTIVE MODE: Entwickle komplett neues induktives Kategoriensystem")
            current_categories = {}  # Leeres induktives System
            deductive_categories = {}  # LEER im inductive mode!
        elif analysis_mode == 'abductive':
            print(f"\n🔄 ABDUCTIVE MODE: Erweitere deduktive Kategorien um Subkategorien")
            current_categories = initial_categories.copy()
            deductive_categories = initial_categories.copy()
        else:  # deductive
            current_categories = initial_categories.copy()
            deductive_categories = initial_categories.copy()

        if batch_size is None:
            batch_size = CONFIG.get('BATCH_SIZE', 5)
        
        total_segments = len(all_segments)
        print(f"Verarbeite {total_segments} Segmente mit Batch-Größe {batch_size}...")

        # Initialisiere ImprovedSaturationController
        saturation_controller = ImprovedSaturationController(analysis_mode)
        
        # HAUPTSCHLEIFE
        batch_count = 0
        use_context = CONFIG.get('CODE_WITH_CONTEXT', False)
        
        while True:
            # Escape-Prüfung
            if self.check_escape_abort():
                print("\n🛑 Abbruch durch Benutzer erkannt...")
                await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                return current_categories, self.coding_results
            
            batch = await self._get_next_batch(all_segments, batch_size)
            if not batch:
                break
                
            batch_count += 1
            material_percentage = (len(self.processed_segments) / total_segments) * 100
            
            print(f"\n{'='*60}")
            print(f"📊 BATCH {batch_count}: {len(batch)} Segmente")
            print(f"📈 Material verarbeitet: {material_percentage:.1f}%")
            print(f"{'='*60}")
            
            batch_start = time.time()
            
            try:
                # 1. ALLGEMEINE RELEVANZPRÜFUNG
                print(f"\n🔍 Schritt 1: Erweiterte Relevanzprüfung für Forschungsfrage...")

                # FIX: Escape-Prüfung vor Relevanzprüfung
                if self.check_escape_abort():
                    print("\n🛑 Abbruch vor Relevanzprüfung erkannt...")
                    await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                    return current_categories, self.coding_results
                
                if analysis_mode == 'deductive':
                    # FIX: Erweiterte Relevanzprüfung für deduktiven Modus
                    extended_relevance_results = await self.relevance_checker.check_relevance_with_category_preselection(
                        batch, current_categories, analysis_mode
                    )
                    
                    # Filtere relevante Segmente und sammle Kategorie-Präferenzen
                    generally_relevant_batch = []
                    category_preselections = {}  # FIX: Neue Variable für Kategorie-Präferenzen
                    
                    for segment_id, text in batch:
                        result = extended_relevance_results.get(segment_id, {})
                        if result.get('is_relevant', False):
                            generally_relevant_batch.append((segment_id, text))
                            # FIX: Speichere Kategorie-Präferenzen für späteren Gebrauch
                            category_preselections[segment_id] = {
                                'preferred_categories': result.get('preferred_categories', []),
                                'relevance_scores': result.get('relevance_scores', {}),
                                'reasoning': result.get('reasoning', '')
                            }
                    
                    print(f"📊 Erweiterte Relevanz: {len(generally_relevant_batch)} von {len(batch)} Segmenten relevant")
                    if category_preselections:
                        preselection_stats = {}
                        for prefs in category_preselections.values():
                            for cat in prefs['preferred_categories']:
                                preselection_stats[cat] = preselection_stats.get(cat, 0) + 1
                        print(f"🎯 Kategorie-Präferenzen: {preselection_stats}")
                        
                else:
                    # FIX: Standard-Relevanzprüfung für andere Modi (unverändert)
                    general_relevance_results = await self.relevance_checker.check_relevance_batch(batch)
                    generally_relevant_batch = [
                        (segment_id, text) for segment_id, text in batch 
                        if general_relevance_results.get(segment_id, False)
                    ]
                    category_preselections = {}  # FIX: Leer für andere Modi
                    print(f"📊 Allgemeine Relevanz: {len(generally_relevant_batch)} von {len(batch)} Segmente relevant")
                
                # Markiere alle Segmente als verarbeitet
                self.processed_segments.update(sid for sid, _ in batch)

                # FIX: Escape-Prüfung nach Relevanzprüfung
                if self.check_escape_abort():
                    print("\n🛑 Abbruch nach Relevanzprüfung erkannt...")
                    await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                    return current_categories, self.coding_results
                
                # 2. INDUKTIVE KATEGORIENENTWICKLUNG
                if not skip_inductive and generally_relevant_batch:
                    print(f"\n🔍 Nächster Schritt: Induktive Kategorienentwicklung...")
                    
                    # FIX: Escape-Prüfung vor Kodierung
                    if self.check_escape_abort():
                        print("\n🛑 Abbruch vor Kodierung erkannt...")
                        await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                        return current_categories, self.coding_results
                    
                    if analysis_mode in ['inductive', 'abductive']:
                        # Standard induktive Kategorienentwicklung
                        new_categories = await self._process_batch_inductively(
                            generally_relevant_batch, 
                            current_categories
                        )
                        
                        if new_categories:
                            before_count = len(current_categories)
                            # FIX: Zähle auch die Subkategorien für bessere Berichterstattung
                            before_subcategories = sum(len(cat.subcategories) for cat in current_categories.values())
                            
                            # Kategorien integrieren
                            current_categories = self._merge_category_systems(
                                current_categories,
                                new_categories
                            )
                            
                            # FIX: Aktualisiere Instanz-Attribut für Validierung
                            self.current_categories = current_categories

                            added_count = len(current_categories) - before_count
                            # FIX: Zähle auch die neuen Subkategorien
                            after_subcategories = sum(len(cat.subcategories) for cat in current_categories.values())
                            added_subcategories = after_subcategories - before_subcategories
                            
                            # FIX: Bessere Ausgabe je nach Analysemodus
                            if analysis_mode == 'abductive':
                                if added_count > 0:
                                    print(f"✅ {added_count} neue Hauptkategorien integriert")
                                if added_subcategories > 0:
                                    print(f"✅ {added_subcategories} neue Subkategorien integriert")
                                if added_count == 0 and added_subcategories == 0:
                                    print("✅ 0 neue Kategorien integriert (wie erwartet im abduktiven Modus)")
                            else:
                                print(f"✅ {added_count} neue Kategorien integriert")
                                if added_subcategories > 0:
                                    print(f"   📝 Zusätzlich {added_subcategories} neue Subkategorien")
                            
                            # Aktualisiere ALLE Kodierer
                            for coder in self.deductive_coders:
                                await coder.update_category_system(current_categories)
                            
                            # FIX: Reset nur bei tatsächlichen Änderungen
                            if added_count > 0 or added_subcategories > 0:
                                saturation_controller.reset_stability_counter()
                        else:
                            saturation_controller.increment_stability_counter()

                    # FIX: Escape-Prüfung nach Kodierung
                    if self.check_escape_abort():
                        print("\n🛑 Abbruch nach Kodierung erkannt...")
                        await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                        return current_categories, self.coding_results
                    
                # 3. DEDUKTIVE KODIERUNG
                print(f"\n🏷️ Nächster Schritt: Deduktive Kodierung aller {len(batch)} Segmente...")

                # FIX: Escape-Prüfung vor Kodierung
                if self.check_escape_abort():
                    print("\n🛑 Abbruch vor Kodierung erkannt...")
                    await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
                    return current_categories, self.coding_results
                
                # Bestimme Kodiersystem je nach Modus
                if analysis_mode == 'inductive':
                    if len(current_categories) == 0:
                        coding_categories = {}
                        print(f"   📝 Inductive Mode: Keine induktiven Kategorien → 'Nicht kodiert'")
                    else:
                        coding_categories = current_categories
                        print(f"   📝 Inductive Mode: Verwende {len(current_categories)} induktive Kategorien")
                elif analysis_mode == 'grounded':
                    # FIX: Im grounded mode nur rein induktive Kategorien verwenden
                    grounded_categories = {}
                    for name, cat in current_categories.items():
                        if name not in DEDUKTIVE_KATEGORIEN:
                            grounded_categories[name] = cat
                    
                    coding_categories = grounded_categories
                    print(f"   📝 Grounded Mode: Verwende {len(grounded_categories)} rein induktive Kategorien")
                    print(f"   🚫 Ausgeschlossen: {len(current_categories) - len(grounded_categories)} deduktive Kategorien")
                else:
                    coding_categories = current_categories
                
                # Führe Kodierung durch
                if use_context:
                    batch_results = await self._code_batch_with_context(
                        batch, 
                        coding_categories,
                        category_preselections=category_preselections  # FIX: Neue Parameter
                    )
                else:
                    batch_results = await self._code_batch_deductively(
                        batch, 
                        coding_categories,
                        category_preselections=category_preselections  # FIX: Neue Parameter
                    )
            
                self.coding_results.extend(batch_results)
                
                # 4. Sättigungsprüfung
                batch_time = time.time() - batch_start
                material_percentage = (len(self.processed_segments) / total_segments) * 100
                total_batches = len(all_segments) / batch_size

                # Normale Sättigungsprüfung
                saturation_status = saturation_controller.assess_saturation(
                    current_categories=current_categories,
                    material_percentage=material_percentage,
                    batch_count=batch_count,
                    total_segments=self._total_segments
                )
            
                print(f"\n📊 Sättigungsstatus:")
                print(f"   🎯 Theoretische Sättigung: {saturation_status['theoretical_saturation']:.1%}")
                print(f"   📈 Materialabdeckung: {saturation_status['material_coverage']:.1%}")
                
                if saturation_status['is_saturated']:
                    print(f"\n🎯 SÄTTIGUNG ERREICHT nach {batch_count} Batches!")
                    break
                
                # Fortschrittsinfo
                print(f"\n📈 Fortschritt:")
                print(f"   - Verarbeitete Segmente: {len(self.processed_segments)}/{total_segments}")
                print(f"   - Aktuelle Kategorien: {len(current_categories)}")
                print(f"   - Kodierungen: {len(self.coding_results)}")
                print(f"   - Batch-Zeit: {batch_time:.2f}s")
                
            except Exception as e:
                print(f"Fehler bei Batch {batch_count}: {str(e)}")
                traceback.print_exc()
                continue

        # Finalisierung
        print(f"\n🏁 FINALISIERUNG ({analysis_mode.upper()} MODE):")

        final_categories = await self._finalize_by_mode(
            analysis_mode, current_categories, deductive_categories, initial_categories
        )

        # FIX: Escape-Prüfung
        if self.check_escape_abort():
            print("\n🛑 Abbruch nach Kodierungsverarbeitung erkannt...")
            await self._export_intermediate_results(chunks, current_categories, deductive_categories, initial_categories)
            return current_categories, self.coding_results
        
        
        # Zeige finale Statistiken
        self._show_final_development_stats(final_categories, initial_categories, batch_count)
        
        # KORRIGIERT: Stelle sicher, dass immer ein Tupel zurückgegeben wird
        return final_categories, self.coding_results
    
    async def _analyze_grounded_mode(self, chunks: Dict[str, List[str]], initial_categories: Dict, 
                                all_segments: List, batch_size: int) -> Tuple[Dict, List]:
        """
        NEUE METHODE: Separate Grounded Mode Analyse
        """
        print("\n🔄 GROUNDED MODE: Starte spezielle Subcode-Sammlung")
        
        if batch_size is None:
            batch_size = CONFIG.get('BATCH_SIZE', 5)
        
        # Initialisiere Grounded-spezifische Variablen
        self.grounded_subcodes_collection = []
        self.grounded_segment_analyses = []
        self.grounded_keywords_collection = []
        self.grounded_batch_history = []
        self.grounded_saturation_counter = 0
        
        batch_count = 0
        use_context = CONFIG.get('CODE_WITH_CONTEXT', False)
        
        # PHASE 1: NUR SUBCODE-SAMMLUNG (KEINE KODIERUNG)
        while True:
            if self.check_escape_abort():
                print("\n🛑 Abbruch durch Benutzer erkannt...")
                break
            
            batch = await self._get_next_batch(all_segments, batch_size)
            if not batch:
                break
                
            batch_count += 1
            material_percentage = (len(self.processed_segments) / len(all_segments)) * 100
            
            print(f"\n{'='*60}")
            print(f"📊 GROUNDED BATCH {batch_count}: {len(batch)} Segmente (NUR SUBCODE-SAMMLUNG)")
            print(f"📈 Material verarbeitet: {material_percentage:.1f}%")
            print(f"{'='*60}")
            
            # 1. Relevanzprüfung
            general_relevance_results = await self.relevance_checker.check_relevance_batch(batch)
            generally_relevant_batch = [
                (segment_id, text) for segment_id, text in batch 
                if general_relevance_results.get(segment_id, False)
            ]
            
            # Markiere Segmente als verarbeitet
            self.processed_segments.update(sid for sid, _ in batch)
            
            # 2. NUR SUBCODE-SAMMLUNG (KEINE KODIERUNG!)
            if generally_relevant_batch:
                relevant_texts = [text for _, text in generally_relevant_batch]
                
                # Grounded-Analyse für Subcodes
                grounded_analysis = await self.inductive_coder.analyze_grounded_batch(
                    segments=relevant_texts,
                    material_percentage=material_percentage
                )
                
                # Sammle Subcodes zentral
                self._collect_grounded_subcodes(grounded_analysis, batch_count)
            
            # 3. Sättigungsprüfung (nur für Subcode-Sammlung)
            if await self._assess_grounded_saturation(batch_count, len(all_segments) / batch_size):
                print(f"\n🛑 GROUNDED SUBCODE-SAMMLUNG abgeschlossen nach {batch_count} Batches!")
                break
        
        print(f"\n🎯 GROUNDED PHASE 1 ABGESCHLOSSEN:")
        print(f"   - Gesammelte Subcodes: {len(self.grounded_subcodes_collection)}")
        print(f"   - Segment-Analysen: {len(self.grounded_segment_analyses)}")
        print(f"   - Keywords: {len(self.grounded_keywords_collection)}")
        
        # PHASE 2: HAUPTKATEGORIEN GENERIEREN
        if len(self.grounded_subcodes_collection) >= 5:
            print(f"\n🔍 PHASE 2: Generiere Hauptkategorien aus Subcodes...")
            
            # Übergebe Subcodes an InductiveCoder
            self.inductive_coder.collected_subcodes = self.grounded_subcodes_collection
            self.inductive_coder.grounded_segment_analyses = self.grounded_segment_analyses
            
            # Generiere Hauptkategorien
            grounded_categories = await self.inductive_coder._generate_main_categories_from_subcodes(initial_categories)
            
            if grounded_categories:
                print(f"✅ {len(grounded_categories)} Hauptkategorien generiert")
                
                # Aktualisiere alle Kodierer mit den neuen Kategorien
                for coder in self.deductive_coders:
                    await coder.update_category_system(grounded_categories)
                
                # PHASE 3: KODIERUNG MIT GROUNDED KATEGORIEN
                print(f"\n🏷️ PHASE 3: Kodiere alle Segmente mit Grounded-Kategorien...")
                coding_results = await self._code_all_segments_with_grounded_categories(
                    all_segments, grounded_categories, use_context
                )
                
                self.coding_results = coding_results
                return grounded_categories, coding_results
            else:
                print("❌ Keine Hauptkategorien generiert - verwende initiale Kategorien")
                return initial_categories, []
        else:
            print(f"⚠️ Zu wenige Subcodes: {len(self.grounded_subcodes_collection)} < 5")
            return initial_categories, []

    def _collect_grounded_subcodes(self, grounded_analysis: Dict, batch_number: int):
        """
        NEUE METHODE: Sammle Subcodes aus Grounded-Analyse
        """
        new_subcodes_count = 0
        
        if grounded_analysis and 'segment_analyses' in grounded_analysis:
            print(f"📝 Verarbeite {len(grounded_analysis['segment_analyses'])} Segment-Analysen")
            
            # Speichere alle Segment-Analysen
            self.grounded_segment_analyses.extend(grounded_analysis['segment_analyses'])
            
            for segment_analysis in grounded_analysis['segment_analyses']:
                subcodes = segment_analysis.get('subcodes', [])
                
                for subcode in subcodes:
                    subcode_name = subcode.get('name', '').strip()
                    if subcode_name:
                        # Prüfe auf Duplikate
                        existing_names = [sc['name'] for sc in self.grounded_subcodes_collection]
                        
                        if subcode_name not in existing_names:
                            # Neuer Subcode
                            subcode_data = {
                                'name': subcode_name,
                                'definition': subcode.get('definition', ''),
                                'keywords': subcode.get('keywords', []),
                                'evidence': subcode.get('evidence', []),
                                'confidence': subcode.get('confidence', 0.7),
                                'batch_number': batch_number,
                                'source_segments': [segment_analysis.get('segment_text', '')[:100]]
                            }
                            
                            self.grounded_subcodes_collection.append(subcode_data)
                            new_subcodes_count += 1
                            
                            # Sammle Keywords
                            self.grounded_keywords_collection.extend(subcode.get('keywords', []))
                            
                            print(f"    ✅ Neuer Subcode: '{subcode_name}'")
                        else:
                            # Erweitere bestehenden Subcode
                            for existing_subcode in self.grounded_subcodes_collection:
                                if existing_subcode['name'] == subcode_name:
                                    # Erweitere ohne Duplikate
                                    existing_keywords = set(existing_subcode['keywords'])
                                    new_keywords = set(subcode.get('keywords', []))
                                    existing_subcode['keywords'] = list(existing_keywords | new_keywords)
                                    
                                    existing_subcode['evidence'].extend(subcode.get('evidence', []))
                                    existing_subcode['source_segments'].append(
                                        segment_analysis.get('segment_text', '')[:100]
                                    )
                                    
                                    # Aktualisiere Konfidenz
                                    old_conf = existing_subcode.get('confidence', 0.7)
                                    new_conf = subcode.get('confidence', 0.7)
                                    existing_subcode['confidence'] = (old_conf + new_conf) / 2
                                    
                                    print(f"    🔄 Subcode erweitert: '{subcode_name}'")
                                    break
        
        # Aktualisiere Sättigungszähler
        if new_subcodes_count == 0:
            self.grounded_saturation_counter += 1
        else:
            self.grounded_saturation_counter = 0
        
        # Speichere Batch-Historie
        batch_info = {
            'batch_number': batch_number,
            'new_subcodes': new_subcodes_count,
            'total_subcodes': len(self.grounded_subcodes_collection),
            'material_percentage': (len(self.processed_segments) / self._total_segments) * 100
        }
        self.grounded_batch_history.append(batch_info)
        
        print(f"✅ SUBCODE-SAMMLUNG BATCH {batch_number}:")
        print(f"   - Neue Subcodes: {new_subcodes_count}")
        print(f"   - Gesamt Subcodes: {len(self.grounded_subcodes_collection)}")
        print(f"   - Sättigungs-Counter: {self.grounded_saturation_counter}")

    async def _code_all_segments_with_grounded_categories(self, all_segments: List, 
                                                        grounded_categories: Dict, 
                                                        use_context: bool) -> List[Dict]:
        """
        NEUE METHODE: Kodiere alle Segmente mit den generierten Grounded-Kategorien
        """
        print(f"🏷️ Kodiere {len(all_segments)} Segmente mit {len(grounded_categories)} Grounded-Kategorien")
        
        coding_results = []
        batch_size = CONFIG.get('BATCH_SIZE', 5)
        
        # Erstelle Batches für die Kodierung
        for i in range(0, len(all_segments), batch_size):
            batch = all_segments[i:i + batch_size]
            print(f"   Kodiere Batch {i//batch_size + 1}: {len(batch)} Segmente")
            
            if use_context:
                batch_results = await self._code_batch_with_context(batch, grounded_categories)
            else:
                batch_results = await self._code_batch_deductively(batch, grounded_categories)
            
            coding_results.extend(batch_results)
            
            # Markiere als verarbeitet (falls noch nicht geschehen)
            for segment_id, _ in batch:
                self.processed_segments.add(segment_id)
        
        print(f"✅ Kodierung abgeschlossen: {len(coding_results)} Kodierungen erstellt")
        return coding_results
    
    async def _recode_segments_with_final_categories(self, final_categories: Dict[str, CategoryDefinition], chunks: Dict[str, List[str]]) -> None:
        """
        GROUNDED MODE: Kodiere alle Segmente nachträglich mit generierten Hauptkategorien
        """
        print(f"\n🔄 GROUNDED MODE: Nachträgliche Kodierung mit {len(final_categories)} Kategorien")
        
        # Aktualisiere ALLE Kodierer mit finalen Kategorien
        for coder in self.deductive_coders:
            success = await coder.update_category_system(final_categories)
            if success:
                print(f"   ✅ Kodierer {coder.coder_id} erfolgreich aktualisiert")
            else:
                print(f"   ❌ Fehler bei Kodierer {coder.coder_id}")
        
        # Rekonstruiere alle Segmente
        all_segments_to_recode = []
        for doc_name, doc_chunks in chunks.items():
            for chunk_id, chunk_text in enumerate(doc_chunks):
                segment_id = f"{doc_name}_chunk_{chunk_id}"
                all_segments_to_recode.append((segment_id, chunk_text))
        
        print(f"📊 Kodiere {len(all_segments_to_recode)} Segmente mit Grounded-Kategorien")
        
        # Kodiere in Batches
        new_codings = []
        batch_size = 5
        
        for i in range(0, len(all_segments_to_recode), batch_size):
            batch = all_segments_to_recode[i:i + batch_size]
            print(f"   Batch {i//batch_size + 1}: {len(batch)} Segmente")
            
            for segment_id, segment_text in batch:
                try:
                    # Kodiere mit dem ersten Kodierer (stellvertretend)
                    coding_result = await self.deductive_coders[0].code_chunk(segment_text, final_categories)
                    
                    if coding_result and coding_result.category != 'Nicht kodiert':
                        new_coding = {
                            'segment_id': segment_id,
                            'coder_id': self.deductive_coders[0].coder_id,
                            'category': coding_result.category,
                            'subcategories': list(coding_result.subcategories),
                            'confidence': coding_result.confidence,
                            'justification': f"[Grounded Theory Nachkodierung] {coding_result.justification}",
                            'text': segment_text,
                            'paraphrase': getattr(coding_result, 'paraphrase', ''),
                            'keywords': getattr(coding_result, 'keywords', ''),
                            'grounded_recoded': True,
                            'multiple_coding_instance': 1,
                            'total_coding_instances': 1,
                            'target_category': '',
                            'category_focus_used': False
                        }
                        new_codings.append(new_coding)
                    else:
                        # "Nicht kodiert" Fallback
                        new_codings.append({
                            'segment_id': segment_id,
                            'coder_id': self.deductive_coders[0].coder_id,
                            'category': 'Nicht kodiert',
                            'subcategories': [],
                            'confidence': {'total': 1.0},
                            'justification': "Nicht relevant für Grounded-Kategorien",
                            'text': segment_text,
                            'grounded_recoded': False
                        })
                        
                except Exception as e:
                    print(f"      ❌ Fehler bei {segment_id}: {str(e)}")
                    continue
        
        # KRITISCH: Ersetze coding_results komplett
        if new_codings:
            print(f"🔄 Ersetze {len(self.coding_results)} alte durch {len(new_codings)} neue Kodierungen")
            self.coding_results = new_codings
            
            # Statistiken
            from collections import Counter
            category_dist = Counter(coding.get('category', 'Unbekannt') for coding in new_codings)
            print(f"\n📈 Kategorienverteilung nach Grounded-Nachkodierung:")
            for cat, count in category_dist.most_common():
                percentage = (count / len(new_codings)) * 100
                print(f"   - {cat}: {count} ({percentage:.1f}%)")
        else:
            print(f"❌ Keine Nachkodierungen erstellt")
    
    def _show_grounded_mode_statistics(self):
        """
        Zeigt detaillierte Statistiken für den Grounded Mode
        """
        if not hasattr(self, 'collected_subcodes'):
            return
            
        print(f"\n📊 GROUNDED MODE STATISTIKEN:")
        print(f"{'='*50}")
        
        # Subcode-Statistiken
        print(f"📝 Subcode-Sammlung:")
        print(f"   - Gesammelte Subcodes: {len(self.collected_subcodes)}")
        print(f"   - Segment-Analysen: {len(self.grounded_segment_analyses)}")
        
        if self.collected_subcodes:
            # Konfidenz-Analyse
            confidences = [sc.get('confidence', 0) for sc in self.collected_subcodes]
            avg_confidence = sum(confidences) / len(confidences)
            print(f"   - Durchschnittliche Konfidenz: {avg_confidence:.2f}")
            
            # Keywords-Analyse
            all_keywords = []
            for sc in self.collected_subcodes:
                all_keywords.extend(sc.get('keywords', []))
            
            keyword_counter = Counter(all_keywords)
            print(f"   - Einzigartige Keywords: {len(set(all_keywords))}")
            print(f"   - Keywords gesamt: {len(all_keywords)}")
            
            # Top Keywords
            top_keywords = keyword_counter.most_common(10)
            print(f"   - Top Keywords: {', '.join([f'{kw}({count})' for kw, count in top_keywords[:5]])}")
            
            # Subcode-Batch-Verteilung
            batch_dist = Counter(sc.get('batch_number', 0) for sc in self.collected_subcodes)
            print(f"   - Verteilung über Batches: {dict(batch_dist)}")

    def _export_grounded_mode_details(self, output_dir: str):
        """
        Exportiert detaillierte Grounded Mode Daten für weitere Analyse
        """
        if not hasattr(self, 'collected_subcodes') or not self.collected_subcodes:
            return
            
        try:
            # Exportiere Subcodes als JSON
            subcodes_path = os.path.join(output_dir, f"grounded_subcodes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            
            export_data = {
                'metadata': {
                    'created': datetime.now().isoformat(),
                    'total_subcodes': len(self.collected_subcodes),
                    'total_segment_analyses': len(self.grounded_segment_analyses),
                    'analysis_mode': 'grounded'
                },
                'subcodes': self.collected_subcodes,
                'segment_analyses': self.grounded_segment_analyses[:100]  # Nur erste 100 für Größe
            }
            
            with open(subcodes_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
                
            print(f"\n📁 Grounded Mode Details exportiert: {subcodes_path}")
            
        except Exception as e:
            print(f"⚠️ Fehler beim Export der Grounded Mode Details: {str(e)}")

    def check_escape_abort(self) -> bool:
        """Prüft ob durch Escape abgebrochen werden soll"""
        return (getattr(self, '_should_abort', False) or 
                getattr(self, '_escape_abort_requested', False) or
                (hasattr(self, 'escape_handler') and self.escape_handler.should_abort()))
    
    async def _export_intermediate_results(self, chunks, current_categories, 
                                         deductive_categories, initial_categories):
        """Exportiert Zwischenergebnisse bei Abbruch"""
        try:
            if not hasattr(self, 'end_time') or self.end_time is None:
                self.end_time = datetime.now()

            print("\n📊 Exportiere Zwischenergebnisse...")
            
            # Erstelle einen speziellen Exporter für Zwischenergebnisse
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            exporter = ResultsExporter(
                output_dir=CONFIG['OUTPUT_DIR'],
                attribute_labels=CONFIG['ATTRIBUTE_LABELS'],
                analysis_manager=self
            )
            
            # Speichere Zwischenkategorien
            if current_categories:
                category_manager = CategoryManager(CONFIG['OUTPUT_DIR'])
                category_manager.save_codebook(
                    categories=current_categories,
                    filename=f"codebook_intermediate_{timestamp}.json"
                )
                print(f"📁 Zwischenkategorien gespeichert: {len(current_categories)} Kategorien")
            
            # Exportiere Zwischenkodierungen falls vorhanden
            if self.coding_results:
                print(f"📊 Exportiere {len(self.coding_results)} Zwischenkodierungen...")
                
                # Revision Manager für Export
                revision_manager = CategoryRevisionManager(
                    output_dir=CONFIG['OUTPUT_DIR'],
                    config=CONFIG
                )
                
                # Berechne eine grobe Reliabilität für Zwischenergebnisse
                reliability = 0.8  # Placeholder
                
                await exporter.export_results(
                    codings=self.coding_results,
                    reliability=reliability,
                    categories=current_categories,
                    chunks=chunks,
                    revision_manager=revision_manager,
                    export_mode="consensus",
                    original_categories=initial_categories,
                    document_summaries=getattr(self, 'document_summaries', None),
                    is_intermediate_export=True  # FIX: Kennzeichnung als Zwischenexport
                )
                
                print("✅ Zwischenergebnisse erfolgreich exportiert!")
                print(f"📂 Dateien im Ordner: {CONFIG['OUTPUT_DIR']}")
                print(f"📄 Export-Datei: QCA-AID_Analysis_INTERMEDIATE_{timestamp}.xlsx")
            else:
                print("⚠️  Keine Kodierungen zum Exportieren vorhanden")
                
        except Exception as e:
            print(f"❌ Fehler beim Export der Zwischenergebnisse: {str(e)}")
            import traceback
            traceback.print_exc()

    def _merge_category_systems(self, 
                            current: Dict[str, CategoryDefinition], 
                            new: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        Führt bestehendes und neues Kategoriensystem zusammen.
        
        Args:
            current: Bestehendes Kategoriensystem
            new: Neue Kategorien
            
        Returns:
            Dict[str, CategoryDefinition]: Zusammengeführtes System
        """
        merged = current.copy()
        
        for name, category in new.items():
            if name not in merged:
                # Komplett neue Kategorie
                merged[name] = category
                print(f"\n🆕 Neue Hauptkategorie hinzugefügt: {name}")
                print(f"   Definition: {category.definition[:100]}...")
                if category.subcategories:
                    print("   Subkategorien:")
                    for sub_name in category.subcategories.keys():
                        print(f"   - {sub_name}")
            else:
                # Bestehende Kategorie aktualisieren
                existing = merged[name]
                
                # Sammle Änderungen für Debug-Ausgabe
                changes = []
                
                # Prüfe auf neue/geänderte Definition
                new_definition = category.definition
                if len(new_definition) > len(existing.definition):
                    changes.append("Definition aktualisiert")
                else:
                    new_definition = existing.definition
                
                # Kombiniere Beispiele
                new_examples = list(set(existing.examples) | set(category.examples))
                if len(new_examples) > len(existing.examples):
                    changes.append(f"{len(new_examples) - len(existing.examples)} neue Beispiele")
                
                # Kombiniere Regeln
                new_rules = list(set(existing.rules) | set(category.rules))
                if len(new_rules) > len(existing.rules):
                    changes.append(f"{len(new_rules) - len(existing.rules)} neue Regeln")
                
                # Neue Subkategorien
                new_subcats = {**existing.subcategories, **category.subcategories}
                if len(new_subcats) > len(existing.subcategories):
                    changes.append(f"{len(new_subcats) - len(existing.subcategories)} neue Subkategorien")
                
                # Erstelle aktualisierte Kategorie
                merged[name] = CategoryDefinition(
                    name=name,
                    definition=new_definition,
                    examples=new_examples,
                    rules=new_rules,
                    subcategories=new_subcats,
                    added_date=existing.added_date,
                    modified_date=datetime.now().strftime("%Y-%m-%d")
                )
                
                if changes:
                    print(f"\n📝 Kategorie '{name}' aktualisiert:")
                    for change in changes:
                        print(f"   - {change}")
        
        return merged

    

    def _prepare_segments(self, chunks: Dict[str, List[str]]) -> List[Tuple[str, str]]:
        """
        Bereitet die Segmente für die Analyse vor.
        
        Args:
            chunks: Dictionary mit Dokumenten-Chunks
            
        Returns:
            List[Tuple[str, str]]: Liste von (segment_id, text) Tupeln
        """
        segments = []
        for doc_name, doc_chunks in chunks.items():
            for chunk_idx, chunk in enumerate(doc_chunks):
                segment_id = f"{doc_name}_chunk_{chunk_idx}"
                segments.append((segment_id, chunk))
        return segments


    def _find_similar_category(self, 
                                category: CategoryDefinition,
                                existing_categories: Dict[str, CategoryDefinition]) -> Optional[str]:
        """
        Findet ähnliche existierende Kategorien basierend auf Namen und Definition.
        
        Args:
            category: Zu prüfende Kategorie
            existing_categories: Bestehendes Kategoriensystem
            
        Returns:
            Optional[str]: Name der ähnlichsten Kategorie oder None
        """
        try:
            best_match = None
            highest_similarity = 0.0
            
            for existing_name, existing_cat in existing_categories.items():
                # Berechne Ähnlichkeit basierend auf verschiedenen Faktoren
                
                # 1. Name-Ähnlichkeit (gewichtet: 0.3)
                name_similarity = self.inductive_coder._calculate_text_similarity(
                    category.name.lower(),
                    existing_name.lower()
                ) * 0.3
                
                # 2. Definitions-Ähnlichkeit (gewichtet: 0.5)
                definition_similarity = self.inductive_coder._calculate_text_similarity(
                    category.definition,
                    existing_cat.definition
                ) * 0.5
                
                # 3. Subkategorien-Überlappung (gewichtet: 0.2)
                subcats1 = set(category.subcategories.keys())
                subcats2 = set(existing_cat.subcategories.keys())
                if subcats1 and subcats2:
                    subcat_overlap = len(subcats1 & subcats2) / len(subcats1 | subcats2)
                else:
                    subcat_overlap = 0
                subcat_similarity = subcat_overlap * 0.2
                
                # Gesamtähnlichkeit
                total_similarity = name_similarity + definition_similarity + subcat_similarity
                
                # Debug-Ausgabe für hohe Ähnlichkeiten
                if total_similarity > 0.5:
                    print(f"\nÄhnlichkeitsprüfung für '{category.name}' und '{existing_name}':")
                    print(f"- Name-Ähnlichkeit: {name_similarity:.2f}")
                    print(f"- Definitions-Ähnlichkeit: {definition_similarity:.2f}")
                    print(f"- Subkategorien-Überlappung: {subcat_similarity:.2f}")
                    print(f"- Gesamt: {total_similarity:.2f}")
                
                # Update beste Übereinstimmung
                if total_similarity > highest_similarity:
                    highest_similarity = total_similarity
                    best_match = existing_name
            
            # Nur zurückgeben wenn Ähnlichkeit hoch genug
            if highest_similarity > 0.7:  # Schwellenwert für Ähnlichkeit
                print(f"\n⚠ Hohe Ähnlichkeit ({highest_similarity:.2f}) gefunden:")
                print(f"- Neue Kategorie: {category.name}")
                print(f"- Existierende Kategorie: {best_match}")
                return best_match
                
            return None
            
        except Exception as e:
            print(f"Fehler bei Ähnlichkeitsprüfung: {str(e)}")
            return None

    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Berechnet die Ähnlichkeit zwischen zwei Texten mit Caching."""
        cache_key = f"{hash(text1)}_{hash(text2)}"
        
        if cache_key in self.similarity_cache:
            return self.similarity_cache[cache_key]
            
        # Konvertiere Texte zu Sets von Wörtern
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        # Berechne Jaccard-Ähnlichkeit
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        similarity = intersection / union if union > 0 else 0.0
        
        # Cache das Ergebnis
        self.similarity_cache[cache_key] = similarity
        self.validation_stats['similarity_calculations'] += 1
        
        return similarity
   
   
    def _prepare_segments(self, chunks: Dict[str, List[str]]) -> List[Tuple[str, str]]:
        """
        Bereitet die Segmente für die Analyse vor.
        
        Args:
            chunks: Dictionary mit Dokumenten-Chunks
            
        Returns:
            List[Tuple[str, str]]: Liste von (segment_id, text) Tupeln
        """
        segments = []
        for doc_name, doc_chunks in chunks.items():
            for chunk_idx, chunk in enumerate(doc_chunks):
                segment_id = f"{doc_name}_chunk_{chunk_idx}"
                segments.append((segment_id, chunk))
        return segments

    def _update_tracking(self,
                        batch_results: List[Dict],
                        processing_time: float,
                        batch_size: int) -> None:
        """
        Aktualisiert die Performance-Metriken.
        
        Args:
            batch_results: Ergebnisse des Batches
            processing_time: Verarbeitungszeit
            batch_size: Größe des Batches
        """
        self.coding_results.extend(batch_results)
        self.performance_metrics['batch_processing_times'].append(processing_time)
        
        # Berechne durchschnittliche Zeit pro Segment
        avg_time_per_segment = processing_time / batch_size
        self.performance_metrics['coding_times'].append(avg_time_per_segment)

    def _log_iteration_status(self,
                            material_percentage: float,
                            saturation_metrics: Dict,
                            num_results: int) -> None:
        """
        Protokolliert den Status der aktuellen Iteration.
        
        Args:
            material_percentage: Prozentsatz des verarbeiteten Materials
            saturation_metrics: Metriken der Sättigungsprüfung
            num_results: Anzahl der Kodierungen
        """
        try:
            # Erstelle Status-Dictionary
            status = {
                'timestamp': datetime.now().isoformat(),
                'material_processed': material_percentage,
                'saturation_metrics': saturation_metrics,
                'results_count': num_results,
                'processing_time': self.performance_metrics['batch_processing_times'][-1] if self.performance_metrics['batch_processing_times'] else 0
            }
            
            # Füge Status zum Log hinzu
            self.analysis_log.append(status)
            
            # Debug-Ausgabe für wichtige Metriken
            print("\nIterations-Status:")
            print(f"- Material verarbeitet: {material_percentage:.1f}%")
            print(f"- Neue Kodierungen: {num_results}")
            print(f"- Verarbeitungszeit: {status['processing_time']:.2f}s")
            if saturation_metrics:
                print("- Sättigungsmetriken:")
                for key, value in saturation_metrics.items():
                    print(f"  • {key}: {value}")
        except Exception as e:
            print(f"Warnung: Fehler beim Logging des Iterationsstatus: {str(e)}")
            # Fehler beim Logging sollte die Hauptanalyse nicht unterbrechen

    def _finalize_analysis(self,
                          final_categories: Dict,
                          initial_categories: Dict) -> Tuple[Dict, List]:
        """
        Schließt die Analyse ab und bereitet die Ergebnisse vor.
        
        Args:
            final_categories: Finales Kategoriensystem
            initial_categories: Initiales Kategoriensystem
            
        Returns:
            Tuple[Dict, List]: (Finales Kategoriensystem, Kodierungsergebnisse)
        """
        # Berechne finale Statistiken
        end_time = datetime.now()
        total_time = (end_time - self.start_time).total_seconds()
        avg_time_per_segment = total_time / len(self.processed_segments)
        
        print("\nAnalyse abgeschlossen:")
        print(f"- Gesamtzeit: {total_time:.2f}s")
        print(f"- Durchschnittliche Zeit pro Segment: {avg_time_per_segment:.2f}s")
        print(f"- Verarbeitete Segmente: {len(self.processed_segments)}")
        print(f"- Finale Kategorien: {len(final_categories)}")
        print(f"- Gesamtanzahl Kodierungen: {len(self.coding_results)}")
        
        return final_categories, self.coding_results

    def get_analysis_report(self) -> Dict:
        """Erstellt einen detaillierten Analysebericht."""
        return {
            'start_time': self.start_time.isoformat() if self.start_time else None,
            'end_time': self.end_time.isoformat() if self.end_time else None,
            'total_segments': len(self.processed_segments),
            'total_codings': len(self.coding_results),
            'performance_metrics': {
                'avg_batch_time': (
                    sum(self.performance_metrics['batch_processing_times']) / 
                    len(self.performance_metrics['batch_processing_times'])
                ) if self.performance_metrics['batch_processing_times'] else 0,
                'total_batches': len(self.performance_metrics['batch_processing_times']),
                'coding_distribution': self._get_coding_distribution()
            },
            'coding_summary': self._get_coding_summary()
        }
    
    def _get_coding_distribution(self) -> Dict:
        """Analysiert die Verteilung der Kodierungen."""
        if not self.coding_results:
            return {}
            
        coders = Counter(coding['coder_id'] for coding in self.coding_results)
        categories = Counter(coding['category'] for coding in self.coding_results)
        
        return {
            'coders': dict(coders),
            'categories': dict(categories)
        }
    
    def _get_coding_summary(self) -> Dict:
        """Erstellt eine Zusammenfassung der Kodierungen."""
        if not self.coding_results:
            return {}
            
        return {
            'total_segments_coded': len(set(coding['segment_id'] for coding in self.coding_results)),
            'avg_codings_per_segment': len(self.coding_results) / len(self.processed_segments) if self.processed_segments else 0,
            'unique_categories': len(set(coding['category'] for coding in self.coding_results))
        }

    def get_progress_report(self) -> Dict:
        """
        Erstellt einen detaillierten Fortschrittsbericht für die laufende Analyse.
        
        Returns:
            Dict: Fortschrittsbericht mit aktuellen Metriken und Status
        """
        current_time = datetime.now()
        elapsed_time = (current_time - self.start_time).total_seconds()
        
        # Berechne durchschnittliche Verarbeitungszeiten
        avg_batch_time = statistics.mean(self.performance_metrics['batch_processing_times']) if self.performance_metrics['batch_processing_times'] else 0
        avg_coding_time = statistics.mean(self.performance_metrics['coding_times']) if self.performance_metrics['coding_times'] else 0
        
                
        # Berechne Verarbeitungsstatistiken
        total_segments = len(self.processed_segments)
        segments_per_hour = (total_segments / elapsed_time) * 3600 if elapsed_time > 0 else 0
        
        return {
            'progress': {
                'processed_segments': total_segments,
                'total_codings': len(self.coding_results),
                'segments_per_hour': round(segments_per_hour, 2),
                'elapsed_time': round(elapsed_time, 2)
            },
            'performance': {
                'avg_batch_processing_time': round(avg_batch_time, 2),
                'avg_coding_time': round(avg_coding_time, 2),
                'last_batch_time': round(self.performance_metrics['batch_processing_times'][-1], 2) if self.performance_metrics['batch_processing_times'] else 0
            },
            'status': {
                'start_time': self.start_time.strftime('%Y-%m-%d %H:%M:%S'),
                'current_time': current_time.strftime('%Y-%m-%d %H:%M:%S'),
                'last_update': self.analysis_log[-1]['timestamp'] if self.analysis_log else None
            }
        }

class ImprovedSaturationController:
    """
    Verbesserte Sättigungskontrolle mit modusabhängigen Kriterien
    """
    
    def __init__(self, analysis_mode: str):
        self.analysis_mode = analysis_mode
        self.stability_counter = 0
        self.saturation_history = []
        
        # Modusabhängige Schwellenwerte
        if analysis_mode == 'inductive':
            self.min_batches = 5
            self.min_material = 0.7
            self.min_stability = 3
            self.min_theoretical = 0.8
        elif analysis_mode == 'abductive':
            self.min_batches = 3
            self.min_material = 0.6
            self.min_stability = 2
            self.min_theoretical = 0.7
        else:  # grounded
            self.min_batches = 4
            self.min_material = 0.8
            self.min_stability = 2
            self.min_theoretical = 0.75

    def assess_saturation(self, current_categories: Dict, material_percentage: float, 
                         batch_count: int, total_segments: int) -> Dict:
        """
        Umfassende Sättigungsbeurteilung
        """
        # Berechne theoretische Sättigung
        theoretical_saturation = self._calculate_theoretical_saturation(current_categories)
        
        # Berechne Kategorienqualität
        category_quality = self._assess_category_quality(current_categories)
        
        # Prüfe alle Kriterien
        criteria = {
            'min_batches': batch_count >= self.min_batches,
            'material_coverage': material_percentage >= (self.min_material * 100),
            'theoretical_saturation': theoretical_saturation >= self.min_theoretical,
            'category_quality': category_quality >= 0.7,
            'stability': self.stability_counter >= self.min_stability,
            'sufficient_categories': len(current_categories) >= 2
        }
        
        is_saturated = all(criteria.values())
        
        # Bestimme Sättigungsgrund
        if is_saturated:
            saturation_reason = "Alle Sättigungskriterien erfüllt"
        else:
            missing = [k for k, v in criteria.items() if not v]
            saturation_reason = f"Fehlende Kriterien: {', '.join(missing)}"
        
        return {
            'is_saturated': is_saturated,
            'theoretical_saturation': theoretical_saturation,
            'material_coverage': material_percentage / 100,
            'stability_batches': self.stability_counter,
            'category_quality': category_quality,
            'saturation_reason': saturation_reason,
            'criteria_met': criteria
        }

    def _calculate_theoretical_saturation(self, categories: Dict) -> float:
        """
        Berechnet theoretische Sättigung
        """
        if not categories:
            return 0.0
        
        # Kategorienreife
        maturity_scores = []
        for cat in categories.values():
            score = 0
            # Definition
            if hasattr(cat, 'definition') and len(cat.definition.split()) >= 15:
                score += 0.4
            # Beispiele
            if hasattr(cat, 'examples') and len(cat.examples) >= 1:
                score += 0.3
            # Subkategorien
            if hasattr(cat, 'subcategories') and len(cat.subcategories) >= 1:
                score += 0.3
            maturity_scores.append(score)
        
        avg_maturity = sum(maturity_scores) / len(maturity_scores) if maturity_scores else 0
        
        # Anzahl-Faktor
        optimal_count = 8 if self.analysis_mode == 'inductive' else 6
        count_factor = min(len(categories) / optimal_count, 1.0)
        
        return (avg_maturity * 0.7) + (count_factor * 0.3)

    def _assess_category_quality(self, categories: Dict) -> float:
        """
        Bewertet Kategorienqualität
        """
        if not categories:
            return 0.0
        
        quality_scores = []
        for cat in categories.values():
            score = 0
            if hasattr(cat, 'definition') and len(cat.definition.split()) >= 10:
                score += 0.5
            if hasattr(cat, 'examples') and len(cat.examples) >= 1:
                score += 0.3
            if hasattr(cat, 'subcategories') and len(cat.subcategories) >= 1:
                score += 0.2
            quality_scores.append(score)
        
        return sum(quality_scores) / len(quality_scores)

    def increment_stability_counter(self):
        """Erhöht Stabilitätszähler"""
        self.stability_counter += 1

    def reset_stability_counter(self):
        """Setzt Stabilitätszähler zurück"""
        self.stability_counter = 0

# --- Klasse: DeductiveCategoryBuilder ---
# Aufgabe: Ableiten deduktiver Kategorien basierend auf theoretischem Vorwissen
class DeductiveCategoryBuilder:
    """
    Baut ein initiales, theoriebasiertes Kategoriensystem auf.
    """
    def load_theoretical_categories(self) -> Dict[str, CategoryDefinition]:
        """
        Lädt die vordefinierten deduktiven Kategorien.
        
        Returns:
            Dict[str, CategoryDefinition]: Dictionary mit Kategorienamen und deren Definitionen
        """
        categories = {}
        today = datetime.now().strftime("%Y-%m-%d")
        
        try:
            # Konvertiere DEDUKTIVE_KATEGORIEN in CategoryDefinition-Objekte
            for name, data in DEDUKTIVE_KATEGORIEN.items():
                # Stelle sicher, dass rules als Liste vorhanden ist
                rules = data.get("rules", [])
                if not isinstance(rules, list):
                    rules = [str(rules)] if rules else []
                
                categories[name] = CategoryDefinition(
                    name=name,
                    definition=data.get("definition", ""),
                    examples=data.get("examples", []),
                    rules=rules,  # Übergebe die validierte rules Liste
                    subcategories=data.get("subcategories", {}),
                    added_date=today,
                    modified_date=today
                )
                
            return categories
            
        except Exception as e:
            print(f"Fehler beim Laden der Kategorien: {str(e)}")
            print("Aktuelle Kategorie:", name)
            print("Kategorie-Daten:", json.dumps(data, indent=2, ensure_ascii=False))
            raise



# --- Klasse: DeductiveCoder ---
# Aufgabe: Automatisches deduktives Codieren von Text-Chunks anhand des Leitfadens
class DeductiveCoder:
    """
    Ordnet Text-Chunks automatisch deduktive Kategorien zu basierend auf dem Kodierleitfaden.
    Nutzt GPT-4-Mini für die qualitative Inhaltsanalyse nach Mayring.
    """
    
    def __init__(self, model_name: str, temperature: str, coder_id: str, skip_inductive: bool = False):
        """
        Initializes the DeductiveCoder with configuration for the GPT model.
        
        Args:
            model_name (str): Name of the GPT model to use
            temperature (float): Controls randomness in the model's output 
                - Lower values (e.g., 0.3) make output more focused and deterministic
                - Higher values (e.g., 0.7) make output more creative and diverse
            coder_id (str): Unique identifier for this coder instance
        """
        self.model_name = model_name
        self.temperature = float(temperature)
        self.coder_id = coder_id
        self.skip_inductive = skip_inductive

        # self.load_theoretical_categories = DeductiveCategoryBuilder.load_theoretical_categories()
        category_builder = DeductiveCategoryBuilder()
        initial_categories = category_builder.load_theoretical_categories()

        # FIX: Initialisiere current_categories modus-abhängig
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'deductive')
        
        if analysis_mode == 'grounded':
            # Im grounded mode starten wir mit leeren Kategorien
            self.current_categories = {}
            print(f"📚 Kodierer {coder_id}: Grounded Mode - startet ohne deduktive Kategorien")
        else:
            # Alle anderen Modi laden deduktive Kategorien
            try:
                self.current_categories = initial_categories
                print(f"📚 Kodierer {coder_id}: {len(self.current_categories)} deduktive Kategorien geladen ({analysis_mode} mode)")
            except Exception as e:
                print(f"⚠️ Fehler beim Laden der Kategorien für Kodierer {coder_id}: {str(e)}")
                self.current_categories = {}
           
        # Hole Provider aus CONFIG
        provider_name = CONFIG.get('MODEL_PROVIDER', 'openai').lower()  # Fallback zu OpenAI
        try:
            # Wir lassen den LLMProvider sich selbst um die Client-Initialisierung kümmern
            self.llm_provider = LLMProviderFactory.create_provider(provider_name)
            print(f"🤖 LLM Provider '{provider_name}' für Kodierer {coder_id} initialisiert")
        except Exception as e:
            print(f"Fehler bei Provider-Initialisierung für {coder_id}: {str(e)}")
            raise
        
        # FIX: Prompt-Handler nur im grounded mode ohne deduktive Kategorien initialisieren
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'deductive')
        if analysis_mode == 'grounded':
            # FIX: Im grounded mode KEINE deduktiven Kategorien verwenden
            self.prompt_handler = QCAPrompts(
                forschungsfrage=FORSCHUNGSFRAGE,
                kodierregeln=KODIERREGELN,
                deduktive_kategorien={}  # FIX: Leeres Dict statt DEDUKTIVE_KATEGORIEN
            )
            print(f"   🔄 Grounded Mode: Kodierer {coder_id} ohne vordefinierte Kategorien initialisiert")
        else:
            # Normale Modi verwenden deduktive Kategorien
            self.prompt_handler = QCAPrompts(
                forschungsfrage=FORSCHUNGSFRAGE,
                kodierregeln=KODIERREGELN,
                deduktive_kategorien=DEDUKTIVE_KATEGORIEN
            )

    async def update_category_system(self, categories: Dict[str, CategoryDefinition]) -> bool:
        """
        Aktualisiert das Kategoriensystem des Kodierers.
        
        Args:
            categories: Neues/aktualisiertes Kategoriensystem
            
        Returns:
            bool: True wenn Update erfolgreich
        """
        try:
            # FIX: Prüfe Analysemodus vor Update
            analysis_mode = CONFIG.get('ANALYSIS_MODE', 'deductive')
            
            if analysis_mode == 'grounded':
                # FIX: Im grounded mode nur mit rein induktiven Kategorien arbeiten
                # Filtere alle deduktiven Kategorien heraus
                grounded_categories = {}
                for name, cat in categories.items():
                    if name not in DEDUKTIVE_KATEGORIEN:
                        grounded_categories[name] = cat
                
                print(f"   🔄 Grounded Mode: Kodierer {self.coder_id} aktualisiert mit {len(grounded_categories)} rein induktiven Kategorien")
                print(f"   📝 Ausgeschlossen: {len(categories) - len(grounded_categories)} deduktive Kategorien")
                
                # Verwende nur die gefilterten Kategorien
                categories_to_use = grounded_categories
            else:
                # Andere Modi verwenden alle Kategorien
                categories_to_use = categories
                print(f"   🔄 {analysis_mode.upper()} Mode: Kodierer {self.coder_id} aktualisiert mit {len(categories_to_use)} Kategorien")

            # Konvertiere CategoryDefinition in serialisierbares Dict
            categories_dict = {
                name: {
                    'definition': cat.definition,
                    'examples': list(cat.examples) if isinstance(cat.examples, set) else cat.examples,
                    'rules': list(cat.rules) if isinstance(cat.rules, set) else cat.rules,
                    'subcategories': dict(cat.subcategories) if isinstance(cat.subcategories, set) else cat.subcategories
                } for name, cat in categories_to_use.items()
            }

            # FIX: Aktualisiere das Kontextwissen des Kodierers modusabhängig
            if analysis_mode == 'grounded':
                prompt = f"""
                GROUNDED THEORY MODUS: Das Kategoriensystem wurde mit rein induktiven Kategorien aktualisiert.
                
                Aktuelle induktive Kategorien:
                {json.dumps(categories_dict, indent=2, ensure_ascii=False)}
                
                WICHTIGE REGELN FÜR GROUNDED KODIERUNG:
                1. Verwende NUR die oben aufgeführten induktiven Kategorien
                2. KEINE deduktiven Kategorien aus dem ursprünglichen Codebook verwenden
                3. Diese Kategorien wurden bottom-up aus den Texten entwickelt
                4. Kodiere nur dann, wenn das Segment eindeutig zu einer dieser Kategorien passt
                5. Bei Unsicherheit: "Nicht kodiert" verwenden
                
                Antworte einfach mit "Verstanden" wenn du die Anweisung erhalten hast.
                """
            else:
                prompt = f"""
                Das Kategoriensystem wurde aktualisiert. Neue Zusammensetzung:
                {json.dumps(categories_dict, indent=2, ensure_ascii=False)}
                
                Berücksichtige bei der Kodierung:
                1. Verwende alle verfügbaren Kategorien entsprechend ihrer Definitionen
                2. Prüfe auch Subkategorien bei der Zuordnung
                3. Kodiere nur bei eindeutiger Zuordnung
                
                Antworte einfach mit "Verstanden" wenn du die Anweisung erhalten hast.
                """

            # Aktualisiere internes Kategoriensystem
            # FIX: Verwende gefilterte Kategorien statt alle
            self.current_categories = categories_to_use
            
            # FIX: Aktualisiere auch den Prompt-Handler mit den richtigen Kategorien
            if analysis_mode == 'grounded':
                # FIX: Im grounded mode deduktive Kategorien leer lassen
                self.prompt_handler.deduktive_kategorien({})
            else:
                # Andere Modi behalten deduktive Kategorien
                self.prompt_handler.deduktive_kategorien(DEDUKTIVE_KATEGORIEN)

            # Test-Anfrage um sicherzustellen, dass das System bereit ist
            response = await self.llm_provider.chat(
                messages=[
                    {"role": "system", "content": "Du bist ein QCA-Kodierer. Antworte auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            if "verstanden" in response.content.lower():
                return True
            else:
                print(f"⚠️ Unerwartete Antwort von Kodierer {self.coder_id}: {response.content}")
                return True  # Fahre trotzdem fort
                
        except Exception as e:
            print(f"❌ Fehler beim Update von Kodierer {self.coder_id}: {str(e)}")
            return False

    async def code_chunk_with_context_switch(self, 
                                         chunk: str, 
                                         categories: Dict[str, CategoryDefinition],
                                         use_context: bool = True,
                                         context_data: Dict = None) -> Dict:
        """
        Wrapper-Methode, die basierend auf dem use_context Parameter
        zwischen code_chunk und code_chunk_with_progressive_context wechselt.
        
        Args:
            chunk: Zu kodierender Text
            categories: Kategoriensystem
            use_context: Ob Kontext verwendet werden soll
            context_data: Kontextdaten (current_summary und segment_info)
            
        Returns:
            Dict: Kodierungsergebnis (und ggf. aktualisiertes Summary)
        """
        if use_context and context_data:
            # Mit progressivem Kontext kodieren
            return await self.code_chunk_with_progressive_context(
                chunk, 
                categories, 
                context_data.get('current_summary', ''),
                context_data.get('segment_info', {})
            )
        else:
            # Klassische Kodierung ohne Kontext
            result = await self.code_chunk(chunk, categories)
            
            # Wandle CodingResult in Dictionary um, wenn nötig
            if result and isinstance(result, CodingResult):
                return {
                    'coding_result': result.to_dict(),
                    'updated_summary': context_data.get('current_summary', '') if context_data else ''
                }
            elif result and isinstance(result, dict):
                return {
                    'coding_result': result,
                    'updated_summary': context_data.get('current_summary', '') if context_data else ''
                }
            else:
                return None
            
    async def code_chunk(self, chunk: str, categories: Optional[Dict[str, CategoryDefinition]] = None, 
                        is_last_segment: bool = False, preferred_cats: Optional[List[str]] = None) -> Optional[CodingResult]:  # FIX: Neuer Parameter hinzugefügt
        """
        Kodiert einen Text-Chunk basierend auf dem aktuellen Kategoriensystem.
        
        Args:
            chunk: Zu kodierender Text
            categories: Kategoriensystem (optional, verwendet current_categories als Fallback)
            is_last_segment: Ob dies das letzte Segment ist
            preferred_cats: Liste bevorzugter Kategorien für gefilterte Kodierung  # FIX: Neue Funktionalität
            
        Returns:
            Optional[CodingResult]: Kodierungsergebnis oder None bei Fehlern
        """
        try:
            # Verwende das interne Kategoriensystem wenn vorhanden, sonst das übergebene
            current_categories = categories
            
            if not current_categories:
                print(f"Fehler: Kein Kategoriensystem für Kodierer {self.coder_id} verfügbar")
                return None

            # FIX: Kategorien-Filterung basierend auf preferred_cats
            if preferred_cats:
                # Filtere Kategorien basierend auf Vorauswahl
                filtered_categories = {
                    name: cat for name, cat in current_categories.items() 
                    if name in preferred_cats
                }
                
                if filtered_categories:
                    print(f"    🎯 Gefilterte Kodierung für {self.coder_id}: {len(filtered_categories)}/{len(current_categories)} Kategorien")
                    print(f"    📋 Fokus auf: {', '.join(preferred_cats)}")
                    effective_categories = filtered_categories
                else:
                    print(f"    ⚠️ Keine der bevorzugten Kategorien {preferred_cats} gefunden - nutze alle Kategorien")
                    effective_categories = current_categories
            else:
                # Standard-Verhalten: alle Kategorien verwenden
                effective_categories = current_categories
            # FIX: Ende der neuen Filterlogik

            # Erstelle formatierte Kategorienübersicht
            categories_overview = []
            for name, cat in effective_categories.items():  # FIX: Nutze gefilterte Kategorien
                category_info = {
                    'name': name,
                    'definition': cat.definition,
                    'examples': list(cat.examples) if isinstance(cat.examples, set) else cat.examples,
                    'rules': list(cat.rules) if isinstance(cat.rules, set) else cat.rules,
                    'subcategories': {}
                }
                
                # Füge Subkategorien hinzu
                for sub_name, sub_def in cat.subcategories.items():
                    category_info['subcategories'][sub_name] = sub_def
                    
                categories_overview.append(category_info)

            # Erstelle Prompt
            prompt = self.prompt_handler.get_deductive_coding_prompt(
                chunk=chunk,
                categories_overview=categories_overview
            )
            
            try:
                token_counter.start_request()

                response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=self.temperature,
                    response_format={"type": "json_object"}
                )
                
                # Verarbeite Response
                llm_response = LLMResponse(response)
                result = json.loads(llm_response.content)
                
                token_counter.track_response(response, self.model_name)
                
                if result and isinstance(result, dict):
                    # FIX: Strikte Subkategorie-Validierung anwenden
                    main_category = result.get('category', '')
                    original_subcats = result.get('subcategories', [])
                    
                    validated_subcats = CategoryValidator.validate_subcategories_for_category(
                        original_subcats, main_category, self.current_categories, warn_only=False
                    )
                    result['subcategories'] = validated_subcats
                    # FIX: Ende der Subkategorie-Validierung
                    
                    coding_result = CodingResult(
                        category=result.get('category', 'Nicht kodiert'),
                        subcategories=tuple(validated_subcats),  # ✅ Tuple statt Set
                        confidence=result.get('confidence', {}),
                        justification=result.get('justification', ''),
                        paraphrase=result.get('paraphrase', ''),
                        keywords=result.get('keywords', ''),
                        text_references=tuple(result.get('text_references', [])),  # ✅ Tuple
                        uncertainties=tuple(result.get('uncertainties', [])) if result.get('uncertainties') else None
                    )
                                        
                    return coding_result
                else:
                    print(f"Ungültige API-Antwort von {self.coder_id}")
                    return None
                    
            except json.JSONDecodeError as e:
                print(f"JSON-Fehler bei {self.coder_id}: {str(e)}")
                print(f"Rohe Antwort: {llm_response.content[:200]}...")
                return None
            except Exception as e:
                print(f"API-Fehler bei {self.coder_id}: {str(e)}")
                return None
                
        except Exception as e:
            print(f"Allgemeiner Fehler bei der Kodierung durch {self.coder_id}: {str(e)}")
            return None
    
    async def code_chunk_with_progressive_context(self, 
                                          chunk: str, 
                                          categories: Dict[str, CategoryDefinition],
                                          current_summary: str,
                                          segment_info: Dict) -> Dict:
        """
        Kodiert einen Text-Chunk und aktualisiert gleichzeitig das Dokument-Summary
        mit einem dreistufigen Reifungsmodell.
        
        Args:
            chunk: Zu kodierender Text
            categories: Kategoriensystem
            current_summary: Aktuelles Dokument-Summary
            segment_info: Zusätzliche Informationen über das Segment
            
        Returns:
            Dict: Enthält sowohl Kodierungsergebnis als auch aktualisiertes Summary
        """
        try:
            
            current_categories = categories
            
            if not current_categories:
                print(f"Fehler: Kein Kategoriensystem für Kodierer {self.coder_id} verfügbar")
                return None

            print(f"\nDeduktiver Kodierer 🧐 **{self.coder_id}** verarbeitet Chunk mit progressivem Kontext...")
            
            # Erstelle formatierte Kategorienübersicht
            categories_overview = []
            for name, cat in current_categories.items():
                category_info = {
                    'name': name,
                    'definition': cat.definition,
                    'examples': list(cat.examples) if isinstance(cat.examples, set) else cat.examples,
                    'rules': list(cat.rules) if isinstance(cat.rules, set) else cat.rules,
                    'subcategories': {}
                }
                
                # Füge Subkategorien hinzu
                for sub_name, sub_def in cat.subcategories.items():
                    category_info['subcategories'][sub_name] = sub_def
                    
                categories_overview.append(category_info)
            
            # Position im Dokument und Fortschritt berechnen
            position_info = f"Segment: {segment_info.get('position', '')}"
            doc_name = segment_info.get('doc_name', 'Unbekanntes Dokument')
            
            # Berechne die relative Position im Dokument (für das Reifungsmodell)
            chunk_id = 0
            total_chunks = 1
            if 'position' in segment_info:
                try:
                    # Extrahiere Chunk-Nummer aus "Chunk X"
                    chunk_id = int(segment_info['position'].split()[-1])
                    
                    # Schätze Gesamtanzahl der Chunks (basierend auf bisherigen Chunks)
                    # Alternative: Tatsächliche Anzahl übergeben, falls verfügbar
                    total_chunks = max(chunk_id * 1.5, 20)  # Schätzung
                    
                    document_progress = chunk_id / total_chunks
                    print(f"Dokumentfortschritt: ca. {document_progress:.1%}")
                except (ValueError, IndexError):
                    document_progress = 0.5  # Fallback
            else:
                document_progress = 0.5  # Fallback
                
            # Bestimme die aktuelle Reifephase basierend auf dem Fortschritt
            if document_progress < 0.3:
                reifephase = "PHASE 1 (Sammlung)"
                max_aenderung = "50%"
            elif document_progress < 0.7:
                reifephase = "PHASE 2 (Konsolidierung)"
                max_aenderung = "30%"
            else:
                reifephase = "PHASE 3 (Präzisierung)"
                max_aenderung = "10%"
                
            print(f"Summary-Reifephase: {reifephase}, max. Änderung: {max_aenderung}")
            
            # Angepasster Prompt basierend auf dem dreistufigen Reifungsmodell
            # Verbesserter summary_update_prompt für die _code_chunk_with_progressive_context Methode

            summary_update_prompt = f"""
            ## AUFGABE 2: SUMMARY-UPDATE ({reifephase}, {int(document_progress*100)}%)

            """

            # Robustere Phasen-spezifische Anweisungen
            if document_progress < 0.3:
                summary_update_prompt += """
            SAMMLUNG (0-30%) - STRUKTURIERTER AUFBAU:
            - SCHLÜSSELINFORMATIONEN: Beginne mit einer LISTE wichtigster Konzepte im Telegrammstil
            - FORMAT: "Thema1: Kernaussage; Thema2: Kernaussage" 
            - SPEICHERSTRUKTUR: Speichere alle Informationen in KATEGORIEN (z.B. Akteure, Prozesse, Faktoren)
            - KEINE EINLEITUNGEN oder narrative Elemente, NUR Fakten und Verbindungen
            - BEHALTE IMMER: Bereits dokumentierte Schlüsselkonzepte müssen bestehen bleiben
            """
            elif document_progress < 0.7:
                summary_update_prompt += """
            KONSOLIDIERUNG (30-70%) - HIERARCHISCHE ORGANISATION:
            - SCHLÜSSELINFORMATIONEN BEWAHREN: Alle bisherigen Hauptkategorien beibehalten
            - NEUE STRUKTUR: Als hierarchische Liste mit Kategorien und Unterpunkten organisieren
            - KOMPRIMIEREN: Details aus gleichen Themenbereichen zusammenführen
            - PRIORITÄTSFORMAT: "Kategorie: Hauptpunkt1; Hauptpunkt2 → Detail"
            - STATT LÖSCHEN: Verwandte Inhalte zusammenfassen, aber KEINE Kategorien eliminieren
            """
            else:
                summary_update_prompt += """
            PRÄZISIERUNG (70-100%) - VERDICHTUNG MIT THESAURUS:
            - THESAURUS-METHODE: Jede Kategorie braucht genau 1-2 Sätze im Telegrammstil
            - HAUPTKONZEPTE STABIL HALTEN: Alle identifizierten Kategorien müssen enthalten bleiben
            - ABSTRAHIEREN: Einzelinformationen innerhalb einer Kategorie verdichten
            - STABILITÄTSPRINZIP: Einmal erkannte wichtige Zusammenhänge dürfen nicht verloren gehen
            - PRIORITÄTSORDNUNG: Wichtigste Informationen IMMER am Anfang jeder Kategorie
            """

            # Allgemeine Kriterien für Stabilität und Komprimierung
            summary_update_prompt += """

            INFORMATIONSERHALTUNGS-SYSTEM:
            - MAXIMUM 80 WÖRTER - Komprimiere alte statt neue Informationen zu verwerfen
            - KATEGORIEBASIERT: Jedes Summary muss immer in 3-5 klare Themenkategorien strukturiert sein
            - SCHLÜSSELPRINZIP: Bilde das Summary als INFORMATIONALE HIERARCHIE:
            1. Stufe: Immer stabile Themenkategorien
            2. Stufe: Zentrale Aussagen zu jeder Kategorie
            3. Stufe: Ergänzende Details (diese können komprimiert werden)
            - STABILITÄTSGARANTIE: Neue Iteration darf niemals vorherige Kategorie-Level-1-Information verlieren
            - KOMPRIMIERUNGSSTRATEGIE: Bei Platzmangel Details (Stufe 3) zusammenfassen statt zu entfernen
            - FORMAT: "Kategorie1: Hauptpunkt; Hauptpunkt. Kategorie2: Hauptpunkt; Detail." (mit Doppelpunkten)
            - GRUNDREGEL: Neue Informationen ergänzen bestehende Kategorien statt sie zu ersetzen
            """
            
            # Prompt mit erweiterter Aufgabe für Summary-Update
            prompt = self.prompt_handler.get_progressive_context_prompt(
                chunk=chunk,
                categories_overview=categories_overview,
                current_summary=current_summary,
                position_info=position_info,
                summary_update_prompt=summary_update_prompt
            )
            
            # API-Call
            token_counter.start_request()

            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                response_format={"type": "json_object"}
            )
            
            # Verarbeite Response mit Wrapper
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)

            
            token_counter.track_response(response, self.model_name)
            
            # Extrahiere relevante Teile
            if result and isinstance(result, dict):
                coding_result = result.get('coding_result', {})
                updated_summary = result.get('updated_summary', current_summary)
                
                # Prüfe Wortlimit beim Summary
                if len(updated_summary.split()) > 80:  # Etwas Spielraum über 70
                    words = updated_summary.split()
                    updated_summary = ' '.join(words[:70])
                    print(f"⚠️ Summary wurde gekürzt: {len(words)} → 70 Wörter")
                
                # Analyse der Veränderungen
                if current_summary:
                    # Berechne Prozent der Änderung
                    old_words = set(current_summary.lower().split())
                    new_words = set(updated_summary.lower().split())
                    
                    if old_words:
                        # Jaccard-Distanz als Maß für Veränderung
                        unchanged = len(old_words.intersection(new_words))
                        total = len(old_words.union(new_words))
                        change_percent = (1 - (unchanged / total)) * 100
                        
                        print(f"Summary Änderung: {change_percent:.1f}% (Ziel: max. {max_aenderung})")
                
                if coding_result:
                    paraphrase = coding_result.get('paraphrase', '')
                    if paraphrase:
                        print(f"\n🗒️  Paraphrase: {paraphrase}")
                    print(f"  ✓ Kodierung von {self.coder_id}: 🏷️  {coding_result.get('category', '')}")
                    print(f"  ✓ Subkategorien von {self.coder_id}: 🏷️  {', '.join(coding_result.get('subcategories', []))}")
                    print(f"  ✓ Keywords von {self.coder_id}: 🏷️  {coding_result.get('keywords', '')}")
                    print(f"\n📝 Summary für {doc_name} aktualisiert ({len(updated_summary.split())} Wörter):")
                    print(f"{updated_summary[:1000]}..." if len(updated_summary) > 100 else f"📄 {updated_summary}")
                    
                    # Kombiniertes Ergebnis zurückgeben
                    return {
                        'coding_result': coding_result,
                        'updated_summary': updated_summary
                    }
                else:
                    print(f"  ✗ Keine gültige Kodierung erhalten")
                    return None
            else:
                print("  ✗ Keine gültige Antwort erhalten")
                return None
                
        except Exception as e:
            print(f"Fehler bei der Kodierung durch {self.coder_id}: {str(e)}")
            print("Details:")
            import traceback
            traceback.print_exc()
            return None

    async def code_chunk_with_focus(self, chunk: str, categories: Dict[str, CategoryDefinition], 
                                focus_category: str, focus_context: Dict) -> Optional[CodingResult]:
        """
        Kodiert einen Text-Chunk mit Fokus auf eine bestimmte Kategorie (für Mehrfachkodierung).
        
        Args:
            chunk: Zu kodierender Text
            categories: Kategoriensystem  
            focus_category: Kategorie auf die fokussiert werden soll
            focus_context: Kontext zur fokussierten Kategorie mit 'justification', 'text_aspects', 'relevance_score'
            
        Returns:
            Optional[CodingResult]: Kodierungsergebnis mit Fokus-Kennzeichnung
        """
        
        try:
            if not categories:
                print(f"Fehler: Kein Kategoriensystem für Kodierer {self.coder_id} verfügbar")
                return None

            # print(f"    🎯 Fokuskodierung für Kategorie: {focus_category} (Relevanz: {focus_context.get('relevance_score', 0):.2f})")

            # Erstelle formatierte Kategorienübersicht mit Fokus-Hervorhebung
            categories_overview = []
            for name, cat in categories.items():  
                # Hebe Fokus-Kategorie hervor
                display_name = name
                if name == focus_category:
                    display_name = name
                
                category_info = {
                    'name': display_name,
                    'definition': cat.definition,
                    'examples': list(cat.examples) if isinstance(cat.examples, set) else cat.examples,
                    'rules': list(cat.rules) if isinstance(cat.rules, set) else cat.rules,
                    'subcategories': {}
                }
                
                # Füge Subkategorien hinzu
                for sub_name, sub_def in cat.subcategories.items():
                    category_info['subcategories'][sub_name] = sub_def
                    
                categories_overview.append(category_info)

            # Erstelle fokussierten Prompt
            prompt = self.prompt_handler.get_focus_coding_prompt(
                chunk=chunk,
                categories_overview=categories_overview,
                focus_category=focus_category,
                focus_context=focus_context
            )
            
            try:
                token_counter.start_request()

                response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=self.temperature,
                    response_format={"type": "json_object"}
                )
                
                # Verarbeite Response mit Wrapper
                llm_response = LLMResponse(response)
                result = json.loads(llm_response.content)

                
                token_counter.track_response(response, self.model_name)
                
                if result and isinstance(result, dict):
                    if result.get('category'):
                        # Verarbeite Paraphrase
                        paraphrase = result.get('paraphrase', '')
                        if paraphrase:
                            print(f"      🗒️  Fokus-Paraphrase: {paraphrase}")

                        # Dokumentiere Fokus-Adherence
                        focus_adherence = result.get('focus_adherence', {})
                        followed_focus = focus_adherence.get('followed_focus', True)
                        focus_icon = "🎯" if followed_focus else "🔄"
                        
                        print(f"      {focus_icon} Fokuskodierung von {self.coder_id}: 🏷️  {result.get('category', '')}")
                        print(f"      ✓ Subkategorien: 🏷️  {', '.join(result.get('subcategories', []))}")
                        print(f"      ✓ Keywords: 🏷️  {result.get('keywords', '')}")
                        
                        if not followed_focus:
                            deviation_reason = focus_adherence.get('deviation_reason', 'Nicht angegeben')
                            print(f"      ⚠️ Fokus-Abweichung: {deviation_reason}")

                        # Debug-Ausgaben für Fokus-Details
                        if focus_adherence:
                            focus_score = focus_adherence.get('focus_category_score', 0)
                            chosen_score = focus_adherence.get('chosen_category_score', 0)
                            print(f"      📊 Fokus-Score: {focus_score:.2f}, Gewählt-Score: {chosen_score:.2f}")

                        # Erweiterte Begründung mit Fokus-Kennzeichnung
                        original_justification = result.get('justification', '')
                        focus_prefix = f"[FOKUS: {focus_category}, Adherence: {'Ja' if followed_focus else 'Nein'}] "
                        enhanced_justification = focus_prefix + original_justification
                    
                        return CodingResult(
                            category=result.get('category', ''),
                            subcategories=tuple(result.get('subcategories', [])),
                            justification=enhanced_justification,
                            confidence=result.get('confidence', {'total': 0.0, 'category': 0.0, 'subcategories': 0.0}),
                            text_references=tuple([chunk[:100]]),
                            uncertainties=None,
                            paraphrase=result.get('paraphrase', ''),
                            keywords=result.get('keywords', '')
                        )
                    else:
                        print("      ✗ Keine passende Kategorie gefunden")
                        return None
                    
            except Exception as e:
                print(f"Fehler bei API Call für fokussierte Kodierung: {str(e)}")
                return None          

        except Exception as e:
            print(f"Fehler bei der fokussierten Kodierung durch {self.coder_id}: {str(e)}")
            return None

    async def code_chunk_with_focus_and_context(self, 
                                            chunk: str, 
                                            categories: Dict[str, CategoryDefinition],
                                            focus_category: str,
                                            focus_context: Dict,
                                            current_summary: str,
                                            segment_info: Dict,
                                            update_summary: bool = False) -> Dict:
        """
        Kodiert einen Text-Chunk mit Fokus auf eine Kategorie UND progressivem Kontext.
        Kombiniert die Funktionalität von code_chunk_with_focus und code_chunk_with_progressive_context.
        
        Args:
            chunk: Zu kodierender Text
            categories: Kategoriensystem
            focus_category: Kategorie auf die fokussiert werden soll
            focus_context: Kontext zur fokussierten Kategorie
            current_summary: Aktuelles Dokument-Summary
            segment_info: Zusätzliche Informationen über das Segment
            update_summary: Ob das Summary aktualisiert werden soll
            
        Returns:
            Dict: Enthält sowohl Kodierungsergebnis als auch aktualisiertes Summary
        """
        try:
            if not categories:
                print(f"Fehler: Kein Kategoriensystem für Kodierer {self.coder_id} verfügbar")
                return None

            print(f"    🎯 Fokuskodierung für Kategorie: {focus_category} (Relevanz: {focus_context.get('relevance_score', 0):.2f})")

            # Erstelle formatierte Kategorienübersicht mit Fokus-Hervorhebung
            categories_overview = []
            for name, cat in categories.items():  
                # Hebe Fokus-Kategorie hervor
                display_name = name
                if name == focus_category:
                    display_name = name
                
                category_info = {
                    'name': display_name,
                    'definition': cat.definition,
                    'examples': list(cat.examples) if isinstance(cat.examples, set) else cat.examples,
                    'rules': list(cat.rules) if isinstance(cat.rules, set) else cat.rules,
                    'subcategories': {}
                }
                
                # Füge Subkategorien hinzu
                for sub_name, sub_def in cat.subcategories.items():
                    category_info['subcategories'][sub_name] = sub_def
                    
                categories_overview.append(category_info)

            # FIX: DEBUG - Prüfe ob gefilterte Kategorien wirklich im Prompt landen
            # print(f"    🔍 DEBUG: Categories Overview für Prompt:")
            # for cat_info in categories_overview:
            #     cat_name = cat_info['name']
            #     subcats = list(cat_info['subcategories'].keys())
            #     print(f"      - {cat_name}: {subcats}")

            # Position im Dokument und Fortschritt berechnen
            position_info = f"Segment: {segment_info.get('position', '')}"
            doc_name = segment_info.get('doc_name', 'Unbekanntes Dokument')
            
            # FIX: Summary-Update-Anweisungen mit dreistufigem Reifungsmodell (gleiche Logik wie in code_chunk_with_progressive_context)
            summary_update_prompt = ""
            if update_summary:
                # Berechne die relative Position im Dokument (für das Reifungsmodell)
                chunk_id = 0
                total_chunks = 1
                if 'position' in segment_info:
                    try:
                        # Extrahiere Chunk-Nummer aus "Chunk X"
                        chunk_id = int(segment_info['position'].split()[-1])
                        
                        # Schätze Gesamtanzahl der Chunks (basierend auf bisherigen Chunks)
                        # Alternative: Tatsächliche Anzahl übergeben, falls verfügbar
                        total_chunks = max(chunk_id * 1.5, 20)  # Schätzung
                        
                        document_progress = chunk_id / total_chunks
                        print(f"        Dokumentfortschritt: ca. {document_progress:.1%}")
                    except (ValueError, IndexError):
                        document_progress = 0.5  # Fallback
                else:
                    document_progress = 0.5  # Fallback
                    
                # Bestimme die aktuelle Reifephase basierend auf dem Fortschritt
                if document_progress < 0.3:
                    reifephase = "PHASE 1 (Sammlung)"
                    max_aenderung = "50%"
                elif document_progress < 0.7:
                    reifephase = "PHASE 2 (Konsolidierung)"
                    max_aenderung = "30%"
                else:
                    reifephase = "PHASE 3 (Präzisierung)"
                    max_aenderung = "10%"
                    
                print(f"        Summary-Reifephase: {reifephase}, max. Änderung: {max_aenderung}")
                
                # Angepasster Prompt basierend auf dem dreistufigen Reifungsmodell
                summary_update_prompt = f"""
                ## AUFGABE 2: SUMMARY-UPDATE ({reifephase}, {int(document_progress*100)}%)

                """

                # Robustere Phasen-spezifische Anweisungen
                if document_progress < 0.3:
                    summary_update_prompt += """
                SAMMLUNG (0-30%) - STRUKTURIERTER AUFBAU:
                - SCHLÜSSELINFORMATIONEN: Beginne mit einer LISTE wichtigster Konzepte im Telegrammstil
                - FORMAT: "Thema1: Kernaussage; Thema2: Kernaussage" 
                - SPEICHERSTRUKTUR: Speichere alle Informationen in KATEGORIEN (z.B. Akteure, Prozesse, Faktoren)
                - KEINE EINLEITUNGEN oder narrative Elemente, NUR Fakten und Verbindungen
                - BEHALTE IMMER: Bereits dokumentierte Schlüsselkonzepte müssen bestehen bleiben
                """
                elif document_progress < 0.7:
                    summary_update_prompt += """
                KONSOLIDIERUNG (30-70%) - HIERARCHISCHE ORGANISATION:
                - SCHLÜSSELINFORMATIONEN BEWAHREN: Alle bisherigen Hauptkategorien beibehalten
                - NEUE STRUKTUR: Als hierarchische Liste mit Kategorien und Unterpunkten organisieren
                - KOMPRIMIEREN: Details aus gleichen Themenbereichen zusammenführen
                - PRIORITÄTSFORMAT: "Kategorie: Hauptpunkt1; Hauptpunkt2 → Detail"
                - STATT LÖSCHEN: Verwandte Inhalte zusammenfassen, aber KEINE Kategorien eliminieren
                """
                else:
                    summary_update_prompt += """
                PRÄZISIERUNG (70-100%) - VERDICHTUNG MIT THESAURUS:
                - THESAURUS-METHODE: Jede Kategorie braucht genau 1-2 Sätze im Telegrammstil
                - HAUPTKONZEPTE STABIL HALTEN: Alle identifizierten Kategorien müssen enthalten bleiben
                - ABSTRAHIEREN: Einzelinformationen innerhalb einer Kategorie verdichten
                - STABILITÄTSPRINZIP: Einmal erkannte wichtige Zusammenhänge dürfen nicht verloren gehen
                - PRIORITÄTSORDNUNG: Wichtigste Informationen IMMER am Anfang jeder Kategorie
                """

                # Allgemeine Kriterien für Stabilität und Komprimierung
                summary_update_prompt += """

                INFORMATIONSERHALTUNGS-SYSTEM:
                - MAXIMUM 80 WÖRTER - Komprimiere alte statt neue Informationen zu verwerfen
                - KATEGORIEBASIERT: Jedes Summary muss immer in 3-5 klare Themenkategorien strukturiert sein
                - SCHLÜSSELPRINZIP: Bilde das Summary als INFORMATIONALE HIERARCHIE:
                1. Stufe: Immer stabile Themenkategorien
                2. Stufe: Zentrale Aussagen zu jeder Kategorie
                3. Stufe: Ergänzende Details (diese können komprimiert werden)
                - STABILITÄTSGARANTIE: Neue Iteration darf niemals vorherige Kategorie-Level-1-Information verlieren
                - KOMPRIMIERUNGSSTRATEGIE: Bei Platzmangel Details (Stufe 3) zusammenfassen statt zu entfernen
                - FORMAT: "Kategorie1: Hauptpunkt; Hauptpunkt. Kategorie2: Hauptpunkt; Detail." (mit Doppelpunkten)
                - GRUNDREGEL: Neue Informationen ergänzen bestehende Kategorien statt sie zu ersetzen
                """

            # Erstelle fokussierten Prompt mit Kontext
            prompt = self.prompt_handler.get_focus_context_coding_prompt(
                chunk=chunk,
                categories_overview=categories_overview,
                focus_category=focus_category,
                focus_context=focus_context,
                current_summary=current_summary,
                position_info=position_info,
                summary_update_prompt=summary_update_prompt,
                update_summary=update_summary
            )
            

            # API-Call
            token_counter.start_request()

            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                response_format={"type": "json_object"}
            )
            
            # Verarbeite Response mit Wrapper
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)

            
            token_counter.track_response(response, self.model_name)
            
            # Extrahiere relevante Teile
            if result and isinstance(result, dict):
                coding_result = result.get('coding_result', {})
                
                # Summary nur aktualisieren wenn angefordert
                if update_summary:
                    updated_summary = result.get('updated_summary', current_summary)
                    
                    # Prüfe Wortlimit beim Summary
                    if len(updated_summary.split()) > 80:
                        words = updated_summary.split()
                        updated_summary = ' '.join(words[:70])
                        print(f"        ⚠️ Summary wurde gekürzt: {len(words)} → 70 Wörter")
                    
                    # FIX: Analyse der Veränderungen (gleich wie in code_chunk_with_progressive_context)
                    if current_summary:
                        # Berechne Prozent der Änderung
                        old_words = set(current_summary.lower().split())
                        new_words = set(updated_summary.lower().split())
                        
                        if old_words:
                            # Jaccard-Distanz als Maß für Veränderung
                            unchanged = len(old_words.intersection(new_words))
                            total = len(old_words.union(new_words))
                            change_percent = (1 - (unchanged / total)) * 100
                            
                            print(f"        Summary Änderung: {change_percent:.1f}% (Ziel: max. {max_aenderung})")
                else:
                    updated_summary = current_summary
                
                if coding_result:
                    paraphrase = coding_result.get('paraphrase', '')
                    if paraphrase:
                        print(f"        🗒️  Fokus-Kontext-Paraphrase: {paraphrase}")

                    # Dokumentiere Fokus-Adherence
                    focus_adherence = coding_result.get('focus_adherence', {})
                    followed_focus = focus_adherence.get('followed_focus', True)
                    focus_icon = "🎯" if followed_focus else "🔄"
                    
                    print(f"        {focus_icon} Fokus-Kontext-Kodierung von {self.coder_id}: 🏷️  {coding_result.get('category', '')}")
                    print(f"        ✓ Subkategorien: 🏷️  {', '.join(coding_result.get('subcategories', []))}")
                    print(f"        ✓ Keywords: 🏷️  {coding_result.get('keywords', '')}")
                    
                    if not followed_focus:
                        deviation_reason = focus_adherence.get('deviation_reason', 'Nicht angegeben')
                        print(f"        ⚠️ Fokus-Abweichung: {deviation_reason}")

                    if update_summary:
                        print(f"        📝 Summary für {doc_name} aktualisiert ({len(updated_summary.split())} Wörter):")
                        print(f"        {updated_summary[:100]}..." if len(updated_summary) > 100 else f"        📄 {updated_summary}")
                    
                    # Kombiniertes Ergebnis zurückgeben
                    return {
                        'coding_result': coding_result,
                        'updated_summary': updated_summary
                    }
                else:
                    print(f"        ✗ Keine gültige Kodierung erhalten")
                    return None
            else:
                print("        ✗ Keine gültige Antwort erhalten")
                return None
                
        except Exception as e:
            print(f"Fehler bei der fokussierten Kontext-Kodierung durch {self.coder_id}: {str(e)}")
            print("Details:")
            import traceback
            traceback.print_exc()
            return None
        
    async def _check_relevance(self, chunk: str) -> bool:
        """
        Prüft die Relevanz eines Chunks für die Forschungsfrage.
        
        Args:
            chunk: Zu prüfender Text
            
        Returns:
            bool: True wenn der Text relevant ist
        """
        try:
            prompt = self.prompt_handler.get_segment_relevance_assessment_prompt(chunk)

            token_counter.start_request()
            
            response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    response_format={"type": "json_object"}
                )
                
            # Verarbeite Response mit Wrapper
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)
            
            
            token_counter.track_response(response, self.model_name)

            # Detaillierte Ausgabe der Relevanzprüfung
            if result.get('is_relevant'):
                print(f"✓ Relevanz bestätigt (Konfidenz: {result.get('confidence', 0):.2f})")
                if result.get('key_aspects'):
                    print("  Relevante Aspekte:")

                    for aspect in result['key_aspects']:
                        print(f"  - {aspect}")
            else:
                print(f"❌ Nicht relevant: {result.get('justification', 'Keine Begründung')}")

            return result.get('is_relevant', False)

        except Exception as e:
            print(f"Fehler bei der Relevanzprüfung: {str(e)}")
            return True  # Im Zweifelsfall als relevant markieren

    
    def _validate_coding(self, result: dict) -> Optional[CodingResult]:
        """Validiert und konvertiert das API-Ergebnis in ein CodingResult-Objekt"""
        try:
            return CodingResult(
                category=result.get('category', ''),
                subcategories=result.get('subcategories', []),
                justification=result.get('justification', ''),
                confidence=result.get('confidence', {'total': 0, 'category': 0, 'subcategories': 0}),
                text_references=result.get('text_references', []),
                uncertainties=result.get('uncertainties', [])
            )
        except Exception as e:
            print(f"Fehler bei der Validierung des Kodierungsergebnisses: {str(e)}")
            return None

# ---
# --- Klasse: InductiveCoder ---
# Aufgabe: Ergänzung deduktiver Kategorien durch induktive Kategorien mittels OpenAI API

class InductiveCoder:
    """
    Vereinfachter induktiver Kodierer mit strikter 2-Phasen-Struktur:
    Phase 1: Kategoriensystem-Aufbau (mit strenger Sättigung)
    Phase 2: Kodierung mit festem System
    """
    
    def __init__(self, model_name: str, history: DevelopmentHistory, output_dir: str, config: dict = None):
        self.model_name = model_name
        self.output_dir = output_dir
        self.history = history
        self.config = config or CONFIG  # KORREKTUR: Speichere config
        
        # LLM Provider (unverändert)
        provider_name = CONFIG.get('MODEL_PROVIDER', 'openai').lower()
        self.llm_provider = LLMProviderFactory.create_provider(provider_name)
        
        # Cache und Tracking (unverändert)
        self.category_cache = {}
        self.analysis_cache = {}
        self.batch_results = []
        self.similarity_cache = {}
        
        # VERBESSERTE Sättigungsschwellen (aus dem verbesserten Code)
        self.MIN_CONFIDENCE = 0.7
        self.MIN_EXAMPLES = 2
        self.MIN_CATEGORY_USAGE = 2
        self.MAX_CATEGORIES_PER_BATCH = 5
        
        # VERSCHÄRFTE Sättigungskriterien (aus dem verbesserten Code)
        self.MIN_BATCHES_BEFORE_SATURATION = 5
        self.MIN_MATERIAL_COVERAGE = 0.8
        self.STABILITY_THRESHOLD = 3
        
        # Theoretische Sättigungsmetriken (aus dem verbesserten Code)
        self.theoretical_saturation_history = []
        self.category_development_phases = []
        
        # Phasen-Management (unverändert)
        self.current_phase = "development"
        self.categories_locked = False
        self.development_complete = False
        
        # Sättigungs-Tracking (unverändert)
        self.batches_without_new_categories = 0
        self.category_usage_history = {}
        self.rejected_categories = []
        
        # Für Grounded Theory Modus (unverändert)
        self.collected_subcodes = []
        self.segment_analyses = []

        self.discovered_aspects = set()
        self.batch_metrics = []
        
        # Prompt-Handler (unverändert)
        self.prompt_handler = QCAPrompts(
            forschungsfrage=FORSCHUNGSFRAGE,
            kodierregeln=KODIERREGELN,
            deduktive_kategorien=DEDUKTIVE_KATEGORIEN
        )

        print(f"\n🔬 Induktive Kodierung initialisiert:")
        print(f"- Min. Batches vor Sättigung: {self.MIN_BATCHES_BEFORE_SATURATION}")
        print(f"- Min. Materialabdeckung: {self.MIN_MATERIAL_COVERAGE:.0%}")
        print(f"- Stabilitätsschwelle: {self.STABILITY_THRESHOLD} Batches")
    
    
    
        
    def _create_proper_batches(self, segments: List[str], batch_size: int) -> List[List[str]]:
        """
        VERBESSERT: Erstellt Batches ohne künstliche Größenreduzierung
        """
        if not segments:
            return []
        
        print(f"📦 Erstelle Batches: {len(segments)} Segmente → Batch-Größe {batch_size}")
        
        batches = []
        for i in range(0, len(segments), batch_size):
            batch = segments[i:i + batch_size]
            batches.append(batch)
        
        print(f"📦 Ergebnis: {len(batches)} gleichmäßige Batches erstellt")
        return batches

    async def _validate_and_integrate_strict(self, candidates: Dict[str, CategoryDefinition], 
                                           existing: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        Validierung und automatische Konsolidierung neuer Kategorien
        """
        validated = {}
        
        for name, category in candidates.items():
            # 1. Ähnlichkeitsprüfung
            similar_existing = self._find_similar_category(category, existing)
            if similar_existing:
                print(f"🔄 '{name}' zu ähnlich zu '{similar_existing}' - wird konsolidiert")
                # Automatische Konsolidierung statt Ablehnung
                consolidated = await self._auto_merge_categories(
                    category, existing[similar_existing], name, similar_existing
                )
                if consolidated:
                    existing[similar_existing] = consolidated
                    # WICHTIG: Nutzung für konsolidierte Kategorie erhöhen
                    self.category_usage_history[similar_existing] = self.category_usage_history.get(similar_existing, 0) + 1
                continue
            
            # 2. Qualitätsprüfung
            if await self._meets_quality_standards(category):
                validated[name] = category
                # WICHTIG: Nutzung für neue Kategorie setzen
                self.category_usage_history[name] = self.category_usage_history.get(name, 0) + 1
                print(f"✅ '{name}' validiert (Nutzung: {self.category_usage_history[name]})")
            else:
                print(f"❌ '{name}' erfüllt Qualitätsstandards nicht")
        
        return validated
    
    async def _consolidate_categories(self, categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        Automatische Konsolidierung ähnlicher Kategorien
        """
        print("\n🔄 Starte automatische Konsolidierung...")
        
        consolidated = categories.copy()
        merge_candidates = []
        
        # Finde Konsolidierungskandidaten
        category_names = list(consolidated.keys())
        for i in range(len(category_names)):
            for j in range(i + 1, len(category_names)):
                name1, name2 = category_names[i], category_names[j]
                if name1 in consolidated and name2 in consolidated:
                    similarity = self._calculate_category_similarity(
                        consolidated[name1], consolidated[name2]
                    )
                    if similarity > 0.7:  # similarity_threshold
                        merge_candidates.append((name1, name2, similarity))
        
        # Sortiere nach Ähnlichkeit
        merge_candidates.sort(key=lambda x: x[2], reverse=True)
        
        # Führe Konsolidierungen durch
        for name1, name2, similarity in merge_candidates[:3]:  # Max 3 Merges pro Runde
            if name1 in consolidated and name2 in consolidated:
                print(f"🔗 Konsolidiere '{name1}' + '{name2}' (Ähnlichkeit: {similarity:.2f})")
                merged = await self._merge_categories_intelligent(
                    consolidated[name1], consolidated[name2], name1, name2
                )
                if merged:
                    # Verwende den besseren Namen
                    better_name = self._choose_better_name(name1, name2)
                    consolidated[better_name] = merged
                    
                    # Entferne die anderen
                    other_name = name2 if better_name == name1 else name1
                    del consolidated[other_name]
                    
                    print(f"✅ Konsolidiert zu '{better_name}'")
        
        return consolidated
    
    async def _finalize_categories(self, categories: Dict[str, CategoryDefinition]) -> Dict[str, CategoryDefinition]:
        """
        Finale Bereinigung des Kategoriensystems
        """
        print("\n🧹 Finale Bereinigung...")
        
        cleaned = {}
        
        for name, category in categories.items():
            # KORRIGIERT: Verwende deutlich niedrigere Schwelle oder überspringe Check
            usage_count = self.category_usage_history.get(name, 0)
            
            # TEMPORÄRER FIX: Akzeptiere alle Kategorien in der Entwicklungsphase
            if self.current_phase == "development":
                print(f"✅ '{name}' übernommen (Entwicklungsphase)")
                cleaned[name] = category
                continue
                
            # KORRIGIERT: Viel niedrigere Schwelle
            min_usage = max(1, self.MIN_CATEGORY_USAGE // 3)  # 1 statt 3
            
            if usage_count >= min_usage:
                # Verbessere Definition falls nötig
                if len(category.definition.split()) < 20:
                    enhanced = await self._enhance_category_definition(category)
                    if enhanced:
                        category = category.replace(definition=enhanced.definition)
                
                cleaned[name] = category
                print(f"✅ '{name}' übernommen (Nutzung: {usage_count})")
            else:
                print(f"❌ '{name}' entfernt (Zu wenig genutzt: {usage_count}, Mindest: {min_usage})")
        
        return cleaned
    
    def _update_usage_history(self, category_names: List[str]) -> None:
        """
        Aktualisiert die Nutzungshistorie für Kategorien
        """
        for name in category_names:
            if name in self.category_usage_history:
                self.category_usage_history[name] += 1
            else:
                self.category_usage_history[name] = 1
        
        print(f"📊 Nutzungshistorie aktualisiert für: {category_names}")
        print(f"    Aktuelle Nutzung: {dict(list(self.category_usage_history.items())[-3:])}")

    def _create_category_definition(self, cat_data: dict) -> CategoryDefinition:
        """
        Erstellt CategoryDefinition aus API-Response Dictionary
        """
        try:
            return CategoryDefinition(
                name=cat_data.get('name', ''),
                definition=cat_data.get('definition', ''),
                examples=cat_data.get('evidence', []),
                rules=[],  # Wird später entwickelt
                subcategories={
                    sub.get('name', ''): sub.get('definition', '')
                    for sub in cat_data.get('subcategories', [])
                },
                added_date=datetime.now().strftime("%Y-%m-%d"),
                modified_date=datetime.now().strftime("%Y-%m-%d")
            )
        except Exception as e:
            print(f"Fehler bei CategoryDefinition-Erstellung: {str(e)}")
            return None
    
    
    
    def _format_existing_categories(self, categories: Dict[str, CategoryDefinition]) -> str:
        """Formatiert bestehende Kategorien für Prompt"""
        if not categories:
            return "Keine bestehenden Kategorien."
        
        formatted = []
        for name, cat in categories.items():
            definition_preview = cat.definition[:100] + "..." if len(cat.definition) > 100 else cat.definition
            formatted.append(f"- {name}: {definition_preview}")
        
        return "\n".join(formatted)

    
    async def develop_category_system(self, segments: List[str], initial_categories: Dict[str, CategoryDefinition] = None) -> Dict[str, CategoryDefinition]:
        """
        VERBESSERTE Kategorienentwicklung mit korrekter Sättigungslogik
        """
        print(f"\n🔍 Starte verbesserte induktive Entwicklung mit {len(segments)} Segmenten")
        
        current_categories = initial_categories.copy() if initial_categories else {}
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'inductive')
        
        print(f"\n📊 Analysemodus: {analysis_mode.upper()}")
        
        # Reset Tracking
        self.theoretical_saturation_history = []
        self.category_development_phases = []
        self.batches_without_new_categories = 0
        
        # VERBESSERTE Batch-Erstellung (keine künstliche Reduzierung)
        print("\n📦 Erstelle optimierte Batches...")

        # Erstelle Batches direkt
        effective_batch_size = min(CONFIG.get('BATCH_SIZE', 5), len(segments))
        batches = self._create_proper_batches(segments, effective_batch_size)
        
        
        print(f"📊 Batch-Konfiguration:")
        print(f"- Relevante Segmente: {len(segments)}")
        print(f"- Batch-Größe: {effective_batch_size}")
        print(f"- Anzahl Batches: {len(batches)}")
        
        working_categories = current_categories.copy()
        
        # HAUPTSCHLEIFE mit verbesserter Sättigungslogik
        for batch_idx, batch in enumerate(batches):
            print(f"\n{'='*60}")
            print(f"📊 BATCH {batch_idx + 1}/{len(batches)} - Kategorienentwicklung")
            print(f"{'='*60}")
            
            # Analysiere Batch
            new_candidates = await self._analyze_batch_improved(batch, working_categories, analysis_mode)
            
            # Validiere und integriere neue Kategorien
            if new_candidates:
                validated_categories = await self._validate_and_integrate_strict(new_candidates, working_categories)
                
                if validated_categories:
                    before_count = len(working_categories)
                    working_categories.update(validated_categories)
                    added_count = len(working_categories) - before_count
                    
                    print(f"✅ {added_count} neue Kategorien integriert")
                    self.batches_without_new_categories = 0
                    self._update_usage_history(list(validated_categories.keys()))
                    
                    # Dokumentiere Entwicklungsphase
                    self.category_development_phases.append({
                        'batch': batch_idx + 1,
                        'new_categories': added_count,
                        'total_categories': len(working_categories),
                        'material_coverage': (batch_idx + 1) / len(batches)
                    })
                else:
                    print("❌ Keine Kategorien haben strenge Validierung bestanden")
                    self.batches_without_new_categories += 1
            else:
                print("ℹ️ Keine neuen Kategorien in diesem Batch")
                self.batches_without_new_categories += 1
            
            # VERBESSERTE Sättigungsprüfung
            saturation_metrics = self._assess_comprehensive_saturation(
                working_categories, 
                batch_idx + 1, 
                len(batches)
            )
            
            print(f"\n📈 SÄTTIGUNGSANALYSE:")
            print(f"- Theoretische Sättigung: {saturation_metrics['theoretical_saturation']:.2f}")
            print(f"- Materialabdeckung: {saturation_metrics['material_coverage']:.1%}")
            print(f"- Stabile Batches: {saturation_metrics['stable_batches']}")
            print(f"- Kategorienqualität: {saturation_metrics['category_quality']:.2f}")
            print(f"- Diversität: {saturation_metrics['category_diversity']:.2f}")
            
            # Speichere Sättigungshistorie
            self.theoretical_saturation_history.append(saturation_metrics)
            
            # Prüfe ALLE Sättigungskriterien
            if self._check_comprehensive_saturation(saturation_metrics, batch_idx + 1, len(batches)):
                print(f"\n🛑 VOLLSTÄNDIGE SÄTTIGUNG erreicht nach Batch {batch_idx + 1}")
                print(f"📊 Sättigungsgrund:")
                for criterion, value in saturation_metrics.items():
                    print(f"   - {criterion}: {value}")
                break
            else:
                print(f"\n⏳ Sättigung noch nicht erreicht - fortsetzen")
                self._log_saturation_progress(saturation_metrics)
            
            # Zwischenkonsolidierung alle 3 Batches
            if (batch_idx + 1) % 3 == 0:
                print(f"\n🔄 Zwischenkonsolidierung nach Batch {batch_idx + 1}")
                working_categories = await self._consolidate_categories(working_categories)
        
        # Finale Bereinigung und Qualitätssicherung
        final_categories = await self._finalize_categories(working_categories)
        
        # Zeige finale Entwicklungsstatistiken
        self._show_development_summary(final_categories, initial_categories)
        
        return final_categories

    def _create_proper_batches(self, segments: List[str], batch_size: int) -> List[List[str]]:
        """
        VERBESSERT: Erstellt Batches ohne künstliche Größenreduzierung
        """
        if not segments:
            return []
        
        print(f"📦 Erstelle Batches: {len(segments)} Segmente → Batch-Größe {batch_size}")
        
        batches = []
        for i in range(0, len(segments), batch_size):
            batch = segments[i:i + batch_size]
            batches.append(batch)
        
        print(f"📦 Ergebnis: {len(batches)} gleichmäßige Batches erstellt")
        return batches

    def _assess_comprehensive_saturation(self, categories: Dict[str, CategoryDefinition], 
                                       current_batch: int, total_batches: int) -> Dict[str, float]:
        """
        VERBESSERTE umfassende Sättigungsbeurteilung
        """
        # 1. Theoretische Sättigung (Kategorienqualität und -vollständigkeit)
        theoretical_saturation = self._calculate_theoretical_saturation(categories)
        
        # 2. Materialabdeckung
        material_coverage = current_batch / total_batches
        
        # 3. Stabilität (Batches ohne neue Kategorien)
        stability_ratio = self.batches_without_new_categories / max(1, current_batch)
        
        # 4. Kategorienqualität (Definition, Beispiele, Subkategorien)
        category_quality = self._assess_category_quality(categories)
        
        # 5. Kategorien-Diversität (thematische Abdeckung)
        category_diversity = self._calculate_category_diversity(categories)
        
        return {
            'theoretical_saturation': theoretical_saturation,
            'material_coverage': material_coverage,
            'stable_batches': self.batches_without_new_categories,
            'stability_ratio': stability_ratio,
            'category_quality': category_quality,
            'category_diversity': category_diversity,
            'total_categories': len(categories)
        }

    def _calculate_theoretical_saturation(self, categories: Dict[str, CategoryDefinition]) -> float:
        """
        Berechnet theoretische Sättigung basierend auf Kategorienreife und Forschungsabdeckung
        """
        if not categories:
            return 0.0
        
        # 1. Kategorienreife (Definition, Beispiele, Subkategorien)
        maturity_scores = []
        for cat in categories.values():
            score = 0
            # Definition (0-0.4)
            def_score = min(len(cat.definition.split()) / 30, 0.4)
            # Beispiele (0-0.3)
            example_score = min(len(cat.examples) / 5, 0.3)
            # Subkategorien (0-0.3)
            subcat_score = min(len(cat.subcategories) / 4, 0.3)
            
            total_score = def_score + example_score + subcat_score
            maturity_scores.append(total_score)
        
        avg_maturity = sum(maturity_scores) / len(maturity_scores)
        
        # 2. Forschungsabdeckung (Anzahl und Diversität der Kategorien)
        # Schätze optimale Kategorienanzahl basierend auf Forschungsfrage
        estimated_optimal = 8  # Typisch für qualitative Analysen
        coverage_ratio = min(len(categories) / estimated_optimal, 1.0)
        
        # 3. Kombinierte theoretische Sättigung
        theoretical_saturation = (avg_maturity * 0.7) + (coverage_ratio * 0.3)
        
        return min(theoretical_saturation, 1.0)

    def _assess_category_quality(self, categories: Dict[str, CategoryDefinition]) -> float:
        """
        Bewertet die durchschnittliche Qualität aller Kategorien
        """
        if not categories:
            return 0.0
        
        quality_scores = []
        for cat in categories.values():
            score = 0
            
            # Definition ausreichend (0-0.4)
            if len(cat.definition.split()) >= 20:
                score += 0.4
            elif len(cat.definition.split()) >= 10:
                score += 0.2
            
            # Beispiele vorhanden (0-0.3)
            if len(cat.examples) >= 3:
                score += 0.3
            elif len(cat.examples) >= 1:
                score += 0.15
            
            # Subkategorien entwickelt (0-0.3)
            if len(cat.subcategories) >= 3:
                score += 0.3
            elif len(cat.subcategories) >= 1:
                score += 0.15
            
            quality_scores.append(score)
        
        return sum(quality_scores) / len(quality_scores)

    def _calculate_category_diversity(self, categories: Dict[str, CategoryDefinition]) -> float:
        """
        Berechnet thematische Diversität der Kategorien
        """
        if not categories:
            return 0.0
        
        # Sammle Schlüsselwörter aus allen Definitionen
        all_keywords = set()
        for cat in categories.values():
            words = cat.definition.lower().split()
            keywords = [w for w in words if len(w) > 4]  # Nur längere Wörter
            all_keywords.update(keywords[:5])  # Top 5 pro Kategorie
        
        # Diversität = Verhältnis von einzigartigen Begriffen zu Kategorien
        diversity = len(all_keywords) / (len(categories) * 3)  # Normalisiert
        return min(diversity, 1.0)

    def _check_comprehensive_saturation(self, saturation_metrics: Dict[str, float], 
                                      current_batch: int, total_batches: int) -> bool:
        """
        VERSCHÄRFTE Sättigungsprüfung mit mehreren Kriterien
        """
        # Mindestkriterien
        min_batches = max(self.MIN_BATCHES_BEFORE_SATURATION, total_batches * 0.3)
        min_material = self.MIN_MATERIAL_COVERAGE
        min_stability = self.STABILITY_THRESHOLD
        
        # Prüfe alle Kriterien
        criteria_met = {
            'min_batches': current_batch >= min_batches,
            'material_coverage': saturation_metrics['material_coverage'] >= min_material,
            'theoretical_saturation': saturation_metrics['theoretical_saturation'] >= 0.8,
            'category_quality': saturation_metrics['category_quality'] >= 0.7,
            'stability': saturation_metrics['stable_batches'] >= min_stability,
            'sufficient_categories': saturation_metrics['total_categories'] >= 3
        }
        
        print(f"\n🔍 Sättigungskriterien:")
        for criterion, met in criteria_met.items():
            status = "✅" if met else "❌"
            print(f"   {status} {criterion}: {met}")
        
        # Sättigung nur wenn ALLE Kriterien erfüllt
        is_saturated = all(criteria_met.values())
        
        if is_saturated:
            print(f"\n🎯 ALLE Sättigungskriterien erfüllt!")
        else:
            missing = [k for k, v in criteria_met.items() if not v]
            print(f"\n⏳ Fehlende Kriterien: {', '.join(missing)}")
        
        return is_saturated

    def _create_inductive_mode_prompt(self, segments_text: str, existing_categories: Dict[str, CategoryDefinition]) -> str:
        """
        Erstellt spezifischen Prompt für INDUCTIVE MODE (vollständige induktive Kategorienentwicklung)
        """
        # Formatiere bestehende induktive Kategorien als Kontext (aber nicht als Einschränkung)
        existing_context = ""
        if existing_categories:
            existing_names = list(existing_categories.keys())
            existing_context = f"""
            BESTEHENDE INDUKTIVE KATEGORIEN (als Kontext, NICHT als Einschränkung):
            {', '.join(existing_names)}
            
            WICHTIG: Entwickle NEUE, EIGENSTÄNDIGE Kategorien, die sich thematisch von den bestehenden unterscheiden.
            Beachte aber die bereits entwickelten Kategorien um Redundanzen zu vermeiden.
            """
        
        return f"""
        INDUCTIVE MODE: Vollständige induktive Kategorienentwicklung

        {existing_context}

        AUFGABE: Entwickle völlig NEUE Hauptkategorien aus den folgenden Textsegmenten.
        Dies ist ein eigenständiges induktives Kategoriensystem, unabhängig von deduktiven Kategorien.

        REGELN FÜR INDUCTIVE MODE:
        - Entwickle 1-{self.MAX_CATEGORIES_PER_BATCH} NEUE Hauptkategorien
        - Jede Kategorie muss mindestens {self.MIN_EXAMPLES} Textbelege haben
        - Konfidenz mindestens {self.MIN_CONFIDENCE}
        - Kategorien müssen thematisch eigenständig und relevant sein
        - Erstelle auch 2-4 Subkategorien pro Hauptkategorie
        - Kategorien sollen neue Aspekte der Forschungsfrage beleuchten
        - Vermeide Redundanzen zu bereits entwickelten Kategorien

        FORSCHUNGSFRAGE: {FORSCHUNGSFRAGE}

        TEXTSEGMENTE:
        {segments_text}

        Antworte NUR mit JSON:
        {{
            "new_categories": [
                {{
                    "name": "Kategorie Name",
                    "definition": "Ausführliche Definition (mindestens 20 Wörter)",
                    "evidence": ["Textbelege aus den Segmenten"],
                    "confidence": 0.0-1.0,
                    "subcategories": [
                        {{
                            "name": "Subkategorie Name", 
                            "definition": "Subkategorie Definition"
                        }}
                    ],
                    "thematic_justification": "Warum diese Kategorie einen eigenständigen Themenbereich abbildet"
                }}
            ],
            "development_assessment": {{
                "categories_developed": 0,
                "theoretical_saturation": 0.0-1.0,
                "new_themes_found": true/false,
                "recommendation": "continue/pause/stop"
            }}
        }}
        """
    
    def _create_abductive_mode_prompt(self, segments_text: str, existing_categories: Dict[str, CategoryDefinition]) -> str:
        """
        Erstellt spezifischen Prompt für ABDUCTIVE MODE (nur Subkategorien)
        """
        categories_context = []
        for cat_name, cat_def in existing_categories.items():
            categories_context.append({
                'name': cat_name,
                'definition': cat_def.definition[:200],
                'existing_subcategories': list(cat_def.subcategories.keys())
            })

        return f"""
        ABDUKTIVER MODUS: Entwickle NUR neue Subkategorien für bestehende Hauptkategorien.

        BESTEHENDE HAUPTKATEGORIEN:
        {json.dumps(categories_context, indent=2, ensure_ascii=False)}

        STRIKTE REGELN FÜR ABDUKTIVEN MODUS:
        - KEINE neuen Hauptkategorien entwickeln
        - NUR neue Subkategorien für bestehende Hauptkategorien
        - Subkategorien müssen neue, relevante Themenaspekte abbilden
        - Mindestens {self.MIN_EXAMPLES} Textbelege pro Subkategorie
        - Konfidenz mindestens {self.MIN_CONFIDENCE}
        - Prüfe JEDE bestehende Hauptkategorie auf mögliche neue Subkategorien
        
        TEXTSEGMENTE:
        {segments_text}
        
        Antworte NUR mit JSON:
        {{
            "extended_categories": {{
                "hauptkategorie_name": {{
                    "new_subcategories": [
                        {{
                            "name": "Subkategorie Name",
                            "definition": "Definition der Subkategorie",
                            "evidence": ["Textbelege"],
                            "confidence": 0.0-1.0,
                            "thematic_novelty": "Warum diese Subkategorie einen neuen Aspekt abbildet"
                        }}
                    ]
                }}
            }},
            "saturation_assessment": {{
                "subcategory_saturation": 0.0-1.0,
                "new_aspects_found": true/false,
                "recommendation": "continue/pause/stop"
            }}
        }}
        """

    def _create_standard_prompt(self, segments_text: str, existing_categories: Dict[str, CategoryDefinition]) -> str:
        """
        Erstellt Standard-Prompt für allgemeine induktive Kategorienentwicklung
        """
        existing_context = ""
        if existing_categories:
            existing_names = list(existing_categories.keys())
            existing_context = f"Bestehende Kategorien: {', '.join(existing_names)}"

        return f"""
        STANDARD INDUKTIVE KATEGORIENENTWICKLUNG

        {existing_context}

        AUFGABE: Entwickle neue Kategorien aus den folgenden Textsegmenten.

        ALLGEMEINE REGELN:
        - Entwickle 1-{self.MAX_CATEGORIES_PER_BATCH} neue Kategorien
        - Jede Kategorie braucht mindestens {self.MIN_EXAMPLES} Textbelege
        - Konfidenz mindestens {self.MIN_CONFIDENCE}
        - Erstelle aussagekräftige Definitionen
        - Füge relevante Subkategorien hinzu

        FORSCHUNGSFRAGE: {FORSCHUNGSFRAGE}

        TEXTSEGMENTE:
        {segments_text}

        Antworte NUR mit JSON:
        {{
            "new_categories": [
                {{
                    "name": "Kategorie Name",
                    "definition": "Kategorie Definition",
                    "evidence": ["Textbelege"],
                    "confidence": 0.0-1.0,
                    "subcategories": [
                        {{"name": "Subkategorie", "definition": "Definition"}}
                    ]
                }}
            ]
        }}
        """
    
    async def _analyze_batch_improved(self, batch: List[str], existing_categories: Dict[str, CategoryDefinition], analysis_mode: str) -> Dict[str, CategoryDefinition]:
        """
        VERBESSERTE Batch-Analyse mit modusabhängiger Logik
        """
        segments_text = "\n\n=== SEGMENT BREAK ===\n\n".join(
            f"SEGMENT {i + 1}:\n{text}" 
            for i, text in enumerate(batch)
        )

        # Modusabhängige Prompt-Erstellung
        if analysis_mode == 'inductive':
            prompt = self._create_inductive_mode_prompt(segments_text, existing_categories)
        elif analysis_mode == 'abductive':
            prompt = self._create_abductive_mode_prompt(segments_text, existing_categories)
        else:
            prompt = self._create_standard_prompt(segments_text, existing_categories)

        try:
            token_counter.start_request()
            
            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Antworte auf deutsch."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                response_format={"type": "json_object"}
            )
            
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)
            
            
            token_counter.track_response(response, self.model_name)
            
            # Verarbeite Ergebnisse
            candidates = {}
            
            for cat_data in result.get('new_categories', []):
                if cat_data.get('confidence', 0) >= self.MIN_CONFIDENCE:
                    candidates[cat_data['name']] = self._create_category_definition(cat_data)
                    print(f"✅ Neuer Kandidat: '{cat_data['name']}' (Konfidenz: {cat_data.get('confidence', 0):.2f})")
            
            return candidates
            
        except Exception as e:
            print(f"Fehler bei verbesserter Batch-Analyse: {str(e)}")
            return {}

    def _create_abductive_mode_prompt(self, segments_text: str, existing_categories: Dict[str, CategoryDefinition]) -> str:
        """
        Erstellt spezifischen Prompt für abduktiven Modus (nur Subkategorien)
        """
        categories_context = []
        for cat_name, cat_def in existing_categories.items():
            categories_context.append({
                'name': cat_name,
                'definition': cat_def.definition,
                'existing_subcategories': list(cat_def.subcategories.keys())
            })

        return f"""
        ABDUKTIVER MODUS: Entwickle NUR neue Subkategorien für bestehende Hauptkategorien.

        BESTEHENDE HAUPTKATEGORIEN:
        {json.dumps(categories_context, indent=2, ensure_ascii=False)}

        STRIKTE REGELN:
        - KEINE neuen Hauptkategorien entwickeln
        - NUR neue Subkategorien für bestehende Hauptkategorien
        - Subkategorien müssen neue, relevante Themenaspekte abbilden
        - Mindestens {self.MIN_EXAMPLES} Textbelege pro Subkategorie
        - Konfidenz mindestens {self.MIN_CONFIDENCE}
        
        TEXTSEGMENTE:
        {segments_text}
        
        Antworte NUR mit JSON:
        {{
            "extended_categories": {{
                "hauptkategorie_name": {{
                    "new_subcategories": [
                        {{
                            "name": "Subkategorie Name",
                            "definition": "Definition der Subkategorie",
                            "evidence": ["Textbelege"],
                            "confidence": 0.0-1.0,
                            "thematic_novelty": "Warum diese Subkategorie einen neuen Aspekt abbildet"
                        }}
                    ]
                }}
            }},
            "saturation_assessment": {{
                "subcategory_saturation": 0.0-1.0,
                "new_aspects_found": true/false,
                "recommendation": "continue/pause/stop"
            }}
        }}
        """

    def _log_saturation_progress(self, saturation_metrics: Dict[str, float]) -> None:
        """
        Protokolliert Sättigungsfortschritt für Benutzer-Feedback
        """
        print(f"\n📊 Sättigungsfortschritt:")
        print(f"   🎯 Theoretische Sättigung: {saturation_metrics['theoretical_saturation']:.1%}")
        print(f"   📈 Materialabdeckung: {saturation_metrics['material_coverage']:.1%}")
        print(f"   🔄 Stabilität: {saturation_metrics['stable_batches']} Batches ohne neue Kategorien")
        print(f"   ⭐ Kategorienqualität: {saturation_metrics['category_quality']:.1%}")
        print(f"   🌈 Diversität: {saturation_metrics['category_diversity']:.1%}")

    def _show_development_summary(self, final_categories: Dict[str, CategoryDefinition], 
                                initial_categories: Dict[str, CategoryDefinition]) -> None:
        """
        Zeigt finale Entwicklungsstatistiken
        """
        print(f"\n{'='*60}")
        print(f"📊 KATEGORIENENTWICKLUNG ABGESCHLOSSEN")
        print(f"{'='*60}")
        
        # Grundstatistiken
        initial_count = len(initial_categories) if initial_categories else 0
        final_count = len(final_categories)
        new_categories = final_count - initial_count
        
        print(f"📈 Kategorien-Bilanz:")
        print(f"   - Initial: {initial_count}")
        print(f"   - Neu entwickelt: {new_categories}")
        print(f"   - Final: {final_count}")
        
        # Sättigungshistorie
        if self.theoretical_saturation_history:
            final_saturation = self.theoretical_saturation_history[-1]
            print(f"\n🎯 Finale Sättigung:")
            print(f"   - Theoretische Sättigung: {final_saturation['theoretical_saturation']:.1%}")
            print(f"   - Kategorienqualität: {final_saturation['category_quality']:.1%}")
            print(f"   - Diversität: {final_saturation['category_diversity']:.1%}")
        
        # Entwicklungsphasen
        if self.category_development_phases:
            print(f"\n📊 Entwicklungsphasen:")
            for phase in self.category_development_phases:
                print(f"   Batch {phase['batch']}: +{phase['new_categories']} → {phase['total_categories']} total")

    
    
    def _format_existing_categories(self, categories: Dict[str, CategoryDefinition]) -> str:
        """Formatiert bestehende Kategorien für Prompt"""
        if not categories:
            return "Keine bestehenden Kategorien."
        
        formatted = []
        for name, cat in categories.items():
            formatted.append(f"- {name}: {cat.definition[:100]}...")
        
        return "\n".join(formatted)

    
    async def _enhance_category_definition(self, category: CategoryDefinition) -> Optional[CategoryDefinition]:
        """Verbessert Kategoriendefinition"""
        try:
            prompt = self.prompt_handler._get_definition_enhancement_prompt({
                'name': category.name,
                'definition': category.definition,
                'examples': category.examples
            })
            
            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            enhanced_def = response.choices[0].message.content.strip()
            
            if len(enhanced_def.split()) >= 20:
                return category.replace(definition=enhanced_def)
            
        except Exception as e:
            print(f"Fehler bei Definition-Verbesserung: {str(e)}")
        
        return None
    
    async def analyze_grounded_batch(self, segments: List[str], material_percentage: float) -> Dict[str, Any]:
        """
        Analysiert einen Batch von Segmenten im 'grounded' Modus.
        Extrahiert Subcodes und Keywords ohne direkte Zuordnung zu Hauptkategorien.
        Sorgt für angemessenen Abstand zwischen Keywords und Subcodes.
        
        Args:
            segments: Liste der Textsegmente
            material_percentage: Prozentsatz des verarbeiteten Materials
            
        Returns:
            Dict[str, Any]: Analyseergebnisse mit Subcodes und Keywords
        """
        try:
            # Cache-Key erstellen
            cache_key = (
                tuple(segments),
                'grounded'
            )
            
            # Prüfe Cache
            if cache_key in self.analysis_cache:
                print("Nutze gecachte Analyse")
                return self.analysis_cache[cache_key]

            # Bestehende Subcodes sammeln
            existing_subcodes = []
            if hasattr(self, 'collected_subcodes'):
                existing_subcodes = [sc.get('name', '') for sc in self.collected_subcodes if isinstance(sc, dict)]
            
            # Definiere JSON-Schema für den grounded Modus
            json_schema = '''{
                "segment_analyses": [
                    {
                        "segment_text": "Textsegment",
                        "subcodes": [
                            {
                                "name": "Subcode-Name",
                                "definition": "Definition des Subcodes",
                                "evidence": ["Textbelege"],
                                "keywords": ["Schlüsselwörter des Subcodes"],
                                "confidence": 0.0-1.0
                            }
                        ],
                        "memo": "Analytische Notizen zum Segment"
                    }
                ],
                "abstraction_quality": {
                    "keyword_subcode_distinction": 0.0-1.0,
                    "comment": "Bewertung der Abstraktionshierarchie"
                },
                "saturation_metrics": {
                    "new_aspects_found": true/false,
                    "coverage": 0.0-1.0,
                    "justification": "Begründung"
                }
            }'''

            # Verbesserter Prompt mit Fokus auf die Abstraktionshierarchie
            prompt = self.prompt_handler.get_grounded_analysis_prompt(
                segments=segments,
                existing_subcodes=existing_subcodes,
                json_schema=json_schema
            )
            
            token_counter.start_request()

            # API-Call
            response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    response_format={"type": "json_object"}
                )
                
            # Verarbeite Response mit Wrapper
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)
            
            
            token_counter.track_response(response, self.model_name)

            # Cache das Ergebnis
            self.analysis_cache[cache_key] = result
            
            # Bewertung der Abstraktionsqualität
            abstraction_quality = result.get('abstraction_quality', {})
            if abstraction_quality and 'keyword_subcode_distinction' in abstraction_quality:
                quality_score = abstraction_quality['keyword_subcode_distinction']
                quality_comment = abstraction_quality.get('comment', '')
                print(f"\nAbstraktionsqualität: {quality_score:.2f}/1.0")
                print(f"Kommentar: {quality_comment}")
            
            # Debug-Ausgabe und verbesserte Fortschrittsanzeige
            segment_count = len(result.get('segment_analyses', []))
            
            # Zähle Subcodes und ihre Keywords
            subcode_count = 0
            keyword_count = 0
            new_subcodes = []
            
            for analysis in result.get('segment_analyses', []):
                subcodes = analysis.get('subcodes', [])
                subcode_count += len(subcodes)
                
                for subcode in subcodes:
                    new_subcodes.append(subcode)
                    keyword_count += len(subcode.get('keywords', []))
                    
                    # Zeige Abstraktionsbeispiele für besseres Monitoring
                    keywords = subcode.get('keywords', [])
                    if keywords and len(keywords) > 0:
                        print(f"\nAbstraktionsbeispiel:")
                        print(f"Keywords: {', '.join(keywords[:3])}" + ("..." if len(keywords) > 3 else ""))
                        print(f"Subcode: {subcode.get('name', '')}")
            
            # Erweiterte Fortschrittsanzeige
            print(f"\nGrounded Analyse für {segment_count} Segmente abgeschlossen:")
            print(f"- {subcode_count} neue Subcodes identifiziert")
            print(f"- {keyword_count} Keywords mit Subcodes verknüpft")
            print(f"- Material-Fortschritt: {material_percentage:.1f}%")
            
            # Progress Bar für Gesamtfortschritt der Subcode-Sammlung
            if hasattr(self, 'collected_subcodes'):
                total_collected = len(self.collected_subcodes) + subcode_count
                # Einfache ASCII Progress Bar
                bar_length = 30
                filled_length = int(bar_length * material_percentage / 100)
                bar = '█' * filled_length + '░' * (bar_length - filled_length)
                
                print(f"\nGesamtfortschritt Grounded-Analyse:")
                print(f"[{bar}] {material_percentage:.1f}%")
                print(f"Bisher gesammelt: {total_collected} Subcodes mit ihren Keywords")
            
            return result
            
        except Exception as e:
            print(f"Fehler bei Grounded-Analyse: {str(e)}")
            print("Details:")
            traceback.print_exc()
            return {}
    
    async def _prefilter_segments(self, segments: List[str]) -> List[str]:
        """
        Filtert Segmente nach Relevanz für Kategorienentwicklung.
        Optimiert durch Parallelverarbeitung und Caching.
        """
        async def check_segment(segment: str) -> Tuple[str, float]:
            cache_key = hash(segment)
            if cache_key in self.category_cache:
                return segment, self.category_cache[cache_key]
            
            relevance = await self._assess_segment_relevance(segment)
            self.category_cache[cache_key] = relevance
            return segment, relevance
        
        # Parallele Relevanzprüfung
        tasks = [check_segment(seg) for seg in segments]
        results = await asyncio.gather(*tasks)
        
        # Filter relevante Segmente
        return [seg for seg, relevance in results if relevance > self.MIN_CONFIDENCE]

    async def _assess_segment_relevance(self, segment: str) -> float:
        """
        Bewertet die Relevanz eines Segments für die Kategorienentwicklung.
        """
        prompt = self.prompt_handler.get_segment_relevance_assessment_prompt(segment)
                
        try:
            token_counter.start_request()

            response = await self.llm_provider.create_completion(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    response_format={"type": "json_object"}
                )
                
            # Verarbeite Response mit Wrapper
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)

            
            token_counter.track_response(response, self.model_name)

            return float(result.get('relevance_score', 0))
            
        except Exception as e:
            print(f"Error in relevance assessment: {str(e)}")
            return 0.0
    
    def _create_batches(self, segments: List[str], batch_size: int = None) -> List[List[str]]:
        """
        Creates batches of segments for processing.
        
        Args:
            segments: List of text segments to process
            batch_size: Optional custom batch size (defaults to self.BATCH_SIZE)
            
        Returns:
            List[List[str]]: List of segment batches
        """
        batch_size = self.BATCH_SIZE
            
        return [
            segments[i:i + batch_size] 
            for i in range(0, len(segments), batch_size)
        ]
    
    async def _generate_main_categories_from_subcodes(self, initial_categories: Dict[str, CategoryDefinition] = None) -> Dict[str, CategoryDefinition]:
        """
        Generiert Hauptkategorien aus den gesammelten Subcodes - VOLLSTÄNDIGE GROUNDED THEORY IMPLEMENTIERUNG
        """
        try:
            # Hole gesammelte Subcodes (mehrere Quellen probieren)
            collected_subcodes = []
            
            if hasattr(self, 'collected_subcodes') and self.collected_subcodes:
                collected_subcodes = self.collected_subcodes
                print(f"📚 Verwende Subcodes aus InductiveCoder: {len(collected_subcodes)}")
            elif hasattr(self, 'analysis_manager') and hasattr(self.analysis_manager, 'collected_subcodes'):
                collected_subcodes = self.analysis_manager.collected_subcodes
                print(f"📚 Verwende Subcodes aus AnalysisManager: {len(collected_subcodes)}")
            else:
                print("⚠️ Keine gesammelten Subcodes gefunden - prüfe verfügbare Attribute:")
                for attr in dir(self):
                    if 'subcode' in attr.lower():
                        print(f"   - {attr}: {getattr(self, attr, 'N/A')}")
                return initial_categories or {}
            
            if len(collected_subcodes) < 5:
                print(f"⚠️ Zu wenige Subcodes für Hauptkategorien-Generierung: {len(collected_subcodes)} < 5")
                return initial_categories or {}
            
            print(f"\n🔍 GROUNDED THEORY: Generiere Hauptkategorien aus {len(collected_subcodes)} Subcodes")
            
            # Bereite Subcodes für LLM-Analyse vor
            subcodes_data = []
            all_keywords = []
            
            for subcode in collected_subcodes:
                subcode_entry = {
                    'name': subcode.get('name', ''),
                    'definition': subcode.get('definition', ''),
                    'keywords': subcode.get('keywords', []),
                    'evidence': subcode.get('evidence', []),
                    'confidence': subcode.get('confidence', 0.7)
                }
                subcodes_data.append(subcode_entry)
                all_keywords.extend(subcode.get('keywords', []))
            
            # Zeige Statistiken
            keyword_counter = Counter(all_keywords)
            top_keywords = keyword_counter.most_common(15)
            avg_confidence = sum(s.get('confidence', 0) for s in collected_subcodes) / len(collected_subcodes)
            
            print(f"\n📊 Subcode-Analyse vor Hauptkategorien-Generierung:")
            print(f"   - Subcodes: {len(subcodes_data)}")
            print(f"   - Einzigartige Keywords: {len(set(all_keywords))}")
            print(f"   - Durchschnittliche Konfidenz: {avg_confidence:.2f}")
            print(f"   - Top Keywords: {', '.join([f'{kw}({count})' for kw, count in top_keywords[:8]])}")
            
            # Erstelle optimierten Prompt für Grounded Theory
            enhanced_prompt = self.prompt_handler.get_main_categories_generation_prompt(
                subcodes_data=subcodes_data,
                top_keywords=top_keywords,
                avg_confidence=avg_confidence
            )
                        
            # LLM-Aufruf
            print("\n⏳ Generiere Hauptkategorien via Grounded Theory Analyse...")
            
            token_counter.start_request()

            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für Grounded Theory und qualitative Inhaltsanalyse. Du antwortest auf deutsch."},
                    {"role": "user", "content": enhanced_prompt}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )
            
            llm_response = LLMResponse(response)
            result = json.loads(llm_response.content)
            
            
            token_counter.track_response(response, self.model_name)
            
            # Verarbeite Ergebnisse zu CategoryDefinition-Objekten
            grounded_categories = {}
            subcode_mapping = result.get('subcode_mappings', {})
            
            print(f"\n✅ Hauptkategorien-Generierung abgeschlossen:")
            
            for i, category_data in enumerate(result.get('main_categories', []), 1):
                name = category_data.get('name', '')
                definition = category_data.get('definition', '')
                
                if name and definition:
                    # Erstelle Subcategories aus zugeordneten Subcodes
                    subcategories = {}
                    assigned_subcodes = []
                    
                    for subcode_data in category_data.get('subcodes', []):
                        subcode_name = subcode_data.get('name', '')
                        subcode_definition = subcode_data.get('definition', '')
                        if subcode_name and subcode_definition:
                            subcategories[subcode_name] = subcode_definition
                            assigned_subcodes.append(subcode_name)
                    
                    # Erstelle CategoryDefinition
                    grounded_categories[name] = CategoryDefinition(
                        name=name,
                        definition=definition,
                        examples=category_data.get('examples', []),
                        rules=category_data.get('rules', []),
                        subcategories=subcategories,
                        added_date=datetime.now().strftime("%Y-%m-%d"),
                        modified_date=datetime.now().strftime("%Y-%m-%d")
                    )
                    
                    # Zeige Details
                    characteristic_keywords = ', '.join(category_data.get('characteristic_keywords', [])[:5])
                    print(f"   {i}. 📁 '{name}': {len(subcategories)} Subcodes zugeordnet")
                    print(f"      Keywords: {characteristic_keywords}")
                    print(f"      Subcodes: {', '.join(assigned_subcodes[:3])}{'...' if len(assigned_subcodes) > 3 else ''}")
            
            # Meta-Analyse Ergebnisse
            meta = result.get('meta_analysis', {})
            if meta:
                print(f"\n📈 Grounded Theory Meta-Analyse:")
                print(f"   - Verarbeitete Subcodes: {meta.get('total_subcodes_processed', len(subcodes_data))}")
                print(f"   - Generierte Hauptkategorien: {len(grounded_categories)}")
                print(f"   - Theoretische Sättigung: {meta.get('theoretical_saturation', 0):.2f}")
                print(f"   - Subcode-Abdeckung: {meta.get('coverage', 0):.2f}")
            
            # Prüfe Subcode-Zuordnung
            mapped_subcodes = set(subcode_mapping.values()) if subcode_mapping else set()
            all_subcode_names = set(s['name'] for s in subcodes_data)
            unmapped_subcodes = all_subcode_names - mapped_subcodes
            
            if unmapped_subcodes:
                print(f"\n⚠️ {len(unmapped_subcodes)} Subcodes wurden nicht zugeordnet:")
                for subcode in list(unmapped_subcodes)[:5]:
                    print(f"   - {subcode}")
                if len(unmapped_subcodes) > 5:
                    print(f"   ... und {len(unmapped_subcodes) - 5} weitere")
            else:
                print(f"\n✅ Alle {len(all_subcode_names)} Subcodes erfolgreich zugeordnet")
            
            # Kombiniere mit initial categories falls vorhanden
            if initial_categories:
                combined_categories = initial_categories.copy()
                for name, category in grounded_categories.items():
                    combined_categories[name] = category
                print(f"\n🔗 Kombiniert mit {len(initial_categories)} initialen Kategorien")
                return combined_categories
            
            return grounded_categories
            
        except Exception as e:
            print(f"❌ Fehler bei Grounded Theory Hauptkategorien-Generierung: {str(e)}")
            import traceback
            traceback.print_exc()
            return initial_categories or {}
        
    def _create_category_definition(self, cat_data: dict) -> CategoryDefinition:
        """
        Erstellt CategoryDefinition aus API-Response Dictionary
        GRUND: Wird für Kategorienentwicklung benötigt
        """
        try:
            return CategoryDefinition(
                name=cat_data.get('name', ''),
                definition=cat_data.get('definition', ''),
                examples=cat_data.get('evidence', []),
                rules=[],  # Wird später entwickelt
                subcategories={
                    sub.get('name', ''): sub.get('definition', '')
                    for sub in cat_data.get('subcategories', [])
                },
                added_date=datetime.now().strftime("%Y-%m-%d"),
                modified_date=datetime.now().strftime("%Y-%m-%d")
            )
        except Exception as e:
            print(f"Fehler bei CategoryDefinition-Erstellung: {str(e)}")
            return None

    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Berechnet die Ähnlichkeit zwischen zwei Texten mit Caching."""
        cache_key = f"{hash(text1)}_{hash(text2)}"
        
        if cache_key in self.similarity_cache:
            return self.similarity_cache[cache_key]
            
        # Konvertiere Texte zu Sets von Wörtern
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        # Berechne Jaccard-Ähnlichkeit
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        similarity = intersection / union if union > 0 else 0.0
        
        # Cache das Ergebnis
        self.similarity_cache[cache_key] = similarity
        
        return similarity

    def _find_similar_category(self, category: CategoryDefinition, existing_categories: Dict[str, CategoryDefinition]) -> Optional[str]:
        """
        Findet ähnliche Kategorie basierend auf Ähnlichkeitsschwelle
        """
        for existing_name, existing_cat in existing_categories.items():
            similarity = self._calculate_category_similarity(category, existing_cat)
            
            if similarity > self.similarity_threshold:
                print(f"🔍 Ähnliche Kategorie gefunden: '{category.name}' ↔ '{existing_name}' ({similarity:.2f})")
                return existing_name
        
        return None

    def _extract_base_segment_id(self, coding: Dict) -> str:
        """
        Extrahiert die Basis-Segment-ID für Reliabilitätsberechnung.
        Behandelt Mehrfachkodierung korrekt.
        
        Args:
            coding: Kodierung mit segment_id
            
        Returns:
            str: Basis-Segment-ID ohne Mehrfachkodierungs-Suffixe
        """
        segment_id = coding.get('segment_id', '')
        
        # Entferne Mehrfachkodierungs-Suffixe
        # Format kann sein: "doc_chunk_5" oder "doc_chunk_5-1" für Mehrfachkodierung
        if '-' in segment_id:
            # Prüfe ob es ein Mehrfachkodierungs-Suffix ist (endet mit -Zahl)
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit():
                base_id = parts[0]
            else:
                base_id = segment_id
        else:
            base_id = segment_id
        
        return base_id
    
    def _document_reliability_results(
        self, 
        alpha: float, 
        total_segments: int, 
        total_coders: int, 
        category_frequencies: dict
    ) -> str:
        """
        Generiert einen detaillierten Bericht über die Intercoder-Reliabilität.

        Args:
            alpha: Krippendorffs Alpha Koeffizient
            total_segments: Gesamtzahl der analysierten Segmente
            total_coders: Gesamtzahl der Kodierer
            category_frequencies: Häufigkeiten der Kategorien
            
        Returns:
            str: Formatierter Bericht als Markdown-Text
        """
        try:
            # Bestimme das Reliabilitätsniveau basierend auf Alpha
            reliability_level = (
                "Excellent" if alpha > 0.8 else
                "Acceptable" if alpha > 0.667 else
                "Poor"
            )
            
            # Erstelle den Bericht
            report = [
                "# Intercoder Reliability Analysis Report",
                f"\n## Overview",
                f"- Number of text segments: {total_segments}",
                f"- Number of coders: {total_coders}",
                f"- Krippendorff's Alpha: {alpha:.3f}",
                f"- Reliability Assessment: {reliability_level}",
                "\n## Category Usage",
                "| Category | Frequency |",
                "|----------|-----------|"
            ]
            
            # Füge Kategorienhäufigkeiten hinzu
            for category, frequency in sorted(category_frequencies.items(), key=lambda x: x[1], reverse=True):
                report.append(f"| {category} | {frequency} |")
            
            # Füge Empfehlungen hinzu
            report.extend([
                "\n## Recommendations",
                "Based on the reliability analysis, the following actions are suggested:"
            ])
            
            if alpha < 0.667:
                report.extend([
                    "1. Review and clarify category definitions",
                    "2. Provide additional coder training",
                    "3. Consider merging similar categories",
                    "4. Add more explicit coding rules"
                ])
            elif alpha < 0.8:
                report.extend([
                    "1. Review cases of disagreement",
                    "2. Refine coding guidelines for ambiguous cases",
                    "3. Consider additional coder calibration"
                ])
            else:
                report.extend([
                    "1. Continue with current coding approach",
                    "2. Document successful coding practices",
                    "3. Consider using this category system as a template for future analyses"
                ])
            
            # Füge detaillierte Analyse hinzu
            report.extend([
                "\n## Detailed Analysis",
                "### Interpretation of Krippendorff's Alpha",
                "- > 0.800: Excellent reliability",
                "- 0.667 - 0.800: Acceptable reliability",
                "- < 0.667: Poor reliability",
                "\n### Category Usage Analysis",
                "- Most frequently used category: " + max(category_frequencies.items(), key=lambda x: x[1])[0],
                "- Least frequently used category: " + min(category_frequencies.items(), key=lambda x: x[1])[0],
                "- Number of categories with single use: " + str(sum(1 for x in category_frequencies.values() if x == 1)),
                "\n### Coder Performance",
                "- Average segments per coder: " + f"{total_segments/total_coders:.1f}" if total_coders > 0 else "N/A"
            ])
            
            # Füge Zeitstempel hinzu
            report.append(f"\nReport generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            return '\n'.join(report)
            
        except Exception as e:
            print(f"Error generating reliability report: {str(e)}")
            import traceback
            traceback.print_exc()
            return "# Reliability Report\n\nError generating report"
    
  
    async def _meets_quality_standards(self, category: CategoryDefinition) -> bool:
        """
        Prüft ob Kategorie strikte Qualitätsstandards erfüllt
        VEREINFACHT für bessere Durchlässigkeit
        """
        # 1. Definition ausreichend lang (weiter reduziert)
        if len(category.definition.split()) < 5:  # reduziert von 10
            print(f"❌ '{category.name}': Definition zu kurz ({len(category.definition.split())} Wörter)")
            return False
        
        # 2. Genügend Beispiele (weiter reduziert) 
        if len(category.examples) < 1:  # reduziert von 2
            print(f"❌ '{category.name}': Zu wenige Beispiele ({len(category.examples)})")
            return False
        
        # 3. Name nicht zu kurz
        if len(category.name) < 3:
            print(f"❌ '{category.name}': Name zu kurz")
            return False
        
        print(f"✅ '{category.name}': Qualitätsstandards erfüllt")
        return True

    async def _auto_merge_categories(self, cat1: CategoryDefinition, cat2: CategoryDefinition, name1: str, name2: str) -> Optional[CategoryDefinition]:
        """
        Automatische intelligente Zusammenführung ähnlicher Kategorien
        """
        print(f"🔗 Automatische Zusammenführung: '{name1}' + '{name2}'")
        
        try:
            # Wähle besseren Namen
            better_name = self._choose_better_name(name1, name2)
            
            # Kombiniere Definitionen intelligent
            combined_definition = await self._merge_definitions_intelligent(cat1.definition, cat2.definition)
            
            # Kombiniere Beispiele (entferne Duplikate)
            combined_examples = list(set(cat1.examples + cat2.examples))
            
            # Kombiniere Regeln
            combined_rules = list(set(cat1.rules + cat2.rules))
            
            # Kombiniere Subkategorien
            combined_subcats = {**cat1.subcategories, **cat2.subcategories}
            
            # Erstelle zusammengeführte Kategorie
            merged = CategoryDefinition(
                name=better_name,
                definition=combined_definition,
                examples=combined_examples,
                rules=combined_rules,
                subcategories=combined_subcats,
                added_date=min(cat1.added_date, cat2.added_date),
                modified_date=datetime.now().strftime("%Y-%m-%d")
            )
            
            print(f"✅ Zusammenführung erfolgreich zu '{better_name}'")
            return merged
            
        except Exception as e:
            print(f"❌ Fehler bei automatischer Zusammenführung: {str(e)}")
            return None

    async def _merge_definitions_intelligent(self, def1: str, def2: str) -> str:
        """
        Intelligente Zusammenführung von Definitionen via LLM
        """
        prompt = self.prompt_handler.get_definition_enhancement_prompt({
            'definition1': def1,
            'definition2': def2
        })
                
        try:
            response = await self.llm_provider.create_completion(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Du bist ein Experte für qualitative Inhaltsanalyse."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            
            merged_def = response.choices[0].message.content.strip()
            
            # Fallback falls LLM-Merge fehlschlägt
            if len(merged_def.split()) < 15:
                return f"{def1} Zusätzlich umfasst dies: {def2}"
            
            return merged_def
            
        except Exception as e:
            print(f"Fehler bei Definition-Merge: {str(e)}")
            return f"{def1} Erweitert um: {def2}"

    def _calculate_category_similarity(self, cat1: CategoryDefinition, cat2: CategoryDefinition) -> float:
        """
        Berechnet Ähnlichkeit zwischen zwei Kategorien basierend auf mehreren Faktoren
        """
        # 1. Name-Ähnlichkeit (30%)
        name_similarity = self._calculate_text_similarity(cat1.name.lower(), cat2.name.lower()) * 0.3
        
        # 2. Definition-Ähnlichkeit (50%)
        def_similarity = self._calculate_text_similarity(cat1.definition, cat2.definition) * 0.5
        
        # 3. Subkategorien-Überlappung (20%)
        subcats1 = set(cat1.subcategories.keys())
        subcats2 = set(cat2.subcategories.keys())
        
        if subcats1 and subcats2:
            subcat_overlap = len(subcats1 & subcats2) / len(subcats1 | subcats2)
        else:
            subcat_overlap = 0
        
        subcat_similarity = subcat_overlap * 0.2
        
        total_similarity = name_similarity + def_similarity + subcat_similarity
        
        return min(total_similarity, 1.0)

    async def _merge_categories_intelligent(self, cat1: CategoryDefinition, cat2: CategoryDefinition, name1: str, name2: str) -> Optional[CategoryDefinition]:
        """
        Intelligente Zusammenführung mit Qualitätsprüfung
        """
        # Verwende die bereits implementierte _auto_merge_categories
        merged = await self._auto_merge_categories(cat1, cat2, name1, name2)
        
        if merged and await self._meets_quality_standards(merged):
            return merged
        
        print(f"❌ Zusammengeführte Kategorie erfüllt Qualitätsstandards nicht")
        return None

    def _choose_better_name(self, name1: str, name2: str) -> str:
        """
        Wählt den besseren Kategorienamen basierend auf Kriterien
        """
        # Kriterien für besseren Namen
        score1 = score2 = 0
        
        # 1. Länge (nicht zu kurz, nicht zu lang)
        if 5 <= len(name1) <= 25:
            score1 += 1
        if 5 <= len(name2) <= 25:
            score2 += 1
        
        # 2. Keine Sonderzeichen/Zahlen
        if name1.replace('_', '').replace('-', '').isalpha():
            score1 += 1
        if name2.replace('_', '').replace('-', '').isalpha():
            score2 += 1
        
        # 3. Keine englischen Wörter
        english_words = {'research', 'development', 'management', 'system', 'process', 'analysis'}
        if not any(word.lower() in english_words for word in name1.split('_')):
            score1 += 1
        if not any(word.lower() in english_words for word in name2.split('_')):
            score2 += 1
        
        # 4. Kürzerer Name bei Gleichstand
        if score1 == score2:
            return name1 if len(name1) <= len(name2) else name2
        
        return name1 if score1 > score2 else name2

    def _update_usage_history(self, category_names: List[str]) -> None:
        """
        Aktualisiert die Nutzungshistorie für Kategorien
        """
        for name in category_names:
            if name in self.category_usage_history:
                self.category_usage_history[name] += 1
            else:
                self.category_usage_history[name] = 1
        
        print(f"📊 Nutzungshistorie aktualisiert für: {category_names}")
        print(f"    Aktuelle Nutzung: {dict(list(self.category_usage_history.items())[-3:])}")
    
    def _log_performance(self, 
                        num_segments: int,
                        num_categories: int,
                        processing_time: float) -> None:
        """
        Protokolliert Performance-Metriken.
        """
        metrics = {
            'timestamp': datetime.now().isoformat(),
            'segments_processed': num_segments,
            'categories_developed': num_categories,
            'processing_time': processing_time,
            'segments_per_second': num_segments / processing_time
        }
        
        print("\nPerformance Metrics:")
        print(f"- Segments processed: {num_segments}")
        print(f"- Categories developed: {num_categories}")
        print(f"- Processing time: {processing_time:.2f}s")
        print(f"- Segments/second: {metrics['segments_per_second']:.2f}")
        
        # Speichere Metriken
        self.batch_results.append(metrics)
    def _find_similar_category(self, category: CategoryDefinition, existing_categories: Dict[str, CategoryDefinition]) -> Optional[str]:
        """
        Findet ähnliche existierende Kategorien basierend auf Namen und Definition.
        """
        try:
            best_match = None
            highest_similarity = 0.0
            
            for existing_name, existing_cat in existing_categories.items():
                # Berechne Ähnlichkeit basierend auf verschiedenen Faktoren
                
                # 1. Name-Ähnlichkeit (gewichtet: 0.3)
                name_similarity = self._calculate_text_similarity(
                    category.name.lower(),
                    existing_name.lower()
                ) * 0.3
                
                # 2. Definitions-Ähnlichkeit (gewichtet: 0.5)
                definition_similarity = self._calculate_text_similarity(
                    category.definition,
                    existing_cat.definition
                ) * 0.5
                
                # 3. Subkategorien-Überlappung (gewichtet: 0.2)
                subcats1 = set(category.subcategories.keys())
                subcats2 = set(existing_cat.subcategories.keys())
                if subcats1 and subcats2:
                    subcat_overlap = len(subcats1 & subcats2) / len(subcats1 | subcats2)
                else:
                    subcat_overlap = 0
                subcat_similarity = subcat_overlap * 0.2
                
                # Gesamtähnlichkeit
                total_similarity = name_similarity + definition_similarity + subcat_similarity
                
                # Update beste Übereinstimmung
                if total_similarity > highest_similarity:
                    highest_similarity = total_similarity
                    best_match = existing_name
            
            # Nur zurückgeben wenn Ähnlichkeit hoch genug
            if highest_similarity > 0.7:  # Schwellenwert für Ähnlichkeit
                print(f"\n⚠ Hohe Ähnlichkeit ({highest_similarity:.2f}) gefunden:")
                print(f"- Neue Kategorie: {category.name}")
                print(f"- Existierende Kategorie: {best_match}")
                return best_match
                
            return None
            
        except Exception as e:
            print(f"Fehler bei Ähnlichkeitsprüfung: {str(e)}")
            return None


# --- Klasse: ManualCoder ---
class ManualCoder:
    def __init__(self, coder_id: str):
        self.coder_id = coder_id
        self.root = None
        self.text_chunk = None
        self.category_listbox = None
        self.categories = {}
        self.current_coding = None
        self._is_processing = False
        self.current_categories = {}
        self.is_last_segment = False
        
        # NEU: Mehrfachkodierung-Support
        self.multiple_selection_enabled = True
        self.category_map = {}  # Mapping von Listbox-Index zu Kategorie-Info

        
    def _on_selection_change(self, event):
        """
        ERWEITERT: Behandelt Änderungen der Kategorienauswahl mit Nummern-Info
        """
        try:
            selected_indices = self.category_listbox.curselection()
            
            if not selected_indices:
                self.selection_info_label.config(text="Keine Auswahl", foreground='gray')
                return
            
            # Analysiere Auswahl mit Nummern
            selected_categories = []
            main_categories = set()
            selected_numbers = []
            
            for idx in selected_indices:
                if idx in self.category_map:
                    cat_info = self.category_map[idx]
                    selected_categories.append(cat_info)
                    main_categories.add(cat_info['main_category'])
                    selected_numbers.append(cat_info.get('number', '?'))
            
            # Erstelle Infotext mit Nummern
            if len(selected_indices) == 1:
                cat_info = selected_categories[0]
                number = cat_info.get('number', '?')
                if cat_info['type'] == 'main':
                    info_text = f"Nr. {number}: Hauptkategorie '{cat_info['name']}'"
                else:
                    info_text = f"Nr. {number}: Subkategorie '{cat_info['name']}' → {cat_info['main_category']}"
                self.selection_info_label.config(text=info_text, foreground='black')
            else:
                numbers_text = ", ".join(selected_numbers)
                if len(main_categories) == 1:
                    info_text = f"Nummern {numbers_text}: {len(selected_indices)} Subkategorien von '{list(main_categories)[0]}'"
                    self.selection_info_label.config(text=info_text, foreground='blue')
                else:
                    info_text = f"Nummern {numbers_text}: Mehrfachkodierung ({len(selected_indices)} Kategorien aus {len(main_categories)} Hauptkategorien)"
                    self.selection_info_label.config(text=info_text, foreground='orange')
        except Exception as e:
            print(f"Fehler bei Auswahlaktualisierung: {str(e)}")

    def get_category_by_number(self, number_input: str) -> dict:
        """
        NEU: Gibt Kategorie-Info basierend auf Nummern-Eingabe zurück
        
        Args:
            number_input: Eingabe wie "1", "1.2", "3" etc.
            
        Returns:
            dict: Kategorie-Information oder None wenn nicht gefunden
        """
        return self.number_to_category_map.get(number_input.strip())

    def update_category_list_enhanced(self):
        """
        ERWEITERT: Aktualisiert die Kategorienliste mit NUMMERIERUNG für bessere Übersicht
        """
        if not self.category_listbox:
            return
            
        self.category_listbox.delete(0, tk.END)
        self.category_map = {}
        self.number_to_category_map = {}  # NEU: Mapping von Nummern zu Kategorien
        
        current_index = 0
        main_category_number = 1
        
        # Sortiere Kategorien alphabetisch für bessere Übersicht
        sorted_categories = sorted(self.categories.items())
        
        for cat_name, cat_def in sorted_categories:
            # Hauptkategorie hinzufügen mit Nummerierung
            main_number = str(main_category_number)
            display_text = f"{main_number}. 📁 {cat_name}"
            self.category_listbox.insert(tk.END, display_text)
            
            # Mapping für Index und Nummer
            self.category_map[current_index] = {
                'type': 'main',
                'name': cat_name,
                'main_category': cat_name,
                'number': main_number
            }
            self.number_to_category_map[main_number] = {
                'type': 'main',
                'name': cat_name,
                'main_category': cat_name
            }
            
            current_index += 1
            sub_category_number = 1
            
            # Subkategorien hinzufügen (eingerückt und nummeriert)
            if hasattr(cat_def, 'subcategories') and cat_def.subcategories:
                sorted_subcats = sorted(cat_def.subcategories.items())
                for sub_name, sub_def in sorted_subcats:
                    sub_number = f"{main_category_number}.{sub_category_number}"
                    display_text = f"    {sub_number} 📄 {sub_name}"
                    self.category_listbox.insert(tk.END, display_text)
                    
                    # Mapping für Index und Nummer
                    self.category_map[current_index] = {
                        'type': 'sub',
                        'name': sub_name,
                        'main_category': cat_name,
                        'definition': sub_def,
                        'number': sub_number
                    }
                    self.number_to_category_map[sub_number] = {
                        'type': 'sub',
                        'name': sub_name,
                        'main_category': cat_name,
                        'definition': sub_def
                    }
                    
                    current_index += 1
                    sub_category_number += 1
            
            main_category_number += 1

        # Scrolle zum Anfang
        if self.category_listbox.size() > 0:
            self.category_listbox.see(0)
        
        print(f"Nummerierte Kategorieliste aktualisiert: {len(self.category_map)} Einträge")
        
        # Zeige Nummern-Referenz in einem Label an
        self._update_number_reference()

    def _update_number_reference(self):
        """
        NEU: Zeigt eine kompakte Nummern-Referenz für schnelle Eingabe
        """
        if not hasattr(self, 'number_reference_label'):
            return
            
        # Erstelle kompakte Referenz
        reference_lines = []
        main_cats = []
        
        for number, info in self.number_to_category_map.items():
            if info['type'] == 'main':
                main_cats.append(f"{number}={info['name'][:15]}")
            elif len(reference_lines) < 5:  # Zeige nur ersten paar Subkategorien
                reference_lines.append(f"{number}={info['name'][:10]}")
        
        # Kompakte Anzeige
        main_line = " | ".join(main_cats)
        sub_line = " | ".join(reference_lines[:3]) + ("..." if len(reference_lines) > 3 else "")
        
        reference_text = f"Hauptkat.: {main_line}\nSubkat.: {sub_line}"
        self.number_reference_label.config(text=reference_text)

    def _safe_code_selection_enhanced(self):
        """
        ERWEITERT: Thread-sichere Kodierungsauswahl mit Mehrfachkodierung-Support
        """
        if not self._is_processing:
            try:
                self._is_processing = True
                
                selected_indices = list(self.category_listbox.curselection())
                if not selected_indices:
                    messagebox.showwarning("Warnung", "Bitte wählen Sie mindestens eine Kategorie aus.")
                    self._is_processing = False
                    return
                
                # Validiere Auswahl
                from QCA_Utils import validate_multiple_selection
                is_valid, error_msg, selected_categories = validate_multiple_selection(
                    selected_indices, self.category_map
                )
                
                if not is_valid:
                    messagebox.showerror("Fehler", error_msg)
                    self._is_processing = False
                    return
                
                # Verarbeite Auswahl
                if len(selected_indices) == 1:
                    # Einzelauswahl - wie bisher
                    self._process_single_selection(selected_indices[0])
                else:
                    # Mehrfachauswahl - neue Logik
                    self._process_multiple_selection(selected_categories)
                
                # Bei letztem Segment Hinweis anzeigen
                if self.is_last_segment:
                    messagebox.showinfo(
                        "Kodierung abgeschlossen",
                        "Die Kodierung des letzten Segments wurde abgeschlossen.\n"
                        "Der manuelle Kodierungsprozess wird beendet."
                    )
                
                self._is_processing = False
                
                # Fenster schließen
                if self.root:
                    try:
                        self.root.destroy()
                        self.root.quit()
                    except Exception as e:
                        print(f"Fehler beim Schließen des Fensters: {str(e)}")
                        
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler bei der Kategorieauswahl: {str(e)}")
                print(f"Fehler bei der Kategorieauswahl: {str(e)}")
                import traceback
                traceback.print_exc()
                self._is_processing = False

    def _process_single_selection(self, index: int):
        """
        Verarbeitet eine Einzelauswahl (als Dictionary statt CodingResult)
        """
        category_info = self.category_map[index]
        
        if category_info['type'] == 'main':
            main_cat = category_info['name']
            sub_cat = None
        else:
            main_cat = category_info['main_category']
            sub_cat = category_info['name']
        
        # Verifiziere Kategorien
        if main_cat not in self.categories:
            messagebox.showerror(
                "Fehler",
                f"Hauptkategorie '{main_cat}' nicht gefunden.\n"
                f"Verfügbare Kategorien: {', '.join(self.categories.keys())}"
            )
            return
            
        if sub_cat and sub_cat not in self.categories[main_cat].subcategories:
            messagebox.showerror(
                "Fehler",
                f"Subkategorie '{sub_cat}' nicht in '{main_cat}' gefunden.\n"
                f"Verfügbare Subkategorien: {', '.join(self.categories[main_cat].subcategories.keys())}"
            )
            return

        # Erstelle Einzelkodierung als Dictionary
        self.current_coding = {
            'category': main_cat,
            'subcategories': [sub_cat] if sub_cat else [],
            'justification': "Manuelle Kodierung",
            'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
            'text_references': [self.text_chunk.get("1.0", tk.END)[:100]],
            'uncertainties': None,
            'paraphrase': "",
            'keywords': "",
            'manual_coding': True,
            'manual_multiple_coding': False,
            'multiple_coding_instance': 1,
            'total_coding_instances': 1,
            'coding_date': datetime.now().isoformat(),
            'coder_id': self.coder_id
        }
        
        print(f"Einzelkodierung erstellt: {main_cat}" + (f" → {sub_cat}" if sub_cat else ""))

    def _process_multiple_selection(self, selected_categories: List[Dict]):
        """
        NEUE METHODE: Verarbeitet Mehrfachauswahl von Kategorien
        """
        # Analysiere Auswahltyp
        main_categories = set(cat['main_category'] for cat in selected_categories)
        
        # Bestätigungsdialog anzeigen
        from QCA_Utils import show_multiple_coding_info
        confirmed = show_multiple_coding_info(
            self.root,
            len(selected_categories),
            list(main_categories)
        )
        
        if not confirmed:
            print("Mehrfachkodierung abgebrochen")
            return
        
        # Erstelle Mehrfachkodierung
        if len(main_categories) == 1:
            # Alle Auswahlen gehören zu einer Hauptkategorie
            main_cat = list(main_categories)[0]
            subcategories = [
                cat['name'] for cat in selected_categories 
                if cat['type'] == 'sub'
            ]
            
            # Füge Hauptkategorie hinzu, wenn sie direkt ausgewählt wurde
            main_cat_selected = any(
                cat['type'] == 'main' for cat in selected_categories
            )
            
            self.current_coding = CodingResult(
                category=main_cat,
                subcategories=tuple(subcategories),
                justification=f"Manuelle Kodierung mit {len(subcategories)} Subkategorien",
                confidence={'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
                text_references=(self.text_chunk.get("1.0", tk.END)[:100],)
            )
            
            print(f"Einzelkodierung mit mehreren Subkategorien: {main_cat} → {', '.join(subcategories)}")
        else:
            # Echte Mehrfachkodierung: verschiedene Hauptkategorien
            from QCA_Utils import create_multiple_coding_results
            
            coding_results = create_multiple_coding_results(
                selected_categories=selected_categories,
                text=self.text_chunk.get("1.0", tk.END),
                coder_id=self.coder_id
            )
            
            self.current_coding = coding_results
            
            main_cat_names = [result.category if hasattr(result, 'category') else result['category'] 
                            for result in coding_results]
            print(f"Mehrfachkodierung erstellt: {len(coding_results)} Kodierungen für {', '.join(main_cat_names)}")

    def _safe_finish_coding_enhanced(self):
        """
        ERWEITERT: Thread-sicherer Abschluss mit Mehrfachkodierung-Support
        """
        if not self._is_processing and self.is_last_segment:
            if messagebox.askyesno(
                "Segment kodieren und abschließen",
                "Möchten Sie das aktuelle Segment kodieren und den manuellen Kodierungsprozess abschließen?"
            ):
                # Verwende die erweiterte Kodierungslogik
                self._safe_code_selection_enhanced()

    def _safe_skip_chunk(self):
        """Thread-sicheres Überspringen (als Dictionary)"""
        if not self._is_processing:
            self.current_coding = {
                'category': "Nicht kodiert",
                'subcategories': [],
                'justification': "Chunk übersprungen",
                'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
                'text_references': [self.text_chunk.get("1.0", tk.END)[:100]],
                'uncertainties': None,
                'paraphrase': "",
                'keywords': "",
                'manual_coding': True,
                'manual_multiple_coding': False,
                'multiple_coding_instance': 1,
                'total_coding_instances': 1,
                'coding_date': datetime.now().isoformat(),
                'coder_id': self.coder_id
            }
            self._is_processing = False
            
            if self.is_last_segment:
                messagebox.showinfo(
                    "Kodierung abgeschlossen",
                    "Die Kodierung des letzten Segments wurde übersprungen.\n"
                    "Der manuelle Kodierungsprozess wird beendet."
                )
            
            # KORRIGIERT: Füge destroy() hinzu um das Fenster komplett zu schließen
            try:
                if self.root and self.root.winfo_exists():
                    self.root.quit()
                    self.root.destroy()  # HINZUGEFÜGT: Zerstört das Fenster komplett
                    self.root = None     # HINZUGEFÜGT: Setze Referenz auf None
            except Exception as e:
                print(f"Info: Fehler beim Schließen des Fensters: {str(e)}")
                # Fallback: Setze root auf None auch bei Fehlern
                self.root = None

    def _safe_abort_coding(self):
        """
        KORRIGIERT: Explizite Abbruch-Funktion (über Button)
        """
        if not self._is_processing:
            if messagebox.askyesno(
                "Kodierung komplett abbrechen",
                "Möchten Sie die gesamte manuelle Kodierung beenden?\n\n"
                "Alle bisher kodierten Segmente werden gespeichert."
            ):
                print("Benutzer hat manuelle Kodierung komplett abgebrochen")
                self.current_coding = "ABORT_ALL"
                self._is_processing = False
                
                try:
                    if self.root and self.root.winfo_exists():
                        self.root.quit()
                except Exception as e:
                    print(f"Info: Abbruch-Bereinigung: {str(e)}")

    def _safe_new_main_category(self):
        """Thread-sichere neue Hauptkategorie (unverändert)"""
        if not self._is_processing:
            from tkinter import simpledialog
            new_cat = simpledialog.askstring(
                "Neue Hauptkategorie",
                "Geben Sie den Namen der neuen Hauptkategorie ein:"
            )
            if new_cat:
                if new_cat in self.categories:
                    messagebox.showwarning("Warnung", "Diese Kategorie existiert bereits.")
                    return
                self.categories[new_cat] = CategoryDefinition(
                    name=new_cat,
                    definition="",
                    examples=[],
                    rules=[],
                    subcategories={},
                    added_date=datetime.now().strftime("%Y-%m-%d"),
                    modified_date=datetime.now().strftime("%Y-%m-%d")
                )
                self.update_category_list_enhanced()

    def _safe_new_sub_category_enhanced(self):
        """
        ERWEITERT: Neue Subkategorie mit Nummern-Eingabe
        """
        if not self._is_processing:
            from tkinter import simpledialog
            
            # Zeige verfügbare Hauptkategorien mit Nummern
            main_cats_info = []
            for number, info in self.number_to_category_map.items():
                if info['type'] == 'main':
                    main_cats_info.append(f"{number} = {info['name']}")
            
            if not main_cats_info:
                messagebox.showwarning("Warnung", "Keine Hauptkategorien verfügbar.")
                return
            
            # Erstelle Eingabedialog mit Nummern-Auswahl
            dialog_text = (
                "Verfügbare Hauptkategorien:\n" + 
                "\n".join(main_cats_info) + 
                "\n\nGeben Sie die Nummer der Hauptkategorie ein:"
            )
            
            main_cat_input = simpledialog.askstring(
                "Hauptkategorie auswählen (per Nummer)",
                dialog_text
            )
            
            if main_cat_input:
                # Prüfe ob Eingabe eine gültige Nummer ist
                main_cat_info = self.number_to_category_map.get(main_cat_input.strip())
                
                if main_cat_info and main_cat_info['type'] == 'main':
                    main_cat_name = main_cat_info['name']
                    
                    # Dialog für neue Subkategorie
                    new_sub = simpledialog.askstring(
                        "Neue Subkategorie",
                        f"Geben Sie den Namen der neuen Subkategorie für\n'{main_cat_name}' (Nr. {main_cat_input}) ein:"
                    )
                    
                    if new_sub:
                        if new_sub in self.categories[main_cat_name].subcategories:
                            messagebox.showwarning("Warnung", "Diese Subkategorie existiert bereits.")
                            return
                            
                        # Füge neue Subkategorie hinzu
                        self.categories[main_cat_name].subcategories[new_sub] = ""
                        
                        # Aktualisiere die Anzeige
                        self.update_category_list_enhanced()
                        
                        # Zeige Erfolg mit neuer Nummer
                        new_number = self._find_number_for_subcategory(main_cat_name, new_sub)
                        messagebox.showinfo(
                            "Subkategorie erstellt", 
                            f"'{new_sub}' wurde als Nr. {new_number} zu '{main_cat_name}' hinzugefügt"
                        )
                        
                elif main_cat_input.strip():
                    # Fallback: Versuche Namen-Eingabe
                    if main_cat_input.strip() in self.categories:
                        main_cat_name = main_cat_input.strip()
                        
                        new_sub = simpledialog.askstring(
                            "Neue Subkategorie",
                            f"Geben Sie den Namen der neuen Subkategorie für '{main_cat_name}' ein:"
                        )
                        
                        if new_sub and new_sub not in self.categories[main_cat_name].subcategories:
                            self.categories[main_cat_name].subcategories[new_sub] = ""
                            self.update_category_list_enhanced()
                            
                            new_number = self._find_number_for_subcategory(main_cat_name, new_sub)
                            messagebox.showinfo(
                                "Subkategorie erstellt", 
                                f"'{new_sub}' wurde als Nr. {new_number} hinzugefügt"
                            )
                    else:
                        messagebox.showwarning(
                            "Warnung", 
                            f"Ungültige Eingabe: '{main_cat_input}'\n\nBitte verwenden Sie die Nummer (z.B. '1') oder den exakten Namen der Hauptkategorie."
                        )

    def _find_number_for_subcategory(self, main_cat_name: str, sub_name: str) -> str:
        """
        NEU: Findet die Nummer einer Subkategorie
        """
        for number, info in self.number_to_category_map.items():
            if (info['type'] == 'sub' and 
                info['main_category'] == main_cat_name and 
                info['name'] == sub_name):
                return number
        return "?"

    def on_closing(self):
        """Sicheres Schließen des Fensters (unverändert)"""
        try:
            if messagebox.askokcancel("Beenden", "Möchten Sie das Kodieren wirklich beenden?"):
                self.current_coding = None
                self._is_processing = False
                
                if hasattr(self, 'root') and self.root:
                    for attr_name in dir(self):
                        attr = getattr(self, attr_name) 
                        if hasattr(attr, '_tk'):
                            delattr(self, attr_name)
                    
                    try:
                        self.root.quit()
                        self.root.destroy()
                        self.root = None
                    except:
                        pass
        except:
            if hasattr(self, 'root') and self.root:
                try:
                    self.root.quit()
                    self.root.destroy()
                    self.root = None
                except:
                    pass

    def _cleanup_tkinter_resources(self):
        """Bereinigt alle Tkinter-Ressourcen (unverändert)"""
        try:
            for attr_name in list(self.__dict__.keys()):
                if attr_name.startswith('_tk_var_'):
                    delattr(self, attr_name)
                    
            self.text_chunk = None
            self.category_listbox = None
            
            import gc
            gc.collect()
            
        except Exception as e:
            print(f"Warnung: Fehler bei der Bereinigung von Tkinter-Ressourcen: {str(e)}") 
        
    async def code_chunk(self, chunk: str, categories: Optional[Dict[str, CategoryDefinition]], is_last_segment: bool = False) -> Optional[Union[Dict, List[Dict]]]:
        """
        KORRIGIERT: Minimale Tkinter-Bereinigung nach MainLoop
        """
        try:
            self.categories = self.current_categories or categories
            self.current_coding = None
            self.is_last_segment = is_last_segment
            
            # Erstelle und starte das Tkinter-Fenster im Hauptthread
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, self._run_enhanced_tk_window, chunk)
            
            # KORRIGIERT: Prüfe auf ABORT_ALL BEVOR weitere Verarbeitung
            if self.current_coding == "ABORT_ALL":
                return "ABORT_ALL"
            
            # KORRIGIERT: Minimale Bereinigung - root sollte bereits None sein
            if hasattr(self, 'root'):
                self.root = None
            
            # Rest der Verarbeitung bleibt gleich...
            if self.current_coding:
                if isinstance(self.current_coding, list):
                    enhanced_codings = []
                    for coding_dict in self.current_coding:
                        enhanced_coding = coding_dict.copy()
                        enhanced_coding['text'] = chunk
                        enhanced_codings.append(enhanced_coding)
                    self.current_coding = enhanced_codings
                else:
                    if isinstance(self.current_coding, dict):
                        self.current_coding['text'] = chunk
                    else:
                        # CodingResult zu Dict konvertieren
                        self.current_coding = {
                            'category': self.current_coding.category,
                            'subcategories': list(self.current_coding.subcategories),
                            'justification': self.current_coding.justification,
                            'confidence': self.current_coding.confidence,
                            'text_references': list(self.current_coding.text_references),
                            'uncertainties': list(self.current_coding.uncertainties) if self.current_coding.uncertainties else None,
                            'paraphrase': getattr(self.current_coding, 'paraphrase', ''),
                            'keywords': getattr(self.current_coding, 'keywords', ''),
                            'text': chunk,
                            'manual_coding': True,
                            'manual_multiple_coding': False,
                            'multiple_coding_instance': 1,
                            'total_coding_instances': 1,
                            'coding_date': datetime.now().isoformat()
                        }
            
            # Debug-Ausgabe
            if isinstance(self.current_coding, list):
                result_status = f"Mehrfachkodierung mit {len(self.current_coding)} Kodierungen erstellt"
            elif self.current_coding == "ABORT_ALL":
                result_status = "Kodierung abgebrochen"
            elif self.current_coding:
                result_status = "Einzelkodierung erstellt"
            else:
                result_status = "Keine Kodierung"
                
            print(f"ManualCoder Ergebnis: {result_status}")
            
            # Finale Bereinigung (nur Ressourcen, nicht Tkinter)
            self._cleanup_tkinter_resources()
            
            return self.current_coding
            
        except Exception as e:
            print(f"Fehler in code_chunk: {str(e)}")
            import traceback
            traceback.print_exc()
            
            # Sichere Bereinigung auch im Fehlerfall
            if hasattr(self, 'root'):
                self.root = None
            self._cleanup_tkinter_resources()
            return None

    def _cleanup_tkinter_safely(self):
        """
        KORRIGIERT: Sichere Bereinigung nur im Hauptthread
        """
        try:
            if hasattr(self, 'root') and self.root:
                # Prüfe ob wir im Hauptthread sind
                if threading.current_thread() is threading.main_thread():
                    try:
                        # Prüfe ob das Fenster noch existiert
                        if self.root.winfo_exists():
                            self.root.quit()
                            self.root.destroy()
                            print("Tkinter-Fenster erfolgreich geschlossen")
                        
                    except tk.TclError:
                        # Fenster wurde bereits zerstört - das ist OK
                        print("Tkinter-Fenster war bereits geschlossen")
                        pass
                    except Exception as e:
                        # Andere Fehler - loggen aber nicht abbrechen
                        print(f"Info: Tkinter-Bereinigung: {str(e)}")
                else:
                    # Wir sind nicht im Hauptthread - nur Referenz entfernen
                    print("Tkinter-Bereinigung übersprungen (nicht im Hauptthread)")
                    
                # Referenz immer entfernen
                self.root = None
                    
        except Exception as e:
            print(f"Info: Tkinter-Bereinigung abgeschlossen: {str(e)}")

    def _run_enhanced_tk_window(self, chunk: str):
        """
        KORRIGIERT: Bessere Thread-Behandlung
        """
        try:
            # Vorherige Fenster sicher schließen (falls vorhanden)
            if hasattr(self, 'root') and self.root:
                try:
                    self.root.quit()
                    self.root.destroy()
                except:
                    pass
                self.root = None
            
            self.root = tk.Tk()
            self.root.title(f"Manueller Coder - {self.coder_id}")
            self.root.geometry("900x700")
            
            # KORRIGIERT: Protokoll für sicheres Schließen
            self.root.protocol("WM_DELETE_WINDOW", self._safe_window_close)
        
            # GUI erstellen
            self._create_enhanced_gui(chunk)
            
            # Fenster in den Vordergrund bringen
            self.root.lift()
            self.root.attributes('-topmost', True)
            self.root.attributes('-topmost', False)
            self.root.focus_force()
            
            # Plattformspezifische Anpassungen
            if platform.system() == "Darwin":  # macOS
                self.root.createcommand('tk::mac::RaiseWindow', self.root.lift)
            
            # KORRIGIERT: MainLoop mit sauberer Beendigung
            try:
                self.root.update()
                self.root.mainloop()
                
                # WICHTIG: Nach mainloop() ist das Fenster bereits geschlossen
                # Setze root auf None ohne weitere quit()/destroy() Aufrufe
                self.root = None
                print("Tkinter MainLoop beendet")
                
            except tk.TclError as tcl_error:
                if "application has been destroyed" in str(tcl_error):
                    # Das ist OK - Fenster wurde ordnungsgemäß geschlossen
                    self.root = None
                    print("Tkinter-Anwendung ordnungsgemäß beendet")
                else:
                    print(f"TclError: {tcl_error}")
                    self.root = None
            
        except Exception as e:
            print(f"Fehler beim Erstellen des Tkinter-Fensters: {str(e)}")
            self.current_coding = None
            self.root = None
            return

    def _create_enhanced_gui(self, chunk: str):
        """
        Erstellt die erweiterte GUI mit Mehrfachauswahl-Support
        """
        # Hauptframe
        main_frame = ttk.Frame(self.root)
        main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Header mit Instruktionen
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Titel
        title_label = ttk.Label(
            header_frame,
            text=f"Manueller Kodierer - {self.coder_id}",
            font=('Arial', 12, 'bold')
        )
        title_label.pack()

        # Instruktionen erweitert um Nummern-Eingabe
        instructions_label = ttk.Label(
            header_frame,
            text="💡 Tipp: Strg+Klick für Mehrfachauswahl • Shift+Klick für Bereichsauswahl\n" +
                "🔢 Neue Subkategorie: Nur Hauptkategorie-Nummer eingeben (z.B. '1' für erste Hauptkategorie)",
            font=('Arial', 9, 'italic'),
            foreground='blue'
        )
        instructions_label.pack(pady=(5, 0))
        
        
        # Fortschrittsinfo bei letztem Segment
        if self.is_last_segment:
            last_segment_label = ttk.Label(
                header_frame,
                text="🏁 LETZTES SEGMENT",
                font=('Arial', 10, 'bold'),
                foreground='red'
            )
            last_segment_label.pack(pady=(5, 0))

        # Textbereich
        text_frame = ttk.LabelFrame(main_frame, text="Textsegment")
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        self.text_chunk = tk.Text(text_frame, height=10, wrap=tk.WORD, font=('Arial', 10))
        text_scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.text_chunk.yview)
        self.text_chunk.config(yscrollcommand=text_scrollbar.set)
        
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_chunk.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.text_chunk.insert(tk.END, chunk)

         # Kategorienbereich mit Nummern-Referenz
        category_frame = ttk.LabelFrame(main_frame, text="🔢 Nummerierte Kategorien (Mehrfachauswahl möglich)")
        category_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # NEU: Nummern-Referenz Label (kompakt, oben)
        reference_frame = ttk.Frame(category_frame)
        reference_frame.pack(fill=tk.X, padx=5, pady=2)
        
        self.number_reference_label = ttk.Label(
            reference_frame,
            text="Lade Nummern-Referenz...",
            font=('Arial', 8),
            foreground='darkblue',
            background='lightgray',
            relief='sunken'
        )
        self.number_reference_label.pack(fill=tk.X, padx=2, pady=2)
        
        
        # Verwende die neue MultiSelectListbox aus QCA_Utils
        from QCA_Utils import MultiSelectListbox
        self.category_listbox = MultiSelectListbox(category_frame, font=('Arial', 10))
        
        cat_scrollbar = ttk.Scrollbar(category_frame, orient=tk.VERTICAL, command=self.category_listbox.yview)
        self.category_listbox.config(yscrollcommand=cat_scrollbar.set)
        
        cat_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.category_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        
        # Auswahlinfo
        selection_info_frame = ttk.Frame(main_frame)
        selection_info_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.selection_info_label = ttk.Label(
            selection_info_frame,
            text="Keine Auswahl",
            font=('Arial', 9),
            foreground='gray'
        )
        self.selection_info_label.pack()
        
        # Binding für Auswahl-Updates
        self.category_listbox.bind('<<ListboxSelect>>', self._on_selection_change)
        
        # Button-Frame mit klareren Beschriftungen
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        

        # Hauptbuttons mit klareren Texten
        if self.is_last_segment:
            ttk.Button(
                button_frame,
                text="Kodieren & Kodierung beenden",  # Klarerer Text
                command=self._safe_finish_coding_enhanced
            ).pack(side=tk.LEFT, padx=(0, 5))
        else:
            ttk.Button(
                button_frame,
                text="Kodieren & Weiter",  # Klarerer Text
                command=self._safe_code_selection_enhanced
            ).pack(side=tk.LEFT, padx=(0, 5))

        ttk.Button(
            button_frame,
            text="Neue Hauptkategorie",
            command=self._safe_new_main_category
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(
            button_frame,
            text="Neue Subkategorie (per Nr.)",  # Aktualisierter Text
            command=self._safe_new_sub_category_enhanced  # Neue Methode
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(
            button_frame,
            text="Segment überspringen",  # Klarerer Text
            command=self._safe_skip_chunk
        ).pack(side=tk.LEFT, padx=(0, 5))
        
        ttk.Button(
            button_frame,
            text="Kodierung komplett beenden",  # Klarerer Text für Abbruch
            command=self._safe_abort_coding
        ).pack(side=tk.RIGHT)
        
        # Kategorien laden
        self.update_category_list_enhanced()

    def _safe_window_close(self):
        """
        KORRIGIERT: Sichere Behandlung des Fenster-Schließens ohne Threading-Warnungen
        """
        try:
            # Bei X-Button-Klick → Kodierung überspringen (nicht abbrechen)
            if not self._is_processing:
                if messagebox.askyesno(
                    "Fenster schließen",
                    "Möchten Sie dieses Segment überspringen und zum nächsten wechseln?\n\n"
                    "Wählen Sie 'Nein' um zum Kodieren zurückzukehren."
                ):
                    # Segment überspringen
                    self.current_coding = {
                        'category': "Nicht kodiert",
                        'subcategories': [],
                        'justification': "Segment übersprungen (Fenster geschlossen)",
                        'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
                        'text_references': [],
                        'uncertainties': None,
                        'paraphrase': "",
                        'keywords': "",
                        'manual_coding': True,
                        'manual_multiple_coding': False,
                        'multiple_coding_instance': 1,
                        'total_coding_instances': 1,
                        'coding_date': datetime.now().isoformat()
                    }
                    
                    # Fenster sicher schließen
                    if self.root and self.root.winfo_exists():
                        self.root.quit()
                    
        except Exception as e:
            print(f"Info: Fenster-Schließung: {str(e)}")
            # Im Fehlerfall: Kodierung überspringen
            self.current_coding = None
            try:
                if self.root:
                    self.root.quit()
            except:
                pass


class ReviewManager:
    """
    KORRIGIERT: Zentrale Verwaltung aller Review-Modi mit kategorie-zentrierter Mehrfachkodierungs-Behandlung
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
    
    def process_coding_review(self, all_codings: List[Dict], export_mode: str) -> List[Dict]:
        """
        Hauptfunktion für alle Review-Modi mit korrekter Mehrfachkodierungs-Behandlung
        
        Args:
            all_codings: Alle ursprünglichen Kodierungen
            export_mode: 'consensus', 'majority', 'manual', etc.
            
        Returns:
            Liste der finalen, reviewten Kodierungen
        """
        print(f"\n=== REVIEW-PROZESS ({export_mode.upper()}) ===")
        
        # 1. FRÜHE SEGMENTIERUNG: Erkenne Mehrfachkodierungen und erstelle kategorie-spezifische Segmente
        category_segments = self._create_category_specific_segments(all_codings)
        
        # 2. REVIEW-PROZESS: Wende gewählten Modus auf kategorie-spezifische Segmente an
        if export_mode == 'manual':
            reviewed_codings = self._manual_review_process(category_segments)
        elif export_mode == 'majority':
            reviewed_codings = self._majority_review_process(category_segments)
        else:  # consensus (default)
            reviewed_codings = self._consensus_review_process(category_segments)
        
        print(f"Review abgeschlossen: {len(reviewed_codings)} finale Kodierungen")
        return reviewed_codings
    
    def _create_category_specific_segments(self, all_codings: List[Dict]) -> List[Dict]:
        """
        KERNFUNKTION: Erstelle kategorie-spezifische Segmente für korrekte Mehrfachkodierungs-Behandlung
        
        Verwandelt:
        - TEDFWI-1: [Akteure, Kontextfaktoren, Legitimation]
        
        In:
        - TEDFWI-1-01: [Akteure] (alle Akteure-Kodierungen für Segment TEDFWI-1)
        - TEDFWI-1-02: [Kontextfaktoren] (alle Kontextfaktoren-Kodierungen für Segment TEDFWI-1)  
        - TEDFWI-1-03: [Legitimation] (alle Legitimation-Kodierungen für Segment TEDFWI-1)
        """
        print("🔄 Erstelle kategorie-spezifische Segmente...")
        
        # Gruppiere nach ursprünglicher Segment-ID
        original_segments = defaultdict(list)
        for coding in all_codings:
            segment_id = coding.get('segment_id', '')
            if segment_id:
                # Extrahiere ursprüngliche Segment-ID (falls bereits erweitert)
                original_id = self._extract_original_segment_id(segment_id)
                original_segments[original_id].append(coding)
        
        # Erstelle kategorie-spezifische Segmente
        category_segments = []
        
        for original_id, codings in original_segments.items():
            # Identifiziere alle Hauptkategorien für dieses Segment
            categories = set()
            for coding in codings:
                category = coding.get('category', '')
                if category and category not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                    categories.add(category)
            
            if len(categories) <= 1:
                # Einfachkodierung oder keine gültigen Kategorien
                if categories:
                    category_segments.append({
                        'segment_id': original_id,
                        'original_segment_id': original_id,
                        'target_category': list(categories)[0],
                        'codings': codings,
                        'is_multiple_coding': False,
                        'instance_info': {'instance_number': 1, 'total_instances': 1}
                    })
                else:
                    # Keine gültigen Kategorien - behalte ursprüngliches Segment
                    category_segments.append({
                        'segment_id': original_id,
                        'original_segment_id': original_id,
                        'target_category': None,
                        'codings': codings,
                        'is_multiple_coding': False,
                        'instance_info': {'instance_number': 1, 'total_instances': 1}
                    })
            else:
                # MEHRFACHKODIERUNG: Erstelle separate Segmente pro Kategorie
                sorted_categories = sorted(categories)  # Konsistente Sortierung für ID-Zuordnung
                total_instances = len(sorted_categories)
                
                for i, category in enumerate(sorted_categories, 1):
                    # Neue Segment-ID mit kategorie-spezifischem Suffix
                    new_segment_id = f"{original_id}-{i:02d}"
                    
                    # Filtere Kodierungen für diese spezifische Kategorie
                    category_codings = [
                        coding for coding in codings 
                        if coding.get('category', '') == category
                    ]
                    
                    category_segments.append({
                        'segment_id': new_segment_id,
                        'original_segment_id': original_id,
                        'target_category': category,
                        'codings': category_codings,
                        'is_multiple_coding': True,
                        'instance_info': {
                            'instance_number': i,
                            'total_instances': total_instances,
                            'category_rank': i,
                            'all_categories': sorted_categories
                        }
                    })
                
                print(f"  📊 Mehrfachkodierung {original_id}: {len(sorted_categories)} Kategorien → {len(sorted_categories)} Segmente")
        
        print(f"✅ {len(category_segments)} kategorie-spezifische Segmente erstellt")
        return category_segments
    
    def _extract_original_segment_id(self, segment_id: str) -> str:
        """
        Extrahiert die ursprüngliche Segment-ID (entfernt Mehrfachkodierungs-Suffixe)
        
        Beispiele:
        - "TEDFWI-1-01" → "TEDFWI-1"
        - "TEDFWI-1" → "TEDFWI-1"
        - "doc_chunk_5-02" → "doc_chunk_5"
        """
        # Prüfe auf Mehrfachkodierungs-Suffix (Format: -XX wo XX eine Zahl ist)
        if '-' in segment_id:
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit() and len(parts[1]) <= 2:
                return parts[0]
        return segment_id
    
    def _consensus_review_process(self, category_segments: List[Dict]) -> List[Dict]:
        """
        Consensus-Review für kategorie-spezifische Segmente
        """
        print("🔍 Führe Consensus-Review durch...")
        reviewed_codings = []
        
        for segment in category_segments:
            codings = segment['codings']
            if len(codings) == 1:
                # Nur eine Kodierung - übernehme direkt
                final_coding = codings[0].copy()
                final_coding['segment_id'] = segment['segment_id']
                reviewed_codings.append(final_coding)
            else:
                # Mehrere Kodierungen - führe Consensus durch
                consensus_coding = self._get_consensus_for_category_segment(segment)
                if consensus_coding:
                    reviewed_codings.append(consensus_coding)
        
        return reviewed_codings
    
    def _majority_review_process(self, category_segments: List[Dict]) -> List[Dict]:
        """
        Majority-Review für kategorie-spezifische Segmente
        """
        print("🗳️ Führe Majority-Review durch...")
        reviewed_codings = []
        
        for segment in category_segments:
            codings = segment['codings']
            if len(codings) == 1:
                # Nur eine Kodierung - übernehme direkt
                final_coding = codings[0].copy()
                final_coding['segment_id'] = segment['segment_id']
                reviewed_codings.append(final_coding)
            else:
                # Mehrere Kodierungen - führe Majority durch
                majority_coding = self._get_majority_for_category_segment(segment)
                if majority_coding:
                    reviewed_codings.append(majority_coding)
        
        return reviewed_codings
    
    def _manual_review_process(self, category_segments: List[Dict]) -> List[Dict]:
        """
        FIX: Korrigierter manueller Review-Prozess ohne Event Loop Konflikt
        Für ReviewManager Klasse - verwendet bestehende Methoden und behält Sortierreihenfolge bei
        """
        print("👤 Führe manuelles Review durch...")
        
        # Identifiziere Segmente, die Review benötigen
        segments_needing_review = []
        for segment in category_segments:
            if len(segment['codings']) > 1:
                # FIX: Verwende bestehende Methode zur Unstimmigkeits-Prüfung
                if self._has_category_disagreement(segment):
                    segments_needing_review.append(segment)
        
        if not segments_needing_review:
            print("✅ Kein manueller Review erforderlich - alle Segmente haben eindeutige Kodierungen")
            return self._consensus_review_process(category_segments)
        
        print(f"🎯 {len(segments_needing_review)} kategorie-spezifische Segmente benötigen Review:")
        for segment in segments_needing_review:
            category = segment.get('category', 'Unbekannt')
            segment_id = segment['segment_id']
            print(f"  📋 {segment_id}: {category} (Teil {segment_id.split('-')[-1]} von {segment_id.rsplit('-', 1)[0]})")
        
        # FIX: Verwende asyncio.create_task() statt loop.run_until_complete()
        print("🎮 Starte GUI-basiertes manuelles Review...")

        try:
            # Konvertiere segments_needing_review zu dem Format, das ManualReviewComponent erwartet
            segment_codings = {}
            for segment in segments_needing_review:
                segment_id = segment['segment_id']
                segment_codings[segment_id] = segment['codings']
            
            # Importiere und verwende ManualReviewComponent für echtes GUI
            from QCA_Utils import ManualReviewGUI, ManualReviewComponent
            import asyncio
            
            # Erstelle ManualReviewComponent
            manual_review_component = ManualReviewComponent(self.output_dir)
            
            # FIX: Verwende asyncio.create_task() für bereits laufende Event Loop
            import concurrent.futures
            
            # Führe GUI-Review in separatem Thread aus
            def run_gui_review():
                try:
                    # Erstelle neue Event Loop für diesen Thread
                    new_loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(new_loop)
                    
                    try:
                        # FIX: Verwende neue Methode, die Unstimmigkeits-Prüfung überspringt
                        review_decisions = new_loop.run_until_complete(
                            manual_review_component.review_discrepancies_direct(segment_codings, skip_discrepancy_check=True)
                        )
                        return review_decisions
                    finally:
                        new_loop.close()
                        
                except Exception as e:
                    print(f"❌ Fehler im GUI-Thread: {e}")
                    return None
            
            # Führe GUI-Review in separatem Thread aus
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(run_gui_review)
                review_decisions = future.result(timeout=300)  # 5 Minuten Timeout
            
            if review_decisions is None:
                raise Exception("GUI-Review fehlgeschlagen")
                
            print(f"✅ GUI-Review abgeschlossen: {len(review_decisions)} Entscheidungen getroffen")
            
        except Exception as e:
            print(f"❌ Fehler beim GUI-Review: {e}")
            print("📝 Verwende automatischen Consensus als Fallback")
            import traceback
            traceback.print_exc()
            
            # FIX: Fallback ohne problematische Coroutine-Aufrufe
            return self._consensus_review_process(category_segments)
        
        # FIX: Kombiniere Review-Entscheidungen mit Segmenten ohne Review IN KORREKTER REIHENFOLGE
        reviewed_codings = []
        review_decisions_dict = {decision['segment_id']: decision for decision in review_decisions}
        
        # Durchlaufe category_segments in ursprünglicher Reihenfolge
        for segment in category_segments:
            segment_id = segment['segment_id']
            
            if segment_id in review_decisions_dict:
                # Verwende manuelle Review-Entscheidung
                reviewed_codings.append(review_decisions_dict[segment_id])
            else:
                # Segment ohne Review - verwende bestehende Logik
                if len(segment['codings']) == 1:
                    final_coding = segment['codings'][0].copy()
                    final_coding['segment_id'] = segment['segment_id']
                    reviewed_codings.append(final_coding)
                else:
                    # FIX: Verwende bestehende Consensus-Methode für Fallback
                    consensus_coding = self._get_consensus_for_category_segment(segment)
                    if consensus_coding:
                        reviewed_codings.append(consensus_coding)
        
        return reviewed_codings
    
    def _has_category_disagreement(self, segment: Dict) -> bool:
        """
        Prüft, ob es echte Unstimmigkeiten innerhalb einer Kategorie gibt
        
        Da alle Kodierungen bereits auf eine Kategorie gefiltert sind,
        prüfen wir hauptsächlich Subkategorien-Unstimmigkeiten
        """
        codings = segment['codings']
        
        # Vergleiche Subkategorien
        subcategory_sets = []
        for coding in codings:
            subcats = set(coding.get('subcategories', []))
            subcategory_sets.append(subcats)
        
        # Prüfe auf Unterschiede in Subkategorien
        if len(set(frozenset(s) for s in subcategory_sets)) > 1:
            return True
        
        return False
    
    def _get_consensus_for_category_segment(self, segment: Dict) -> Optional[Dict]:
        """
        Ermittelt Consensus für ein kategorie-spezifisches Segment
        """
        codings = segment['codings']
        target_category = segment['target_category']
        
        if not codings:
            return None
        
        # Da alle Kodierungen bereits die gleiche Hauptkategorie haben,
        # konzentrieren wir uns auf Subkategorien-Consensus
        
        # Sammle alle Subkategorien für diese Kategorie
        all_subcategories = []
        for coding in codings:
            subcats = coding.get('subcategories', [])
            all_subcategories.extend(subcats)
        
        # Zähle Subkategorien-Häufigkeiten
        from collections import Counter
        subcat_counts = Counter(all_subcategories)
        
        # Consensus-Subkategorien: Nur die, die von der Mehrheit gewählt wurden
        total_coders = len(codings)
        consensus_threshold = total_coders // 2 + 1  # Mehr als die Hälfte
        
        consensus_subcats = []
        for subcat, count in subcat_counts.items():
            if count >= consensus_threshold:
                consensus_subcats.append(subcat)
        
        # Wähle beste Kodierung als Basis
        best_coding = max(codings, key=lambda x: self._extract_confidence_value(x))
        consensus_coding = best_coding.copy()
        
        # Aktualisiere mit Consensus-Informationen
        consensus_coding.update({
            'segment_id': segment['segment_id'],
            'category': target_category,  # Bereits gefiltert
            'subcategories': consensus_subcats,
            'consensus_info': {
                'total_coders': total_coders,
                'selection_type': 'consensus',
                'subcat_consensus_threshold': consensus_threshold,
                'original_segment_id': segment['original_segment_id'],
                'is_multiple_coding_instance': segment['is_multiple_coding'],
                'instance_info': segment['instance_info']
            }
        })
        
        return consensus_coding
    
    def _get_majority_for_category_segment(self, segment: Dict) -> Optional[Dict]:
        """
        Ermittelt Majority für ein kategorie-spezifisches Segment
        """
        codings = segment['codings']
        target_category = segment['target_category']
        
        if not codings:
            return None
        
        # Wähle beste Kodierung als Basis (höchste Konfidenz)
        best_coding = max(codings, key=lambda x: self._extract_confidence_value(x))
        majority_coding = best_coding.copy()
        
        # Aktualisiere mit Majority-Informationen
        majority_coding.update({
            'segment_id': segment['segment_id'],
            'category': target_category,
            'consensus_info': {
                'total_coders': len(codings),
                'selection_type': 'majority',
                'confidence_based_selection': True,
                'original_segment_id': segment['original_segment_id'],
                'is_multiple_coding_instance': segment['is_multiple_coding'],
                'instance_info': segment['instance_info']
            }
        })
        
        return majority_coding
    
    def _extract_confidence_value(self, coding: Dict) -> float:
        """
        Extrahiert Konfidenzwert aus Kodierung
        """
        confidence = coding.get('confidence', {})
        if isinstance(confidence, dict):
            return confidence.get('total', 0.0)
        elif isinstance(confidence, (int, float)):
            return float(confidence)
        return 0.0


# ============================
# KORRIGIERTE KRIPPENDORFF'S ALPHA BERECHNUNG
# ============================

class ReliabilityCalculator:
    """
    FIX: Einheitliche Krippendorff's Alpha Berechnung nach Krippendorff (2011)
    Alle Reliabilitäts-Berechnungen laufen über diese Klasse
    """
    
    def __init__(self):
        self.debug = True
    
    def _extract_base_segment_id(self, coding: Dict) -> str:
        """
        FIX: Extrahiert Basis-Segment-ID ohne Mehrfachkodierungs-Suffixe
        Für ReliabilityCalculator Klasse
        """
        segment_id = coding.get('segment_id', '')
        
        # Entferne Mehrfachkodierungs-Suffixe
        # Format kann sein: "doc_chunk_5" oder "doc_chunk_5-1" für Mehrfachkodierung
        if '-' in segment_id:
            # Prüfe ob es ein Mehrfachkodierungs-Suffix ist (endet mit -Zahl)
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit():
                base_id = parts[0]
            else:
                base_id = segment_id
        else:
            base_id = segment_id
        
        return base_id
    
    def calculate_comprehensive_reliability(self, codings: List[Dict]) -> dict:
        """
        FIX: Aktualisierte Hauptmethode mit robusteren Berechnungen
        Für ReliabilityCalculator Klasse
        """
        print("\n📊 Umfassende Krippendorff's Alpha Analyse...")
        
        # FIX: Robusterer Filter
        original_codings = self._filter_original_codings(codings)
        
        if len(original_codings) < 2:
            print("⚠️ Zu wenige ursprüngliche Kodierungen für Reliabilitätsanalyse")
            return self._create_empty_reliability_report()
        
        # Basis-Statistiken (mit Fallback)
        statistics = self._calculate_basic_statistics(original_codings)
        
        # 1. Overall Alpha (kombinierte Sets) - extrahiere Float-Wert
        overall_alpha = self._calculate_combined_sets_alpha(original_codings)

        # 2. Hauptkategorien Alpha - extrahiere Float-Wert
        main_categories_alpha = self._calculate_main_categories_alpha(original_codings)

        # 3. FIX: Subkategorien Alpha mit partieller Übereinstimmung - bereits Float
        subcategories_alpha = self._calculate_subcategories_alpha(original_codings)
        
        # 4. Detaillierte Übereinstimmungsanalyse
        agreement_analysis = self._calculate_detailed_agreement_analysis(original_codings)
        
        reliability_report = {
            'overall_alpha': overall_alpha,
            'main_categories_alpha': main_categories_alpha,
            'subcategories_alpha': subcategories_alpha,
            'agreement_analysis': agreement_analysis,
            'statistics': statistics
        }
        
        self._print_reliability_summary(reliability_report)
        
        return reliability_report
    
    def calculate_reliability(self, all_codings: List[Dict]) -> float:
        """
        FIX: Hauptmethode - gibt Overall Alpha als Float zurück (für Rückwärtskompatibilität)
        """
        report = self.calculate_comprehensive_reliability(all_codings)
        overall_alpha = report['overall_alpha']
        # FIX: Stelle sicher, dass es ein Float ist
        if isinstance(overall_alpha, dict):
            return overall_alpha.get('alpha', 0.0)
        return float(overall_alpha)
    
    def _filter_original_codings(self, codings: List[Dict]) -> List[Dict]:
        """
        FIX: Robusterer Filter für ursprüngliche Kodierungen
        Für ReliabilityCalculator Klasse
        """
        original_codings = []
        
        # print(f"🔍 Debug Filter - Input: {len(codings)} Kodierungen")
        
        for i, coding in enumerate(codings):
            coder_id = coding.get('coder_id', '')
            consensus_info = coding.get('consensus_info', {})
            manual_review = coding.get('manual_review', False)
            selection_type = consensus_info.get('selection_type', '')
            
            # FIX: Debug-Ausgabe für erste 3 Kodierungen
            # if i < 3:
            #     print(f"  Kodierung {i}: coder_id='{coder_id}', manual_review={manual_review}, selection_type='{selection_type}'")
            
            # FIX: Weniger strenger Filter - akzeptiere mehr Kodierungen
            is_excluded = (
                coder_id in ['consensus', 'majority', 'review'] or
                manual_review == True or
                selection_type in ['consensus', 'majority', 'manual_consensus']
            )
            
            if not is_excluded:
                original_codings.append(coding)
            elif i < 3:
                print(f"    -> Ausgeschlossen")
        
        print(f"🔍 Gefilterte ursprüngliche Kodierungen: {len(original_codings)}")
        
        # FIX: Falls zu wenige gefunden, weniger streng filtern
        if len(original_codings) < 2:
            print("⚠️ Zu wenige gefunden - verwende weniger strengen Filter...")
            original_codings = []
            
            for coding in codings:
                coder_id = coding.get('coder_id', '')
                # FIX: Nur explizite Review-Resultate ausschließen
                if coder_id not in ['consensus', 'majority', 'review']:
                    original_codings.append(coding)
            
            print(f"🔍 Mit weniger strengem Filter: {len(original_codings)} Kodierungen")
        
        return original_codings

    def _calculate_combined_sets_alpha(self, codings: List[Dict]) -> float:
        """
        FIX: Overall Alpha mit Jaccard-Ähnlichkeit (konsistent mit Subkategorien-Behandlung)
        Für ReliabilityCalculator Klasse
        """
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        all_comparisons = 0
        
        for segment_id, coders_data in segment_data.items():
            if len(coders_data) < 2:
                continue
            
            coders = list(coders_data.keys())
            
            for i in range(len(coders)):
                for j in range(i + 1, len(coders)):
                    coder1, coder2 = coders[i], coders[j]
                    
                    # FIX: Kombiniere Haupt- und Subkategorien zu Sets
                    set1 = set()
                    set2 = set()
                    
                    for coding in coders_data[coder1]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set1.add(main_cat)
                        
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set1.update(subcats)
                    
                    for coding in coders_data[coder2]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set2.add(main_cat)
                        
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set2.update(subcats)
                    
                    all_comparisons += 1
                    
                    # FIX: Jaccard-Ähnlichkeit statt exakter Gleichheit
                    if len(set1) == 0 and len(set2) == 0:
                        # Beide haben keine Kategorien - perfekte Übereinstimmung
                        overlap_score = 1.0
                    elif len(set1) == 0 or len(set2) == 0:
                        # Einer hat keine, der andere schon - keine Übereinstimmung
                        overlap_score = 0.0
                    else:
                        # Jaccard-Koeffizient: |Schnittmenge| / |Vereinigungsmenge|
                        intersection = len(set1.intersection(set2))
                        union = len(set1.union(set2))
                        overlap_score = intersection / union if union > 0 else 0.0
                    
                    all_overlap_scores.append(overlap_score)
        
        if all_comparisons == 0:
            return 0.0
        
        # Durchschnittliche Übereinstimmung
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores)
        
        # Erwartete Übereinstimmung für kombinierte Sets
        expected_agreement = 0.25  # Konservative Schätzung
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        print(f"📊 Overall Alpha Details:")
        print(f"   • Durchschnittliche Jaccard-Übereinstimmung: {observed_agreement:.3f}")
        print(f"   • Erwartete Zufallsübereinstimmung: {expected_agreement:.3f}")
        print(f"   • Overall Alpha (Jaccard-basiert): {alpha:.3f}")
        
        return max(0.0, alpha)
    
    def _calculate_main_categories_alpha(self, codings: List[Dict]) -> float:
        """
        FIX: Hauptkategorien Alpha mit Jaccard-Ähnlichkeit (für Konsistenz)
        Für ReliabilityCalculator Klasse
        """
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        all_comparisons = 0
        
        for segment_id, coders_data in segment_data.items():
            if len(coders_data) < 2:
                continue
            
            coders = list(coders_data.keys())
            
            for i in range(len(coders)):
                for j in range(i + 1, len(coders)):
                    coder1, coder2 = coders[i], coders[j]
                    
                    # Nur Hauptkategorien sammeln
                    set1 = set()
                    set2 = set()
                    
                    for coding in coders_data[coder1]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set1.add(main_cat)
                    
                    for coding in coders_data[coder2]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set2.add(main_cat)
                    
                    all_comparisons += 1
                    
                    # Jaccard-Ähnlichkeit
                    if len(set1) == 0 and len(set2) == 0:
                        overlap_score = 1.0
                    elif len(set1) == 0 or len(set2) == 0:
                        overlap_score = 0.0
                    else:
                        intersection = len(set1.intersection(set2))
                        union = len(set1.union(set2))
                        overlap_score = intersection / union if union > 0 else 0.0
                    
                    all_overlap_scores.append(overlap_score)
        
        if all_comparisons == 0:
            return 0.0
        
        # Durchschnittliche Übereinstimmung
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores)
        
        # Erwartete Übereinstimmung
        expected_agreement = 0.20  # Für Hauptkategorien
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        print(f"📊 Hauptkategorien Alpha Details:")
        print(f"   • Durchschnittliche Übereinstimmung: {observed_agreement:.3f}")
        print(f"   • Hauptkategorien Alpha (Jaccard): {alpha:.3f}")
        
        return max(0.0, alpha)

    
    def _calculate_subcategories_alpha_old(self, codings: List[Dict]) -> dict:
        """
        FIX: Korrigierte Subkategorien Alpha Berechnung - behandelt Dictionary-Struktur korrekt
        Für ReliabilityCalculator Klasse
        """
        # FIX: Korrekte Gruppierung nach Segment-ID und Kodierer
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        total_comparisons = 0
        total_agreements = 0
        
        # FIX: Iteriere über die Dictionary-Struktur korrekt
        for original_id, coder_data in segment_data.items():
            # Extrahiere alle Kodierungen für dieses Segment von verschiedenen Kodierern
            all_segment_codings = []
            for coder_id, coder_codings in coder_data.items():
                all_segment_codings.extend(coder_codings)
            
            # Mindestens 2 Kodierungen pro Segment nötig für Vergleich
            if len(all_segment_codings) < 2:
                continue
                
            # FIX: Paarweise Vergleiche zwischen allen Kodierungen dieses Segments
            for i in range(len(all_segment_codings)):
                for j in range(i + 1, len(all_segment_codings)):
                    subcats1 = all_segment_codings[i].get('subcategories', [])
                    subcats2 = all_segment_codings[j].get('subcategories', [])
                    
                    # Normalisiere zu Sets
                    if isinstance(subcats1, (list, tuple)):
                        set1 = set(subcats1)
                    else:
                        set1 = set()
                    
                    if isinstance(subcats2, (list, tuple)):
                        set2 = set(subcats2)
                    else:
                        set2 = set()
                    
                    total_comparisons += 1
                    
                    # Jaccard-Ähnlichkeit für Sets
                    if len(set1) == 0 and len(set2) == 0:
                        overlap = 1.0  # Beide leer = perfekte Übereinstimmung
                        total_agreements += 1
                    elif len(set1.union(set2)) == 0:
                        overlap = 0.0
                    else:
                        overlap = len(set1.intersection(set2)) / len(set1.union(set2))
                        if overlap == 1.0:
                            total_agreements += 1
                    
                    # FIX: Nur numerische Overlap-Werte hinzufügen, KEINE String-Werte
                    all_overlap_scores.append(overlap)
                    
                    # FIX: Entfernt - keine String-Subkategorien zu all_overlap_scores hinzufügen
                    # all_overlap_scores.extend(list(set1))  # ENTFERNT
                    # all_overlap_scores.extend(list(set2))  # ENTFERNT
        
        if total_comparisons == 0:
            return {'alpha': 0.0, 'observed_agreement': 0.0, 'expected_agreement': 0.25, 'comparisons': 0}
        
        # FIX: Durchschnittliche Jaccard-Ähnlichkeit als beobachtete Übereinstimmung
        # Jetzt sind alle Werte in all_overlap_scores garantiert Float-Werte
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores) if len(all_overlap_scores) > 0 else 0.0
        expected_agreement = 0.25  # Vereinfacht
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        return {
            'alpha': max(0.0, alpha),
            'observed_agreement': observed_agreement,
            'expected_agreement': expected_agreement,
            'comparisons': total_comparisons
        }
    
    def _calculate_subcategories_alpha(self, codings: List[Dict]) -> float:
        """
        FIX: Subkategorien Alpha mit partieller Übereinstimmung (Jaccard-ähnlich)
        Für ReliabilityCalculator Klasse
        
        Behandelt: "subcat1, subcat2" vs. "subcat1, subcat3" als partielle Übereinstimmung
        statt als komplette Nicht-Übereinstimmung
        """
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        all_comparisons = 0
        
        print(f"🔍 Debug Subkategorien: Analysiere {len(segment_data)} Segmente")
        
        for segment_id, coders_data in segment_data.items():
            if len(coders_data) < 2:
                continue
            
            coders = list(coders_data.keys())
            
            for i in range(len(coders)):
                for j in range(i + 1, len(coders)):
                    coder1, coder2 = coders[i], coders[j]
                    
                    # Sammle Subkategorien für beide Kodierer
                    set1 = set()
                    set2 = set()
                    
                    for coding in coders_data[coder1]:
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set1.update(subcats)
                    
                    for coding in coders_data[coder2]:
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set2.update(subcats)
                    
                    all_comparisons += 1
                    
                    # FIX: Jaccard-Ähnlichkeit statt exakter Gleichheit
                    if len(set1) == 0 and len(set2) == 0:
                        # Beide haben keine Subkategorien - perfekte Übereinstimmung
                        overlap_score = 1.0
                    elif len(set1) == 0 or len(set2) == 0:
                        # Einer hat keine, der andere schon - keine Übereinstimmung
                        overlap_score = 0.0
                    else:
                        # Jaccard-Koeffizient: |Schnittmenge| / |Vereinigungsmenge|
                        intersection = len(set1.intersection(set2))
                        union = len(set1.union(set2))
                        overlap_score = intersection / union if union > 0 else 0.0
                    
                    all_overlap_scores.append(overlap_score)
                    
                    # Debug für erste 3 Vergleiche
                    if all_comparisons <= 3:
                        print(f"  Vergleich {all_comparisons}: {list(set1)} vs {list(set2)} → {overlap_score:.3f}")
        
        if all_comparisons == 0:
            print("⚠️ Keine Subkategorien-Vergleiche möglich")
            return 0.0
        
        # Durchschnittliche Übereinstimmung
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores)
        
        # FIX: Erwartete Zufallsübereinstimmung für partielle Übereinstimmung
        # Vereinfachte Berechnung: Bei zufälliger Verteilung würde man etwa 0.2-0.3 erwarten
        expected_agreement = 0.25  # Konservative Schätzung
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        print(f"📊 Subkategorien-Alpha Details:")
        print(f"   • Vergleiche durchgeführt: {all_comparisons}")
        print(f"   • Durchschnittliche Übereinstimmung: {observed_agreement:.3f}")
        print(f"   • Erwartete Zufallsübereinstimmung: {expected_agreement:.3f}")
        print(f"   • Partielle Übereinstimmungs-Alpha: {alpha:.3f}")
        
        return max(0.0, alpha)
    
    def _calculate_detailed_agreement_analysis(self, codings: List[Dict]) -> dict:
        """
        FIX: Korrigierte detaillierte Übereinstimmungsanalyse nach Kategorien
        Für ReliabilityCalculator Klasse
        """
        # FIX: Korrekte Gruppierung nach Segment-ID und Kodierer
        segment_data = self._group_by_original_segments(codings)
        
        agreement_stats = {
            'Vollständige Übereinstimmung': 0,
            'Hauptkategorie gleich, Subkat. unterschiedlich': 0,
            'Hauptkategorie unterschiedlich': 0
        }
        
        # FIX: Iteriere über die Dictionary-Struktur korrekt
        for original_id, coder_data in segment_data.items():
            # Extrahiere alle Kodierungen für dieses Segment von verschiedenen Kodierern
            all_segment_codings = []
            for coder_id, coder_codings in coder_data.items():
                all_segment_codings.extend(coder_codings)
            
            # Mindestens 2 Kodierungen pro Segment nötig für Vergleich
            if len(all_segment_codings) < 2:
                continue
                
            # FIX: Paarweise Vergleiche zwischen allen Kodierungen dieses Segments
            for i in range(len(all_segment_codings)):
                for j in range(i + 1, len(all_segment_codings)):
                    coding1 = all_segment_codings[i]
                    coding2 = all_segment_codings[j]
                    
                    main_cat1 = coding1.get('category', '')
                    main_cat2 = coding2.get('category', '')
                    subcats1 = set(coding1.get('subcategories', []))
                    subcats2 = set(coding2.get('subcategories', []))
                    
                    if main_cat1 == main_cat2 and subcats1 == subcats2:
                        agreement_stats['Vollständige Übereinstimmung'] += 1
                    elif main_cat1 == main_cat2:
                        agreement_stats['Hauptkategorie gleich, Subkat. unterschiedlich'] += 1
                    else:
                        agreement_stats['Hauptkategorie unterschiedlich'] += 1
        
        return agreement_stats
    
    def _group_by_original_segments(self, codings: List[Dict]) -> dict:
        """
        FIX: Gruppiert Kodierungen nach ursprünglicher Segment-ID für ReliabilityCalculator
        Rückgabe: {segment_id: {coder_id: [codings]}}
        Für ReliabilityCalculator Klasse
        """
        segment_data = {}
        
        for coding in codings:
            # Extrahiere ursprüngliche Segment-ID
            original_id = self._extract_base_segment_id(coding)
            
            if original_id not in segment_data:
                segment_data[original_id] = {}
            
            coder_id = coding.get('coder_id', 'unknown')
            if coder_id not in segment_data[original_id]:
                segment_data[original_id][coder_id] = []
            
            segment_data[original_id][coder_id].append(coding)
        
        return segment_data
    
    def _calculate_alpha_from_sets(self, agreements: int, comparisons: int, category_sets: List[set]) -> float:
        """
        FIX: Berechnet Krippendorff's Alpha aus Set-Daten
        """
        if comparisons == 0:
            return 0.0
        
        observed_agreement = agreements / comparisons
        expected_agreement = self._calculate_expected_set_agreement(category_sets)
        
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        return max(0.0, alpha)
    
    def _calculate_expected_set_agreement(self, all_category_sets: List[set]) -> float:
        """
        FIX: Berechnet erwartete Set-Übereinstimmung nach Krippendorff (2011)
        """
        if not all_category_sets:
            return 0.0
        
        # Sammle alle individuellen Kategorien
        all_individual_categories = []
        for category_set in all_category_sets:
            all_individual_categories.extend(list(category_set))
        
        if not all_individual_categories:
            return 0.0
        
        # Häufigkeitsverteilung
        from collections import Counter
        category_frequencies = Counter(all_individual_categories)
        total_instances = len(all_individual_categories)
        
        # Vereinfachte erwartete Übereinstimmung für Set-Variable
        # (Krippendorff 2011 empfiehlt komplexere Berechnung, aber das ist eine praktikable Näherung)
        expected_agreement = 0.0
        unique_sets = list(set(frozenset(s) for s in all_category_sets))
        
        for unique_set in unique_sets:
            # Wahrscheinlichkeit dieses Set-Typs
            set_probability = 1.0
            for category in unique_set:
                cat_prob = category_frequencies[category] / total_instances
                set_probability *= cat_prob
            
            expected_agreement += set_probability ** 2
        
        return min(expected_agreement, 0.99)  # Verhindere Division durch 0
    
    def _extract_original_segment_id(self, coding: Dict) -> str:
        """
        FIX: Extrahiert ursprüngliche Segment-ID
        """
        # Erst consensus_info prüfen
        consensus_info = coding.get('consensus_info', {})
        if consensus_info.get('original_segment_id'):
            return consensus_info['original_segment_id']
        
        # Fallback: Aus segment_id ableiten
        segment_id = coding.get('segment_id', '')
        if '-' in segment_id:
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit() and len(parts[1]) <= 2:
                return parts[0]
        
        return segment_id
    
    def _calculate_basic_statistics(self, codings: List[Dict]) -> dict:
        """
        FIX: Robuste Basis-Statistiken mit Fallback
        Für ReliabilityCalculator Klasse
        """
        if not codings:
            return {
                'total_codings': 0,
                'vergleichbare_segmente': 0,
                'total_segmente': 0,
                'anzahl_kodierer': 0,
                'mittelwert_kodierungen': 0.0
            }
        
        segment_data = self._group_by_original_segments(codings)
        vergleichbare_segmente = sum(1 for data in segment_data.values() if len(data) >= 2)
        
        coders = set()
        for coding in codings:
            coder_id = coding.get('coder_id', 'unknown')
            if coder_id:  # FIX: Leere coder_ids ignorieren
                coders.add(coder_id)
        
        anzahl_kodierer = len(coders) if coders else 1  # FIX: Mindestens 1 um Division durch 0 zu vermeiden
        
        return {
            'total_codings': len(codings),
            'vergleichbare_segmente': vergleichbare_segmente,
            'total_segmente': len(segment_data),
            'anzahl_kodierer': anzahl_kodierer,
            'mittelwert_kodierungen': len(codings) / anzahl_kodierer
        }
    
    def _create_empty_reliability_report(self) -> dict:
        """
        FIX: Erstellt vollständigen leeren Bericht
        Für ReliabilityCalculator Klasse
        """
        return {
            'overall_alpha': 0.0,
            'main_categories_alpha': 0.0,
            'subcategories_alpha': 0.0,
            'agreement_analysis': {
                'Vollständige Übereinstimmung': 0,
                'Hauptkategorie gleich, Subkat. unterschiedlich': 0,
                'Hauptkategorie unterschiedlich': 0,
                'Gesamt': 0
            },
            'statistics': {
                'total_codings': 0,
                'vergleichbare_segmente': 0,
                'total_segmente': 0,
                'anzahl_kodierer': 0,
                'mittelwert_kodierungen': 0.0
            }
        }
    
    def _calculate_combined_sets_alpha(self, codings: List[Dict]) -> float:
        """
        FIX: Overall Alpha mit Jaccard-Ähnlichkeit (konsistent mit Subkategorien-Behandlung)
        Für ReliabilityCalculator Klasse
        """
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        all_comparisons = 0
        
        for segment_id, coders_data in segment_data.items():
            if len(coders_data) < 2:
                continue
            
            coders = list(coders_data.keys())
            
            for i in range(len(coders)):
                for j in range(i + 1, len(coders)):
                    coder1, coder2 = coders[i], coders[j]
                    
                    # FIX: Kombiniere Haupt- und Subkategorien zu Sets
                    set1 = set()
                    set2 = set()
                    
                    for coding in coders_data[coder1]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set1.add(main_cat)
                        
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set1.update(subcats)
                    
                    for coding in coders_data[coder2]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set2.add(main_cat)
                        
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            set2.update(subcats)
                    
                    all_comparisons += 1
                    
                    # FIX: Jaccard-Ähnlichkeit statt exakter Gleichheit
                    if len(set1) == 0 and len(set2) == 0:
                        # Beide haben keine Kategorien - perfekte Übereinstimmung
                        overlap_score = 1.0
                    elif len(set1) == 0 or len(set2) == 0:
                        # Einer hat keine, der andere schon - keine Übereinstimmung
                        overlap_score = 0.0
                    else:
                        # Jaccard-Koeffizient: |Schnittmenge| / |Vereinigungsmenge|
                        intersection = len(set1.intersection(set2))
                        union = len(set1.union(set2))
                        overlap_score = intersection / union if union > 0 else 0.0
                    
                    all_overlap_scores.append(overlap_score)
        
        if all_comparisons == 0:
            return 0.0
        
        # Durchschnittliche Übereinstimmung
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores)
        
        # Erwartete Übereinstimmung für kombinierte Sets
        expected_agreement = 0.25  # Konservative Schätzung
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        print(f"📊 Overall Alpha Details:")
        print(f"   • Durchschnittliche Jaccard-Übereinstimmung: {observed_agreement:.3f}")
        print(f"   • Erwartete Zufallsübereinstimmung: {expected_agreement:.3f}")
        print(f"   • Overall Alpha (Jaccard-basiert): {alpha:.3f}")
        
        return max(0.0, alpha)

    def _calculate_main_categories_alpha(self, codings: List[Dict]) -> float:
        """
        FIX: Hauptkategorien Alpha mit Jaccard-Ähnlichkeit (für Konsistenz)
        Für ReliabilityCalculator Klasse
        """
        segment_data = self._group_by_original_segments(codings)
        
        all_overlap_scores = []
        all_comparisons = 0
        
        for segment_id, coders_data in segment_data.items():
            if len(coders_data) < 2:
                continue
            
            coders = list(coders_data.keys())
            
            for i in range(len(coders)):
                for j in range(i + 1, len(coders)):
                    coder1, coder2 = coders[i], coders[j]
                    
                    # Nur Hauptkategorien sammeln
                    set1 = set()
                    set2 = set()
                    
                    for coding in coders_data[coder1]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set1.add(main_cat)
                    
                    for coding in coders_data[coder2]:
                        main_cat = coding.get('category', '')
                        if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                            set2.add(main_cat)
                    
                    all_comparisons += 1
                    
                    # Jaccard-Ähnlichkeit
                    if len(set1) == 0 and len(set2) == 0:
                        overlap_score = 1.0
                    elif len(set1) == 0 or len(set2) == 0:
                        overlap_score = 0.0
                    else:
                        intersection = len(set1.intersection(set2))
                        union = len(set1.union(set2))
                        overlap_score = intersection / union if union > 0 else 0.0
                    
                    all_overlap_scores.append(overlap_score)
        
        if all_comparisons == 0:
            return 0.0
        
        # Durchschnittliche Übereinstimmung
        observed_agreement = sum(all_overlap_scores) / len(all_overlap_scores)
        
        # Erwartete Übereinstimmung
        expected_agreement = 0.20  # Für Hauptkategorien
        
        # Krippendorff's Alpha
        if expected_agreement >= 1.0:
            alpha = 1.0 if observed_agreement >= 1.0 else 0.0
        else:
            alpha = (observed_agreement - expected_agreement) / (1 - expected_agreement)
        
        print(f"📊 Hauptkategorien Alpha Details:")
        print(f"   • Durchschnittliche Übereinstimmung: {observed_agreement:.3f}")
        print(f"   • Hauptkategorien Alpha (Jaccard): {alpha:.3f}")
        
        return max(0.0, alpha)

    def _print_reliability_summary(self, report: dict):
        """
        FIX: Erweiterte Zusammenfassung mit Konsistenz-Prüfung
        Für ReliabilityCalculator Klasse
        """
        print(f"\n📈 Krippendorff's Alpha Reliabilitäts-Analyse:")
        print(f"=" * 60)
        print(f"Overall Alpha (Jaccard-basiert):       {report['overall_alpha']:.3f}")
        print(f"Hauptkategorien Alpha (Jaccard):       {report['main_categories_alpha']:.3f}")
        print(f"Subkategorien Alpha (Jaccard):         {report['subcategories_alpha']:.3f}")
        print(f"Vergleichbare Segmente:                {report['statistics']['vergleichbare_segmente']}")
        print(f"Anzahl Kodierer:                       {report['statistics']['anzahl_kodierer']}")
        
        # FIX: Konsistenz-Prüfung
        overall = report['overall_alpha']
        main_alpha = report['main_categories_alpha']
        sub_alpha = report['subcategories_alpha']
        
        print(f"\n💡 Methodik:")
        print(f"   • Alle Alpha-Werte verwenden Jaccard-Ähnlichkeit")
        print(f"   • Konsistente Set-basierte Berechnung")
        print(f"   • Overall sollte zwischen Haupt- und Sub-Alpha liegen")
        
        # Konsistenz-Check
        min_component = min(main_alpha, sub_alpha)
        max_component = max(main_alpha, sub_alpha)
        
        if min_component <= overall <= max_component:
            print(f"   ✅ Mathematische Konsistenz: {min_component:.3f} ≤ {overall:.3f} ≤ {max_component:.3f}")
        else:
            print(f"   ⚠️ Mathematische Inkonsistenz: Overall liegt außerhalb der Komponenten!")
            print(f"      Bereich: {min_component:.3f} - {max_component:.3f}, Overall: {overall:.3f}")
        
        # Bewertung
        rating = "Exzellent" if overall > 0.8 else "Akzeptabel" if overall > 0.667 else "Unzureichend"
        print(f"\nBewertung Overall Alpha:               {rating}")
        
        if overall < 0.667:
            print(f"⚠️  Reliabilität unter Schwellenwert - Kategoriensystem überarbeiten")


# --- Klasse: ResultsExporter ---
# Aufgabe: Export der kodierten Daten und des finalen Kategoriensystems
class ResultsExporter:
    """
    Exports the results of the analysis in various formats (JSON, CSV, Excel, Markdown).
    Supports the documentation of coded chunks and the category system.
    """
    
    def __init__(self, output_dir: str, attribute_labels: Dict[str, str], inductive_coder: InductiveCoder = None,
                 analysis_manager: 'IntegratedAnalysisManager' = None):
        self.output_dir = output_dir
        self.attribute_labels = attribute_labels
        self.inductive_coder = inductive_coder
        self.analysis_manager = analysis_manager
        self.relevance_checker = analysis_manager.relevance_checker if analysis_manager else None
        self.category_colors = {}
        self.current_categories = {}  # Wird von main() gesetzt

        os.makedirs(output_dir, exist_ok=True)

        # Importierte Funktionen als Instanzmethoden verfügbar machen
        self._sanitize_text_for_excel = _sanitize_text_for_excel
        self._generate_pastel_colors = _generate_pastel_colors
        self._format_confidence = _format_confidence

    def _get_consensus_coding(self, segment_codes: List[Dict]) -> Dict:
        """
        FIX: Einheitliche Consensus-Bildung mit robuster Subkategorien-Validierung
        Für IntegratedAnalysisManager Klasse
        """
        if not segment_codes:
            return {}

        # FIX: Debug-Info über Eingabe-Kodierungen
        categories = [coding.get('category', 'UNKNOWN') for coding in segment_codes]
        unique_categories = list(set(categories))
        
        print(f"🔍 DEBUG Consensus: {len(segment_codes)} Kodierungen, Kategorien: {unique_categories}")
        
        # FIX: Detaillierte Analyse der Subkategorien VOR Consensus
        for i, coding in enumerate(segment_codes):
            cat = coding.get('category', 'UNKNOWN')
            subcats = coding.get('subcategories', [])
            print(f"   Kodierung {i+1}: {cat} → {subcats}")
        
        # Wenn alle dieselbe Hauptkategorie haben, normale Konsensbildung
        if len(unique_categories) == 1:
            return self._get_single_consensus_coding(segment_codes)
        
        # Mehrfachkodierung: Erstelle präzises Kategorie-Subkategorie-Mapping
        print(f"🔀 Mehrfachkodierung erkannt mit Kategorien: {unique_categories}")
        
        best_coding = None
        highest_confidence = 0
                
        for coding in segment_codes:
            category = coding.get('category', '')
            subcats = coding.get('subcategories', [])
            confidence = self._extract_confidence_value(coding)
            
            print(f"   Prüfe Kodierung: {category} (Subkat: {len(subcats)}, Konfidenz: {confidence:.2f})")
                                   
            # Globale beste Kodierung
            if confidence > highest_confidence:
                highest_confidence = confidence
                best_coding = coding

        if best_coding:
            consensus_coding = best_coding.copy()
            
            # FIX: Konsistente Subkategorien-Behandlung mit detailliertem Logging
            main_category = consensus_coding.get('category', '')
            original_subcats = best_coding.get('subcategories', [])
            
            print(f"🎯 Beste Kodierung gewählt: {main_category}")
            print(f"   Original Subkategorien: {original_subcats}")
            
            # FIX: Verwende IMMER die robuste CategoryValidator-Methode
            try:
                validated_subcats = CategoryValidator.validate_subcategories_for_category(
                    original_subcats, main_category, self.current_categories, warn_only=False
                )
                
                print(f"   Nach Validierung: {validated_subcats}")
                
                # FIX: Dokumentiere Validierungsaktionen
                removed_subcats = set(original_subcats) - set(validated_subcats)
                if removed_subcats:
                    print(f"   🔧 ENTFERNT: {removed_subcats}")
                    # FIX: Füge Validierungs-Info zur Begründung hinzu
                    original_justification = consensus_coding.get('justification', '')
                    consensus_coding['justification'] = f"{original_justification} [FIX: Subkategorien-Validierung entfernte: {list(removed_subcats)}]"
                
                consensus_coding['subcategories'] = validated_subcats
                
            except Exception as e:
                print(f"⚠️ FEHLER bei Subkategorien-Validierung: {str(e)}")
                print(f"   Fallback: Verwende ursprüngliche Subkategorien ohne Validierung")
                consensus_coding['subcategories'] = original_subcats
            
            # FIX: Füge Validierungs-Metadaten hinzu
            consensus_coding['validation_applied'] = True
            consensus_coding['original_subcategory_count'] = len(original_subcats)
            consensus_coding['validated_subcategory_count'] = len(consensus_coding['subcategories'])
            
            return consensus_coding
        
        # Fallback: Erste Kodierung verwenden
        print("⚠️ FALLBACK: Verwende erste verfügbare Kodierung")
        fallback_coding = segment_codes[0] if segment_codes else {}
        
        # FIX: Auch Fallback-Kodierung validieren
        if fallback_coding:
            try:
                main_cat = fallback_coding.get('category', '')
                orig_subcats = fallback_coding.get('subcategories', [])
                
                validated_subcats = CategoryValidator.validate_subcategories_for_category(
                    orig_subcats, main_cat, self.current_categories, warn_only=False
                )
                
                fallback_coding['subcategories'] = validated_subcats
                fallback_coding['validation_applied'] = True
                fallback_coding['is_fallback_coding'] = True
                
            except Exception as e:
                print(f"⚠️ Fallback-Validierung fehlgeschlagen: {str(e)}")
        
        return fallback_coding

    def _get_majority_coding(self, segment_codes: List[Dict]) -> Optional[Dict]:
        """
        VEREINFACHT: Nutzt dieselbe Logik wie Schlüsselwörter - nimmt aus bester Kodierung
        """
        if not segment_codes:
            return None

        print(f"\nMehrheitsentscheidung für Segment mit {len(segment_codes)} Kodierungen...")

        # 1. Zähle Hauptkategorien
        category_counts = Counter(coding['category'] for coding in segment_codes)
        total_coders = len(segment_codes)
        
        # Finde häufigste Hauptkategorie(n)
        max_count = max(category_counts.values())
        majority_categories = [
            category for category, count in category_counts.items()
            if count == max_count
        ]
        
        print(f"  Kategorieverteilung: {dict(category_counts)}")
        print(f"  Häufigste Kategorie(n): {majority_categories} ({max_count}/{total_coders})")
        
        # 2. Bei eindeutiger Mehrheit
        if len(majority_categories) == 1:
            majority_category = majority_categories[0]
            print(f"  ✓ Eindeutige Mehrheit für: '{majority_category}'")
        else:
            # 3. Bei Gleichstand: Wähle nach höchster Konfidenz
            print(f"  Gleichstand zwischen {len(majority_categories)} Kategorien")
            
            # Sammle Kodierungen für die gleichstehenden Kategorien
            tied_codings = [
                coding for coding in segment_codes
                if coding['category'] in majority_categories
            ]
            
            # Finde die Kodierung mit der höchsten Konfidenz
            highest_confidence = -1
            best_coding = None
            
            for coding in tied_codings:
                # Extrahiere Konfidenzwert
                confidence = 0.0
                if isinstance(coding.get('confidence'), dict):
                    confidence = float(coding['confidence'].get('total', 0))
                    if confidence == 0:  # Fallback auf category-Konfidenz
                        confidence = float(coding['confidence'].get('category', 0))
                elif isinstance(coding.get('confidence'), (int, float)):
                    confidence = float(coding['confidence'])
                
                if confidence > highest_confidence:
                    highest_confidence = confidence
                    best_coding = coding
                    majority_category = coding['category']
            
            print(f"  ✓ Tie-Breaking durch Konfidenz: '{majority_category}' (Konfidenz: {highest_confidence:.2f})")
        
        # 4. Sammle alle Kodierungen für die gewählte Mehrheitskategorie
        matching_codings = [
            coding for coding in segment_codes
            if coding['category'] == majority_category
        ]
        
        # VEREINFACHT: Wähle beste Kodierung und nutze ihre Subkategorien direkt
        base_coding = max(
            matching_codings,
            key=lambda x: self._extract_confidence_value(x)
        )
        
        # VEREINFACHT: Keine komplexe Subkategorien-Sammlung
        majority_coding = base_coding.copy()
        main_category = majority_coding.get('category', '')
        original_subcats = base_coding.get('subcategories', [])
        validated_subcats = CategoryValidator.validate_subcategories_for_category(
                original_subcats, main_category, self.current_categories, warn_only=False
            )
        majority_coding['subcategories'] = validated_subcats
        
        # Kombiniere Begründungen (bleibt gleich)
        all_justifications = []
        for coding in matching_codings:
            justification = coding.get('justification', '')
            if justification and justification not in all_justifications:
                all_justifications.append(justification)  
        
        if all_justifications:
            majority_coding['justification'] = f"[Mehrheit aus {len(matching_codings)} Kodierern] " + " | ".join(all_justifications[:3])
        
        # Rest der Dokumentation bleibt gleich
        majority_coding['consensus_info'] = {
            'total_coders': total_coders,
            'category_votes': max_count,
            'category_agreement': max_count / total_coders,
            'tied_categories': majority_categories if len(majority_categories) > 1 else [],
            'source_codings': len(matching_codings),
            'selection_type': 'majority',
            'tie_broken_by_confidence': len(majority_categories) > 1
        }
        
        print(f"  ✓ Mehrheits-Kodierung erstellt: '{majority_category}' mit {len(majority_coding['subcategories'])} Subkategorien direkt übernommen: {', '.join(majority_coding['subcategories'])}")
        
        return majority_coding


    def _get_single_consensus_coding(self, segment_codes: List[Dict]) -> Optional[Dict]:
        """
        Ermittelt die Konsens-Kodierung für ein Segment basierend auf einem mehrstufigen Prozess.
        KORRIGIERT: Präzise Subkategorien-Zuordnung ohne Vermischung zwischen Hauptkategorien
        """
        if not segment_codes:
            return None

        # 1. Analyse der Hauptkategorien
        category_counts = Counter(coding['category'] for coding in segment_codes)
        total_coders = len(segment_codes)
        
        # Finde häufigste Hauptkategorie(n)
        max_count = max(category_counts.values())
        majority_categories = [
            category for category, count in category_counts.items()
            if count == max_count
        ]
        
        # Prüfe ob es eine klare Mehrheit gibt (>50%)
        if max_count <= total_coders / 2:
            print(f"Keine Mehrheit für Hauptkategorie gefunden: {dict(category_counts)}")
            
            # Suche nach Kodierung mit höchster Konfidenz
            highest_confidence = -1
            best_coding = None
            
            for coding in segment_codes:
                confidence = self._extract_confidence_value(coding)
                
                if confidence > highest_confidence:
                    highest_confidence = confidence
                    best_coding = coding
            
            # Minimalschwelle für Konfidenz (kann angepasst werden)
            confidence_threshold = 0.7
            
            if highest_confidence >= confidence_threshold:
                # Verwende die Kodierung mit der höchsten Konfidenz
                result_coding = best_coding.copy()
                
                # KORRIGIERT: Behalte nur Subkategorien der gewählten Hauptkategorie
                result_coding['subcategories'] = best_coding.get('subcategories', [])
                
                # Füge Hinweis zur konfidenzbedingten Auswahl hinzu
                result_coding['justification'] = (f"[Konfidenzbasierte Auswahl: {highest_confidence:.2f}] " + 
                                                result_coding.get('justification', ''))
                
                # Dokumentiere den Konsensprozess
                result_coding['consensus_info'] = {
                    'total_coders': total_coders,
                    'category_distribution': dict(category_counts),
                    'max_agreement': max_count / total_coders,
                    'selection_type': 'confidence_based',
                    'confidence': highest_confidence
                }
                
                print(f"  Konfidenzbasierte Auswahl: '{result_coding['category']}' mit {len(result_coding.get('subcategories', []))} Subkategorien")
                return result_coding
            else:
                # Erstelle "Kein Kodierkonsens"-Eintrag
                base_coding = segment_codes[0].copy()
                base_coding['category'] = "Kein Kodierkonsens"
                base_coding['subcategories'] = []  # Keine Subkategorien bei fehlendem Konsens
                base_coding['justification'] = (f"Keine Mehrheit unter den Kodierern und keine Kodierung " +
                                            f"mit ausreichender Konfidenz (max: {highest_confidence:.2f}). " +
                                            f"Kategorien: {', '.join(category_counts.keys())}")
                
                # Dokumentiere den Konsensprozess
                base_coding['consensus_info'] = {
                    'total_coders': total_coders,
                    'category_distribution': dict(category_counts),
                    'max_agreement': max_count / total_coders,
                    'selection_type': 'no_consensus',
                    'max_confidence': highest_confidence
                }
                
                print(f"  Kein Konsens und unzureichende Konfidenz (max: {highest_confidence:.2f})")
                return base_coding

        # 2. Wenn es mehrere gleichhäufige Hauptkategorien gibt, verwende Tie-Breaking
        if len(majority_categories) > 1:
            print(f"Gleichstand zwischen Kategorien: {majority_categories}")
            # Sammle alle Kodierungen für die Mehrheitskategorien
            candidate_codings = [
                coding for coding in segment_codes
                if coding['category'] in majority_categories
            ]
            
            # Wähle basierend auf höchster durchschnittlicher Konfidenz
            highest_avg_confidence = -1
            selected_category = None
            
            for category in majority_categories:
                category_codings = [c for c in candidate_codings if c['category'] == category]
                total_confidence = 0.0
                
                for coding in category_codings:
                    confidence = self._extract_confidence_value(coding)
                    total_confidence += confidence
                    
                avg_confidence = total_confidence / len(category_codings) if category_codings else 0
                
                if avg_confidence > highest_avg_confidence:
                    highest_avg_confidence = avg_confidence
                    selected_category = category
                    
            majority_category = selected_category
            print(f"  Kategorie '{majority_category}' durch höchste Konfidenz ({highest_avg_confidence:.2f}) gewählt")
        else:
            majority_category = majority_categories[0]

        # 3. KORRIGIERT: Sammle nur Kodierungen für die gewählte Mehrheitskategorie
        matching_codings = [
            coding for coding in segment_codes
            if coding['category'] == majority_category
        ]
        
        # VEREINFACHT: Wähle beste Kodierung und nutze ihre Subkategorien direkt
        base_coding = max(
            matching_codings,
            key=lambda x: self._extract_confidence_value(x)
        )
        
        # VEREINFACHT: Keine komplexe Subkategorien-Sammlung mehr
        consensus_coding = base_coding.copy()
        main_category = consensus_coding.get('category', '')
        original_subcats = base_coding.get('subcategories', [])
        # FIX: Prüfe ob wir Zugriff auf das vollständige Kategoriensystem haben
        categories_for_validation = getattr(self, 'current_categories', {})
        
        if categories_for_validation and main_category in categories_for_validation:
            validated_subcats = CategoryValidator.validate_subcategories_for_category(
                original_subcats, main_category, categories_for_validation, warn_only=False
            )
            
            # FIX: Debug-Ausgabe der Validierung
            if len(original_subcats) != len(validated_subcats):
                removed_subcats = set(original_subcats) - set(validated_subcats)
                print(f"🔧 FIX: Consensus-Validierung entfernte {len(removed_subcats)} Subkategorien: {removed_subcats}")
                
                # FIX: Dokumentiere Änderung in Begründung
                original_justification = consensus_coding.get('justification', '')
                consensus_coding['justification'] = f"{original_justification} [FIX: Consensus-Validierung entfernte ungültige Subkategorien: {list(removed_subcats)}]"
                
        else:
            # FIX: Fallback ohne Validierung
            print(f"⚠️ WARNUNG: Keine Kategorie-Validierung möglich für '{main_category}' - verwende ursprüngliche Subkategorien")
            validated_subcats = original_subcats
        
        consensus_coding['subcategories'] = validated_subcats
        
        # Kombiniere nur Begründungen der matching codings
        all_justifications = []
        for coding in matching_codings:
            justification = coding.get('justification', '')
            if justification and justification not in all_justifications:
                all_justifications.append(justification)
        
        if all_justifications:
            consensus_coding['justification'] = f"[Konsens aus {len(matching_codings)} Kodierern] " + " | ".join(all_justifications[:3])
        
        consensus_coding['consensus_info'] = {
            'total_coders': total_coders,
            'category_agreement': max_count / total_coders,
            'source_codings': len(matching_codings),
            'selection_type': 'consensus'
        }
        
        print(f"\nKonsens-Kodierung erstellt:")
        print(f"- Hauptkategorie: {consensus_coding['category']} ({max_count}/{total_coders} Kodierer)")
        print(f"- Subkategorien: {len(consensus_coding['subcategories'])} direkt übernommen: {', '.join(consensus_coding['subcategories'])}")
        print(f"- Übereinstimmung: {(max_count/total_coders)*100:.1f}%")
        
        return consensus_coding

    def _create_category_specific_codings(self, segment_codes: List[Dict], segment_id: str) -> List[Dict]:
        """
        KORRIGIERT: Präzise Subkategorien-Zuordnung OHNE Mehrfachkodierung zu verhindern
        """
        # Gruppiere Kodierungen nach Hauptkategorien
        category_groups = {}
        
        for coding in segment_codes:
            main_cat = coding.get('category', '')
            if main_cat and main_cat not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                if main_cat not in category_groups:
                    category_groups[main_cat] = []
                category_groups[main_cat].append(coding)
        
        result_codings = []
        
        for i, (main_cat, codings_for_cat) in enumerate(category_groups.items(), 1):
            # print(f"DEBUG: Verarbeite Hauptkategorie '{main_cat}' mit {len(codings_for_cat)} Kodierungen")
            
            # Wähle die beste Kodierung für diese Kategorie als Basis
            best_coding = max(codings_for_cat, key=lambda x: self._extract_confidence_value(x))
            
            # KRITISCH: Sammle NUR Subkategorien, die für DIESE Hauptkategorie kodiert wurden
            relevant_subcats = []
            
            for coding in codings_for_cat:
                # Prüfe ob diese Kodierung wirklich für die aktuelle Hauptkategorie ist
                if coding.get('category') == main_cat:
                    # Prüfe ob es eine fokussierte Kodierung war
                    target_category = coding.get('target_category', '')
                    
                    if target_category == main_cat or not target_category:
                        # Diese Kodierung war für diese Hauptkategorie bestimmt
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            relevant_subcats.extend(subcats)
                        elif isinstance(subcats, str) and subcats:
                            subcat_list = [s.strip() for s in subcats.split(',') if s.strip()]
                            relevant_subcats.extend(subcat_list)
            
            # Entferne Duplikate
            final_subcats = list(set(relevant_subcats))
            
            # OPTIONAL: Validiere gegen Kategoriensystem (aber nur als Warnung)
            if hasattr(self, 'current_categories') and main_cat in self.current_categories:
                valid_subcats_for_main = set(self.current_categories[main_cat].subcategories.keys())
                invalid_subcats = [sub for sub in final_subcats if sub not in valid_subcats_for_main]
                
                if invalid_subcats:
                    print(f"  WARNUNG: Ungültige Subkategorien für '{main_cat}' gefunden: {invalid_subcats}")
                    print(f"  Gültige Subkategorien: {list(valid_subcats_for_main)}")
                    # NICHT entfernen, nur warnen!
            
            # print(f"  Finale Subkategorien für '{main_cat}': {final_subcats}")
            
            # Erstelle konsolidierte Kodierung
            consolidated_coding = best_coding.copy()
            consolidated_coding['category'] = main_cat
            consolidated_coding['subcategories'] = final_subcats  # Nur relevante Subkategorien
            consolidated_coding['multiple_coding_instance'] = i
            consolidated_coding['total_coding_instances'] = len(category_groups)
            consolidated_coding['target_category'] = main_cat
            consolidated_coding['category_focus_used'] = True
            
            # Erweiterte Begründung
            original_justification = consolidated_coding.get('justification', '')
            consolidated_coding['justification'] = f"[Mehrfachkodierung - Kategorie {i}/{len(category_groups)}] {original_justification}"
            
            result_codings.append(consolidated_coding)
        
        # print(f"DEBUG: Erstellt {len(result_codings)} kategorie-spezifische Kodierungen für {segment_id}")
        return result_codings
   
   
    # Zusätzliche Methode für ResultsExporter Klasse
    def debug_export_process(self, codings: List[Dict]) -> None:
        """
        Öffentliche Debug-Methode für Export-Prozess
        Kann vor dem eigentlichen Export aufgerufen werden
        """
        print(f"\n🔍 STARTE EXPORT-DEBUG für {len(codings)} Kodierungen")
        self._debug_export_preparation(codings)
        
        # Zusätzliche Checks
        segments_with_issues = []
        
        for coding in codings:
            segment_id = coding.get('segment_id', '')
            category = coding.get('category', '')
            subcats = coding.get('subcategories', [])
            
            # Prüfe auf leere Subkategorien bei kategorisierten Segmenten
            if category and category not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                if not subcats or (isinstance(subcats, list) and len(subcats) == 0):
                    segments_with_issues.append({
                        'segment_id': segment_id,
                        'category': category,
                        'issue': 'Keine Subkategorien trotz Kategorisierung'
                    })
        
        if segments_with_issues:
            print(f"\n⚠ GEFUNDENE PROBLEME: {len(segments_with_issues)} Segmente mit fehlenden Subkategorien")
            for issue in segments_with_issues[:3]:
                print(f"  - {issue['segment_id']}: {issue['category']} -> {issue['issue']}")
            if len(segments_with_issues) > 3:
                print(f"  ... und {len(segments_with_issues) - 3} weitere")
        else:
            print(f"\n✅ Keine offensichtlichen Subkategorien-Probleme gefunden")
        
        print(f"\n🔍 EXPORT-DEBUG ABGESCHLOSSEN")

    def _extract_confidence_value(self, coding: Dict) -> float:
        """
        Hilfsmethode zum Extrahieren des Konfidenzwerts aus einer Kodierung.
        
        Args:
            coding: Kodierung mit Konfidenzinformation
            
        Returns:
            float: Konfidenzwert zwischen 0 und 1
        """
        try:
            if isinstance(coding.get('confidence'), dict):
                confidence = float(coding['confidence'].get('total', 0))
                if confidence == 0:  # Fallback auf category-Konfidenz
                    confidence = float(coding['confidence'].get('category', 0))
            elif isinstance(coding.get('confidence'), (int, float)):
                confidence = float(coding['confidence'])
            else:
                confidence = 0.0
            return confidence
        except (ValueError, TypeError):
            return 0.0

    def _get_manual_priority_coding(self, segment_codes: List[Dict]) -> Optional[Dict]:
        """
        Bevorzugt manuelle Kodierungen vor automatischen Kodierungen.
        KORRIGIERT: Subkategorien werden korrekt verarbeitet und zusammengeführt.
        
        Args:
            segment_codes: Liste der Kodierungen für ein Segment von verschiedenen Kodierern
                
        Returns:
            Optional[Dict]: Priorisierte Kodierung mit korrekten Subkategorien
        """
        if not segment_codes:
            return None

        print(f"\nManuelle Priorisierung für Segment mit {len(segment_codes)} Kodierungen...")

        # 1. Trenne manuelle und automatische Kodierungen
        manual_codings = []
        auto_codings = []
        
        for coding in segment_codes:
            coder_id = coding.get('coder_id', '')
            # Erkenne manuelle Kodierungen anhand der coder_id
            if 'human' in coder_id.lower() or 'manual' in coder_id.lower() or coding.get('manual_coding', False):
                manual_codings.append(coding)
            else:
                auto_codings.append(coding)
        
        print(f"  Gefunden: {len(manual_codings)} manuelle, {len(auto_codings)} automatische Kodierungen")
        
        # 2. Wenn manuelle Kodierungen vorhanden sind, bevorzuge diese
        if manual_codings:
            print("  ✓ Verwende manuelle Kodierungen mit Priorität")
            
            if len(manual_codings) == 1:
                # Nur eine manuelle Kodierung - verwende diese direkt
                selected_coding = manual_codings[0].copy()
                
                # WICHTIG: Subkategorien beibehalten!
                if 'subcategories' in manual_codings[0]:
                    selected_coding['subcategories'] = manual_codings[0]['subcategories']
                
                selected_coding['consensus_info'] = {
                    'total_coders': len(segment_codes),
                    'manual_coders': len(manual_codings),
                    'auto_coders': len(auto_codings),
                    'selection_type': 'single_manual',
                    'priority_reason': 'Einzige manuelle Kodierung verfügbar'
                }
                print(f"    Einzige manuelle Kodierung: '{selected_coding['category']}' mit {len(selected_coding.get('subcategories', []))} Subkategorien")
                
            else:
                # Mehrere manuelle Kodierungen - suche Konsens unter diesen
                print(f"    Suche Konsens unter {len(manual_codings)} manuellen Kodierungen")
                
                # Prüfe ob alle dieselbe Hauptkategorie haben
                manual_categories = [c['category'] for c in manual_codings]
                if len(set(manual_categories)) == 1:
                    # Alle haben dieselbe Hauptkategorie - konsolidiere Subkategorien
                    main_category = manual_categories[0]
                    
                    # Sammle alle Subkategorien von manuellen Kodierungen
                    all_manual_subcats = []
                    for coding in manual_codings:
                        subcats = coding.get('subcategories', [])
                        if isinstance(subcats, (list, tuple)):
                            all_manual_subcats.extend(subcats)
                        elif isinstance(subcats, str) and subcats:
                            subcats_list = [s.strip() for s in subcats.split(',') if s.strip()]
                            all_manual_subcats.extend(subcats_list)
                    
                    # Finde Konsens-Subkategorien (mindestens von der Hälfte verwendet)
                    subcat_counts = Counter(all_manual_subcats)
                    min_votes = len(manual_codings) / 2
                    consensus_subcats = [
                        subcat for subcat, count in subcat_counts.items()
                        if count >= min_votes
                    ]
                    
                    # Wähle beste manuelle Kodierung als Basis
                    selected_coding = max(
                        manual_codings,
                        key=lambda x: self._extract_confidence_value(x)
                    ).copy()
                    
                    # Setze konsolidierte Subkategorien
                    selected_coding['subcategories'] = consensus_subcats
                    
                    # Kombiniere Begründungen
                    manual_justifications = [c.get('justification', '')[:100] for c in manual_codings if c.get('justification', '')]
                    if manual_justifications:
                        selected_coding['justification'] = f"[Konsens aus {len(manual_codings)} manuellen Kodierungen] " + " | ".join(manual_justifications[:3])
                    
                    selected_coding['consensus_info'] = {
                        'total_coders': len(segment_codes),
                        'manual_coders': len(manual_codings),
                        'auto_coders': len(auto_codings),
                        'selection_type': 'manual_consensus',
                        'priority_reason': 'Konsens unter manuellen Kodierungen',
                        'subcategory_distribution': dict(subcat_counts)
                    }
                    print(f"    Konsens bei manuellen Kodierungen: '{selected_coding['category']}' mit {len(consensus_subcats)} Subkategorien: {', '.join(consensus_subcats)}")
                else:
                    # Verschiedene Hauptkategorien - wähle nach Konfidenz
                    selected_coding = max(
                        manual_codings,
                        key=lambda x: self._extract_confidence_value(x)
                    ).copy()

                    # VEREINFACHT: Direkte Übernahme
                    # selected_coding['subcategories'] = selected_coding.get('subcategories', [])  # DIREKT
                    main_category = selected_coding.get('category', '')
                    original_subcats = selected_coding.get('subcategories', [])
                    validated_subcats = CategoryValidator.validate_subcategories_for_category(
                        original_subcats, main_category, self.current_categories, warn_only=False
                    )
                    selected_coding['subcategories'] = validated_subcats
    
        else:
            # 3. Keine manuellen Kodierungen - verwende automatische mit Konsens
            print("  Keine manuellen Kodierungen - verwende automatische Kodierungen")
            
            # Verwende die bestehende Konsens-Logik für automatische Kodierungen
            consensus_coding = self._get_consensus_coding(auto_codings)
            
            
            if consensus_coding:
                selected_coding = consensus_coding.copy()
                # Aktualisiere consensus_info
                selected_coding['consensus_info'] = selected_coding.get('consensus_info', {})
                selected_coding['consensus_info'].update({
                    'total_coders': len(segment_codes),
                    'manual_coders': 0,
                    'auto_coders': len(auto_codings),
                    'selection_type': 'auto_consensus',
                    'priority_reason': 'Keine manuellen Kodierungen verfügbar - automatischer Konsens'
                })
                print(f"    Automatischer Konsens: '{selected_coding['category']}' mit {len(selected_coding.get('subcategories', []))} Subkategorien")
            else:
                # Fallback: Wähle automatische Kodierung mit höchster Konfidenz
                selected_coding = max(
                    auto_codings,
                    key=lambda x: self._extract_confidence_value(x)
                ).copy()
                
                # VEREINFACHT: Direkte Übernahme
            selected_coding['subcategories'] = selected_coding.get('subcategories', [])  # DIREKT
    
        return selected_coding


    def _calculate_coding_quality(self, coding: Dict, consensus_subcats: List[str]) -> float:
        """
        Berechnet einen Qualitätsscore für eine Kodierung.
        Berücksichtigt mehrere Faktoren:
        - Konfidenz der Kodierung
        - Übereinstimmung mit Konsens-Subkategorien
        - Qualität der Begründung

        Args:
            coding: Einzelne Kodierung
            consensus_subcats: Liste der Konsens-Subkategorien

        Returns:
            float: Qualitätsscore zwischen 0 und 1
        """
        try:
            # Hole Konfidenzwert (gesamt oder Hauptkategorie)
            if isinstance(coding.get('confidence'), dict):
                confidence = float(coding['confidence'].get('total', 0))
            else:
                confidence = float(coding.get('confidence', 0))

            # Berechne Übereinstimmung mit Konsens-Subkategorien
            coding_subcats = set(coding.get('subcategories', []))
            consensus_subcats_set = set(consensus_subcats)
            if consensus_subcats_set:
                subcat_overlap = len(coding_subcats & consensus_subcats_set) / len(consensus_subcats_set)
            else:
                subcat_overlap = 1.0  # Volle Punktzahl wenn keine Konsens-Subkategorien

            # Bewerte Qualität der Begründung
            justification = coding.get('justification', '')
            if isinstance(justification, str):
                justification_score = min(len(justification.split()) / 20, 1.0)  # Max bei 20 Wörtern
            else:
                justification_score = 0.0  # Keine Begründung vorhanden oder ungültiger Typ

            # Gewichtete Kombination der Faktoren
            quality_score = (
                confidence * 0.5 +          # 50% Konfidenz
                subcat_overlap * 0.3 +      # 30% Subkategorien-Übereinstimmung
                justification_score * 0.2   # 20% Begründungsqualität
            )

            return quality_score

        except Exception as e:
            print(f"Fehler bei der Berechnung der Codierungsqualität: {str(e)}")
            return 0.0  # Rückgabe eines neutralen Scores im Fehlerfall
    
    def export_optimization_analysis(self, 
                                original_categories: Dict[str, CategoryDefinition],
                                optimized_categories: Dict[str, CategoryDefinition],
                                optimization_log: List[Dict]):
        """Exportiert eine detaillierte Analyse der Kategorienoptimierungen."""
        
        analysis_path = os.path.join(self.output_dir, 
                                    f'category_optimization_{datetime.now().strftime("%Y%m%d_%H%M%S")}.md')
        
        with open(analysis_path, 'w', encoding='utf-8') as f:
            f.write("# Analyse der Kategorienoptimierungen\n\n")
            
            f.write("## Übersicht\n")
            f.write(f"- Ursprüngliche Kategorien: {len(original_categories)}\n")
            f.write(f"- Optimierte Kategorien: {len(optimized_categories)}\n")
            f.write(f"- Anzahl der Optimierungen: {len(optimization_log)}\n\n")
            
            f.write("## Detaillierte Optimierungen\n")
            for entry in optimization_log:
                if entry['type'] == 'merge':
                    f.write(f"\n### Zusammenführung zu: {entry['result_category']}\n")
                    f.write(f"- Ursprüngliche Kategorien: {', '.join(entry['original_categories'])}\n")
                    f.write(f"- Zeitpunkt: {entry['timestamp']}\n\n")
                    
                    f.write("#### Ursprüngliche Definitionen:\n")
                    for cat in entry['original_categories']:
                        if cat in original_categories:
                            f.write(f"- {cat}: {original_categories[cat].definition}\n")
                    f.write("\n")
                    
                    if entry['result_category'] in optimized_categories:
                        f.write("#### Neue Definition:\n")
                        f.write(f"{optimized_categories[entry['result_category']].definition}\n\n")
                        
                # Weitere Optimierungstypen hier...
            
            f.write("\n## Statistiken\n")
            f.write(f"- Kategorienreduktion: {(1 - len(optimized_categories) / len(original_categories)) * 100:.1f}%\n")
            
            # Zähle Optimierungstypen
            optimization_types = Counter(entry['type'] for entry in optimization_log)
            f.write("\nOptimierungstypen:\n")
            for opt_type, count in optimization_types.items():
                f.write(f"- {opt_type}: {count}\n")
        
        print(f"Optimierungsanalyse exportiert nach: {analysis_path}")
  
    
    def _validate_export_data(self, export_data: List[dict]) -> bool:
        """
        Validiert die zu exportierenden Daten.
        
        Args:
            export_data: Liste der aufbereiteten Export-Daten
            
        Returns:
            bool: True wenn Daten valide sind
        """
        required_columns = {
            'Dokument', 'Chunk_Nr', 'Text', 'Kodiert', 
            'Hauptkategorie', 'Kategorietyp', 'Subkategorien', 
            'Begründung', 'Konfidenz', 'Mehrfachkodierung'
        }
        
        try:
            if not export_data:
                print("Warnung: Keine Daten zum Exportieren vorhanden")
                return False
                
            # Prüfe ob alle erforderlichen Spalten vorhanden sind
            for entry in export_data:
                missing_columns = required_columns - set(entry.keys())
                if missing_columns:
                    print(f"Warnung: Fehlende Spalten in Eintrag: {missing_columns}")
                    return False
                    
                # Prüfe Kodiert-Status
                if entry['Kodiert'] not in {'Ja', 'Nein'}:
                    print(f"Warnung: Ungültiger Kodiert-Status: {entry['Kodiert']}")
                    return False
                    
            print("Validierung der Export-Daten erfolgreich")
            return True
            
        except Exception as e:
            print(f"Fehler bei der Validierung der Export-Daten: {str(e)}")
            return False
            
    def _extract_metadata(self, filename: str) -> tuple:
        """Extrahiert Metadaten aus dem Dateinamen"""
        from pathlib import Path
        tokens = Path(filename).stem.split("_")
        attribut1 = tokens[0] if len(tokens) >= 1 else ""
        attribut2 = tokens[1] if len(tokens) >= 2 else ""
        attribut3 = tokens[2] if len(tokens) >= 3 else "" 
        return attribut1, attribut2, attribut3

    def _initialize_category_colors(self, df: pd.DataFrame) -> None:
        """
        Initialisiert die Farbzuordnung für alle Kategorien einmalig.
        
        Args:
            df: DataFrame mit einer 'Hauptkategorie' Spalte
        """
        if not self.category_colors:  # Nur initialisieren wenn noch nicht geschehen
            # Hole alle eindeutigen Hauptkategorien außer 'Nicht kodiert'
            categories = sorted([cat for cat in df['Hauptkategorie'].unique() 
                              if cat != 'Nicht kodiert'])
            
            # Generiere Pastellfarben
            colors = self._generate_pastel_colors(len(categories))
            
            # Erstelle Mapping in alphabetischer Reihenfolge
            self.category_colors = {
                category: color for category, color in zip(categories, colors)
            }
            
            # Füge 'Nicht kodiert' mit grauer Farbe hinzu
            if 'Nicht kodiert' in df['Hauptkategorie'].unique():
                self.category_colors['Nicht kodiert'] = 'CCCCCC'
            
            print("\nFarbzuordnung initialisiert:")
            for cat, color in self.category_colors.items():
                print(f"- {cat}: {color}")


    async def export_results(self,
                            codings: List[Dict],
                            reliability: float,
                            categories: Dict[str, CategoryDefinition],
                            chunks: Dict[str, List[str]],
                            revision_manager: 'CategoryRevisionManager',
                            export_mode: str = "consensus",
                            original_categories: Dict[str, CategoryDefinition] = None,
                            original_codings: List[Dict] = None,  
                            inductive_coder: 'InductiveCoder' = None,
                            document_summaries: Dict[str, str] = None,
                            is_intermediate_export: bool = False) -> None: 
        """
        FIX: Exportiert mit korrekten ursprünglichen Kodierungen für Reliabilität
        Für ResultsExporter Klasse
        
        Args:
            codings: Finale/Review-Kodierungen für Export
            categories: Finale Kategorien
            original_categories: Ursprüngliche Kategorien
            document_summaries: Document Summaries
            revision_manager: Revision Manager
            original_codings: FIX: Ursprüngliche Kodierungen für Reliabilitätsberechnung
            export_mode: Export-Modus
            reliability: FIX: Bereits berechnete Reliabilität aus main()
        """
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            analysis_mode = CONFIG.get('ANALYSIS_MODE', 'deductive')
            # FIX: Unterschiedliche Dateinamen für normale und Zwischenexporte
            if is_intermediate_export:
                filename = f"QCA-AID_Analysis_INTERMEDIATE_{analysis_mode}_{timestamp}.xlsx"
                print(f"📊 Exportiere Zwischenergebnisse bei Abbruch...")
            else:
                filename = f"QCA-AID_Analysis_{analysis_mode}_{timestamp}.xlsx"
                print(f"📊 Exportiere finale Ergebnisse mit {export_mode}-Modus...")
            
            filepath = os.path.join(self.output_dir, filename)
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                print(f"📊 Exportiere umfassende Ergebnisse mit {export_mode}-Modus...")
                
                # 1. HAUPT-SHEET: Kodierungsergebnisse (finale Kodierungen)
                print("📋 Exportiere Hauptergebnisse...")
                self._export_main_results(writer, codings, original_categories)
                
                # 2. FIX: HÄUFIGKEITEN-SHEET (finale Kodierungen)
                print("📊 Exportiere Häufigkeiten...")
                df_coded = self._prepare_dataframe_for_frequency_analysis(codings)
                if not df_coded.empty:
                    attribut1_label = self.attribute_labels.get('attribut1', 'Attribut1')
                    attribut2_label = self.attribute_labels.get('attribut2', 'Attribut2')
                    self._export_frequency_analysis(writer, df_coded, attribut1_label, attribut2_label)
                
                # 3. FIX: INTERCODER-BERICHT mit ursprünglichen Kodierungen
                print("📊 Exportiere IntercoderBericht...")
                if original_codings and reliability is not None:
                    # FIX: Verwende bereits berechnete Reliabilität und ursprüngliche Kodierungen
                    self._export_intercoder_bericht(writer, original_codings, reliability)
                    print(f"✅ IntercoderBericht mit Alpha={reliability:.3f} erstellt")
                else:
                    print("⚠️ Keine ursprünglichen Kodierungen oder Reliabilität verfügbar")
                    self._create_empty_intercoder_sheet(writer)
                
                # 4. KATEGORIEN-ÜBERSICHT
                if categories:
                    print("📂 Exportiere Kategorien-Übersicht...")
                    self._export_categories_sheet_formatted(writer, categories, original_categories)
                
                # 5. PROGRESSIVE SUMMARIES (falls vorhanden)
                if document_summaries:
                    print("📝 Exportiere Progressive Summaries...")
                    self._export_progressive_summaries(writer, document_summaries)
                
                # 6. REVIEW-STATISTIKEN
                print("🎯 Exportiere Review-Statistiken...")
                review_stats = self._calculate_review_statistics(codings, export_mode, original_codings)
                self._export_review_statistics(writer, review_stats, export_mode)
                
                # 7. REVISIONSHISTORIE (falls verfügbar)
                if revision_manager and hasattr(revision_manager, 'changes'):
                    print("📜 Exportiere Revisionshistorie...")
                    revision_manager._export_revision_history(writer, revision_manager.changes)
                else:
                    print("ℹ️ Keine Revisionshistorie verfügbar")
                
                # 8. KONFIGURATION-SHEET 
                print("⚙️ Exportiere Konfiguration...")
                self._export_configuration(writer, export_mode)
                
                if is_intermediate_export:
                    print(f"✅ Zwischenergebnisse erfolgreich exportiert!")
                else:
                    print(f"✅ Export erfolgreich: {filename}")
                
            # FIX: Nur bei normalem Export Dateiinfo anzeigen
            if not is_intermediate_export:
                print(f"📂 Dateien im Ordner: {self.output_dir}")
                print(f"📄 Export-Datei: {filename}")
            
            # FIX: Return filename am Ende hinzugefügt
            return filename
            
        except Exception as e:
            print(f"❌ Fehler beim Export: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    

    def _prepare_dataframe_for_frequency_analysis(self, codings: List[Dict]) -> pd.DataFrame:
        """
        NEUE METHODE: Bereitet DataFrame für Häufigkeitsanalyse vor
        """
        try:
            # Erstelle DataFrame aus Kodierungen
            data = []
            
            for coding in codings:
                doc_name = coding.get('document', '')
                
                # Extrahiere Attribute aus Dokumentname
                attribut1, attribut2 = self._extract_attributes_from_document(doc_name)
                
                # Grunddaten
                row_data = {
                    'Dokument': doc_name,
                    self.attribute_labels.get('attribut1', 'Attribut1'): attribut1,
                    self.attribute_labels.get('attribut2', 'Attribut2'): attribut2,
                    'Chunk_Nr': coding.get('chunk_id', coding.get('segment_id', '')),
                    'Hauptkategorie': coding.get('category', ''),
                    'Kodiert': 'Ja' if coding.get('category') and coding.get('category') not in ['Nicht kodiert', 'Kein Kodierkonsens'] else 'Nein',
                    'Subkategorien': ', '.join(coding.get('subcategories', [])),
                    'Konfidenz': self._extract_confidence_from_coding(coding)
                }
                
                data.append(row_data)
            
            if data:
                df = pd.DataFrame(data)
                print(f"📊 DataFrame erstellt: {len(df)} Zeilen für Häufigkeitsanalyse")
                return df
            else:
                print("⚠️ Keine Daten für DataFrame verfügbar")
                return pd.DataFrame()
                
        except Exception as e:
            print(f"❌ Fehler beim Erstellen des DataFrames: {str(e)}")
            return pd.DataFrame()
    
    def _extract_confidence_from_coding(self, coding: Dict) -> float:
        """
        HILFSMETHODE: Extrahiert Konfidenzwert aus Kodierung
        """
        confidence = coding.get('confidence', {})
        if isinstance(confidence, dict):
            return confidence.get('total', 0.0)
        elif isinstance(confidence, (int, float)):
            return float(confidence)
        return 0.0
    
    def _export_main_results(self, writer, codings: List[Dict], original_categories: Dict):
        """
        Exportiert Haupt-Kodierungsergebnisse
        FIX: Übernimmt alle fehlenden Spalten und Logik aus _prepare_coding_for_export
        """
        
        # DataFrame erstellen
        data = []
        
        for coding in codings:
            # FIX: Übernehme Dokumentname-Extraktion aus _prepare_coding_for_export
            doc_name = coding.get('document', '')
            if not doc_name:
                # Fallback: Extrahiere aus segment_id
                segment_id = coding.get('segment_id', '')
                if segment_id and '_chunk_' in segment_id:
                    doc_name = segment_id.split('_chunk_')[0]
                else:
                    doc_name = 'Unbekanntes_Dokument'
            
            # FIX: Übernehme _extract_metadata Logik für drei Attribute
            if hasattr(self, '_extract_metadata'):
                attribut1, attribut2, attribut3 = self._extract_metadata(doc_name)
            else:
                # Fallback: Extrahiere aus Dokumentname
                from pathlib import Path
                tokens = Path(doc_name).stem.split("_")
                attribut1 = tokens[0] if len(tokens) >= 1 else ""
                attribut2 = tokens[1] if len(tokens) >= 2 else ""
                attribut3 = tokens[2] if len(tokens) >= 3 else ""

            # FIX: Erstelle eindeutigen Präfix für Chunk-Nr mit bis zu 6 Buchstaben pro Attribut
            chunk_prefix = ""
            if attribut1 and attribut2:
                # FIX: Extrahiere bis zu 6 Buchstaben pro Attribut (oder alle verfügbaren)
                import re
                attr1_letters = re.sub(r'[^a-zA-Z0-9]', '', attribut1)[:6]
                attr2_letters = re.sub(r'[^a-zA-Z0-9]', '', attribut2)[:6]
                attr3_letters = re.sub(r'[^a-zA-Z0-9]', '', attribut3)[:6]
                chunk_prefix = (attr1_letters + "_" + attr2_letters + "_" + attr3_letters).upper()
            else:
                chunk_prefix = doc_name[:5].upper()
            
            # FIX: Korrekte chunk_id Extraktion
            chunk_id = 0
            segment_id = coding.get('segment_id', '')
            
            # Versuche chunk_id direkt zu bekommen
            if 'chunk_id' in coding and isinstance(coding['chunk_id'], int):
                chunk_id = coding['chunk_id']
            elif segment_id and '_chunk_' in segment_id:
                # Extrahiere aus segment_id
                try:
                    chunk_part = segment_id.split('_chunk_')[1]
                    # Falls Mehrfachkodierung: "123-1" -> nimm nur "123"
                    if '-' in chunk_part:
                        chunk_id = int(chunk_part.split('-')[0])
                    else:
                        chunk_id = int(chunk_part)
                except (ValueError, IndexError):
                    chunk_id = 0
            
            # Haupt-Spalten
            text = coding.get('text', '')
            paraphrase = coding.get('paraphrase', '')
            
            # FIX: Hole Kategorie und behandle leere Kategorien explizit
            category = coding.get('category', '')
            
            # FIX: Zentrale Behandlung von leeren/fehlenden Kategorien
            if not category or str(category).strip() == "":
                display_category = "Nicht kodiert"
                is_coded = 'Nein'
                category_type = 'unkodiert'
            elif category == "Kein Kodierkonsens":
                display_category = "Kein Kodierkonsens"
                is_coded = 'Nein'
                category_type = 'unkodiert'
            else:
                display_category = category
                is_coded = 'Ja'
                # Kategorietyp bestimmen für gültige Kategorien
                if hasattr(self, '_determine_category_type'):
                    category_type = self._determine_category_type(category, original_categories or {})
                else:
                    # Fallback-Logik
                    if original_categories and category in original_categories:
                        category_type = 'deduktiv'
                    else:
                        category_type = 'induktiv'
            # FIX: Ende
            
            # FIX: Subkategorien verarbeiten wie in _prepare_coding_for_export
            subcategories = coding.get('subcategories', [])
            subcats_text = ""
            if subcategories:
                if isinstance(subcategories, str):
                    subcats_text = subcategories.strip()
                elif isinstance(subcategories, (list, tuple)):
                    clean_subcats = []
                    for subcat in subcategories:
                        if subcat and str(subcat).strip():
                            clean_text = str(subcat).strip()
                            clean_text = clean_text.replace('[', '').replace(']', '').replace("'", "").replace('"', '')
                            if clean_text:
                                clean_subcats.append(clean_text)
                    subcats_text = ', '.join(clean_subcats)
                elif isinstance(subcategories, dict):
                    clean_subcats = []
                    for key in subcategories.keys():
                        clean_key = str(key).strip()
                        if clean_key:
                            clean_subcats.append(clean_key)
                    subcats_text = ', '.join(clean_subcats)
                else:
                    subcats_text = str(subcategories).strip()
                    subcats_text = subcats_text.replace('[', '').replace(']', '').replace("'", "").replace('"', '')
            
            # Zusätzliche Bereinigung
            subcats_text = subcats_text.replace('[', '').replace(']', '').replace("'", "")
            
            # FIX: Keywords verarbeiten wie in _prepare_coding_for_export
            raw_keywords = coding.get('keywords', '')
            if isinstance(raw_keywords, list):
                formatted_keywords = [kw.strip() for kw in raw_keywords]
            else:
                formatted_keywords = raw_keywords.replace("[", "").replace("]", "").replace("'", "").split(",")
                formatted_keywords = [kw.strip() for kw in formatted_keywords if kw.strip()]
            keywords_text = ', '.join(formatted_keywords)
            
            # FIX: VERBESSERTE BEGRÜNDUNGSVERARBEITUNG - KERNFIX
            justification = ""
            
            # Prioritätssystem für Begründungen:
            # 1. Normale justification (höchste Priorität)
            if coding.get('justification') and coding.get('justification').strip():
                justification = coding.get('justification')
            # 2. reasoning Feld (oft von RelevanceChecker)
            elif coding.get('reasoning') and coding.get('reasoning').strip() and coding.get('reasoning') != 'NICHT VORHANDEN':
                justification = coding.get('reasoning')
            # 3. original_justification (Backup)
            elif coding.get('original_justification') and coding.get('original_justification').strip() and coding.get('original_justification') != 'NICHT VORHANDEN':
                justification = coding.get('original_justification')
            # 4. Fallback für "Nicht kodiert" basierend auf Analyse
            else:
                if category in ["Nicht kodiert", ""] or display_category == "Nicht kodiert":
                    # FIX: Hole Details vom RelevanceChecker falls verfügbar
                    segment_id = coding.get('segment_id', '')
                    if hasattr(self, 'relevance_checker') and segment_id:
                        relevance_details = self.relevance_checker.get_relevance_details(segment_id)
                        if relevance_details:
                            if relevance_details.get('reasoning') and relevance_details['reasoning'] != 'Keine Begründung verfügbar':
                                justification = relevance_details['reasoning']
                            elif relevance_details.get('justification') and relevance_details['justification'] != 'Keine Begründung verfügbar':
                                justification = relevance_details['justification']
                    
                    # FIX: Intelligente Fallback-Begründungen basierend auf Textanalyse
                    if not justification:
                        text_content = text.lower() if text else ""
                        text_length = len(text_content.strip())
                        
                        if text_length < 20:
                            justification = "Segment zu kurz für sinnvolle Kodierung"
                        elif any(pattern in text_content for pattern in ['seite ', 'page ', 'copyright', '©', 'inhaltsverzeichnis', 'table of contents']):
                            justification = "Segment als Metadaten (z.B. Seitenzahl, Copyright) identifiziert"
                        elif any(pattern in text_content for pattern in ['abstract', 'zusammenfassung', 'einleitung']):
                            justification = "Segment außerhalb des Analysebereichs der Forschungsfrage"
                        elif text_length < 100:
                            justification = "Segment enthält zu wenig Substanz für thematische Kodierung"
                        else:
                            justification = "Segment nicht relevant für die definierten Analysekategorien"
                
                # FIX: Fallback für andere Kategorien ohne Begründung
                elif not justification:
                    justification = "Kodierung ohne spezifische Begründung dokumentiert"
            
            # FIX: Debug-Output für Problemdiagnose
            if display_category == "Nicht kodiert":
                segment_id = coding.get('segment_id', 'unknown')
                # print(f"🔧 FIX DEBUG Segment {segment_id}:")
                # print(f"   - Finale justification: '{justification}'")
                # print(f"   - Original justification: '{coding.get('justification', 'LEER')}'")
                # print(f"   - Reasoning: '{coding.get('reasoning', 'LEER')}'")
                # print(f"   - Original_justification: '{coding.get('original_justification', 'LEER')}'")
                # FIX: ZUSÄTZLICHER CHECK - Direkt vom RelevanceChecker holen
                if hasattr(self, 'relevance_checker') and self.relevance_checker:
                    # print(f"   - RelevanceChecker verfügbar: JA")
                    try:
                        relevance_details = self.relevance_checker.get_relevance_details(segment_id)
                        print(f"   - RelevanceChecker Details: {relevance_details}")
                        
                        # FIX: Verwende RelevanceChecker-Daten wenn vorhanden
                        if relevance_details and justification == "Keine Begründung verfügbar":
                            if relevance_details.get('reasoning') and relevance_details['reasoning'] != 'Keine Begründung verfügbar':
                                justification = relevance_details['reasoning']
                                # print(f"   - ✅ Begründung aus RelevanceChecker geholt: '{justification}'")
                            elif relevance_details.get('justification') and relevance_details['justification'] != 'Keine Begründung verfügbar':
                                justification = relevance_details['justification'] 
                                # print(f"   - ✅ Justification aus RelevanceChecker geholt: '{justification}'")
                    except Exception as e:
                        print(f"   - ❌ Fehler beim RelevanceChecker-Zugriff: {e}")
                else:
                    print(f"   - RelevanceChecker verfügbar: NEIN")
                    # FIX: Fallback auf analysis_manager falls self.relevance_checker nicht da ist
                    if hasattr(self, 'analysis_manager') and self.analysis_manager and hasattr(self.analysis_manager, 'relevance_checker'):
                        # print(f"   - Analysis Manager RelevanceChecker verfügbar: JA")
                        try:
                            relevance_details = self.analysis_manager.relevance_checker.get_relevance_details(segment_id)
                            # print(f"   - Analysis Manager RelevanceChecker Details: {relevance_details}")
                            
                            # FIX: Verwende analysis_manager RelevanceChecker-Daten
                            if relevance_details and justification == "Keine Begründung verfügbar":
                                if relevance_details.get('reasoning') and relevance_details['reasoning'] != 'Keine Begründung verfügbar':
                                    justification = relevance_details['reasoning']
                                    # print(f"   - ✅ Begründung aus Analysis Manager RelevanceChecker geholt: '{justification}'")
                                elif relevance_details.get('justification') and relevance_details['justification'] != 'Keine Begründung verfügbar':
                                    # justification = relevance_details['justification']
                                    print(f"   - ✅ Justification aus Analysis Manager RelevanceChecker geholt: '{justification}'")
                        except Exception as e:
                            print(f"   - ❌ Fehler beim Analysis Manager RelevanceChecker-Zugriff: {e}")
                    else:
                        print(f"   - Analysis Manager RelevanceChecker verfügbar: NEIN")
                
                # FIX: Wenn immer noch keine Begründung, verwende intelligente Fallbacks
                if not justification or justification == "Keine Begründung verfügbar":
                    text_content = text.lower() if text else ""
                    text_length = len(text_content.strip())
                    
                    if text_length < 20:
                        justification = "Segment zu kurz für sinnvolle Kodierung"
                        # print(f"   - ✅ Fallback: Zu kurz")
                    elif any(pattern in text_content for pattern in ['seite ', 'page ', 'copyright', '©', 'inhaltsverzeichnis']):
                        justification = "Segment als Metadaten identifiziert"
                        # print(f"   - ✅ Fallback: Metadaten")
                    elif text_length < 100:
                        justification = "Segment enthält zu wenig Substanz für Kodierung"
                        # print(f"   - ✅ Fallback: Zu wenig Substanz")
                    else:
                        justification = "Segment nicht relevant für Analysekategorien"
                        # print(f"   - ✅ Fallback: Nicht relevant")
                
                #  print(f"   - 🎯 FINAL justification: '{justification}'")
            
            # FIX: Konfidenz korrekt extrahieren
            confidence = coding.get('confidence', {})
            if isinstance(confidence, dict):
                confidence_value = confidence.get('total', 0.0)
            elif isinstance(confidence, (int, float)):
                confidence_value = float(confidence)
            else:
                confidence_value = 0.0
            
            # FIX: Mehrfachkodierungs-Info wie in _prepare_coding_for_export
            consensus_info = coding.get('consensus_info', {})
            original_chunk_id = consensus_info.get('original_segment_id', chunk_id)
            is_multiple = consensus_info.get('is_multiple_coding_instance', False)
            instance_info = consensus_info.get('instance_info', {})
            
            # FIX: Erstelle eindeutige Chunk-ID mit Mehrfachkodierungs-Suffix
            if coding.get('total_coding_instances', 1) > 1:
                unique_chunk_id = f"{chunk_prefix}-{chunk_id}-{coding.get('multiple_coding_instance', 1)}"
                mehrfachkodierung_status = 'Ja'
            else:
                unique_chunk_id = f"{chunk_prefix}-{chunk_id}"
                mehrfachkodierung_status = 'Nein'
            
            # FIX: Zeile erstellen mit ALLEN Spalten aus _prepare_coding_for_export
            row_data = {
                'Dokument': _sanitize_text_for_excel(doc_name),  # FIX: Korrekte Parameterübergabe
                self.attribute_labels.get('attribut1', 'Attribut1'): _sanitize_text_for_excel(attribut1),
                self.attribute_labels.get('attribut2', 'Attribut2'): _sanitize_text_for_excel(attribut2),
            }
            
            # FIX: Füge attribut3 hinzu, wenn es definiert ist
            if 'attribut3' in self.attribute_labels and self.attribute_labels['attribut3']:
                row_data[self.attribute_labels['attribut3']] = _sanitize_text_for_excel(attribut3)
            
            if coding.get('manual_review', False):
                review_typ = 'manual'
            else:
                review_typ = consensus_info.get('selection_type', 'single')
            
            # FIX: Alle weiteren Spalten hinzufügen
            additional_fields = {
                'Chunk_Nr': unique_chunk_id,
                'Text': _sanitize_text_for_excel(text),  # FIX: Korrekte Funktionsaufrufe
                'Paraphrase': _sanitize_text_for_excel(paraphrase),
                'Kodiert': is_coded,  # FIX: Verwende korrekten is_coded Wert
                'Hauptkategorie': _sanitize_text_for_excel(display_category),  # FIX: Korrigiere Leerzeichen-Fehlern
                'Kategorietyp': category_type,  # FIX: Verwende korrekten category_type
                'Subkategorien': _sanitize_text_for_excel(subcats_text),
                'Schlüsselwörter': _sanitize_text_for_excel(keywords_text),
                'Begründung': _sanitize_text_for_excel(justification),
                'Konfidenz': f"{confidence_value:.2f}",  # FIX: Korrekte Konfidenz-Formatierung
                'Original_Chunk_Nr': original_chunk_id,  # FIX: Hinzugefügt
                'Mehrfachkodierung': mehrfachkodierung_status,  # FIX: Korrekte Mehrfachkodierung-Logik
                'Instanz_Nr': instance_info.get('instance_number', 1) if is_multiple else 1,  # FIX: Hinzugefügt
                'Gesamt_Instanzen': instance_info.get('total_instances', 1) if is_multiple else 1,  # FIX: Hinzugefügt
                'Review_Typ': review_typ,
                'Kodierer': coding.get('coder_id', 'Unbekannt'),  # FIX: Hinzugefügt               
                'Fokus_Kategorie': _sanitize_text_for_excel(coding.get('target_category', '')),
                'Fokus_verwendet': 'Ja' if coding.get('category_focus_used', False) else 'Nein',
                'Original_Chunk_ID': f"{chunk_prefix}-{chunk_id}"
            }
            
            row_data.update(additional_fields)
            
            # FIX: Kontext-bezogene Felder hinzufügen, wenn vorhanden
            if 'context_summary' in coding and coding['context_summary']:
                row_data['Progressive_Context'] = _sanitize_text_for_excel(coding.get('context_summary', ''))
            
            if 'context_influence' in coding and coding['context_influence']:
                row_data['Context_Influence'] = _sanitize_text_for_excel(coding.get('context_influence', ''))
            
            data.append(row_data)
        
        # Als DataFrame exportieren
        if data:
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='Kodierungsergebnisse', index=False)
            
            # FIX: Aktiviere Tabellenfunktionalität explizit
            worksheet = writer.sheets['Kodierungsergebnisse']
            # print(f"FIX: Formatiere Worksheet 'Kodierungsergebnisse' mit {len(df)} Zeilen und {len(df.columns)} Spalten")
            # print(f"FIX: Spalten: {list(df.columns)}")
            self._format_worksheet(worksheet, as_table=True)  # FIX: Explizit as_table=True
            
        else:
            print("⚠️ Keine Hauptergebnisse zum Exportieren")
    
    def _extract_document_from_segment_id(self, segment_id: str) -> str:
        """
        Extrahiert Dokumentnamen aus segment_id falls document-Feld fehlt
        """
        if not segment_id:
            return 'Unbekanntes_Dokument'
        
        # Segment-ID Format: "dokument_chunk_X" oder "dokument_chunk_X-Y"
        parts = segment_id.split('_chunk_')
        if len(parts) >= 2:
            return parts[0]
        
        # Fallback: Nimm alles vor dem letzten Unterstrich
        parts = segment_id.rsplit('_', 1)
        return parts[0] if len(parts) > 1 else segment_id
    
    def _extract_three_attributes_from_document(self, doc_name: str) -> tuple:
        """
        PUNKT 3: Korrekte Extraktion von 3 Attributen aus Dokumentname
        
        Erwartet Format: attribut1_attribut2_attribut3_rest.txt
        """
        # Entferne Dateierweiterung
        clean_name = doc_name
        for ext in ['.txt', '.docx', '.pdf', '.doc']:
            if clean_name.lower().endswith(ext.lower()):
                clean_name = clean_name[:-len(ext)]
                break
        
        # Teile am Unterstrich
        parts = clean_name.split('_')
        
        attribut1 = parts[0] if len(parts) > 0 else ''
        attribut2 = parts[1] if len(parts) > 1 else ''
        attribut3 = parts[2] if len(parts) > 2 else ''
        
        return attribut1, attribut2, attribut3
    
    def _generate_correct_chunk_id(self, coding: Dict, attribut1: str, attribut2: str, attribut3: str) -> str:
        """
        PUNKT 4: Generiert korrekte Chunk-ID im Format AABBCC-01-01
        
        Format: [Erste 2 Buchstaben Attr1][Erste 2 Buchstaben Attr2][Erste 2 Buchstaben Attr3]-[Segment-Nr]-[Mehrfachkodierungs-Nr]
        """
        # Extrahiere erste 2 Buchstaben aus jedem Attribut
        aa = self._extract_first_letters(attribut1, 2)
        bb = self._extract_first_letters(attribut2, 2)
        cc = self._extract_first_letters(attribut3, 2)
        
        # Basis-Präfix
        prefix = f"{aa}{bb}{cc}".upper()
        
        # Segment-Nummer aus ursprünglicher segment_id extrahieren
        segment_id = coding.get('segment_id', '')
        original_segment_id = coding.get('consensus_info', {}).get('original_segment_id', segment_id)
        
        # Extrahiere Segment-Nummer
        segment_nr = self._extract_segment_number(original_segment_id)
        
        # Mehrfachkodierungs-Nummer
        consensus_info = coding.get('consensus_info', {})
        instance_info = consensus_info.get('instance_info', {})
        multiple_nr = instance_info.get('instance_number', 1)
        
        # Zusammenbauen: AABBCC-01-01
        chunk_id = f"{prefix}-{segment_nr:02d}-{multiple_nr:02d}"
        
        return chunk_id
    
    def _extract_segment_number(self, segment_id: str) -> int:
        """
        Extrahiert Segment-Nummer aus segment_id
        """
        if not segment_id:
            return 1
        
        # Suche nach Zahlen in der segment_id
        numbers = re.findall(r'\d+', segment_id)
        
        if numbers:
            # Nimm die erste gefundene Zahl
            return int(numbers[0])
        
        return 1
    
    def _extract_first_letters(self, text: str, count: int) -> str:
        """
        Extrahiert die ersten N Buchstaben aus einem Text
        """
        if not text:
            return 'XX'[:count]
        
        # Nur Buchstaben extrahieren
        letters = re.sub(r'[^a-zA-Z]', '', text)
        if not letters:
            return 'XX'[:count]
        
        # Erste N Buchstaben
        result = letters[:count].upper()
        
        # Auffüllen falls zu kurz
        while len(result) < count:
            result += 'X'
        
        return result
    
    def _determine_category_type(self, category: str, original_categories: dict) -> str:
        """
        BESTEHENDE METHODE: Bestimmt Kategorietyp (deduktiv/induktiv)
        """
        if not category or category in ['Nicht kodiert', 'Kein Kodierkonsens']:
            return ''
        
        if category in original_categories:
            return 'Deduktiv'
        else:
            return 'Induktiv'
    
    def _format_worksheet(self, worksheet, as_table: bool = False) -> None:
        """
        Formatiert das Detail-Worksheet mit flexibler Farbkodierung und adaptiven Spaltenbreiten
        FIX: Korrigiert AutoFilter/Tabellen-Konflikt
        """
        try:
            if worksheet.max_row <= 1:
                print(f"⚠️ Worksheet '{worksheet.title}' enthält keine Daten")
                return

            # FIX: Hole DataFrame für Farbinitialisierung
            df_data = []
            headers = [cell.value for cell in worksheet[1]]
            
            for row in worksheet.iter_rows(min_row=2, values_only=True):
                if any(cell is not None for cell in row):
                    df_data.append(row)
            
            if df_data and 'Hauptkategorie' in headers:
                # Erstelle DataFrame für Farbinitialisierung
                df = pd.DataFrame(df_data, columns=headers)
                
                # FIX: Initialisiere Kategorie-Farben wenn noch nicht vorhanden
                if not hasattr(self, 'category_colors') or not self.category_colors:
                    self._initialize_category_colors(df)
                
                print(f"Verfügbare Kategorie-Farben: {list(self.category_colors.keys())}")

            # Bestimme Spaltenbreiten adaptiv
            column_widths = []
            for col in range(1, worksheet.max_column + 1):
                max_length = 0
                column = get_column_letter(col)
                
                for row in range(1, min(worksheet.max_row + 1, 101)):
                    cell = worksheet[f"{column}{row}"]
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                
                # Spaltenbreite setzen (min 8, max 50)
                width = min(max(max_length + 2, 8), 50)
                column_widths.append(width)
                worksheet.column_dimensions[column].width = width

            # Header und Datenformatierung
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            # FIX: Finde Hauptkategorie-Spalten-Index korrekt
            hauptkategorie_idx = None
            headers = [cell.value for cell in worksheet[1]]
            for i, header in enumerate(headers, 1):
                if header == 'Hauptkategorie':
                    hauptkategorie_idx = i
                    break
            
            print(f"Hauptkategorie-Spalte gefunden bei Index: {hauptkategorie_idx}")
            
            # Header formatieren
            for cell in worksheet[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
                cell.border = thin_border

            # FIX: Verbesserte Datenformatierung mit Farbkodierung
            for row in worksheet.iter_rows(min_row=2):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=False, vertical='top')
                    cell.border = thin_border

                    # FIX: Farbkodierung für Hauptkategorien-Spalte
                    if (hauptkategorie_idx and 
                        cell.column == hauptkategorie_idx and 
                        cell.value and 
                        hasattr(self, 'category_colors') and 
                        cell.value in self.category_colors):
                        
                        color = self.category_colors[cell.value]
                        cell.fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
                        
                        # print(f"Farbe angewendet für '{cell.value}': {color} in Zeile {cell.row}")

            # FIX: Excel-Tabelle oder AutoFilter erstellen - NICHT BEIDES gleichzeitig
            if as_table:
                try:
                    # Entferne vorhandene Tabellen sicher
                    table_names = list(worksheet.tables.keys()).copy()
                    for table_name in table_names:
                        del worksheet.tables[table_name]
                    
                    # FIX: Entferne auch eventuelle AutoFilter vor Tabellenerstellung
                    worksheet.auto_filter.ref = None
                    
                    # Sichere Bestimmung der letzten Spalte und Zeile
                    last_col_index = worksheet.max_column
                    last_col_letter = get_column_letter(last_col_index)
                    last_row = worksheet.max_row
                    
                    # Generiere eindeutigen Tabellennamen
                    safe_table_name = f"Table_{worksheet.title.replace(' ', '_').replace('-', '_')}"
                    
                    # Tabellenverweis generieren
                    table_ref = f"A1:{last_col_letter}{last_row}"
                    
                    # FIX: Erstelle NUR Excel-Tabelle (nicht AutoFilter)
                    tab = Table(displayName=safe_table_name, ref=table_ref)
                    style = TableStyleInfo(
                        name="TableStyleMedium9", 
                        showFirstColumn=False,
                        showLastColumn=False, 
                        showRowStripes=True, 
                        showColumnStripes=False
                    )
                    tab.tableStyleInfo = style
                    worksheet.add_table(tab)
                    
                    print(f"Excel-Tabelle '{safe_table_name}' erfolgreich erstellt")
                    
                except Exception as table_error:
                    print(f"Warnung bei Tabellenerstellung: {str(table_error)}")
                    # FIX: Fallback zu AutoFilter wenn Tabelle fehlschlägt
                    try:
                        last_col_index = worksheet.max_column
                        last_col_letter = get_column_letter(last_col_index)
                        last_row = worksheet.max_row
                        filter_range = f"A1:{last_col_letter}{last_row}"
                        worksheet.auto_filter.ref = filter_range
                        print(f"AutoFilter als Fallback erstellt: {filter_range}")
                    except Exception as filter_error:
                        print(f"Auch AutoFilter-Fallback fehlgeschlagen: {str(filter_error)}")

            print(f"Worksheet '{worksheet.title}' erfolgreich formatiert" + 
                (f" mit Farbkodierung für Hauptkategorien (Spalte {hauptkategorie_idx})" if hauptkategorie_idx else ""))
            
        except Exception as e:
            print(f"Fehler bei der Formatierung von {worksheet.title}: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _apply_category_color_coding(self, worksheet, df, category_colors):
        """
        Wendet Farbkodierung auf Hauptkategorien-Spalte an
        """
        try:
            # Finde Spalten-Index für Hauptkategorie
            hauptkat_col = None
            for i, col_name in enumerate(df.columns, 1):
                if col_name == 'Hauptkategorie':
                    hauptkat_col = i
                    break
            
            if not hauptkat_col:
                print("⚠️ Hauptkategorie-Spalte nicht gefunden für Farbkodierung")
                return
            
            # Wende Farben auf Zellen an (ab Zeile 2, da Zeile 1 Header ist)
            for row in range(2, len(df) + 2):
                cell = worksheet.cell(row=row, column=hauptkat_col)
                category = cell.value
                
                if category and category in category_colors:
                    color = category_colors[category]
                    cell.fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
            
            print(f"✅ Farbkodierung angewendet für {len(category_colors)} Kategorien")
            
        except Exception as e:
            print(f"⚠️ Fehler bei Farbkodierung: {str(e)}")

    def _format_main_worksheet(self, worksheet, num_columns: int, num_rows: int):
        """
        BESTEHENDE METHODE: Formatiert das Haupt-Arbeitsblatt
        """
        # Spaltenbreite anpassen
        column_widths = {
            1: 20,   # Dokument
            2: 15,   # Attribut1
            3: 15,   # Attribut2  
            4: 15,   # Chunk_Nr
            5: 50,   # Text - WICHTIG: Breit für Text
            6: 30,   # Paraphrase - WICHTIG: Breit für Paraphrase
            7: 10,   # Kodiert
            8: 20,   # Hauptkategorie - WICHTIG
            9: 15,   # Kategorietyp
            10: 25,  # Subkategorien - WICHTIG
            11: 20,  # Schlüsselwörter - WICHTIG
            12: 40,  # Begründung - WICHTIG: Breit für Begründung
            13: 10,  # Konfidenz - WICHTIG
            14: 15,  # Original_Chunk_Nr
            15: 12,  # Mehrfachkodierung
            16: 10,  # Instanz_Nr
            17: 12,  # Gesamt_Instanzen
            18: 12,  # Review_Typ
            19: 15   # Kodierer
        }
        
        for col, width in column_widths.items():
            if col <= num_columns:
                worksheet.column_dimensions[get_column_letter(col)].width = width
        
        # Text-Wrapping für wichtige Spalten
        for row in range(2, num_rows + 2):
            # Text-Spalte (5)
            worksheet.cell(row=row, column=5).alignment = Alignment(wrap_text=True, vertical='top')
            # Paraphrase-Spalte (6) 
            worksheet.cell(row=row, column=6).alignment = Alignment(wrap_text=True, vertical='top')
            # Begründung-Spalte (12)
            worksheet.cell(row=row, column=12).alignment = Alignment(wrap_text=True, vertical='top')
    
    def _create_reliability_sheet(self, workbook, reliability: float, export_mode: str):
        """
        ERWEITERTE METHODE: Erstellt Reliabilitäts-Sheet mit Mehrfachkodierungs-Hinweisen
        """
        ws = workbook.create_sheet("Reliabilität")
        
        # Header
        ws.cell(row=1, column=1, value="Intercoder-Reliabilität").font = Font(bold=True, size=14)
        
        # Krippendorff's Alpha
        ws.cell(row=3, column=1, value="Krippendorff's Alpha:")
        ws.cell(row=3, column=2, value=f"{reliability:.3f}")
        ws.cell(row=3, column=2).font = Font(bold=True, size=12)
        
        # Farbkodierung
        if reliability > 0.8:
            fill_color = '90EE90'  # Grün
        elif reliability > 0.667:
            fill_color = 'FFFF90'  # Gelb
        else:
            fill_color = 'FFB6C1'  # Rot
        
        ws.cell(row=3, column=2).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
        
        # Bewertung
        ws.cell(row=4, column=1, value="Bewertung:")
        rating = ("Exzellent" if reliability > 0.8 else 
                 "Akzeptabel" if reliability > 0.667 else 
                 "Schwach" if reliability > 0.5 else "Unzureichend")
        ws.cell(row=4, column=2, value=rating)
        
        # NEUE INFORMATION: Mehrfachkodierungs-Behandlung
        ws.cell(row=6, column=1, value="Mehrfachkodierungs-Behandlung:").font = Font(bold=True)
        ws.cell(row=7, column=1, value="• Reliabilität basiert auf ursprünglichen Segment-IDs")
        ws.cell(row=8, column=1, value="• Mehrfachkodierungen werden als Set-Variable behandelt")
        ws.cell(row=9, column=1, value="• Kategorie-spezifische Segmentierung für Review")
        ws.cell(row=10, column=1, value=f"• Review-Modus: {export_mode}")
    
    def _create_categories_sheet(self, workbook, categories: Dict, original_categories: Dict):
        """
        BESTEHENDE METHODE: Erstellt Kategorien-Übersicht (bleibt unverändert)
        """
        ws = workbook.create_sheet("Kategorien")
        
        # Header
        headers = ['Hauptkategorie', 'Typ', 'Definition', 'Anzahl Subkategorien', 'Subkategorien']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
        
        # Daten
        row = 2
        for cat_name, cat_obj in categories.items():
            cat_type = 'Deduktiv' if cat_name in original_categories else 'Induktiv'
            definition = getattr(cat_obj, 'definition', '')
            subcats = getattr(cat_obj, 'subcategories', [])
            
            ws.cell(row=row, column=1, value=cat_name)
            ws.cell(row=row, column=2, value=cat_type)
            ws.cell(row=row, column=3, value=definition)
            ws.cell(row=row, column=4, value=len(subcats))
            ws.cell(row=row, column=5, value=', '.join(subcats))
            
            row += 1
    
    def _create_multiple_coding_analysis_sheet(self, workbook, codings: List[Dict]):
        """
        NEUE METHODE: Erstellt detaillierte Mehrfachkodierungs-Analyse
        """
        ws = workbook.create_sheet("Mehrfachkodierungs_Analyse")
        
        # Analysiere Mehrfachkodierungen
        original_segments = defaultdict(list)
        for coding in codings:
            consensus_info = coding.get('consensus_info', {})
            original_id = consensus_info.get('original_segment_id', coding.get('segment_id', ''))
            original_segments[original_id].append(coding)
        
        # Übersicht
        ws.cell(row=1, column=1, value="Mehrfachkodierungs-Analyse").font = Font(bold=True, size=14)
        
        # Statistiken
        total_original = len(original_segments)
        multiple_coded = sum(1 for segs in original_segments.values() if len(segs) > 1)
        single_coded = total_original - multiple_coded
        expansion_factor = len(codings) / total_original if total_original > 0 else 1
        
        ws.cell(row=3, column=1, value="Gesamtstatistik:").font = Font(bold=True)
        ws.cell(row=4, column=1, value="• Ursprüngliche Segmente:")
        ws.cell(row=4, column=2, value=total_original)
        ws.cell(row=5, column=1, value="• Einzelkodierungen:")
        ws.cell(row=5, column=2, value=single_coded)
        ws.cell(row=6, column=1, value="• Mehrfachkodierungen:")
        ws.cell(row=6, column=2, value=multiple_coded)
        ws.cell(row=7, column=1, value="• Finale Segmente:")
        ws.cell(row=7, column=2, value=len(codings))
        ws.cell(row=8, column=1, value="• Expansionsfaktor:")
        ws.cell(row=8, column=2, value=f"{expansion_factor:.2f}")
        
        # Detaillierte Liste der Mehrfachkodierungen
        ws.cell(row=10, column=1, value="Mehrfachkodierungs-Details:").font = Font(bold=True)
        
        headers = ['Original_Segment', 'Anzahl_Kategorien', 'Kategorien', 'Neue_Segment_IDs']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=11, column=col, value=header)
            cell.font = Font(bold=True)
        
        row = 12
        for original_id, segment_codings in original_segments.items():
            if len(segment_codings) > 1:
                categories = [c.get('category', '') for c in segment_codings]
                unique_categories = list(set(categories))
                segment_ids = [c.get('segment_id', '') for c in segment_codings]
                
                ws.cell(row=row, column=1, value=original_id)
                ws.cell(row=row, column=2, value=len(unique_categories))
                ws.cell(row=row, column=3, value=', '.join(unique_categories))
                ws.cell(row=row, column=4, value=', '.join(segment_ids))
                
                row += 1
    
    async def _export_json_results(self, codings: List[Dict], reliability: float, 
                                 categories: Dict, export_mode: str):
        """
        BESTEHENDE METHODE: JSON-Export (bleibt unverändert)
        """
        import json
        from pathlib import Path
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"qca_ergebnisse_{export_mode}_{timestamp}.json"
        output_path = Path(self.output_dir) / filename
        
        export_data = {
            'export_info': {
                'timestamp': datetime.now().isoformat(),
                'export_mode': export_mode,
                'total_codings': len(codings),
                'reliability': reliability
            },
            'codings': codings,
            'categories': {name: {
                'definition': getattr(cat, 'definition', ''),
                'subcategories': getattr(cat, 'subcategories', [])
            } for name, cat in categories.items()}
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        print(f"📄 JSON-Export erstellt: {output_path}")
       
    def _extract_doc_and_chunk_id(self, segment_id: str) -> tuple:
        """
        Extrahiert Dokumentnamen und Chunk-ID aus einer Segment-ID.
        
        Args:
            segment_id: Segment-ID im Format "dokument_chunk_X"
                
        Returns:
            tuple: (doc_name, chunk_id)
        """
        parts = segment_id.split('_chunk_')
        if len(parts) == 2:
            return parts[0], parts[1]
        else:
            # Fallback für ungültige Segment-IDs
            return segment_id, "unknown"

    def _get_base_segment_id(self, coding: Dict) -> str:
        """
        FIX: Extrahiert Basis-Segment-ID ohne Mehrfachkodierungs-Suffix
        Für ResultsExporter Klasse
        """
        segment_id = coding.get('segment_id', '')
        
        # Entferne Mehrfachkodierungs-Suffixe (Format: "doc_chunk_5-1")
        if '-' in segment_id:
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit():
                return parts[0]
        
        return segment_id
    
    def _export_reliability_sheet(self, writer, reliability: float, export_mode: str):
        """
        PUNKT 7: Formatiertes Reliabilitäts-Sheet
        """
        # Erstelle formatierte Reliabilitätsdaten
        reliability_data = [
            ['Metrik', 'Wert', 'Bewertung'],
            ['Krippendorff\'s Alpha', f'{reliability:.3f}', self._get_reliability_rating(reliability)],
            ['Export-Modus', export_mode, ''],
            ['Berechnung', 'Set-basiert für Mehrfachkodierungen', ''],
            ['Basis', 'Ursprüngliche Segment-IDs', '']
        ]
        
        # Als DataFrame exportieren
        df_rel = pd.DataFrame(reliability_data[1:], columns=reliability_data[0])
        df_rel.to_excel(writer, sheet_name='Reliabilität', index=False)
        
        # Formatierung anwenden
        worksheet = writer.sheets['Reliabilität']
        self._apply_professional_formatting(worksheet, df_rel)
        
        # Spezielle Farbkodierung für Alpha-Wert
        alpha_cell = worksheet.cell(row=2, column=2)  # Alpha-Wert
        if reliability > 0.8:
            alpha_cell.fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
        elif reliability > 0.667:
            alpha_cell.fill = PatternFill(start_color='FFFF90', end_color='FFFF90', fill_type='solid')
        else:
            alpha_cell.fill = PatternFill(start_color='FFB6C1', end_color='FFB6C1', fill_type='solid')
    
    def _get_reliability_rating(self, reliability: float) -> str:
        """Bestimmt Reliabilitäts-Bewertung"""
        if reliability > 0.8:
            return "Exzellent"
        elif reliability > 0.667:
            return "Akzeptabel"
        elif reliability > 0.5:
            return "Schwach"
        else:
            return "Unzureichend"
    
    def _export_categories_sheet_formatted(self, writer, categories: Dict, original_categories: Dict):
        """
        Formatiertes Kategorien-Sheet
        """
        cat_data = []
        
        for cat_name, cat_obj in categories.items():
            cat_type = 'Deduktiv' if cat_name in (original_categories or {}) else 'Induktiv'
            definition = getattr(cat_obj, 'definition', '')
            subcats = getattr(cat_obj, 'subcategories', [])
            
            cat_data.append({
                'Hauptkategorie': cat_name,
                'Typ': cat_type,
                'Definition': definition,
                'Anzahl_Subkategorien': len(subcats),
                'Subkategorien': ', '.join(subcats)
            })
        
        if cat_data:
            df_cats = pd.DataFrame(cat_data)
            df_cats.to_excel(writer, sheet_name='Kategorien', index=False)
            
            # Formatierung
            worksheet = writer.sheets['Kategorien']
            self._apply_professional_formatting(worksheet, df_cats)
            
            # Spaltenbreiten anpassen
            worksheet.column_dimensions['A'].width = 25  # Hauptkategorie
            worksheet.column_dimensions['B'].width = 12  # Typ
            worksheet.column_dimensions['C'].width = 50  # Definition
            worksheet.column_dimensions['D'].width = 15  # Anzahl
            worksheet.column_dimensions['E'].width = 40  # Subkategorien
    
    def _export_frequency_analysis(self, writer, df_coded: pd.DataFrame, attribut1_label: str, attribut2_label: str) -> None:
        try:
            # Hole alle Datensätze, auch "Nicht kodiert"
            df_all = df_coded.copy()
            
            # Hole eindeutige Hauptkategorien, inkl. "Nicht kodiert"
            main_categories = df_all['Hauptkategorie'].unique()
            category_colors = {cat: color for cat, color in zip(main_categories, self._generate_pastel_colors(len(main_categories)))}

            if 'Häufigkeitsanalysen' not in writer.sheets:
                writer.book.create_sheet('Häufigkeitsanalysen')
            
            worksheet = writer.sheets['Häufigkeitsanalysen']
            worksheet.delete_rows(1, worksheet.max_row)  # Bestehende Daten löschen

            current_row = 1
            
            from openpyxl.styles import Font
            title_font = Font(bold=True, size=12)

            # 1. Hauptkategorien nach Dokumenten
            cell = worksheet.cell(row=current_row, column=1, value="1. Verteilung der Hauptkategorien")
            cell.font = title_font
            current_row += 2

            # FIX: Verwende dieselbe einfache Methode wie für Subkategorien
            pivot_main = pd.pivot_table(
                df_all,
                index=['Hauptkategorie'],
                columns='Dokument',
                values='Chunk_Nr',
                aggfunc='count',
                margins=True,
                margins_name='Gesamt',
                fill_value=0
            )

            # FIX: Konvertiere zu DataFrame für einfache Ausgabe
            temp_df_main = pivot_main.copy().reset_index()
            
            # Formatierte Spaltenbezeichnungen
            formatted_columns = []
            for col in pivot_main.columns:
                if isinstance(col, tuple):
                    col_parts = [str(part) for part in col if part and part != '']
                    formatted_columns.append(' - '.join(col_parts))
                else:
                    formatted_columns.append(str(col))
            
            # FIX: Erstelle Header-Zeile mit korrekten Spaltennamen wie bei Subkategorien
            headers = ['Hauptkategorie'] + formatted_columns
            for col_idx, header in enumerate(headers):
                worksheet.cell(row=current_row, column=col_idx+1, value=header)
            
            # FIX: Exportiere Daten-Zeilen (beginne bei current_row + 1)
            for row_idx, (index, row) in enumerate(temp_df_main.iterrows()):
                for col_idx, value in enumerate(row):
                    worksheet.cell(row=current_row + 1 + row_idx, column=col_idx+1, value=value)
            
            # Formatiere den Bereich (Header + Daten)
            self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(temp_df_main) + 1, len(headers))
            
            # Zusätzliche Farbkodierung für Hauptkategorien
            for row_idx in range(1, len(temp_df_main) + 1):
                kategorie = temp_df_main.iloc[row_idx-1]['Hauptkategorie']
                if kategorie != 'Gesamt' and kategorie in self.category_colors:
                    color = self.category_colors[kategorie]
                    from openpyxl.styles import PatternFill
                    worksheet.cell(row=current_row + row_idx, column=1).fill = PatternFill(start_color=color, end_color=color, fill_type='solid')
            
            current_row += len(temp_df_main) + 3
            # FIX: Ende - einfache Methode wie bei Subkategorien

            # 2. Subkategorien-Hierarchie (nur für kodierte Segmente)
            cell = worksheet.cell(row=current_row, column=1, value="2. Subkategorien nach Hauptkategorien")
            cell.font = title_font
            current_row += 2

            # Filtere "Nicht kodiert" für Subkategorien-Analyse aus
            df_sub = df_all[df_all['Hauptkategorie'] != "Nicht kodiert"].copy()
            df_sub['Subkategorie'] = df_sub['Subkategorien'].str.split(', ')
            df_sub = df_sub.explode('Subkategorie')
            
            # Erstelle Pivot-Tabelle
            pivot_sub = pd.pivot_table(
                df_sub,
                index=['Hauptkategorie', 'Subkategorie'],
                columns='Dokument',
                values='Chunk_Nr',
                aggfunc='count',
                margins=True,
                margins_name='Gesamt',
                fill_value=0
            )

            # DataFrame für Subkategorien mit korrekten Spaltennamen
            temp_df_sub = pivot_sub.copy().reset_index()
            
            # Formatierte Spaltenbezeichnungen
            formatted_columns = []
            for col in pivot_sub.columns:
                if isinstance(col, tuple):
                    col_parts = [str(part) for part in col if part and part != '']
                    formatted_columns.append(' - '.join(col_parts))
                else:
                    formatted_columns.append(str(col))
            
            # Erstelle Header-Zeile mit korrekten Spaltennamen
            headers = ['Hauptkategorie', 'Subkategorie'] + formatted_columns
            for col_idx, header in enumerate(headers):
                worksheet.cell(row=current_row, column=col_idx+1, value=header)
            
            # Exportiere Daten-Zeilen (beginne bei current_row + 1)
            for row_idx, (index, row) in enumerate(temp_df_sub.iterrows()):
                for col_idx, value in enumerate(row):
                    worksheet.cell(row=current_row + 1 + row_idx, column=col_idx+1, value=value)
            
            # Formatiere den Bereich (Header + Daten)
            self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(temp_df_sub) + 1, len(headers))
            current_row += len(temp_df_sub) + 2

            # 3. Attribut-Analysen
            cell = worksheet.cell(row=current_row, column=1, value="3. Verteilung nach Attributen")
            cell.font = title_font
            current_row += 2

            # 3.1 Attribut 1
            cell = worksheet.cell(row=current_row, column=1, value=f"3.1 Verteilung nach {attribut1_label}")
            cell.font = title_font
            current_row += 1

            attr1_counts = df_coded[attribut1_label].value_counts()
            attr1_counts['Gesamt'] = attr1_counts.sum()
            
            attr1_data = [[attribut1_label, 'Anzahl']] + [[idx, value] for idx, value in attr1_counts.items()]
            
            for row_idx, row_data in enumerate(attr1_data):
                for col_idx, value in enumerate(row_data):
                    worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
            
            self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(attr1_data), 2)
            current_row += len(attr1_data) + 3

            # 3.2 Attribut 2
            cell = worksheet.cell(row=current_row, column=1, value=f"3.2 Verteilung nach {attribut2_label}")
            cell.font = title_font
            current_row += 1

            attr2_counts = df_coded[attribut2_label].value_counts()
            attr2_counts['Gesamt'] = attr2_counts.sum()
            
            attr2_data = [[attribut2_label, 'Anzahl']] + [[idx, value] for idx, value in attr2_counts.items()]
            
            for row_idx, row_data in enumerate(attr2_data):
                for col_idx, value in enumerate(row_data):
                    worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
            
            self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(attr2_data), 2)
            current_row += len(attr2_data) + 3

            # 3.3 Attribut 3 (nur wenn definiert)
            attribut3_label = self.attribute_labels.get('attribut3', '')
            if attribut3_label and attribut3_label in df_coded.columns:
                cell = worksheet.cell(row=current_row, column=1, value=f"3.3 Verteilung nach {attribut3_label}")
                cell.font = title_font
                current_row += 1

                attr3_counts = df_coded[attribut3_label].value_counts()
                attr3_counts['Gesamt'] = attr3_counts.sum()
                
                attr3_data = [[attribut3_label, 'Anzahl']] + [[idx, value] for idx, value in attr3_counts.items()]
                
                for row_idx, row_data in enumerate(attr3_data):
                    for col_idx, value in enumerate(row_data):
                        worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
                
                self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(attr3_data), 2)
                current_row += len(attr3_data) + 3

            # FIX: 3.4 Kreuztabellen der Attribute wieder hinzugefügt
            cell = worksheet.cell(row=current_row, column=1, value="3.4 Kreuztabelle der Attribute")
            cell.font = title_font
            current_row += 1

            # Kreuztabelle 1-2
            cross_tab = pd.crosstab(
                df_coded[attribut1_label], 
                df_coded[attribut2_label],
                margins=True,
                margins_name='Gesamt'
            )
            
            cross_tab_df = cross_tab.copy().reset_index()
            cross_tab_df.columns.name = None
            
            for row_idx, (index, row) in enumerate(cross_tab_df.iterrows()):
                for col_idx, value in enumerate(row):
                    worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
            
            self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(cross_tab_df), len(cross_tab_df.columns))
            current_row += len(cross_tab_df) + 3

            # Weitere Kreuztabellen für Attribut 3, wenn vorhanden
            if attribut3_label and attribut3_label in df_coded.columns:
                # Kreuztabelle 1-3
                cross_tab_1_3 = pd.crosstab(
                    df_coded[attribut1_label], 
                    df_coded[attribut3_label],
                    margins=True,
                    margins_name='Gesamt'
                )
                
                cross_tab_1_3_df = cross_tab_1_3.copy().reset_index()
                cross_tab_1_3_df.columns.name = None
                
                for row_idx, (index, row) in enumerate(cross_tab_1_3_df.iterrows()):
                    for col_idx, value in enumerate(row):
                        worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
                
                self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(cross_tab_1_3_df), len(cross_tab_1_3_df.columns))
                current_row += len(cross_tab_1_3_df) + 3
                
                # Kreuztabelle 2-3
                cross_tab_2_3 = pd.crosstab(
                    df_coded[attribut2_label], 
                    df_coded[attribut3_label],
                    margins=True,
                    margins_name='Gesamt'
                )
                
                cross_tab_2_3_df = cross_tab_2_3.copy().reset_index()
                cross_tab_2_3_df.columns.name = None
                
                for row_idx, (index, row) in enumerate(cross_tab_2_3_df.iterrows()):
                    for col_idx, value in enumerate(row):
                        worksheet.cell(row=current_row + row_idx, column=col_idx+1, value=value)
                
                self._apply_professional_formatting_to_range(worksheet, current_row, 1, len(cross_tab_2_3_df), len(cross_tab_2_3_df.columns))
                current_row += len(cross_tab_2_3_df) + 3
            
            print("✅ Häufigkeitsanalysen erfolgreich mit standardisierter Formatierung exportiert")
            
        except Exception as e:
            print(f"❌ Fehler bei Häufigkeitsanalysen: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _apply_professional_formatting_to_range(self, worksheet, start_row: int, start_col: int, num_rows: int, num_cols: int) -> None:
        """
        FIX: Hilfsmethode für die Formatierung eines bestimmten Bereichs ohne StyleProxy-Probleme
        """
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        # Header formatieren (erste Zeile des Bereichs)
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        header_alignment = Alignment(horizontal='center')
        
        for col in range(start_col, start_col + num_cols):
            cell = worksheet.cell(row=start_row, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Rahmen für alle Zellen
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in range(start_row, start_row + num_rows):
            for col in range(start_col, start_col + num_cols):
                cell = worksheet.cell(row=row, column=col)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=False)
        
        # Abwechselnde Zeilenfärbung (ab zweiter Zeile)
        alternate_fill = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
        for row_num in range(start_row + 1, start_row + num_rows):
            if (row_num - start_row) % 2 == 0:
                for col in range(start_col, start_col + num_cols):
                    cell = worksheet.cell(row=row_num, column=col)
                    if not cell.fill or cell.fill.start_color.rgb == 'FFFFFF':
                        cell.fill = alternate_fill
    
    def _create_formatted_attribute_analysis(self, writer, df_coded, attribute_label):
        """Erstellt formatierte Attribut-Analyse"""
        try:
            # Kreuztabelle erstellen
            crosstab = pd.crosstab(df_coded[attribute_label], df_coded['Hauptkategorie'], margins=True)
            
            if not crosstab.empty:
                sheet_name = f'Analyse_{attribute_label}'[:31]  # Excel Sheet-Name Limit
                crosstab.to_excel(writer, sheet_name=sheet_name)
                
                # Formatierung
                worksheet = writer.sheets[sheet_name]
                self._apply_crosstable_formatting(worksheet, crosstab)
                
        except Exception as e:
            print(f"❌ Fehler bei {attribute_label}-Analyse: {str(e)}")
    
    def _create_formatted_crosstable(self, writer, df_coded, attr1_label, attr2_label):
        """Erstellt formatierte Kreuztabelle"""
        try:
            crosstab = pd.crosstab(df_coded[attr1_label], df_coded[attr2_label], margins=True)
            
            if not crosstab.empty:
                crosstab.to_excel(writer, sheet_name='Kreuztabelle_Attribute')
                
                # Formatierung
                worksheet = writer.sheets['Kreuztabelle_Attribute']
                self._apply_crosstable_formatting(worksheet, crosstab)
                
        except Exception as e:
            print(f"❌ Fehler bei Kreuztabelle: {str(e)}")
    
    def _export_progressive_summaries(self, writer, document_summaries):
        """
        PUNKT 7: Formatierte Progressive Summaries
        """
        try:
            print("📝 Erstelle formatierte Progressive Summaries...")
            
            if not document_summaries:
                print("ℹ️ Keine Document-Summaries verfügbar")
                return
            
            summary_data = []
            for doc_name, summary in document_summaries.items():
                # FIX: Verwende 'summary' statt 'self.summary'
                clean_summary = summary.replace('\n', ' ').replace('\r', ' ').strip()
                
                
                summary_data.append({
                    'Dokument': doc_name,
                    'Finales_Summary': clean_summary,
                    'Wortanzahl': len(summary.split()),
                    'Zeichenanzahl': len(summary),
                    'Durchschnittliche_Wortlänge': f"{len(summary)/len(summary.split()):.1f}" if summary.split() else "0"
                })
            
            if summary_data:
                df_summaries = pd.DataFrame(summary_data)
                df_summaries.to_excel(writer, sheet_name='Progressive_Summaries', index=False)
                
                # Formatierung
                worksheet = writer.sheets['Progressive_Summaries']
                self._apply_professional_formatting(worksheet, df_summaries)
                
                # Spaltenbreiten anpassen
                worksheet.column_dimensions['A'].width = 25  # Dokument
                worksheet.column_dimensions['B'].width = 80  # Summary (breit aber ohne wrapping)
                worksheet.column_dimensions['C'].width = 12  # Wortanzahl
                worksheet.column_dimensions['D'].width = 15  # Zeichenanzahl
                worksheet.column_dimensions['E'].width = 20  # Durchschnittliche Wortlänge
                
                print(f"✅ {len(summary_data)} formatierte Document-Summaries exportiert")
            
        except Exception as e:
            print(f"❌ Fehler bei formatierten Progressive Summaries: {str(e)}")
    
    def _calculate_review_statistics(self, codings: List[Dict], export_mode: str, original_codings: List[Dict] = None) -> Dict[str, int]:
        """
        # FIX: Erweiterte Berechnung der Review-Statistiken mit Vor-/Nach-Review Vergleich
        Berechnet umfassende Statistiken des Review-Prozesses.
        
        Args:
            codings: Liste der finalen Kodierungen nach Review
            export_mode: Verwendeter Review-Modus
            original_codings: Liste der ursprünglichen Kodierungen vor Review (optional)
            
        Returns:
            Dict[str, int]: Erweiterte Statistiken des Review-Prozesses
        """
        stats = {
            'consensus_found': 0,
            'majority_found': 0,
            'manual_priority': 0,
            'no_consensus': 0,
            'single_coding': 0,
            'multiple_coding_consolidated': 0,
            # FIX: Neue Statistiken für bessere Übersicht
            'segments_before_review': 0,
            'segments_after_review': 0,
            'segments_with_conflicts': 0,
            'segments_resolved': 0,
            'total_original_codings': 0,
            'categories_involved': 0
        }
        
        # FIX: Analysiere ursprüngliche Kodierungen vor Review
        if original_codings:
            stats['total_original_codings'] = len(original_codings)
            
            # Gruppiere ursprüngliche Kodierungen nach Segmenten
            from collections import defaultdict
            original_segments = defaultdict(list)
            all_categories = set()
            
            for coding in original_codings:
                segment_id = coding.get('segment_id', '')
                category = coding.get('category', '')
                if segment_id and category not in ['Nicht kodiert', 'Kein Kodierkonsens']:
                    original_segments[segment_id].append(coding)
                    all_categories.add(category)
            
            stats['segments_before_review'] = len(original_segments)
            stats['categories_involved'] = len(all_categories)
            
            # Zähle Segmente mit Konflikten (mehrere verschiedene Kategorien)
            for segment_id, segment_codings in original_segments.items():
                categories = [c.get('category', '') for c in segment_codings]
                unique_categories = set(categories)
                if len(unique_categories) > 1:
                    stats['segments_with_conflicts'] += 1
        # FIX: Ende
        
        stats['segments_after_review'] = len(codings)
        
        for coding in codings:
            # Bestimme den Typ der Kodierung basierend auf verfügbaren Informationen
            if coding.get('manual_review', False):
                stats['manual_priority'] += 1
            elif coding.get('consolidated_from_multiple', False):
                stats['multiple_coding_consolidated'] += 1
            elif coding.get('consensus_info', {}).get('selection_type') == 'consensus':
                stats['consensus_found'] += 1
                stats['segments_resolved'] += 1  # FIX: Zähle als aufgelöst
            elif coding.get('consensus_info', {}).get('selection_type') == 'majority':
                stats['majority_found'] += 1
                stats['segments_resolved'] += 1  # FIX: Zähle als aufgelöst
            elif coding.get('consensus_info', {}).get('selection_type') == 'no_consensus':
                stats['no_consensus'] += 1
            elif coding.get('category') == 'Kein Kodierkonsens':
                stats['no_consensus'] += 1
            else:
                stats['single_coding'] += 1
        
        return stats

    def _export_review_statistics(self, writer, review_stats: Dict, export_mode: str):
        """
        # FIX: Erweiterte Exportfunktion für Review-Statistiken mit detaillierteren Informationen
        Exportiert umfassende Statistiken des Review-Prozesses in ein separates Excel-Sheet.
        
        Args:
            writer: Excel Writer Objekt
            review_stats: Erweiterte Statistiken des Review-Prozesses
            export_mode: Verwendeter Review-Modus
        """
        try:
            if 'Review_Statistiken' not in writer.sheets:
                writer.book.create_sheet('Review_Statistiken')
                
            worksheet = writer.sheets['Review_Statistiken']
            current_row = 1
            
            # Titel
            worksheet.cell(row=current_row, column=1, value=f"Review-Statistiken ({export_mode}-Modus)")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True, size=14)
            current_row += 2
            
            # FIX: Übersichts-Statistiken vor Detail-Aufschlüsselung
            worksheet.cell(row=current_row, column=1, value="ÜBERSICHT")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True, size=12)
            current_row += 1
            
            # Grundlegende Zahlen
            basic_stats = [
                ("Ursprüngliche Kodierungen", review_stats.get('total_original_codings', 0)),
                ("Segmente vor Review", review_stats.get('segments_before_review', 0)),
                ("Segmente nach Review", review_stats.get('segments_after_review', 0)),
                ("Segmente mit Konflikten", review_stats.get('segments_with_conflicts', 0)),
                ("Segmente aufgelöst", review_stats.get('segments_resolved', 0)),
                ("Involvierte Kategorien", review_stats.get('categories_involved', 0))
            ]
            
            for label, value in basic_stats:
                if value > 0:  # Nur anzeigen wenn Werte vorhanden
                    worksheet.cell(row=current_row, column=1, value=label)
                    worksheet.cell(row=current_row, column=2, value=value)
                    current_row += 1
            
            current_row += 1
            # FIX: Ende Übersichts-Sektion
            
            # Header für Detail-Aufschlüsselung
            worksheet.cell(row=current_row, column=1, value="REVIEW-TYPEN AUFSCHLÜSSELUNG")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True, size=12)
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Review-Typ")
            worksheet.cell(row=current_row, column=2, value="Anzahl")
            worksheet.cell(row=current_row, column=3, value="Prozent")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True)
            worksheet.cell(row=current_row, column=2).font = Font(bold=True)
            worksheet.cell(row=current_row, column=3).font = Font(bold=True)
            current_row += 1
            
            # FIX: Berechne Gesamtsumme nur der Review-spezifischen Kategorien
            review_specific_stats = {
                'consensus_found': review_stats.get('consensus_found', 0),
                'majority_found': review_stats.get('majority_found', 0),
                'manual_priority': review_stats.get('manual_priority', 0),
                'no_consensus': review_stats.get('no_consensus', 0),
                'single_coding': review_stats.get('single_coding', 0),
                'multiple_coding_consolidated': review_stats.get('multiple_coding_consolidated', 0)
            }
            
            total_reviewed = sum(review_specific_stats.values())
            
            for stat_name, count in review_specific_stats.items():
                if count > 0:
                    # Übersetze Statistik-Namen
                    german_names = {
                        'consensus_found': 'Konsens gefunden',
                        'majority_found': 'Mehrheit gefunden', 
                        'manual_priority': 'Manuelle Priorität',
                        'no_consensus': 'Kein Konsens',
                        'single_coding': 'Einzelkodierung',
                        'multiple_coding_consolidated': 'Mehrfachkodierung konsolidiert'
                    }
                    
                    display_name = german_names.get(stat_name, stat_name)
                    worksheet.cell(row=current_row, column=1, value=display_name)
                    worksheet.cell(row=current_row, column=2, value=count)
                    
                    # Prozentangabe
                    if total_reviewed > 0:
                        percentage = (count / total_reviewed) * 100
                        worksheet.cell(row=current_row, column=3, value=f"{percentage:.1f}%")
                    
                    current_row += 1
            
            # Gesamtsumme
            worksheet.cell(row=current_row, column=1, value="GESAMT")
            worksheet.cell(row=current_row, column=2, value=total_reviewed)
            worksheet.cell(row=current_row, column=3, value="100.0%")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True)
            worksheet.cell(row=current_row, column=2).font = Font(bold=True)
            worksheet.cell(row=current_row, column=3).font = Font(bold=True)
            
            # FIX: Zusätzliche Analyse-Sektion
            current_row += 2
            worksheet.cell(row=current_row, column=1, value="REVIEW-EFFIZIENZ")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True, size=12)
            current_row += 1
            
            # Effizienz-Berechnung
            conflicts = review_stats.get('segments_with_conflicts', 0)
            resolved = review_stats.get('segments_resolved', 0)
            
            if conflicts > 0:
                resolution_rate = (resolved / conflicts) * 100
                worksheet.cell(row=current_row, column=1, value="Konflikt-Auflösungsrate")
                worksheet.cell(row=current_row, column=2, value=f"{resolution_rate:.1f}%")
                current_row += 1
            
            # Reduzierungsrate berechnen
            original_count = review_stats.get('total_original_codings', 0)
            final_count = review_stats.get('segments_after_review', 0)
            
            if original_count > 0 and final_count > 0:
                reduction_rate = ((original_count - final_count) / original_count) * 100
                worksheet.cell(row=current_row, column=1, value="Kodierungs-Reduzierung")
                worksheet.cell(row=current_row, column=2, value=f"{reduction_rate:.1f}%")
                current_row += 1
            # FIX: Ende Zusätzliche Analyse
            
            # Spaltenbreiten anpassen
            worksheet.column_dimensions['A'].width = 25
            worksheet.column_dimensions['B'].width = 15
            worksheet.column_dimensions['C'].width = 15
            
            print("✅ Erweiterte Review-Statistiken erfolgreich exportiert")
            
        except Exception as e:
            print(f"❌ Fehler beim Export der Review-Statistiken: {str(e)}")
            import traceback
            traceback.print_exc()

    
    def _export_detailed_intercoder_analysis(self, writer, codings: List[Dict], reliability: float):
        """
        PUNKT 6: Detaillierte Intercoder-Analyse mit Subkategorien-Unstimmigkeiten
        """
        try:
            print("👥 Erstelle detaillierte Intercoder-Analyse...")
            
            # Gruppiere Kodierungen nach ursprünglicher Segment-ID
            from collections import defaultdict
            
            original_segments = defaultdict(list)
            for coding in codings:
                # Extrahiere ursprüngliche Segment-ID
                consensus_info = coding.get('consensus_info', {})
                original_id = consensus_info.get('original_segment_id', coding.get('segment_id', ''))
                if original_id:
                    original_segments[original_id].append(coding)
            
            # Analysiere Unstimmigkeiten
            disagreement_data = []
            
            for original_id, segment_codings in original_segments.items():
                if len(segment_codings) < 2:
                    continue  # Keine Unstimmigkeit möglich bei nur einem Kodierer
                
                # Analysiere Hauptkategorien-Unstimmigkeiten
                main_categories = [c.get('category', '') for c in segment_codings]
                unique_main_cats = set(main_categories)
                
                # Analysiere Subkategorien-Unstimmigkeiten
                for main_cat in unique_main_cats:
                    if main_cat in ['Nicht kodiert', 'Kein Kodierkonsens', '']:
                        continue
                    
                    # Finde alle Kodierungen für diese Hauptkategorie
                    cat_codings = [c for c in segment_codings if c.get('category') == main_cat]
                    
                    if len(cat_codings) < 2:
                        continue
                    
                    # Subkategorien-Analyse
                    subcat_sets = []
                    coder_info = []
                    
                    for coding in cat_codings:
                        subcats = set(coding.get('subcategories', []))
                        subcat_sets.append(subcats)
                        coder_info.append({
                            'coder': coding.get('coder_id', 'Unbekannt'),
                            'subcats': list(subcats),
                            'confidence': self._extract_confidence_from_coding(coding)
                        })
                    
                    # Prüfe auf Subkategorien-Unstimmigkeiten
                    all_subcats_identical = all(s == subcat_sets[0] for s in subcat_sets)
                    
                    if not all_subcats_identical or len(unique_main_cats) > 1:
                        # Unstimmigkeit gefunden
                        disagreement_data.append({
                            'Segment_ID': original_id,
                            'Hauptkategorie': main_cat,
                            'Anzahl_Kodierer': len(cat_codings),
                            'Hauptkat_Konsens': 'Ja' if len(unique_main_cats) == 1 else 'Nein',
                            'Subkat_Konsens': 'Ja' if all_subcats_identical else 'Nein',
                            'Kodierer_1': coder_info[0]['coder'] if len(coder_info) > 0 else '',
                            'Subkats_1': ', '.join(coder_info[0]['subcats']) if len(coder_info) > 0 else '',
                            'Konfidenz_1': f"{coder_info[0]['confidence']:.2f}" if len(coder_info) > 0 else '',
                            'Kodierer_2': coder_info[1]['coder'] if len(coder_info) > 1 else '',
                            'Subkats_2': ', '.join(coder_info[1]['subcats']) if len(coder_info) > 1 else '',
                            'Konfidenz_2': f"{coder_info[1]['confidence']:.2f}" if len(coder_info) > 1 else '',
                            'Kodierer_3': coder_info[2]['coder'] if len(coder_info) > 2 else '',
                            'Subkats_3': ', '.join(coder_info[2]['subcats']) if len(coder_info) > 2 else '',
                            'Konfidenz_3': f"{coder_info[2]['confidence']:.2f}" if len(coder_info) > 2 else '',
                            'Unstimmigkeits_Typ': self._classify_disagreement_type(unique_main_cats, all_subcats_identical)
                        })
            
            # Exportiere Unstimmigkeits-Analyse
            if disagreement_data:
                df_disagreements = pd.DataFrame(disagreement_data)
                df_disagreements.to_excel(writer, sheet_name='Intercoder_Unstimmigkeiten', index=False)
                
                # Formatierung anwenden
                worksheet = writer.sheets['Intercoder_Unstimmigkeiten']
                self._format_intercoder_sheet(worksheet, df_disagreements)
                
                print(f"✅ {len(disagreement_data)} Intercoder-Unstimmigkeiten analysiert")
            else:
                # Leeres Sheet mit Info erstellen
                empty_data = [{'Info': 'Keine Intercoder-Unstimmigkeiten gefunden'}]
                df_empty = pd.DataFrame(empty_data)
                df_empty.to_excel(writer, sheet_name='Intercoder_Unstimmigkeiten', index=False)
                print("ℹ️ Keine Intercoder-Unstimmigkeiten gefunden")
            
            # Zusätzlich: Übersichts-Statistiken
            self._create_intercoder_summary(writer, original_segments, reliability)
            
        except Exception as e:
            print(f"❌ Fehler bei Intercoder-Analyse: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _classify_disagreement_type(self, unique_main_cats, all_subcats_identical):
        """
        Klassifiziert den Typ der Unstimmigkeit
        """
        if len(unique_main_cats) > 1:
            return 'Hauptkategorie-Konflikt'
        elif not all_subcats_identical:
            return 'Subkategorie-Konflikt'
        else:
            return 'Andere Unstimmigkeit'
    
    def _format_intercoder_sheet(self, worksheet, df):
        """
        PUNKT 7: Formatierung des Intercoder-Sheets
        """
        # Header formatieren
        for cell in worksheet[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.alignment = Alignment(horizontal='center')
        
        # Rahmen für alle Zellen
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in worksheet.iter_rows(min_row=1, max_row=len(df)+1, min_col=1, max_col=len(df.columns)):
            for cell in row:
                cell.border = thin_border
        
        # Spaltenbreiten anpassen
        column_widths = {
            'A': 15,  # Segment_ID
            'B': 20,  # Hauptkategorie
            'C': 12,  # Anzahl_Kodierer
            'D': 15,  # Hauptkat_Konsens
            'E': 15,  # Subkat_Konsens
            'F': 12,  # Kodierer_1
            'G': 25,  # Subkats_1
            'H': 10,  # Konfidenz_1
            # ... weitere Spalten
        }
        
        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width
    
    def _export_intercoder_bericht(self, writer, original_codings: List[Dict], reliability: float):
        """
        FIX: Intercoder-Bericht mit bereits berechneter Reliabilität
        Für ResultsExporter Klasse
        """
        try:
            print("📊 Erstelle IntercoderBericht mit ursprünglichen Daten...")
            
            worksheet = writer.book.create_sheet("IntercoderBericht")
            current_row = 1
            
            # Titel
            title_cell = worksheet.cell(row=current_row, column=1, value="Intercoder-Reliabilitäts-Bericht")
            title_cell.font = Font(bold=True, size=14)
            current_row += 2
            
            # FIX: Verwende bereits berechnete Reliabilität (aus main())
            print(f"📊 Verwende bereits berechnete Reliabilität: {reliability:.3f}")
            
            # FIX: Berechne zusätzliche Statistiken nur für Display
            reliability_calc = ReliabilityCalculator()
            
            # FIX: Berechne vollständigen Bericht für detaillierte Anzeige
            comprehensive_report = reliability_calc.calculate_comprehensive_reliability(original_codings)
            
            statistics = reliability_calc._calculate_basic_statistics(original_codings)
            agreement_analysis = reliability_calc._calculate_detailed_agreement_analysis(original_codings)
            
            # 1. Reliabilitäts-Übersicht
            worksheet.cell(row=current_row, column=1, value="1. Reliabilitäts-Übersicht")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True)
            current_row += 1
            
            # FIX: Overall Alpha fett gedruckt und prominent
            worksheet.cell(row=current_row, column=1, value="Overall Alpha (Jaccard-basiert):")
            alpha_cell = worksheet.cell(row=current_row, column=2, value=f"{comprehensive_report['overall_alpha']:.3f}")
            alpha_cell.font = Font(bold=True, size=12)
            current_row += 1
            
            # FIX: Bewertung hinzufügen
            worksheet.cell(row=current_row, column=1, value="Bewertung:")
            overall_alpha = comprehensive_report['overall_alpha']
            rating = "Exzellent" if overall_alpha > 0.8 else "Akzeptabel" if overall_alpha > 0.667 else "Unzureichend"
            rating_cell = worksheet.cell(row=current_row, column=2, value=rating)
            rating_cell.font = Font(bold=True)
            
            # FIX: Farbkodierung für die Bewertung
            if overall_alpha > 0.8:
                rating_cell.fill = PatternFill(start_color='90EE90', end_color='90EE90', fill_type='solid')
            elif overall_alpha > 0.667:
                rating_cell.fill = PatternFill(start_color='FFFF90', end_color='FFFF90', fill_type='solid')
            else:
                rating_cell.fill = PatternFill(start_color='FFB6C1', end_color='FFB6C1', fill_type='solid')
            current_row += 2
            
            # FIX: Zusätzliche Alpha-Werte aus dem comprehensive report
            worksheet.cell(row=current_row, column=1, value="Hauptkategorien Alpha (Jaccard):")
            worksheet.cell(row=current_row, column=2, value=f"{comprehensive_report['main_categories_alpha']:.3f}")
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Subkategorien Alpha (Jaccard):")
            worksheet.cell(row=current_row, column=2, value=f"{comprehensive_report['subcategories_alpha']:.3f}")
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Vergleichbare Segmente:")
            worksheet.cell(row=current_row, column=2, value=comprehensive_report['statistics']['vergleichbare_segmente'])
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Anzahl Kodierer:")
            worksheet.cell(row=current_row, column=2, value=comprehensive_report['statistics']['anzahl_kodierer'])
            current_row += 2
            
            # FIX: Methodik-Informationen hinzufügen
            worksheet.cell(row=current_row, column=1, value="2. Methodik")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True)
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Alle Alpha-Werte verwenden:")
            worksheet.cell(row=current_row, column=2, value="Jaccard-Ähnlichkeit")
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Berechnung:")
            worksheet.cell(row=current_row, column=2, value="Set-basierte Berechnung")
            current_row += 1
            
            worksheet.cell(row=current_row, column=1, value="Konsistenz:")
            worksheet.cell(row=current_row, column=2, value="Overall zwischen Haupt- und Sub-Alpha")
            current_row += 2
            
            # Spaltenbreiten anpassen
            worksheet.column_dimensions['A'].width = 35
            worksheet.column_dimensions['B'].width = 20
            
            print("✅ IntercoderBericht mit ursprünglichen Daten erstellt")
            
        except Exception as e:
            print(f"❌ Fehler beim IntercoderBericht: {str(e)}")
            import traceback
            traceback.print_exc()

    def _create_empty_intercoder_sheet(self, writer):
        """
        FIX: Erstellt Info-Sheet wenn keine Reliabilitätsdaten verfügbar
        Für ResultsExporter Klasse
        """
        worksheet = writer.book.create_sheet("IntercoderBericht")
        
        worksheet.cell(row=1, column=1, value="Intercoder-Reliabilitäts-Bericht").font = Font(bold=True, size=14)
        worksheet.cell(row=3, column=1, value="⚠️ Keine ursprünglichen Kodierungen für Reliabilitätsberechnung verfügbar")
        worksheet.cell(row=4, column=1, value="Reliabilität muss vor dem Review-Prozess berechnet werden")
        
        worksheet.column_dimensions['A'].width = 60


    def _create_intercoder_summary(self, writer, original_segments, reliability):
        """
        FIX: Erstellt erweiterte Intercoder-Übersicht mit detaillierten Alpha-Informationen
        Für ResultsExporter Klasse
        """
        try:
            print("📊 Erstelle erweiterte Intercoder-Übersicht...")
            
            # FIX: Berechne detaillierte Reliabilitätsinformationen
            reliability_calc = ReliabilityCalculator()
            
            # Extrahiere ursprüngliche Kodierungen aus den Segment-Daten
            original_codings = []
            for segment_codings in original_segments.values():
                original_codings.extend(segment_codings)
            
            # FIX: Berechne umfassende Reliabilität mit Details
            if original_codings:
                comprehensive_report = reliability_calc.calculate_comprehensive_reliability(original_codings)
                overall_alpha = comprehensive_report['overall_alpha']
                main_categories_alpha = comprehensive_report['main_categories_alpha']
                subcategories_alpha = comprehensive_report['subcategories_alpha']
                agreement_analysis = comprehensive_report['agreement_analysis']
                statistics = comprehensive_report['statistics']
            else:
                # Fallback wenn keine Daten
                overall_alpha = reliability
                main_categories_alpha = 0.0
                subcategories_alpha = 0.0
                agreement_analysis = {'Vollständige Übereinstimmung': 0, 'Hauptkategorie gleich, Subkat. unterschiedlich': 0, 'Hauptkategorie unterschiedlich': 0, 'Gesamt': 0}
                statistics = {'vergleichbare_segmente': 0, 'anzahl_kodierer': 0}
            
            # Berechne Übersichtsstatistiken
            total_segments = len(original_segments)
            segments_with_multiple_coders = sum(1 for segs in original_segments.values() if len(segs) > 1)
            segments_single_coder = total_segments - segments_with_multiple_coders
            
            # Kodierer-Statistiken
            all_coders = set()
            for segment_codings in original_segments.values():
                for coding in segment_codings:
                    coder_id = coding.get('coder_id', 'Unbekannt')
                    all_coders.add(coder_id)
            
            # FIX: Erweiterte Daten mit allen Alpha-Details
            summary_data = [
                ['Metrik', 'Wert'],
                ['Gesamte Segmente', total_segments],
                ['Segmente mit mehreren Kodierern', segments_with_multiple_coders],
                ['Segmente mit einem Kodierer', segments_single_coder],
                ['Anteil Mehrfachkodierung', f"{(segments_with_multiple_coders/total_segments)*100:.1f}%" if total_segments > 0 else "0%"],
                ['Anzahl Kodierer', len(all_coders)],
                ['Kodierer', ', '.join(sorted(all_coders))],
                ['', ''],  # Leerzeile als Trenner
                ['--- KRIPPENDORFF\'S ALPHA DETAILS ---', ''],
                ['Overall Alpha (Jaccard-basiert)', f"{overall_alpha:.3f}"],
                ['Hauptkategorien Alpha (Jaccard)', f"{main_categories_alpha:.3f}"],
                ['Subkategorien Alpha (Jaccard)', f"{subcategories_alpha:.3f}"],
                ['Vergleichbare Segmente', statistics.get('vergleichbare_segmente', 0)],
                ['', ''],  # Leerzeile als Trenner
                ['--- ÜBEREINSTIMMUNGSANALYSE ---', ''],
                ['Vollständige Übereinstimmung', agreement_analysis.get('Vollständige Übereinstimmung', 0)],
                ['Hauptkategorie gleich, Subkat. unterschiedlich', agreement_analysis.get('Hauptkategorie gleich, Subkat. unterschiedlich', 0)],
                ['Hauptkategorie unterschiedlich', agreement_analysis.get('Hauptkategorie unterschiedlich', 0)],
                ['Gesamt analysiert', agreement_analysis.get('Gesamt', 0)],
                ['', ''],  # Leerzeile als Trenner
                ['--- BEWERTUNG ---', ''],
                ['Overall Alpha Bewertung', self._get_reliability_rating(overall_alpha)],
                ['Hauptkategorien Bewertung', self._get_reliability_rating(main_categories_alpha)],
                ['Subkategorien Bewertung', self._get_reliability_rating(subcategories_alpha)]
            ]
            
            df_summary = pd.DataFrame(summary_data[1:], columns=summary_data[0])
            df_summary.to_excel(writer, sheet_name='Intercoder_Übersicht', index=False)
            
            # Formatierung
            worksheet = writer.sheets['Intercoder_Übersicht']
            self._apply_professional_formatting(worksheet, df_summary)
            
            # FIX: Spezielle Formatierung für Trennzeilen
            for row_idx in range(1, len(summary_data)):
                cell_value = summary_data[row_idx][0]
                if cell_value.startswith('---') and cell_value.endswith('---'):
                    # Fette Formatierung für Sektions-Header
                    worksheet.cell(row=row_idx + 1, column=1).font = Font(bold=True, size=11)
                elif cell_value == '':
                    # Leerzeilen
                    worksheet.cell(row=row_idx + 1, column=1).value = ''
            
            # Spaltenbreiten
            worksheet.column_dimensions['A'].width = 40
            worksheet.column_dimensions['B'].width = 25
            
            print("✅ Erweiterte Intercoder-Übersicht mit Alpha-Details erstellt")
            
        except Exception as e:
            print(f"❌ Fehler bei erweiterter Intercoder-Übersicht: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _apply_professional_formatting(self, worksheet, df):
        """
        Wendet professionelle Formatierung an
        """
        # Header formatieren
        for cell in worksheet[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
            cell.alignment = Alignment(horizontal='center')
        
        # Rahmen für alle Zellen
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in worksheet.iter_rows(min_row=1, max_row=len(df)+1, min_col=1, max_col=len(df.columns)):
            for cell in row:
                cell.border = thin_border
                # PUNKT 1: Keine Text-Wrapping
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=False)
        
        # Abwechselnde Zeilenfärbung
        for row_num in range(2, len(df) + 2):
            if row_num % 2 == 0:
                for col in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=row_num, column=col)
                    cell.fill = PatternFill(start_color='F8F9FA', end_color='F8F9FA', fill_type='solid')
    
    def _apply_crosstable_formatting(self, worksheet, crosstab):
        """
        PUNKT 7: Spezielle Formatierung für Kreuztabellen
        """
        # Header-Zeile und -Spalte formatieren
        for cell in worksheet[1]:  # Erste Zeile
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='B0C4DE', end_color='B0C4DE', fill_type='solid')
        
        for row in range(1, len(crosstab) + 2):  # Erste Spalte
            cell = worksheet.cell(row=row, column=1)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='B0C4DE', end_color='B0C4DE', fill_type='solid')
        
        # Rahmen
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for row in worksheet.iter_rows(min_row=1, max_row=len(crosstab)+1, min_col=1, max_col=len(crosstab.columns)+1):
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Marginal-Zeilen/-Spalten hervorheben (falls vorhanden)
        if 'All' in crosstab.index:
            last_row = len(crosstab) + 1
            for col in range(1, len(crosstab.columns) + 2):
                cell = worksheet.cell(row=last_row, column=col)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='FFE4B5', end_color='FFE4B5', fill_type='solid')
        
        if 'All' in crosstab.columns:
            last_col = len(crosstab.columns) + 1
            for row in range(1, len(crosstab) + 2):
                cell = worksheet.cell(row=row, column=last_col)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='FFE4B5', end_color='FFE4B5', fill_type='solid')
    
    def _extract_confidence_from_coding(self, coding: Dict) -> float:
        """
        Extrahiert Konfidenzwert aus Kodierung
        """
        confidence = coding.get('confidence', {})
        if isinstance(confidence, dict):
            return confidence.get('total', 0.0)
        elif isinstance(confidence, (int, float)):
            return float(confidence)
        return 0.0
    
    def _determine_category_type(self, category: str, original_categories: dict) -> str:
        """
        # FIX: Erweiterte Bestimmung des Kategorietyps (deduktiv/induktiv/grounded)
        Bestimmt Kategorietyp (deduktiv/induktiv/grounded) basierend auf Analysemodus und Kategoriendefinition
        
        Args:
            category: Name der Kategorie
            original_categories: Dictionary der ursprünglichen Kategorien mit CategoryDefinition-Objekten
            
        Returns:
            str: 'Deduktiv', 'Induktiv', 'Grounded' oder '' (für nicht kodiert)
        """
        if not category or category in ['Nicht kodiert', 'Kein Kodierkonsens']:
            return ''
        
        # FIX: Prüfe zuerst ob es eine deduktive Kategorie ist
        if category in original_categories:
            return 'Deduktiv'
        
        # FIX: Prüfe ob die Analyse im grounded mode durchgeführt wurde
        analysis_mode = CONFIG.get('ANALYSIS_MODE', 'deductive')
        if analysis_mode == 'grounded':
            return 'Grounded'
        
        # FIX: Alternative: Prüfe ob in der CategoryDefinition ein development_type gespeichert ist
        if original_categories and category in original_categories:
            category_obj = original_categories[category]
            if hasattr(category_obj, 'development_type'):
                dev_type = getattr(category_obj, 'development_type', '')
                if dev_type == 'grounded':
                    return 'Grounded'
                elif dev_type == 'deductive':
                    return 'Deduktiv'
                elif dev_type == 'inductive':
                    return 'Induktiv'
        
        # FIX: Fallback: Alle anderen Kategorien sind induktiv
        return 'Induktiv'
    
    def _extract_attributes_from_document(self, doc_name: str) -> tuple:
        """
        Extrahiert erste 2 Attribute (für Rückwärtskompatibilität)
        """
        attr1, attr2, _ = self._extract_three_attributes_from_document(doc_name)
        return attr1, attr2
    
    def _prepare_dataframe_for_frequency_analysis(self, codings: List[Dict]) -> pd.DataFrame:
        """
        Bereitet DataFrame für Häufigkeitsanalyse vor
        """
        try:
            data = []
            
            for coding in codings:
                doc_name = coding.get('document', '')
                if not doc_name:
                    doc_name = self._extract_document_from_segment_id(coding.get('segment_id', ''))
                
                attribut1, attribut2, attribut3 = self._extract_three_attributes_from_document(doc_name)
                
                row_data = {
                    'Dokument': doc_name,
                    self.attribute_labels.get('attribut1', 'Attribut1'): attribut1,
                    self.attribute_labels.get('attribut2', 'Attribut2'): attribut2,
                    self.attribute_labels.get('attribut3', 'Attribut3'): attribut3,
                    'Chunk_Nr': coding.get('chunk_id', coding.get('segment_id', '')),
                    'Hauptkategorie': coding.get('category', ''),
                    'Kodiert': 'Ja' if coding.get('category') and coding.get('category') not in ['Nicht kodiert', 'Kein Kodierkonsens'] else 'Nein',
                    'Subkategorien': ', '.join(coding.get('subcategories', [])),
                    'Konfidenz': self._extract_confidence_from_coding(coding)
                }
                
                data.append(row_data)
            
            return pd.DataFrame(data) if data else pd.DataFrame()
            
        except Exception as e:
            print(f"❌ Fehler beim DataFrame erstellen: {str(e)}")
            return pd.DataFrame()
    
    def export_annotated_pdfs(self, 
                             codings: List[Dict], 
                             chunks: Dict[str, List[str]], 
                             data_dir: str) -> List[str]:
        """
        FIX: Neue Methode für ResultsExporter Klasse
        Exportiert annotierte PDFs für alle gefundenen PDF-Eingabedateien
        
        Args:
            codings: Liste der finalen Kodierungen
            chunks: Dictionary mit chunk_id -> text mapping
            data_dir: Input-Verzeichnis mit Original-PDF-Dateien (aus CONFIG['DATA_DIR'])
            
        Returns:
            List[str]: Liste der Pfade zu erstellten annotierten PDFs
        """
        print(f"\n🎨 Beginne PDF-Annotations-Export...")
        
        try:
            # FIX: Importiere PDF-Annotator (nur wenn benötigt)
            from QCA_Utils import PDFAnnotator
        except ImportError:
            print("   ❌ PyMuPDF nicht verfügbar - PDF-Annotation übersprungen")
            print("   💡 Installieren Sie mit: pip install PyMuPDF")
            return []
        
        # FIX: Initialisiere PDF-Annotator
        pdf_annotator = PDFAnnotator(self)
        
        # FIX: Nutze os.path (wie im Rest von QCA-AID) statt Path
        pdf_files = []
        
        # FIX: Finde PDF-Dateien mit der gleichen Logik wie DocumentReader
        try:
            if not os.path.exists(data_dir):
                print(f"   ⚠️ Verzeichnis {data_dir} existiert nicht")
                return []

            for file in os.listdir(data_dir):
                file_path = os.path.join(data_dir, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                if os.path.isfile(file_path) and file_ext == '.pdf':
                    pdf_files.append((file, file_path))
                    
        except Exception as e:
            print(f"   ❌ Fehler beim Durchsuchen des Verzeichnisses: {e}")
            return []
        
        if not pdf_files:
            print("   ℹ️ Keine PDF-Dateien im Input-Verzeichnis gefunden")
            return []
        
        print(f"   📄 {len(pdf_files)} PDF-Dateien gefunden")
        annotated_files = []
        
        # FIX: Annotiere jede PDF-Datei
        for filename, file_path in pdf_files:
            print(f"\n   📄 Verarbeite: {filename}")
            
            # FIX: Filtere nur konsolidierte/Review-Kodierungen für diese Datei
            file_stem = os.path.splitext(filename)[0]
            file_codings = []
            
            for coding in codings:
                # FIX: Nur Kodierungen nach Review/Consensus nehmen
                is_review_coding = (
                    coding.get('consensus_info') is not None or          # Hat Consensus-Info
                    coding.get('review_decision') is not None or         # Hat Review-Entscheidung  
                    coding.get('selection_type') in ['consensus', 'majority', 'manual_priority'] or  # Ist Review-Ergebnis
                    len([c for c in codings if c.get('segment_id') == coding.get('segment_id')]) == 1  # Einzige Kodierung für Segment
                )
                
                # FIX: Prüfe ob Kodierung zu dieser Datei gehört
                matches_file = (
                    file_stem in coding.get('document', '') or 
                    file_stem in coding.get('segment_id', '')
                )
                
                if is_review_coding and matches_file:
                    file_codings.append(coding)
                    
            print(f"      📋 {len(file_codings)} Review-Kodierungen für {filename} gefunden")
            
            if not file_codings:
                print(f"      ⚠️ Keine Kodierungen für {filename} gefunden")
                continue
            
            # FIX: Erstelle Ausgabepfad mit os.path
            output_filename = f"{file_stem}_QCA_annotiert.pdf"
            output_file = os.path.join(self.output_dir, output_filename)
            
            # FIX: Annotiere PDF
            try:
                result_path = pdf_annotator.annotate_pdf_with_codings(
                    file_path,
                    file_codings,
                    chunks,
                    output_file
                )
                
                if result_path:
                    annotated_files.append(result_path)
                    print(f"      ✅ Erstellt: {os.path.basename(result_path)}")
                
            except Exception as e:
                print(f"      ❌ Fehler bei {filename}: {e}")
                continue
        
        print(f"\n✅ PDF-Annotation abgeschlossen: {len(annotated_files)} Dateien erstellt")
        return annotated_files
    
    def export_annotated_pdfs_all_formats(self, 
                                         codings: List[Dict], 
                                         chunks: Dict[str, List[str]], 
                                         data_dir: str) -> List[str]:
        """
        FIX: Neue Methode für ResultsExporter Klasse - Erweiterte PDF-Annotation für alle Formate
        
        Args:
            codings: Liste der finalen Kodierungen
            chunks: Dictionary mit chunk_id -> text mapping
            data_dir: Input-Verzeichnis mit Original-Dateien
            
        Returns:
            List[str]: Liste der Pfade zu erstellten annotierten PDFs
        """
        print(f"\n🎨 Beginne erweiterte PDF-Annotations-Export für alle Formate...")
        
        try:
            # FIX: Importiere benötigte Klassen
            from QCA_Utils import PDFAnnotator, DocumentToPDFConverter
        except ImportError:
            print("   ❌ Benötigte Bibliotheken nicht verfügbar")
            print("   💡 Installieren Sie mit: pip install PyMuPDF reportlab")
            return []
        
        # FIX: Initialisiere Konverter und Annotator
        try:
            pdf_converter = DocumentToPDFConverter(self.output_dir)
            pdf_annotator = PDFAnnotator(self)
        except Exception as e:
            print(f"   ❌ Fehler beim Initialisieren: {e}")
            return []
        
        # FIX: Finde alle unterstützten Dateien
        supported_extensions = ['.pdf', '.txt', '.docx', '.doc']
        input_files = []
        
        try:
            if not os.path.exists(data_dir):
                print(f"   ⚠️ Verzeichnis {data_dir} existiert nicht")
                return []

            for file in os.listdir(data_dir):
                file_path = os.path.join(data_dir, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                if os.path.isfile(file_path) and file_ext in supported_extensions:
                    input_files.append((file, file_path, file_ext))
                    
        except Exception as e:
            print(f"   ❌ Fehler beim Durchsuchen des Verzeichnisses: {e}")
            return []
        
        if not input_files:
            print("   ℹ️ Keine unterstützten Dateien im Input-Verzeichnis gefunden")
            return []
        
        print(f"   📁 {len(input_files)} Dateien gefunden:")
        for filename, _, ext in input_files:
            print(f"      • {filename} ({ext})")
        
        annotated_files = []
        
        # FIX: Verarbeite jede Datei
        for filename, file_path, file_ext in input_files:
            print(f"\n   📄 Verarbeite: {filename}")
            
            # FIX: Filtere Review-Kodierungen für diese Datei
            file_stem = os.path.splitext(filename)[0]
            file_codings = []
            
            for coding in codings:
                # FIX: Nur Kodierungen nach Review/Consensus nehmen
                is_review_coding = (
                    coding.get('consensus_info') is not None or          
                    coding.get('review_decision') is not None or         
                    coding.get('selection_type') in ['consensus', 'majority', 'manual_priority'] or  
                    len([c for c in codings if c.get('segment_id') == coding.get('segment_id')]) == 1  
                )
                
                # FIX: Prüfe ob Kodierung zu dieser Datei gehört
                matches_file = (
                    file_stem in coding.get('document', '') or 
                    file_stem in coding.get('segment_id', '')
                )
                
                if is_review_coding and matches_file:
                    file_codings.append(coding)
            
            if not file_codings:
                print(f"      ⚠️ Keine Review-Kodierungen für {filename} gefunden")
                continue
            
            print(f"      📋 {len(file_codings)} Review-Kodierungen gefunden")
            
            # FIX: Konvertiere zu PDF falls nötig
            if file_ext == '.pdf':
                pdf_path = file_path
                print(f"      ✅ Bereits PDF")
            else:
                print(f"      🔄 Konvertiere {file_ext.upper()} zu PDF...")
                pdf_path = pdf_converter.convert_document_to_pdf(file_path)
                
                if not pdf_path:
                    print(f"      ❌ Konvertierung fehlgeschlagen")
                    continue
                
                print(f"      ✅ PDF erstellt: {os.path.basename(pdf_path)}")
            
            # FIX: Annotiere PDF
            try:
                output_filename = f"{file_stem}_QCA_annotiert.pdf"
                output_file = os.path.join(self.output_dir, output_filename)
                
                result_path = pdf_annotator.annotate_pdf_with_codings(
                    pdf_path,
                    file_codings,
                    chunks,
                    output_file
                )
                
                if result_path:
                    annotated_files.append(result_path)
                    print(f"      ✅ Annotiert: {os.path.basename(result_path)}")
                else:
                    print(f"      ❌ Annotation fehlgeschlagen")
                
            except Exception as e:
                print(f"      ❌ Fehler bei Annotation: {e}")
                continue
        
        # FIX: Bereinige temporäre Dateien
        try:
            pdf_converter.cleanup_temp_pdfs()
        except Exception as e:
            print(f"   ⚠️ Fehler bei Bereinigung: {e}")
        
        print(f"\n✅ Erweiterte PDF-Annotation abgeschlossen: {len(annotated_files)} Dateien erstellt")
        return annotated_files
    
    def _export_configuration(self, writer, export_mode: str):
        """
        # FIX: Fehlende Methode _export_configuration hinzugefügt
        Exportiert die Konfiguration in ein separates Excel-Sheet.
        Für ResultsExporter Klasse
        
        Args:
            writer: Excel Writer Objekt
            export_mode: Verwendeter Export-Modus
        """
        try:
            from openpyxl.styles import Font, PatternFill
            
            if 'Konfiguration' not in writer.sheets:
                writer.book.create_sheet('Konfiguration')
                
            worksheet = writer.sheets['Konfiguration']
            current_row = 1
            
            # Titel
            worksheet.cell(row=current_row, column=1, value="QCA-AID Konfiguration")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True, size=14)
            current_row += 2
            
            # Header für die Konfigurationstabelle
            worksheet.cell(row=current_row, column=1, value="Parameter")
            worksheet.cell(row=current_row, column=2, value="Wert")
            worksheet.cell(row=current_row, column=1).font = Font(bold=True)
            worksheet.cell(row=current_row, column=2).font = Font(bold=True)
            worksheet.cell(row=current_row, column=1).fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
            worksheet.cell(row=current_row, column=2).fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
            current_row += 1
            
            # FIX: Alle wichtigen Konfigurationsparameter exportieren
            config_params = [
                ('MODEL_PROVIDER', CONFIG.get('MODEL_PROVIDER', 'OpenAI')),
                ('MODEL_NAME', CONFIG.get('MODEL_NAME', 'gpt-4o-mini')),
                ('CHUNK_SIZE', CONFIG.get('CHUNK_SIZE', 2000)),
                ('CHUNK_OVERLAP', CONFIG.get('CHUNK_OVERLAP', 200)),
                ('BATCH_SIZE', CONFIG.get('BATCH_SIZE', 5)),
                ('CODE_WITH_CONTEXT', CONFIG.get('CODE_WITH_CONTEXT', False)),
                ('ANALYSIS_MODE', CONFIG.get('ANALYSIS_MODE', 'deductive')),
                ('REVIEW_MODE', export_mode),
                ('EXPORT_ANNOTATED_PDFS', CONFIG.get('EXPORT_ANNOTATED_PDFS', True)),
                ('ATTRIBUT1_LABEL', CONFIG.get('ATTRIBUTE_LABELS', {}).get('attribut1', 'Attribut1')),
                ('ATTRIBUT2_LABEL', CONFIG.get('ATTRIBUTE_LABELS', {}).get('attribut2', 'Attribut2')),
                ('TIMESTAMP', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            ]
            
            for param_name, param_value in config_params:
                worksheet.cell(row=current_row, column=1, value=param_name)
                worksheet.cell(row=current_row, column=2, value=str(param_value))
                current_row += 1
            
            # Spaltenbreiten anpassen
            worksheet.column_dimensions['A'].width = 25
            worksheet.column_dimensions['B'].width = 40
            
            print("⚙️ Konfiguration erfolgreich exportiert")
            
        except Exception as e:
            print(f"❌ Fehler beim Export der Konfiguration: {str(e)}")
            import traceback
            traceback.print_exc()
                
# --- Klasse: CategoryRevisionManager ---
# Aufgabe: Verwaltung der iterativen Kategorienrevision
class CategoryRevisionManager:
    """
    Manages the iterative refinement of categories during qualitative content analysis.
    
    This class implements Mayring's approach to category revision, ensuring systematic
    documentation of changes, validation of the category system, and maintenance of
    methodological quality standards throughout the analysis process.
    """
    
    def __init__(self, output_dir: str, config: Dict):
        """
        Initializes the CategoryRevisionManager with necessary tracking mechanisms.
        
        Args:
            output_dir (str): Directory for storing revision documentation
        """
        self.config = config
        self.output_dir = output_dir
        self.changes: List[CategoryChange] = []
        self.revision_log_path = os.path.join(output_dir, "category_revisions.json")
        
        # Define quality thresholds for category system
        # self.validator = CategoryValidator(config)

        # Load existing revision history if available
        self._load_revision_history()

        # Prompt-Handler hinzufügen
        self.prompt_handler = QCAPrompts(
            forschungsfrage=FORSCHUNGSFRAGE,
            kodierregeln=KODIERREGELN,
            deduktive_kategorien=DEDUKTIVE_KATEGORIEN
        )

    def _load_revision_history(self) -> None:
        """
        Loads the existing revision history from the JSON file if available.
        This method is called during initialization to restore previous category changes.
        """
        try:
            if os.path.exists(self.revision_log_path):
                with open(self.revision_log_path, 'r', encoding='utf-8') as f:
                    history = json.load(f)
                    # Convert the loaded dictionary data back into CategoryChange objects
                    self.changes = [
                        CategoryChange(
                            category_name=change['category_name'],
                            change_type=change['change_type'],
                            description=change['description'],
                            timestamp=change['timestamp'],
                            old_value=change.get('old_value'),
                            new_value=change.get('new_value'),
                            affected_codings=change.get('affected_codings'),
                            justification=change.get('justification', '')
                        )
                        for change in history
                    ]
                print(f"Loaded revision history: {len(self.changes)} changes")
            else:
                print("No existing revision history found - starting fresh")
                self.changes = []
                
        except Exception as e:
            print(f"Error loading revision history: {str(e)}")
            print("Starting with empty revision history")
            self.changes = []

    def _export_revision_history(self, writer, changes: List['CategoryChange']) -> None:
        """
        Exportiert die Revisionshistorie in ein separates Excel-Sheet.
        
        Args:
            writer: Excel Writer Objekt
            changes: Liste der Kategorieänderungen
        """
        try:
            # Erstelle DataFrame für Revisionshistorie
            revision_data = []
            for change in changes:
                # Erstelle lesbares Änderungsdatum
                change_date = datetime.fromisoformat(change.timestamp)
                
                # Bereite die betroffenen Kodierungen auf
                affected_codings = (
                    ', '.join(change.affected_codings)
                    if change.affected_codings
                    else 'Keine'
                )
                
                # Sammle Änderungsdetails
                if change.old_value and change.new_value:
                    details = []
                    for key in set(change.old_value.keys()) | set(change.new_value.keys()):
                        old = change.old_value.get(key, 'Nicht vorhanden')
                        new = change.new_value.get(key, 'Nicht vorhanden')
                        if old != new:
                            details.append(f"{key}: {old} → {new}")
                    details_str = '\n'.join(details)
                else:
                    details_str = ''

                revision_data.append({
                    'Datum': change_date.strftime('%Y-%m-%d %H:%M'),
                    'Kategorie': change.category_name,
                    'Art der Änderung': change.change_type,
                    'Beschreibung': change.description,
                    'Details': details_str,
                    'Begründung': change.justification,
                    'Betroffene Kodierungen': affected_codings
                })
            
            if revision_data:
                df_revisions = pd.DataFrame(revision_data)
                sheet_name = 'Revisionshistorie'
                
                # Erstelle Sheet falls nicht vorhanden
                if sheet_name not in writer.sheets:
                    writer.book.create_sheet(sheet_name)
                    
                # Exportiere Daten
                df_revisions.to_excel(writer, sheet_name=sheet_name, index=False)
                
                # Formatiere Worksheet
                worksheet = writer.sheets[sheet_name]
                
                # Setze Spaltenbreiten
                column_widths = {
                    'A': 20,  # Datum
                    'B': 25,  # Kategorie
                    'C': 15,  # Art der Änderung
                    'D': 50,  # Beschreibung
                    'E': 50,  # Details
                    'F': 50,  # Begründung
                    'G': 40   # Betroffene Kodierungen
                }
                
                for col, width in column_widths.items():
                    worksheet.column_dimensions[col].width = width
                    
                # Formatiere Überschriften
                for cell in worksheet[1]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='EEEEEE', end_color='EEEEEE', fill_type='solid')
                
                # Aktiviere Zeilenumbruch für lange Texte
                for row in worksheet.iter_rows(min_row=2):
                    for cell in row:
                        cell.alignment = Alignment(wrap_text=True, vertical='top')
                
                print(f"Revisionshistorie mit {len(revision_data)} Einträgen exportiert")
                
            else:
                print("Keine Revisionshistorie zum Exportieren vorhanden")
                
        except Exception as e:
            print(f"Warnung: Fehler beim Export der Revisionshistorie: {str(e)}")
            import traceback
            traceback.print_exc()




# --- Hilfsfunktionen ---

async def perform_manual_coding(chunks, categories, manual_coders):
    """
    KORRIGIERT: Behandelt Abbruch und Mehrfachkodierung korrekt
    """
    manual_codings = []
    total_segments = sum(len(chunks[doc]) for doc in chunks)
    processed_segments = 0
    
    # Erstelle eine flache Liste aller zu kodierenden Segmente
    all_segments = []
    for document_name, document_chunks in chunks.items():
        for chunk_id, chunk in enumerate(document_chunks):
            all_segments.append((document_name, chunk_id, chunk))
    
    print(f"\nManuelles Kodieren: Insgesamt {total_segments} Segmente zu kodieren")
    
    try:
        # Verarbeite alle Segmente
        for idx, (document_name, chunk_id, chunk) in enumerate(all_segments):
            processed_segments += 1
            progress_percentage = (processed_segments / total_segments) * 100
            
            print(f"\nManuelles Codieren: Dokument {document_name}, "
                  f"Chunk {chunk_id + 1}/{len(chunks[document_name])} "
                  f"(Gesamt: {processed_segments}/{total_segments}, {progress_percentage:.1f}%)")
            
            # Prüfe, ob es das letzte Segment ist
            last_segment = (processed_segments == total_segments)
            
            for coder_idx, manual_coder in enumerate(manual_coders):
                try:
                    # Informiere den Benutzer über den Fortschritt
                    if last_segment:
                        print(f"Dies ist das letzte zu kodierende Segment!")
                    
                    # Übergabe des last_segment Parameters an die code_chunk Methode
                    coding_result = await manual_coder.code_chunk(chunk, categories, is_last_segment=last_segment)
                    
                    # KORRIGIERT: Prüfe auf ABORT_ALL
                    if coding_result == "ABORT_ALL":
                        print("Manuelles Kodieren wurde vom Benutzer abgebrochen.")
                        
                        # Schließe alle verbliebenen GUI-Fenster
                        for coder in manual_coders:
                            if hasattr(coder, 'root') and coder.root:
                                try:
                                    coder.root.quit()
                                    coder.root.destroy()
                                except:
                                    pass
                        
                        return manual_codings  # Gebe bisher gesammelte Kodierungen zurück
                    
                    # KORRIGIERT: Behandle sowohl Liste als auch einzelne Kodierungen
                    if coding_result:
                        if isinstance(coding_result, list):
                            # Mehrfachkodierung: Verarbeite jede Kodierung in der Liste
                            print(f"Mehrfachkodierung erkannt: {len(coding_result)} Kodierungen")
                            
                            for i, single_coding in enumerate(coding_result, 1):
                                # Erstelle Dictionary-Eintrag für jede Kodierung
                                coding_entry = {
                                    'segment_id': f"{document_name}_chunk_{chunk_id}",
                                    'coder_id': manual_coder.coder_id,
                                    'category': single_coding.get('category', ''),
                                    'subcategories': single_coding.get('subcategories', []),
                                    'confidence': single_coding.get('confidence', {'total': 1.0}),
                                    'justification': single_coding.get('justification', ''),
                                    'text': chunk,
                                    'document_name': document_name,
                                    'chunk_id': chunk_id,
                                    'manual_coding': True,
                                    'manual_multiple_coding': True,
                                    'multiple_coding_instance': i,
                                    'total_coding_instances': len(coding_result),
                                    'coding_date': datetime.now().isoformat()
                                }
                                
                                # Füge weitere Attribute hinzu falls vorhanden
                                for attr in ['paraphrase', 'keywords', 'text_references', 'uncertainties']:
                                    if attr in single_coding:
                                        coding_entry[attr] = single_coding[attr]
                                
                                manual_codings.append(coding_entry)
                                print(f"  ✓ Mehrfachkodierung {i}/{len(coding_result)}: {coding_entry['category']}")
                        
                        else:
                            # Einzelkodierung (Dictionary)
                            coding_entry = {
                                'segment_id': f"{document_name}_chunk_{chunk_id}",
                                'coder_id': manual_coder.coder_id,
                                'category': coding_result.get('category', ''),
                                'subcategories': coding_result.get('subcategories', []),
                                'confidence': coding_result.get('confidence', {'total': 1.0}),
                                'justification': coding_result.get('justification', ''),
                                'text': chunk,
                                'document_name': document_name,
                                'chunk_id': chunk_id,
                                'manual_coding': True,
                                'manual_multiple_coding': False,
                                'multiple_coding_instance': 1,
                                'total_coding_instances': 1,
                                'coding_date': datetime.now().isoformat()
                            }
                            
                            # Füge weitere Attribute hinzu falls vorhanden
                            for attr in ['paraphrase', 'keywords', 'text_references', 'uncertainties']:
                                if attr in coding_result:
                                    coding_entry[attr] = coding_result[attr]
                            
                            manual_codings.append(coding_entry)
                            print(f"✓ Manuelle Kodierung erfolgreich: {coding_entry['category']}")
                    else:
                        print("⚠ Manuelle Kodierung übersprungen")
                        
                except Exception as e:
                    print(f"Fehler bei manuellem Kodierer {manual_coder.coder_id}: {str(e)}")
                    import traceback
                    traceback.print_exc()
                    continue  # Fahre mit dem nächsten Chunk fort
                    
                # Kurze Pause zwischen den Chunks
                await asyncio.sleep(0.5)
    
        print("\n✅ Manueller Kodierungsprozess abgeschlossen")
        print(f"- {len(manual_codings)}/{total_segments} Segmente erfolgreich kodiert")
        
        # Sicherstellen, dass alle Fenster geschlossen sind
        for coder in manual_coders:
            if hasattr(coder, 'root') and coder.root:
                try:
                    coder.root.quit()
                    coder.root.destroy()
                    coder.root = None
                except:
                    pass
    
    except Exception as e:
        print(f"Fehler im manuellen Kodierungsprozess: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Versuche, alle Fenster zu schließen, selbst im Fehlerfall
        for coder in manual_coders:
            if hasattr(coder, 'root') and coder.root:
                try:
                    coder.root.quit()
                    coder.root.destroy()
                    coder.root = None
                except:
                    pass
    
    return manual_codings

# ============================ 
# 5. Hauptprogramm
# ============================ 

# Aufgabe: Zusammenführung aller Komponenten, Steuerung des gesamten Analyseprozesses
async def main() -> None:
    try:
        # FIX: Console Logging initialisieren
        console_logger = ConsoleLogger(CONFIG['OUTPUT_DIR'])
        console_logger.start_logging()

        print("=== Qualitative Inhaltsanalyse nach Mayring ===")

        # 1. Konfiguration laden
        script_dir = os.path.dirname(os.path.abspath(__file__))
        config_loader = ConfigLoader(script_dir)
        
        if config_loader.load_codebook():
            config = config_loader.get_config() 
            config_loader.update_script_globals(globals())
            print("\nKonfiguration erfolgreich geladen")
            # DEBUG: Prüfe die finale Konfiguration
            # print(f"🔍 DEBUG: MULTIPLE_CODINGS nach Config-Update: {CONFIG.get('MULTIPLE_CODINGS')}")
            # print(f"🔍 DEBUG: MULTIPLE_CODING_THRESHOLD nach Config-Update: {CONFIG.get('MULTIPLE_CODING_THRESHOLD')}")
        else:
            print("Verwende Standard-Konfiguration")
            config = CONFIG

        # Mehrfachkodierungs-Konfiguration anzeigen
        print(f"""
╔══════════════════════════════════════════════════════════════╗
║                    MEHRFACHKODIERUNG                         ║
╠══════════════════════════════════════════════════════════════╣
║ Status: {'✓ AKTIVIERT' if CONFIG.get('MULTIPLE_CODINGS', True) else '✗ DEAKTIVIERT'}                                   ║
║ Schwellenwert: {CONFIG.get('MULTIPLE_CODING_THRESHOLD', 0.67):.1%} Relevanz                        ║
║ Verhalten: Segmente werden mehrfach kodiert wenn sie         ║
║           >= {CONFIG.get('MULTIPLE_CODING_THRESHOLD', 0.7):.0%} Relevanz für verschiedene Hauptkategorien   ║
║           haben                                              ║
╚══════════════════════════════════════════════════════════════╝""")
       

        # 2. Kategoriensystem initialisieren
        print("\n1. Initialisiere Kategoriensystem...")
        category_builder = DeductiveCategoryBuilder()
        initial_categories = category_builder.load_theoretical_categories()
        
        # 3. Manager und History initialisieren
        development_history = DevelopmentHistory(CONFIG['OUTPUT_DIR'])
        revision_manager = CategoryRevisionManager(
            output_dir=CONFIG['OUTPUT_DIR'],
            config=config
        )
        
        # Initiale Kategorien dokumentieren
        for category_name in initial_categories.keys():
            revision_manager.changes.append(CategoryChange(
                category_name=category_name,
                change_type='add',
                description="Initiale deduktive Kategorie",
                timestamp=datetime.now().isoformat(),
                justification="Teil des ursprünglichen deduktiven Kategoriensystems"
            ))

        # 4. Dokumente einlesen
        print("\n2. Lese Dokumente ein...")
        reader = DocumentReader(CONFIG['DATA_DIR'])  # Import aus QCA_Utils
        documents = await reader.read_documents()

        if not documents:
            print("\nKeine Dokumente zum Analysieren gefunden.")
            return

        # 4b. Abfrage zur induktiven Kodierung
        print("\n3. Induktive Kodierung konfigurieren...")

        # Prüfe ob ein induktives Codebook existiert
        codebook_path = os.path.join(CONFIG['OUTPUT_DIR'], "codebook_inductive.json")
        skip_inductive = False

        if os.path.exists(codebook_path):
            print("\nGespeichertes induktives Codebook gefunden.")
            print("Automatische Fortführung in 10 Sekunden...")
            
            use_saved = get_input_with_timeout(
                "\nMöchten Sie das gespeicherte erweiterte Kodesystem laden? (j/N)",
                timeout=10
            )
            
            if use_saved.lower() == 'j':
                try:
                    with open(codebook_path, 'r', encoding='utf-8') as f:
                        saved_categories = json.load(f)
                        
                    if 'categories' in saved_categories:
                        # Konvertiere JSON zurück in CategoryDefinition Objekte
                        for name, cat_data in saved_categories['categories'].items():
                            initial_categories[name] = CategoryDefinition(
                                name=name,
                                definition=cat_data['definition'],
                                examples=cat_data.get('examples', []),
                                rules=cat_data.get('rules', []),
                                subcategories=cat_data.get('subcategories', {}),
                                added_date=cat_data.get('added_date', datetime.now().strftime("%Y-%m-%d")),
                                modified_date=cat_data.get('modified_date', datetime.now().strftime("%Y-%m-%d"))
                            )
                        print(f"\n✓ {len(saved_categories['categories'])} Kategorien aus Codebook geladen")
                        skip_inductive = True
                    else:
                        print("\nWarnung: Ungültiges Codebook-Format")
                        
                except Exception as e:
                    print(f"\nFehler beim Laden des Codebooks: {str(e)}")
                    print("Fahre mit Standard-Kategorien fort")

        if not skip_inductive:
            default_mode = CONFIG['ANALYSIS_MODE']
            print("\nAktueller Analysemodus aus Codebook: {default_mode}")
            print("Sie haben 10 Sekunden Zeit für die Eingabe.")
            print("Optionen:")
            print("1 = inductive (volle induktive Analyse)")
            print("2 = abductive (nur Subkategorien entwickeln)")
            print("3 = deductive (nur deduktiv)")
            print("4 = grounded (Subkategorien sammeln, später Hauptkategorien generieren)")

            analysis_mode = get_input_with_timeout(
                f"\nWelchen Analysemodus möchten Sie verwenden? [1/2/3/4] (Standard: {CONFIG['ANALYSIS_MODE']})", 
                timeout=10
            )

            # Mapping von Zahlen zu Modi
            mode_mapping = {
                '1': 'inductive',
                '2': 'abductive',
                '3': 'deductive',
                '4': 'grounded'
            }

            # Verarbeite Zahlen oder direkte Modusangaben, behalte Default wenn leere oder ungültige Eingabe
            if analysis_mode:  # Nur wenn etwas eingegeben wurde
                if analysis_mode in mode_mapping:
                    CONFIG['ANALYSIS_MODE'] = mode_mapping[analysis_mode]
                elif analysis_mode.lower() in mode_mapping.values():
                    CONFIG['ANALYSIS_MODE'] = analysis_mode.lower()
                else:
                    print(f"\nUngültiger Modus '{analysis_mode}'. Verwende Default-Modus '{default_mode}'.")
                    # Keine Änderung an CONFIG['ANALYSIS_MODE'], Default bleibt bestehen
            else:
                print(f"Keine Eingabe. Verwende Default-Modus '{default_mode}'.")

            # Bestimme, ob induktive Analyse übersprungen wird
            skip_inductive = CONFIG['ANALYSIS_MODE'] == 'deductive'

            print(f"\nAnalysemodus: {CONFIG['ANALYSIS_MODE']} {'(Skip induktiv)' if skip_inductive else ''}")

            # Bei Modus 'grounded' zusätzliche Informationen anzeigen
            if CONFIG['ANALYSIS_MODE'] == 'grounded':
                print("""
            Grounded Theory Modus ausgewählt:
            - Zunächst werden Subcodes und Keywords gesammelt, ohne Hauptkategorien zu bilden
            - Erst nach Abschluss aller Segmente werden die Hauptkategorien generiert
            - Die Subcodes werden anhand ihrer Keywords zu thematisch zusammenhängenden Hauptkategorien gruppiert
            - Im Export werden diese als 'grounded' (statt 'induktiv') markiert
            """)

        # 5. Kodierer konfigurieren
        print("\n4. Konfiguriere Kodierer...")
        # Automatische Kodierer
        auto_coders = [
            DeductiveCoder(
                model_name=CONFIG['MODEL_NAME'],
                temperature=coder_config['temperature'],
                coder_id=coder_config['coder_id']
            )
            for coder_config in CONFIG['CODER_SETTINGS']
        ]

        # Manuelle Kodierung konfigurieren
        print("\nKonfiguriere manuelle Kodierung...")
        print("Sie haben 10 Sekunden Zeit für die Eingabe.")
        print("Drücken Sie 'j' für manuelle Kodierung oder 'n' zum Überspringen.")

        manual_coders = []
        user_input = get_input_with_timeout(
            "\nMöchten Sie manuell kodieren? (j/N)",
            timeout=10
        )
        
        if user_input.lower() == 'j':
            manual_coders.append(ManualCoder(coder_id="human_1"))
            print("\n✓ Manueller Kodierer wurde hinzugefügt")
        else:
            print("\nℹ Keine manuelle Kodierung - nur automatische Kodierung wird durchgeführt")

        # 6. Material vorbereiten
        print("\n5. Bereite Material vor...")
        loader = MaterialLoader(
            data_dir=CONFIG['DATA_DIR'],
            chunk_size=CONFIG['CHUNK_SIZE'],
            chunk_overlap=CONFIG['CHUNK_OVERLAP']
        )
        chunks = {}
        for doc_name, doc_text in documents.items():
            chunks[doc_name] = loader.chunk_text(doc_text)
            print(f"- {doc_name}: {len(chunks[doc_name])} Chunks erstellt")

        # 7. Manuelle Kodierung durchführen
        manual_codings = []
        if manual_coders:
            print("\n6. Starte manuelle Kodierung...")
            
            # Verwende die verbesserte perform_manual_coding Funktion
            manual_coding_result = await perform_manual_coding(
                chunks=chunks, 
                categories=initial_categories,
                manual_coders=manual_coders
            )
            
            if manual_coding_result == "ABORT_ALL":
                print("Manuelle Kodierung abgebrochen. Beende Programm.")
                return
                
            manual_codings = manual_coding_result
            print(f"Manuelle Kodierung abgeschlossen: {len(manual_codings)} Kodierungen")
            
            # Stelle sicher, dass alle Kodierer-Fenster geschlossen sind
            for coder in manual_coders:
                if hasattr(coder, 'root') and coder.root:
                    try:
                        coder.root.quit()
                        coder.root.destroy()
                        coder.root = None
                    except:
                        pass


        # 8. Integrierte Analyse starten
        print("\n7. Starte integrierte Analyse...")

        # Zeige Kontext-Modus an
        print(f"\nKodierungsmodus: {'Mit progressivem Kontext' if CONFIG.get('CODE_WITH_CONTEXT', True) else 'Ohne Kontext'}")
        
        analysis_manager = IntegratedAnalysisManager(CONFIG)

        # Initialisiere Fortschrittsüberwachung
        progress_task = asyncio.create_task(
            monitor_progress(analysis_manager)
        )

        try:
            # Starte die Hauptanalyse
            final_categories, coding_results = await analysis_manager.analyze_material(
                chunks=chunks,
                initial_categories=initial_categories,
                skip_inductive=skip_inductive
            )

            # Beende Fortschrittsüberwachung
            progress_task.cancel()
            await progress_task

            # Kombiniere alle Kodierungen
            all_codings = []
            if coding_results and len(coding_results) > 0:
                print(f"\nFüge {len(coding_results)} automatische Kodierungen hinzu")
                for coding in coding_results:
                    if isinstance(coding, dict) and 'segment_id' in coding:
                        all_codings.append(coding)
                    else:
                        print(f"Überspringe ungültige Kodierung: {coding}")

            # Füge manuelle Kodierungen hinzu
            if manual_codings and len(manual_codings) > 0:
                print(f"Füge {len(manual_codings)} manuelle Kodierungen hinzu")
                all_codings.extend(manual_codings)

            print(f"\nGesamtzahl Kodierungen: {len(all_codings)}")


            # 8.  Intercoder-Reliabilität mit kategorie-spezifischer Berechnung
            if all_codings:
                print("\n8. Berechne korrekte Intercoder-Reliabilität...")
                
                # FIX: SICHER ursprüngliche Kodierungen BEVOR Review-Prozess
                original_codings_for_reliability = all_codings.copy()  # Kopie der ursprünglichen Kodierungen
                
                # NEUE LOGIK: Verwende korrigierte ReliabilityCalculator
                reliability_calculator = ReliabilityCalculator()
                reliability = reliability_calculator.calculate_reliability(original_codings_for_reliability)
                
                print(f"📊 Krippendorff's Alpha (korrigiert für Mehrfachkodierungen): {reliability:.3f}")
            else:
                print("\nKeine Kodierungen für Reliabilitätsberechnung")
                reliability = 0.0
                original_codings_for_reliability = []

            # 9. Review-Behandlung mit kategorie-zentrierter Mehrfachkodierungs-Logik
            print(f"\n9. Führe kategorie-zentrierten Review-Prozess durch...")

            # Gruppiere Kodierungen nach Segmenten für Review
            segment_codings = {}
            for coding in all_codings:
                segment_id = coding.get('segment_id')
                if segment_id:
                    if segment_id not in segment_codings:
                        segment_codings[segment_id] = []
                    segment_codings[segment_id].append(coding)
            
            # Erkenne manuelle Kodierer
            manual_coders = set()
            for coding in all_codings:
                coder_id = coding.get('coder_id', '')
                if 'manual' in coder_id.lower() or 'human' in coder_id.lower():
                    manual_coders.add(coder_id)
            
            # Bestimme Review-Modus
            review_mode = CONFIG.get('REVIEW_MODE', 'consensus')

            if manual_coders:
                print(f"🎯 Manuelle Kodierung erkannt von {len(manual_coders)} Kodierern")
                if review_mode == 'manual':
                    print("   Manueller Review-Modus aus CONFIG aktiviert")
                else:
                    print(f"   CONFIG-Einstellung '{review_mode}' wird verwendet (nicht automatisch auf 'manual' geändert)")
            else:
                if review_mode == 'manual':
                    print("   Manueller Review-Modus aus CONFIG aktiviert (auch ohne manuelle Kodierer)")
            
            print(f"📋 Review-Modus: {review_mode}")
            print(f"📊 Eingabe: {len(all_codings)} ursprüngliche Kodierungen")
            
            review_manager = ReviewManager(CONFIG['OUTPUT_DIR'])
            
            try:
                # Führe kategorie-zentrierten Review durch
                reviewed_codings = review_manager.process_coding_review(all_codings, review_mode)
                
                print(f"✅ Review abgeschlossen: {len(reviewed_codings)} finale Kodierungen")
                
                # Überschreibe all_codings mit den reviewten Ergebnissen
                all_codings = reviewed_codings
                
                # Setze Export-Modus
                export_mode = review_mode

                if 'console_logger' in locals():
                    console_logger.stop_logging()    

            except Exception as e:
                print(f"❌ Fehler beim Review-Prozess: {str(e)}")
                print("📝 Verwende ursprüngliche Kodierungen ohne Review")
                # all_codings bleibt unverändert
                export_mode = review_mode
                import traceback
                traceback.print_exc()
                if 'console_logger' in locals():
                    console_logger.stop_logging() 
            

            # 10. Speichere induktiv erweitertes Codebook
            # Hier die Zusammenfassung der finalen Kategorien vor dem Speichern:
            print("\nFinales Kategoriensystem komplett:")
            print(f"- Insgesamt {len(final_categories)} Hauptkategorien")
            print(f"- Davon {len(final_categories) - len(initial_categories)} neu entwickelt")
            
            # Zähle Subkategorien für zusammenfassende Statistik
            total_subcats = sum(len(cat.subcategories) for cat in final_categories.values())
            print(f"- Insgesamt {total_subcats} Subkategorien")
            
            # 10. Speichere induktiv erweitertes Codebook
            if final_categories:
                category_manager = CategoryManager(CONFIG['OUTPUT_DIR'])
                category_manager.save_codebook(
                    categories=final_categories,
                    filename="codebook_inductive.json"
                )
                print(f"\nCodebook erfolgreich gespeichert mit {len(final_categories)} Hauptkategorien und {total_subcats} Subkategorien")

            # 11. Export der Ergebnisse
            print("\n10. Exportiere Ergebnisse...")
            if all_codings:
                exporter = ResultsExporter(
                    output_dir=CONFIG['OUTPUT_DIR'],
                    attribute_labels=CONFIG['ATTRIBUTE_LABELS'],
                    analysis_manager=analysis_manager,
                    inductive_coder=reliability_calculator
                )

                exporter.current_categories = final_categories 
                
                # FIX: Store original codings in exporter for reliability calculation
                exporter.original_codings_for_reliability = original_codings_for_reliability
                
                # Exportiere Ergebnisse mit Document-Summaries, wenn vorhanden
                summary_arg = analysis_manager.document_summaries if CONFIG.get('CODE_WITH_CONTEXT', True) else None

                # FIX: VERWENDE den bereits bestimmten export_mode, lade NICHT nochmal aus CONFIG
                # ENTFERNT: export_mode = CONFIG.get('REVIEW_MODE', 'consensus') 


                # Validiere und mappe den Export-Modus
                if export_mode == 'auto':
                    export_mode = 'consensus'  # 'auto' ist ein Alias für 'consensus'
                elif export_mode not in ['consensus', 'majority', 'manual_priority', 'manual']:
                    print(f"Warnung: Unbekannter export_mode '{export_mode}', verwende 'consensus'")
                    export_mode = 'consensus'
                
                # FIX: Mappe 'manual' auf 'manual_priority' für Export
                if export_mode == 'manual':
                    export_mode = 'manual_priority'

                print(f"Export wird mit Modus '{export_mode}' durchgeführt")

                await exporter.export_results(
                    codings=all_codings,  # Review-Ergebnisse für Export  
                    reliability=reliability,  # Bereits berechnete Reliabilität
                    categories=final_categories,
                    chunks=chunks,  
                    revision_manager=revision_manager,
                    export_mode=export_mode,
                    original_categories=initial_categories,
                    inductive_coder=reliability_calculator,  
                    document_summaries=summary_arg,
                    original_codings=original_codings_for_reliability,
                    is_intermediate_export=False
                )

                # Ausgabe der finalen Summaries, wenn vorhanden
                if CONFIG.get('CODE_WITH_CONTEXT', True) and analysis_manager.document_summaries:
                    print("\nFinale Document-Summaries:")
                    for doc_name, summary in analysis_manager.document_summaries.items():
                        print(f"\n📄 {doc_name}:")
                        print(f"  {summary}")

                # FIX: Korrekte Prüfung von EXPORT_ANNOTATED_PDFS
                export_pdfs_enabled = CONFIG.get('EXPORT_ANNOTATED_PDFS', True)
                print(f"DEBUG: EXPORT_ANNOTATED_PDFS Wert: {export_pdfs_enabled} (Typ: {type(export_pdfs_enabled)})")
                
                if export_pdfs_enabled is False or str(export_pdfs_enabled).lower() in ['false', '0', 'no', 'nein', 'off']:
                    print("\n   ℹ️ PDF-Annotation deaktiviert (EXPORT_ANNOTATED_PDFS=False)")
                elif not pdf_annotation_available:
                    print("\n   ℹ️ PDF-Annotation nicht verfügbar (PyMuPDF/ReportLab fehlt)")
                    print("   💡 Installieren Sie mit: pip install PyMuPDF reportlab")
                else:
                    # PDF-Annotation ist aktiviert und verfügbar
                    try:
                        print("\n🎨 Exportiere annotierte PDFs für alle Dateiformate...")
                        
                        # FIX: Verwende erweiterte Methode für alle Formate
                        annotated_pdfs = exporter.export_annotated_pdfs_all_formats(
                            codings=all_codings,
                            chunks=chunks,
                            data_dir=CONFIG['DATA_DIR']
                        )
                        
                        if annotated_pdfs:
                            print(f"📄 {len(annotated_pdfs)} annotierte PDFs erstellt:")
                            for pdf_path in annotated_pdfs:
                                print(f"   • {os.path.basename(pdf_path)}")
                        else:
                            print("   ℹ️ Keine Dateien für Annotation gefunden")
                            
                    except Exception as e:
                        print(f"   ❌ Fehler bei erweiterter PDF-Annotation: {e}")
                        print("   💡 PDF-Annotation übersprungen, normaler Export fortgesetzt")

                print("Export erfolgreich abgeschlossen")

            else:
                print("Keine Kodierungen zum Exportieren vorhanden")


            # 12. Zeige finale Statistiken
            print("\nAnalyse abgeschlossen:")
            print(analysis_manager.get_analysis_report())

            if CONFIG.get('MULTIPLE_CODINGS', True):
                # Verwende die ursprünglichen Kodierungen für Mehrfachkodierungs-Statistiken
                codings_for_stats = original_codings_for_reliability if original_codings_for_reliability else all_codings
                
                if codings_for_stats:
                    multiple_coding_stats = _calculate_multiple_coding_stats(codings_for_stats)
                    
                    # FIX: ZeroDivisionError bei Division durch Null verhindern
                    auto_coder_ids = set(c.get('coder_id', '') for c in codings_for_stats if c.get('coder_id', '').startswith('auto'))
                    num_auto_coders = len(auto_coder_ids) if auto_coder_ids else 1  # FIX: Mindestens 1 für Division
                    
                    print(f"""
                    Mehrfachkodierungs-Statistiken:
                    - Segmente mit Mehrfachkodierung: {multiple_coding_stats['segments_with_multiple']}
                    - Durchschnittliche Kodierungen pro Segment: {multiple_coding_stats['avg_codings_per_segment']:.2f}
                    - Häufigste Kategorie-Kombinationen: {', '.join(multiple_coding_stats['top_combinations'][:3]) if multiple_coding_stats['top_combinations'] else 'Keine'}
                    - Fokus-Adherence Rate: {multiple_coding_stats['focus_adherence_rate']:.1%}
                    - Mehrfachkodierungs-Faktor: {multiple_coding_stats['avg_codings_per_segment'] / num_auto_coders:.2f}x""")  # FIX: Verwende num_auto_coders statt direkter Division
                else:
                    print("\n                    Mehrfachkodierungs-Statistiken: Keine Kodierungen für Analyse verfügbar")
            else:
                print("\n                    Mehrfachkodierungs-Statistiken: DEAKTIVIERT")
            
            # Token-Statistiken
            print("\nToken-Nutzung:")
            print(token_counter.get_report())
            
            # Relevanz-Statistiken
            relevance_stats = analysis_manager.relevance_checker.get_statistics()
            print("\nRelevanz-Statistiken:")
            print(f"- Segmente analysiert: {relevance_stats['total_segments']}")
            print(f"- Relevante Segmente: {relevance_stats['relevant_segments']}")
            print(f"- Relevanzrate: {relevance_stats['relevance_rate']*100:.1f}%")
            print(f"- API-Calls gespart: {relevance_stats['total_segments'] - relevance_stats['api_calls']}")
            print(f"- Cache-Nutzung: {relevance_stats['cache_size']} Einträge")

            if 'console_logger' in locals():
                console_logger.stop_logging() 

        except asyncio.CancelledError:
            print("\nAnalyse wurde abgebrochen.")
            if 'console_logger' in locals():
                console_logger.stop_logging() 
        finally:
            # Stelle sicher, dass die Fortschrittsüberwachung beendet wird
            if not progress_task.done():
                progress_task.cancel()
                try:
                    await progress_task
                    if 'console_logger' in locals():
                        console_logger.stop_logging() 
                except asyncio.CancelledError:
                    pass
                    if 'console_logger' in locals():
                        console_logger.stop_logging() 

        if 'console_logger' in locals():
            console_logger.stop_logging() 

    except Exception as e:
        import traceback
        print(f"Fehler in der Hauptausführung: {str(e)}")
        traceback.print_exc()
        if 'console_logger' in locals():
            console_logger.stop_logging() 

        try:
            if 'analysis_manager' in locals() and hasattr(analysis_manager, 'coding_results'):
                print("\nVersuche Zwischenergebnisse zu exportieren...")
                await analysis_manager._export_intermediate_results(
                    chunks=chunks if 'chunks' in locals() else {},
                    current_categories=final_categories if 'final_categories' in locals() else {},
                    deductive_categories=initial_categories if 'initial_categories' in locals() else {},
                    initial_categories=initial_categories if 'initial_categories' in locals() else {}
                )
            if 'console_logger' in locals():
                console_logger.stop_logging() 

        except Exception as export_error:
            print(f"Fehler beim Export der Zwischenergebnisse: {str(export_error)}")
            if 'console_logger' in locals():
                console_logger.stop_logging() 


async def monitor_progress(analysis_manager: IntegratedAnalysisManager):
    """
    Überwacht und zeigt den Analysefortschritt an.
    """
    try:
        while True:
            progress = analysis_manager.get_progress_report()
            
            # Formatiere Fortschrittsanzeige
            print("\n--- Analysefortschritt ---")
            print(f"Verarbeitet: {progress['progress']['processed_segments']} Segmente")
            print(f"Geschwindigkeit: {progress['progress']['segments_per_hour']:.1f} Segmente/Stunde")
            print("------------------------")
            
            await asyncio.sleep(30)  # Update alle 30 Sekunden
            
    except asyncio.CancelledError:
        print("\nFortschrittsüberwachung beendet.")

_patch_tkinter_for_threaded_exit()

if __name__ == "__main__":
    try:
        # Windows-spezifische Event Loop Policy setzen
        if os.name == 'nt':
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
        
        # Hauptprogramm ausführen
        asyncio.run(main())
        
    except KeyboardInterrupt:
        print("\nProgramm durch Benutzer beendet")
    except Exception as e:
        print(f"Fehler im Hauptprogramm: {str(e)}")
        raise

