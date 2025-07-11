import re
import asyncio
import sys
import time
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from collections import defaultdict
import os
import json
import pandas as pd
from openpyxl import load_workbook
import tkinter as tk
from tkinter import ttk, messagebox
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

CONFIG = { } # hier leer, in Hauptskript integriert 


# FIX: Kompletter Ersatz der TokenCounter-Klasse durch präzisen TokenTracker
class TokenTracker:
    def __init__(self):
        self.session_stats = {'input': 0, 'output': 0, 'requests': 0, 'cost': 0.0}
        self.daily_stats = self.load_daily_stats()
        self.model_prices = {
            # === CLAUDE MODELLE (bereits korrekt mit vielen Nachkommastellen) ===
            'claude-sonnet-4-20250514': {'input': 0.000015, 'output': 0.000075},     # $15/$75 per 1M tokens
            'claude-opus-4-20241022': {'input': 0.000075, 'output': 0.000375},       # $75/$375 per 1M tokens
            'claude-3-5-sonnet-20241022': {'input': 0.000003, 'output': 0.000015},   # $3/$15 per 1M tokens
            'claude-3-5-haiku-20241022': {'input': 0.00000025, 'output': 0.00000125}, # $0.25/$1.25 per 1M tokens
            
            # === OPENAI GPT-4.1 SERIE ===
            'gpt-4.1': {'input': 0.000003, 'output': 0.000015},                     # $3/$15 per 1M tokens
            'gpt-4.1-mini': {'input': 0.000001, 'output': 0.000004},                # $1/$4 per 1M tokens  
            'gpt-4.1-nano': {'input': 0.0000001, 'output': 0.0000004},              # $0.10/$0.40 per 1M tokens
            
            # === OPENAI GPT-4O SERIE ===
            'gpt-4o': {'input': 0.000003, 'output': 0.00001},                       # $3/$10 per 1M tokens
            'gpt-4o-2024-11-20': {'input': 0.000003, 'output': 0.00001},
            'gpt-4o-mini': {'input': 0.00000015, 'output': 0.0000006},              # $0.15/$0.60 per 1M tokens
            'gpt-4o-mini-2024-07-18': {'input': 0.00000015, 'output': 0.0000006},
            
            # === GPT-4O AUDIO/REALTIME ===
            'gpt-4o-realtime-preview': {'input': 0.000005, 'output': 0.00002},       # $5/$20 per 1M tokens (text)
            'gpt-4o-audio-preview': {'input': 0.000005, 'output': 0.00002},
            
            # === LEGACY GPT-4 MODELLE ===
            'gpt-4': {'input': 0.00003, 'output': 0.00006},                         # $30/$60 per 1M tokens
            'gpt-4-turbo': {'input': 0.00001, 'output': 0.00003},                   # $10/$30 per 1M tokens
            'gpt-4-turbo-2024-04-09': {'input': 0.00001, 'output': 0.00003},
            'gpt-4-1106-preview': {'input': 0.00001, 'output': 0.00003},
            'gpt-4-vision-preview': {'input': 0.00001, 'output': 0.00003},
            
            # === GPT-3.5 TURBO ===
            'gpt-3.5-turbo': {'input': 0.000001, 'output': 0.000002},               # $1/$2 per 1M tokens
            'gpt-3.5-turbo-0125': {'input': 0.000001, 'output': 0.000002},
            'gpt-3.5-turbo-instruct': {'input': 0.0000015, 'output': 0.000002},
            
            # === BATCH API PREISE (50% Rabatt) ===
            'gpt-4o-batch': {'input': 0.0000015, 'output': 0.000005},               # 50% günstiger
            'gpt-4o-mini-batch': {'input': 0.000000075, 'output': 0.0000003},       # 50% günstiger
            'gpt-4-turbo-batch': {'input': 0.000005, 'output': 0.000015},           # 50% günstiger
        }
        
        self.request_start_time = None
        self.request_start_time = None
        
        # FIX: Debug-Tracking hinzufügen
        self.debug_calls = []
    
    def get_model_price(self, model_name):
        """
        Ermittelt den Preis für ein Modell mit Fallback-Logik.
        """
        # Exakte Übereinstimmung
        if model_name in self.model_prices:
            return self.model_prices[model_name]
        
        # Fallback-Logik für ähnliche Modelle
        model_lower = model_name.lower()
        
        # GPT-4.1 Familie
        if 'gpt-4.1' in model_lower:
            if 'nano' in model_lower:
                return self.model_prices['gpt-4.1-nano']
            elif 'mini' in model_lower:
                return self.model_prices['gpt-4.1-mini']
            else:
                return self.model_prices['gpt-4.1']
        
        # GPT-4o Familie
        elif 'gpt-4o' in model_lower:
            if 'mini' in model_lower:
                return self.model_prices['gpt-4o-mini']
            elif 'batch' in model_lower:
                return self.model_prices['gpt-4o-batch']
            else:
                return self.model_prices['gpt-4o']
        
        # GPT-4 Familie
        elif 'gpt-4' in model_lower:
            if 'turbo' in model_lower:
                return self.model_prices['gpt-4-turbo']
            else:
                return self.model_prices['gpt-4']
        
        # GPT-3.5 Familie
        elif 'gpt-3.5' in model_lower:
            return self.model_prices['gpt-3.5-turbo']
        
        # Claude Familie
        elif 'claude' in model_lower:
            if 'sonnet-4' in model_lower:
                return self.model_prices['claude-sonnet-4-20250514']
            elif 'opus-4' in model_lower:
                return self.model_prices['claude-opus-4-20241022']
            else:
                return self.model_prices['claude-3-5-sonnet-20241022']
        
        # Default Fallback (GPT-4o-mini - günstig aber leistungsfähig)
        print(f"⚠️ Unbekanntes Modell '{model_name}' - verwende GPT-4o-mini Preise als Fallback")
        return self.model_prices['gpt-4o-mini']
    
    def load_daily_stats(self):
        """Lade Tagesstatistiken aus Datei"""
        from datetime import date
        import json
        import os
        
        today = str(date.today())
        try:
            if os.path.exists('token_stats.json'):
                with open('token_stats.json', 'r') as f:
                    data = json.load(f)
                    if data.get('date') == today:
                        return data.get('stats', {'input': 0, 'output': 0, 'requests': 0, 'cost': 0.0})
        except:
            pass
        return {'input': 0, 'output': 0, 'requests': 0, 'cost': 0.0}
    
    def save_daily_stats(self):
        """Speichere Tagesstatistiken in Datei"""
        from datetime import date
        import json
        
        data = {
            'date': str(date.today()),
            'stats': self.daily_stats
        }
        try:
            with open('token_stats.json', 'w') as f:
                json.dump(data, f)
        except:
            pass
    
    def start_request(self):
        """Markiere Start einer Anfrage"""
        import time
        self.request_start_time = time.time()
        # FIX: Debug-Log für start_request
        # print(f"🟢 DEBUG: start_request() aufgerufen um {time.time()}")
        
    def add_tokens(self, input_tokens: int, output_tokens: int = 0):
        """
        Kompatibilitätsmethode für Legacy-Code mit ausführlichem Debug.
        """
        import time
        
        # FIX: Ausführliches Debug-Logging
        # print(f"🔵 DEBUG: add_tokens() aufgerufen")
        # print(f"   📨 Input: {input_tokens}, 📤 Output: {output_tokens}")
        # print(f"   🕐 Zeit: {time.time()}")
        
        # Session-Stats aktualisieren
        old_session_total = self.session_stats['input'] + self.session_stats['output']
        
        self.session_stats['input'] += input_tokens
        self.session_stats['output'] += output_tokens
        self.session_stats['requests'] += 1
        
        new_session_total = self.session_stats['input'] + self.session_stats['output']
        
        # Daily-Stats aktualisieren  
        self.daily_stats['input'] += input_tokens
        self.daily_stats['output'] += output_tokens
        self.daily_stats['requests'] += 1
        self.save_daily_stats()
        
        # Debug-Info
        print(f"   ✅ Session-Total: {old_session_total} → {new_session_total}")
        print(f"   📊 Session-Stats: {self.session_stats}")
        
        # Debug-Call protokollieren
        self.debug_calls.append({
            'method': 'add_tokens',
            'input': input_tokens,
            'output': output_tokens,
            'time': time.time(),
            'session_total_after': new_session_total
        })
    
    def track_response(self, response_data, model):
        """Verfolge Token-Verbrauch mit aktualisierter Preislogik"""
        import time
        
        try:
            # Usage-Daten finden
            usage_data = None
            if hasattr(response_data, 'usage'):
                usage_data = response_data.usage
            elif isinstance(response_data, dict) and 'usage' in response_data:
                usage_data = response_data['usage']
            else:
                return
            
            if usage_data:
                input_tokens = 0
                output_tokens = 0
                
                # Multi-Provider Token-Extraktion
                if hasattr(usage_data, 'prompt_tokens'):
                    input_tokens = usage_data.prompt_tokens
                elif hasattr(usage_data, 'input_tokens'):
                    input_tokens = usage_data.input_tokens
                elif isinstance(usage_data, dict):
                    input_tokens = usage_data.get('prompt_tokens', 0) or usage_data.get('input_tokens', 0)
                
                if hasattr(usage_data, 'completion_tokens'):
                    output_tokens = usage_data.completion_tokens
                elif hasattr(usage_data, 'output_tokens'):
                    output_tokens = usage_data.output_tokens
                elif isinstance(usage_data, dict):
                    output_tokens = usage_data.get('completion_tokens', 0) or usage_data.get('output_tokens', 0)
                
                if input_tokens > 0 or output_tokens > 0:
                    # FIX: Verwende neue Preislogik
                    price = self.get_model_price(model)
                    cost = (input_tokens * price['input']) + (output_tokens * price['output'])
                    
                    # Stats aktualisieren
                    self.session_stats['input'] += input_tokens
                    self.session_stats['output'] += output_tokens
                    self.session_stats['requests'] += 1
                    self.session_stats['cost'] += cost
                    
                    self.daily_stats['input'] += input_tokens
                    self.daily_stats['output'] += output_tokens
                    self.daily_stats['requests'] += 1
                    self.daily_stats['cost'] += cost
                    self.save_daily_stats()
                    
                    # Console Output
                    duration = time.time() - self.request_start_time if self.request_start_time else 0
                    # self.print_stats(input_tokens, output_tokens, cost, model, duration)
                    
                    # Rate Limit Warnung
                    self.check_rate_limits(input_tokens, output_tokens)
                    
        except Exception as e:
            print(f"Token-Tracking-Fehler: {e}")
    
    def print_debug_summary(self):
        """Zeige Debug-Zusammenfassung"""
        print(f"\n🔍 DEBUG-ZUSAMMENFASSUNG:")
        print(f"   📞 Gesamt Debug-Calls: {len(self.debug_calls)}")
        
        add_tokens_calls = [c for c in self.debug_calls if c['method'] == 'add_tokens']
        track_response_calls = [c for c in self.debug_calls if c['method'] == 'track_response']
        
        print(f"   🔵 add_tokens aufrufe: {len(add_tokens_calls)}")
        print(f"   🟡 track_response aufrufe: {len(track_response_calls)}")
        
        if add_tokens_calls:
            total_add_tokens = sum(c['input'] + c['output'] for c in add_tokens_calls)
            print(f"   📊 add_tokens Gesamt: {total_add_tokens}")
            
        if track_response_calls:
            total_track_tokens = sum(c['input'] + c['output'] for c in track_response_calls)
            print(f"   📊 track_response Gesamt: {total_track_tokens}")
    
    def print_stats(self, input_tokens, output_tokens, cost, model, duration):
        """Zeige detaillierte Statistiken in der Konsole"""
        print(f"\n🔢 Token-Statistiken:")
        print(f"   📨 Input: {input_tokens:,} | 📤 Output: {output_tokens:,}")
        print(f"   💰 Kosten: ${cost:.6f} | ⏱️ Dauer: {duration:.2f}s")
        print(f"   🤖 Modell: {model}")
        print(f"   📊 Session: {self.session_stats['input']:,}in + {self.session_stats['output']:,}out = ${self.session_stats['cost']:.4f}")
        print(f"   📅 Heute: {self.daily_stats['input']:,}in + {self.daily_stats['output']:,}out = ${self.daily_stats['cost']:.4f}")
    
    def check_rate_limits(self, input_tokens, output_tokens):
        """Prüfe und warne vor Rate-Limits"""
        total_tokens = input_tokens + output_tokens
        
        # Warnung bei hohem Token-Verbrauch
        if total_tokens > 50000:
            print(f"⚠️  WARNUNG: Hoher Token-Verbrauch ({total_tokens:,} Tokens)")
        
        # Warnung bei hohen Tageskosten
        if self.daily_stats['cost'] > 10.0:
            print(f"⚠️  WARNUNG: Hohe Tageskosten (${self.daily_stats['cost']:.2f})")
        
        # Warnung bei vielen Requests
        if self.daily_stats['requests'] > 1000:
            print(f"⚠️  WARNUNG: Viele Requests heute ({self.daily_stats['requests']})")
    
    # FIX: Abwärtskompatibilität zur alten get_report Methode
    def get_report(self):
        """
        Kompatibilitätsmethode für detaillierten Report.
        Erweitert um Session- und Tagesstatistiken sowie Kosten.
        """
        return (f"📊 DETAILLIERTE TOKEN-STATISTIKEN\n"
                f"{'='*50}\n"
                f"🎯 Session-Statistiken:\n"
                f"   📨 Input Tokens: {self.session_stats['input']:,}\n"
                f"   📤 Output Tokens: {self.session_stats['output']:,}\n"
                f"   🔢 Gesamt Tokens: {self.session_stats['input'] + self.session_stats['output']:,}\n"
                f"   🔄 Requests: {self.session_stats['requests']}\n"
                f"   💰 Session-Kosten: ${self.session_stats['cost']:.4f}\n"
                f"\n📅 Tages-Statistiken:\n"
                f"   📨 Input Tokens: {self.daily_stats['input']:,}\n"
                f"   📤 Output Tokens: {self.daily_stats['output']:,}\n"
                f"   🔢 Gesamt Tokens: {self.daily_stats['input'] + self.daily_stats['output']:,}\n"
                f"   🔄 Requests: {self.daily_stats['requests']}\n"
                f"   💰 Tages-Kosten: ${self.daily_stats['cost']:.4f}\n"
                f"{'='*50}")
    
    def get_session_cost(self):
        """Gibt die Session-Kosten zurück"""
        return self.session_stats['cost']
    
    def get_daily_cost(self):
        """Gibt die Tages-Kosten zurück"""
        return self.daily_stats['cost']
    
    def reset_session(self):
        """Setzt Session-Statistiken zurück"""
        self.session_stats = {'input': 0, 'output': 0, 'requests': 0, 'cost': 0.0}
        print("✅ Session-Statistiken zurückgesetzt")


# FIX: Globale Instanz anpassen (ersetzt den alten TokenCounter)
token_counter = TokenTracker()

# --- Klasse: TokenCounter ---
class TokenCounter:
    def __init__(self):
        self.input_tokens = 0
        self.output_tokens = 0

    def add_tokens(self, input_tokens: int, output_tokens: int = 0):
        """
        Zählt Input- und Output-Tokens.
        
        Args:
            input_tokens: Anzahl der Input-Tokens
            output_tokens: Anzahl der Output-Tokens (optional, Standard 0)
        """
        self.input_tokens += input_tokens
        self.output_tokens += output_tokens

    def get_report(self):
        return f"Gesamte Token-Nutzung:\n" \
               f"Input Tokens: {self.input_tokens}\n" \
               f"Output Tokens: {self.output_tokens}\n" \
               f"Gesamt Tokens: {self.input_tokens + self.output_tokens}"

    def estimate_tokens(self,text: str) -> int:
        """
        Schätzt die Anzahl der Tokens in einem Text mit verbesserter Genauigkeit.
        Berücksichtigt verschiedene Zeichentypen, nicht nur Wortgrenzen.
        
        Args:
            text: Zu schätzender Text
            
        Returns:
            int: Geschätzte Tokenanzahl
        """
        if not text:
            return 0
            
        # Grundlegende Schätzung: 1 Token ≈ 4 Zeichen für englischen Text
        # Für deutsche Texte (mit längeren Wörtern) etwas anpassen: 1 Token ≈ 4.5 Zeichen
        char_per_token = 4.5
        
        # Anzahl der Sonderzeichen, die oft eigene Tokens bilden
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
        
        # Anzahl der Wörter
        words = len(text.split())
        
        # Gewichtete Berechnung
        estimated_tokens = int(
            (len(text) / char_per_token) * 0.7 +  # Zeichenbasierte Schätzung (70% Gewichtung)
            (words + special_chars) * 0.3          # Wort- und Sonderzeichenbasierte Schätzung (30% Gewichtung)
        )
        
        return max(1, estimated_tokens)  # Mindestens 1 Token

#token_counter = TokenCounter()



# Hilfsfunktion zur Token-Schätzung


def _safe_speed_calculation(count: int, time_elapsed: float) -> str:
    """
    Sichere Geschwindigkeitsberechnung mit Division-by-Zero Schutz.
    
    Args:
        count: Anzahl der verarbeiteten Elemente
        time_elapsed: Verstrichene Zeit in Sekunden
        
    Returns:
        str: Formatierte Geschwindigkeitsangabe
    """
    if time_elapsed > 0:
        speed = count / time_elapsed
        return f"{speed:.1f} Elemente/Sekunde"
    else:
        return f"{count} Elemente in <0.01s (sehr schnell)"
    
def get_input_with_timeout(prompt: str, timeout: int = 30) -> str:
    """
    Fragt nach Benutzereingabe mit Timeout.
    
    Args:
        prompt: Anzuzeigender Text
        timeout: Timeout in Sekunden
        
    Returns:
        str: Benutzereingabe oder 'n' bei Timeout
    """
    import threading
    import sys
    import time
    from threading import Event
    
    # Plattformspezifische Imports
    if sys.platform == 'win32':
        import msvcrt
    else:
        import select

    answer = {'value': None}
    stop_event = Event()
    
    def input_thread():
        try:
            # Zeige Countdown
            remaining_time = timeout
            while remaining_time > 0 and not stop_event.is_set():
                sys.stdout.write(f'\r{prompt} ({remaining_time}s): ')
                sys.stdout.flush()
                
                # Plattformspezifische Eingabeprüfung
                if sys.platform == 'win32':
                    if msvcrt.kbhit():
                        answer['value'] = msvcrt.getche().decode().strip().lower()
                        sys.stdout.write('\n')
                        stop_event.set()
                        return
                else:
                    if select.select([sys.stdin], [], [], 1)[0]:
                        answer['value'] = sys.stdin.readline().strip().lower()
                        stop_event.set()
                        return
                
                time.sleep(1)
                remaining_time -= 1
            
            # Bei Timeout
            if not stop_event.is_set():
                sys.stdout.write('\n')
                sys.stdout.flush()
                
        except (KeyboardInterrupt, EOFError):
            stop_event.set()
    
    # Starte Input-Thread
    thread = threading.Thread(target=input_thread)
    thread.daemon = True
    thread.start()
    
    # Warte auf Antwort oder Timeout
    thread.join(timeout)
    stop_event.set()
    
    if answer['value'] is None:
        print(f"\nKeine Eingabe innerhalb von {timeout} Sekunden - verwende 'n'")
        return 'n'
        
    return answer['value']

def _calculate_multiple_coding_stats(all_codings: List[Dict]) -> Dict:
    """
    Berechnung der Mehrfachkodierungs-Statistiken
    
    Berücksichtigt sowohl:
    1. Mehrfachkodierung durch verschiedene Kodierer (deduktiver Modus)
    2. Echte Mehrfachkodierung (verschiedene Kategorien für gleichen Text)
    
    Args:
        all_codings: Liste aller Kodierungen
        
    Returns:
        Dict: Statistiken zur Mehrfachkodierung
    """
    from collections import defaultdict, Counter
    
    # Gruppiere nach Segment-ID und analysiere die verschiedenen Arten von Mehrfachkodierung
    segment_codings = defaultdict(list)
    focus_adherence = []
    category_combinations = []
    
    for coding in all_codings:
        segment_id = coding.get('segment_id', '')
        segment_codings[segment_id].append(coding)
        
        # Focus adherence tracking
        if coding.get('category_focus_used', False):
            focus_adherence.append(coding.get('target_category', '') == coding.get('category', ''))
    
    # Analysiere verschiedene Arten von Mehrfachkodierung
    segments_with_multiple_coders = 0
    segments_with_multiple_categories = 0
    segments_with_true_multiple_coding = 0
    
    for segment_id, codings in segment_codings.items():
        if len(codings) > 1:
            # Verschiedene Kodierer für gleiches Segment
            unique_coders = set(c.get('coder_id', '') for c in codings)
            if len(unique_coders) > 1:
                segments_with_multiple_coders += 1
            
            # Verschiedene Kategorien für gleiches Segment
            unique_categories = set(c.get('category', '') for c in codings)
            if len(unique_categories) > 1:
                segments_with_multiple_categories += 1
                # Sammle Kategorie-Kombinationen
                category_combinations.append(' + '.join(sorted(unique_categories)))
            
            # Echte Mehrfachkodierung (verschiedene Instanzen)
            multiple_instances = any(c.get('total_coding_instances', 1) > 1 for c in codings)
            if multiple_instances:
                segments_with_true_multiple_coding += 1
    
    # Bestimme die dominante Art der Mehrfachkodierung für Statistik-Output
    if segments_with_multiple_coders > segments_with_true_multiple_coding:
        # Deduktiver Modus: Zähle Segmente mit mehreren Kodierern als Mehrfachkodierung
        segments_with_multiple = segments_with_multiple_coders
    else:
        # Echter Mehrfachkodierungs-Modus: Zähle nur echte Mehrfachkodierungen
        segments_with_multiple = segments_with_true_multiple_coding
    
    total_segments = len(segment_codings)
    total_codings = len(all_codings)
    
    combination_counter = Counter(category_combinations)
    
    return {
        'segments_with_multiple': segments_with_multiple,
        'total_segments': total_segments,
        'avg_codings_per_segment': total_codings / total_segments if total_segments > 0 else 0,
        'top_combinations': [combo for combo, _ in combination_counter.most_common(5)],
        'focus_adherence_rate': sum(focus_adherence) / len(focus_adherence) if focus_adherence else 0,
        # Zusätzliche Details für erweiterte Ausgaben (optional)
        'segments_with_multiple_coders': segments_with_multiple_coders,
        'segments_with_multiple_categories': segments_with_multiple_categories,
        'segments_with_true_multiple_coding': segments_with_true_multiple_coding
    }


def _patch_tkinter_for_threaded_exit():
    """
    Patcht die Tkinter Variable.__del__ Methode, um den RuntimeError beim Beenden zu vermeiden.
    """
    import tkinter
    
    # Originale __del__ Methode speichern
    original_del = tkinter.Variable.__del__
    
    # Neue __del__ Methode definieren, die Ausnahmen abfängt
    def safe_del(self):
        try:
            # Nur aufrufen, wenn _tk existiert und es sich um ein valides Tkinter-Objekt handelt
            if hasattr(self, '_tk') and self._tk:
                original_del(self)
        except (RuntimeError, TypeError, AttributeError):
            # Diese Ausnahmen stillschweigend ignorieren
            pass
    
    # Die ursprüngliche Methode ersetzen
    tkinter.Variable.__del__ = safe_del
    print("Tkinter für sicheres Beenden gepatcht.")


# Neue Ergänzungen für QCA_Utils.py - Mehrfachkodierung im manuellen Modus

import tkinter as tk
from tkinter import ttk, messagebox
from typing import Dict, List, Optional, Tuple
from datetime import datetime

class MultiSelectListbox(tk.Listbox):
    """
    Erweiterte Listbox mit Mehrfachauswahl per Ctrl+Klick
    für manuelle Mehrfachkodierung
    """
    
    def __init__(self, parent, **kwargs):
        # Aktiviere erweiterte Mehrfachauswahl
        kwargs['selectmode'] = tk.EXTENDED
        super().__init__(parent, **kwargs)
        
        # Bindings für Mehrfachauswahl
        self.bind('<Button-1>', self._on_single_click)
        self.bind('<Control-Button-1>', self._on_ctrl_click)
        self.bind('<Shift-Button-1>', self._on_shift_click)
        
        # Speichere ursprüngliche Auswahl für Ctrl-Klick
        self._last_selection = set()
        
    def _on_single_click(self, event):
        """Normale Einzelauswahl"""
        # Lasse normale Behandlung durch tkinter zu
        self.after_idle(self._update_last_selection)
        
    def _on_ctrl_click(self, event):
        """Ctrl+Klick für Mehrfachauswahl"""
        index = self.nearest(event.y)
        
        if index in self.curselection():
            # Deselektiere wenn bereits ausgewählt
            self.selection_clear(index)
        else:
            # Füge zur Auswahl hinzu
            self.selection_set(index)
            
        self._update_last_selection()
        return "break"  # Verhindert normale Behandlung
        
    def _on_shift_click(self, event):
        """Shift+Klick für Bereichsauswahl"""
        if not self._last_selection:
            return
            
        index = self.nearest(event.y)
        last_indices = list(self._last_selection)
        
        if last_indices:
            start = min(last_indices[0], index)
            end = max(last_indices[0], index)
            
            # Wähle Bereich aus
            for i in range(start, end + 1):
                self.selection_set(i)
                
        self._update_last_selection()
        return "break"
        
    def _update_last_selection(self):
        """Aktualisiert die gespeicherte Auswahl"""
        self._last_selection = set(self.curselection())

class ManualMultipleCodingDialog:
    """
    Dialog für die Bestätigung und Konfiguration von Mehrfachkodierungen
    """
    
    def __init__(self, parent, selected_categories: List[Dict], segment_text: str):
        self.parent = parent
        self.selected_categories = selected_categories
        self.segment_text = segment_text
        self.result = None
        self.dialog = None
        
    def show_dialog(self) -> Optional[List[Dict]]:
        """
        Zeigt den Bestätigungsdialog für Mehrfachkodierung
        
        Returns:
            Optional[List[Dict]]: Liste der bestätigten Kodierungen oder None bei Abbruch
        """
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("Mehrfachkodierung bestätigen")
        self.dialog.geometry("600x500")
        self.dialog.transient(self.parent)
        self.dialog.grab_set()
        
        # Zentriere Dialog
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - (600 // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (500 // 2)
        self.dialog.geometry(f"600x500+{x}+{y}")
        
        self._create_widgets()
        
        # Warte auf Schließung des Dialogs
        self.dialog.wait_window()
        
        return self.result
    
    def _create_widgets(self):
        """Erstellt die Dialog-Widgets"""
        main_frame = ttk.Frame(self.dialog)
        main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        # Titel
        title_label = ttk.Label(
            main_frame, 
            text="Mehrfachkodierung erkannt",
            font=('Arial', 12, 'bold')
        )
        title_label.pack(pady=(0, 10))
        
        # Informationstext
        info_text = f"Sie haben {len(self.selected_categories)} Kategorien/Subkategorien ausgewählt.\n"
        
        # Analysiere Auswahltyp
        main_categories = set()
        for cat in self.selected_categories:
            main_categories.add(cat['main_category'])
            
        if len(main_categories) == 1:
            info_text += f"Alle gehören zur Hauptkategorie '{list(main_categories)[0]}'.\n"
            info_text += "Dies wird als eine Kodierung mit mehreren Subkategorien behandelt."
        else:
            info_text += f"Sie umfassen {len(main_categories)} verschiedene Hauptkategorien.\n"
            info_text += "Dies wird als Mehrfachkodierung behandelt (mehrere Zeilen im Export)."
        
        info_label = ttk.Label(main_frame, text=info_text, wraplength=550)
        info_label.pack(pady=(0, 15))
        
        # Textsegment anzeigen
        text_frame = ttk.LabelFrame(main_frame, text="Textsegment")
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        text_widget = tk.Text(text_frame, height=8, wrap=tk.WORD, state=tk.DISABLED)
        text_widget.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
        text_widget.config(state=tk.NORMAL)
        text_widget.insert(tk.END, self.segment_text[:1000] + ("..." if len(self.segment_text) > 1000 else ""))
        text_widget.config(state=tk.DISABLED)
        
        # Ausgewählte Kategorien anzeigen
        selection_frame = ttk.LabelFrame(main_frame, text="Ihre Auswahl")
        selection_frame.pack(fill=tk.X, pady=(0, 15))
        
        selection_text = tk.Text(selection_frame, height=6, wrap=tk.WORD, state=tk.DISABLED)
        selection_text.pack(padx=5, pady=5, fill=tk.X)
        
        selection_text.config(state=tk.NORMAL)
        for i, cat in enumerate(self.selected_categories, 1):
            if cat['type'] == 'main':
                selection_text.insert(tk.END, f"{i}. Hauptkategorie: {cat['name']}\n")
            else:
                selection_text.insert(tk.END, f"{i}. Subkategorie: {cat['name']} (→ {cat['main_category']})\n")
        selection_text.config(state=tk.DISABLED)
        
        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        ttk.Button(
            button_frame,
            text="Bestätigen",
            command=self._confirm_selection
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(
            button_frame,
            text="Ändern",
            command=self._modify_selection
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(
            button_frame,
            text="Abbrechen",
            command=self._cancel_selection
        ).pack(side=tk.RIGHT)
    
    def _confirm_selection(self):
        """Bestätigt die aktuelle Auswahl"""
        self.result = self.selected_categories
        self.dialog.destroy()
    
    def _modify_selection(self):
        """Schließt Dialog zum Ändern der Auswahl"""
        self.result = "MODIFY"
        self.dialog.destroy()
    
    def _cancel_selection(self):
        """Bricht die Mehrfachkodierung ab"""
        self.result = None
        self.dialog.destroy()

def create_multiple_coding_results(selected_categories: List[Dict], 
                                 text: str, 
                                 coder_id: str) -> List:
    """
    Erstellt CodingResult-Objekte für Mehrfachkodierung
    
    Args:
        selected_categories: Liste der ausgewählten Kategorien
        text: Zu kodierender Text
        coder_id: ID des Kodierers
        
    Returns:
        List: Liste von Kodierungs-Dictionaries (da CodingResult nicht importierbar)
    """
    from datetime import datetime
    
    # Gruppiere nach Hauptkategorien
    main_category_groups = {}
    for cat in selected_categories:
        main_cat = cat['main_category']
        if main_cat not in main_category_groups:
            main_category_groups[main_cat] = {
                'main_category': main_cat,
                'subcategories': []
            }
        
        if cat['type'] == 'sub':
            main_category_groups[main_cat]['subcategories'].append(cat['name'])
    
    # Erstelle Kodierungsergebnisse als Dictionaries
    coding_results = []
    total_instances = len(main_category_groups)
    
    for instance_num, (main_cat, group_data) in enumerate(main_category_groups.items(), 1):
        # Erstelle immer Dictionary (einfacher und zuverlässiger)
        coding_result = {
            'category': main_cat,
            'subcategories': tuple(group_data['subcategories']),
            'justification': f"Manuelle Mehrfachkodierung (Instanz {instance_num}/{total_instances})",
            'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0},
            'text_references': (text[:100],),
            'uncertainties': None,
            'paraphrase': "",
            'keywords': "",
            'multiple_coding_instance': instance_num,
            'total_coding_instances': total_instances,
            'manual_multiple_coding': True,
            'coding_date': datetime.now().isoformat(),
            'coder_id': coder_id,
            'text': text
        }
        
        coding_results.append(coding_result)
    
    return coding_results

def show_multiple_coding_info(parent, num_categories: int, main_categories: List[str]) -> bool:
    """
    Zeigt Informationen über erkannte Mehrfachkodierung
    
    Args:
        parent: Parent-Widget
        num_categories: Anzahl ausgewählter Kategorien
        main_categories: Liste der Hauptkategorien
        
    Returns:
        bool: True wenn fortgefahren werden soll
    """
    if len(main_categories) == 1:
        message = (f"Sie haben {num_categories} Subkategorien der Hauptkategorie "
                  f"'{main_categories[0]}' ausgewählt.\n\n"
                  f"Dies wird als eine Kodierung mit mehreren Subkategorien behandelt.")
        title = "Mehrere Subkategorien"
    else:
        message = (f"Sie haben Kategorien aus {len(main_categories)} verschiedenen "
                  f"Hauptkategorien ausgewählt:\n"
                  f"{', '.join(main_categories)}\n\n"
                  f"Dies wird als Mehrfachkodierung behandelt - das Segment erscheint "
                  f"mehrmals im Export (einmal pro Hauptkategorie).")
        title = "Mehrfachkodierung erkannt"
    
    return messagebox.askyesno(
        title,
        message + "\n\nMöchten Sie fortfahren?",
        parent=parent
    )

# Zusätzliche Hilfsfunktionen für die GUI

def setup_manual_coding_window_enhanced(root, categories, chunk_text, coder_id, is_last_segment=False):
    """
    Erweiterte Einrichtung des manuellen Kodierungsfensters mit Mehrfachauswahl
    
    Args:
        root: Tkinter Root-Fenster
        categories: Verfügbare Kategorien
        chunk_text: Zu kodierender Text
        coder_id: ID des Kodierers
        is_last_segment: Ob es das letzte Segment ist
        
    Returns:
        Dict: GUI-Komponenten für erweiterte Funktionalität
    """
    # Hauptframe
    main_frame = ttk.Frame(root)
    main_frame.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)
    
    # Instruktionen für Mehrfachauswahl
    instructions = ttk.Label(
        main_frame,
        text="💡 Tipp: Verwenden Sie Strg+Klick für Mehrfachauswahl von Kategorien",
        font=('Arial', 9, 'italic'),
        foreground='blue'
    )
    instructions.pack(pady=(0, 10))
    
    # Fortschrittsinfo
    if is_last_segment:
        progress_label = ttk.Label(
            main_frame,
            text="🏁 LETZTES SEGMENT - Kodierung wird nach diesem Segment abgeschlossen",
            font=('Arial', 10, 'bold'),
            foreground='red'
        )
        progress_label.pack(pady=(0, 10))
    
    return main_frame

def validate_multiple_selection(selected_indices: List[int], 
                              category_map: Dict[int, Dict]) -> Tuple[bool, str, List[Dict]]:
    """
    Validiert eine Mehrfachauswahl von Kategorien
    
    Args:
        selected_indices: Liste der ausgewählten Indizes
        category_map: Mapping von Index zu Kategorie-Info
        
    Returns:
        Tuple[bool, str, List[Dict]]: (Gültig, Fehlermeldung, Kategorien-Liste)
    """
    if not selected_indices:
        return False, "Keine Kategorie ausgewählt", []
    
    if len(selected_indices) == 1:
        # Einzelauswahl - normal verarbeiten
        return True, "", []
    
    # Mehrfachauswahl validieren
    selected_categories = []
    main_categories = set()
    
    for idx in selected_indices:
        if idx not in category_map:
            return False, f"Ungültiger Index: {idx}", []
        
        cat_info = category_map[idx]
        selected_categories.append(cat_info)
        main_categories.add(cat_info['main_category'])
    
    # Prüfe auf gültige Kombinationen
    if len(main_categories) > 3:
        return False, "Maximal 3 verschiedene Hauptkategorien können gleichzeitig ausgewählt werden", []
    
    return True, "", selected_categories


# --- Klasse: ConfigLoader ---
class ConfigLoader:
    def __init__(self, script_dir):
        self.script_dir = script_dir
        self.excel_path = os.path.join(script_dir, "QCA-AID-Codebook.xlsx")
        self.config = {
            'FORSCHUNGSFRAGE': "",
            'KODIERREGELN': {},
            'DEDUKTIVE_KATEGORIEN': {},
            'CONFIG': {}
        }
        
    # ConfigLoader Klasse
    def load_codebook(self):
        print(f"Versuche Konfiguration zu laden von: {self.excel_path}")
        if not os.path.exists(self.excel_path):
            print(f"Excel-Datei nicht gefunden: {self.excel_path}")
            return False

        try:
            # FIX: Erweiterte Optionen für das Öffnen der Excel-Datei, auch wenn sie bereits geöffnet ist
            print("\nÖffne Excel-Datei...")
            wb = load_workbook(
                self.excel_path, 
                read_only=True, 
                data_only=True,
                keep_vba=False  # FIX: Verhindert Probleme mit VBA-Code
            )
            print(f"Excel-Datei erfolgreich geladen. Verfügbare Sheets: {wb.sheetnames}")
            
            # Prüfe DEDUKTIVE_KATEGORIEN Sheet
            if 'DEDUKTIVE_KATEGORIEN' in wb.sheetnames:
                print("\nLese DEDUKTIVE_KATEGORIEN Sheet...")
                sheet = wb['DEDUKTIVE_KATEGORIEN']
            
            # Lade die verschiedenen Komponenten
            self._load_research_question(wb)
            self._load_coding_rules(wb)
            self._load_config(wb)
            
            # Debug-Ausgabe vor dem Laden der Kategorien
            print("\nStarte Laden der deduktiven Kategorien...")
            kategorien = self._load_deduktive_kategorien(wb)
            
            # Prüfe das Ergebnis
            if kategorien:
                print("\nGeladene Kategorien:")
                for name, data in kategorien.items():
                    print(f"\n{name}:")
                    print(f"- Definition: {len(data['definition'])} Zeichen")
                    print(f"- Beispiele: {len(data['examples'])}")
                    print(f"- Regeln: {len(data['rules'])}")
                    print(f"- Subkategorien: {len(data['subcategories'])}")
                
                # Speichere in Config
                self.config['DEDUKTIVE_KATEGORIEN'] = kategorien
                print("\nKategorien erfolgreich in Config gespeichert")
                return True
            else:
                print("\nKeine Kategorien geladen!")
                return False

        except PermissionError:
            # FIX: Spezifische Behandlung für geöffnete Dateien
            print("\n⚠️ Datei ist bereits geöffnet (z.B. in Excel).")
            print("Versuche alternative Lademethode...")
            try:
                # FIX: Versuche mit keep_links=False für bessere Kompatibilität
                wb = load_workbook(
                    self.excel_path, 
                    read_only=True, 
                    data_only=True,
                    keep_vba=False,
                    keep_links=False  # FIX: Entfernt externe Links die Probleme verursachen können
                )
                
                # Lade die verschiedenen Komponenten
                self._load_research_question(wb)
                self._load_coding_rules(wb)
                self._load_config(wb)
                kategorien = self._load_deduktive_kategorien(wb)
                
                if kategorien:
                    self.config['DEDUKTIVE_KATEGORIEN'] = kategorien
                    print("\n✓ Konfiguration trotz geöffneter Datei erfolgreich geladen!")
                    return True
                else:
                    print("\n✗ Keine Kategorien geladen, auch bei alternativer Methode")
                    return False
                    
            except Exception as alt_e:
                print(f"\n✗ Auch alternative Lademethode fehlgeschlagen: {str(alt_e)}")
                print("\nLösung: Schließen Sie die Excel-Datei und versuchen Sie erneut,")
                print("oder deaktivieren Sie das automatische Speichern in OneDrive.")
                return False

        except Exception as e:
            print(f"Fehler beim Lesen der Excel-Datei: {str(e)}")
            print("Details:")
            import traceback
            traceback.print_exc()
            return False
        
    def _load_research_question(self, wb):
        if 'FORSCHUNGSFRAGE' in wb.sheetnames:
            sheet = wb['FORSCHUNGSFRAGE']
            value = sheet['B1'].value
            # print(f"Geladene Forschungsfrage: {value}")  # Debug-Ausgabe
            self.config['FORSCHUNGSFRAGE'] = value


    def _load_coding_rules(self, wb):
        """Lädt Kodierregeln aus dem Excel-Codebook."""
        if 'KODIERREGELN' in wb.sheetnames:
            df = pd.read_excel(self.excel_path, sheet_name='KODIERREGELN', header=0)
            
            # Initialisiere Regelkategorien
            rules = {
                'general': [],       # Allgemeine Kodierregeln
                'format': [],        # Formatregeln
                'exclusion': []      # Neue Kategorie für Ausschlussregeln
            }
            
            # Verarbeite jede Spalte
            for column in df.columns:
                rules_list = df[column].dropna().tolist()
                
                if 'Allgemeine' in column:
                    rules['general'].extend(rules_list)
                elif 'Format' in column:
                    rules['format'].extend(rules_list)
                elif 'Ausschluss' in column:  # Neue Spalte für Ausschlussregeln
                    rules['exclusion'].extend(rules_list)
            
            print("\nKodierregeln geladen:")
            print(f"- Allgemeine Regeln: {len(rules['general'])}")
            print(f"- Formatregeln: {len(rules['format'])}")
            print(f"- Ausschlussregeln: {len(rules['exclusion'])}")
            
            self.config['KODIERREGELN'] = rules

    def _load_deduktive_kategorien(self, wb):
        try:
            if 'DEDUKTIVE_KATEGORIEN' not in wb.sheetnames:
                print("Warnung: Sheet 'DEDUKTIVE_KATEGORIEN' nicht gefunden")
                return {}

            print("\nLade deduktive Kategorien...")
            sheet = wb['DEDUKTIVE_KATEGORIEN']
            
            # Initialisiere Kategorien
            kategorien = {}
            current_category = None
            
            # Hole Header-Zeile
            headers = []
            for cell in sheet[1]:
                headers.append(cell.value)
            # print(f"Gefundene Spalten: {headers}")
            
            # Indizes für Spalten finden
            key_idx = headers.index('Key') if 'Key' in headers else None
            sub_key_idx = headers.index('Sub-Key') if 'Sub-Key' in headers else None
            sub_sub_key_idx = headers.index('Sub-Sub-Key') if 'Sub-Sub-Key' in headers else None
            value_idx = headers.index('Value') if 'Value' in headers else None
            
            if None in [key_idx, sub_key_idx, value_idx]:
                print("Fehler: Erforderliche Spalten fehlen!")
                return {}
                
            # Verarbeite Zeilen
            for row_idx, row in enumerate(sheet.iter_rows(min_row=2), 2):
                try:
                    key = row[key_idx].value
                    sub_key = row[sub_key_idx].value if row[sub_key_idx].value else None
                    sub_sub_key = row[sub_sub_key_idx].value if sub_sub_key_idx is not None and row[sub_sub_key_idx].value else None
                    value = row[value_idx].value if row[value_idx].value else None
                    
                    
                    # Neue Hauptkategorie
                    if key and isinstance(key, str):
                        key = key.strip()
                        if key not in kategorien:
                            print(f"\nNeue Hauptkategorie: {key}")
                            current_category = key
                            kategorien[key] = {
                                'definition': '',
                                'rules': [],
                                'examples': [],
                                'subcategories': {}
                            }
                    
                    # Verarbeite Unterkategorien und Werte
                    if current_category and sub_key:
                        sub_key = sub_key.strip()
                        if isinstance(value, str):
                            value = value.strip()
                            
                        if sub_key == 'definition':
                            kategorien[current_category]['definition'] = value
                            print(f"  Definition hinzugefügt: {len(value)} Zeichen")
                            
                        elif sub_key == 'rules':
                            if value:
                                kategorien[current_category]['rules'].append(value)
                                print(f"  Regel hinzugefügt: {value[:50]}...")
                                
                        elif sub_key == 'examples':
                            if value:
                                kategorien[current_category]['examples'].append(value)
                                print(f"  Beispiel hinzugefügt: {value[:50]}...")
                                
                        elif sub_key == 'subcategories' and sub_sub_key:
                            kategorien[current_category]['subcategories'][sub_sub_key] = value
                            print(f"  Subkategorie hinzugefügt: {sub_sub_key}")
                                
                except Exception as e:
                    print(f"Fehler in Zeile {row_idx}: {str(e)}")
                    continue

            # Validierung der geladenen Daten
            print("\nValidiere geladene Kategorien:")
            # for name, kat in kategorien.items():
            #     print(f"\nKategorie: {name}")
            #     print(f"- Definition: {len(kat['definition'])} Zeichen")
            #     if not kat['definition']:
            #         print("  WARNUNG: Keine Definition!")
            #     print(f"- Regeln: {len(kat['rules'])}")
            #     print(f"- Beispiele: {len(kat['examples'])}")
            #     print(f"- Subkategorien: {len(kat['subcategories'])}")
            #     for sub_name, sub_def in kat['subcategories'].items():
            #         print(f"  • {sub_name}: {sub_def[:50]}...")

            # Ergebnis
            if kategorien:
                print(f"\nErfolgreich {len(kategorien)} Kategorien geladen")
                return kategorien
            else:
                print("\nKeine Kategorien gefunden!")
                return {}

        except Exception as e:
            print(f"Fehler beim Laden der Kategorien: {str(e)}")
            print("Details:")
            import traceback
            traceback.print_exc()
            return {}
        
    def _load_config(self, wb):
        if 'CONFIG' in wb.sheetnames:
            df = pd.read_excel(self.excel_path, sheet_name='CONFIG')
            config = {}
            
            for _, row in df.iterrows():
                key = row['Key']
                sub_key = row['Sub-Key']
                sub_sub_key = row['Sub-Sub-Key']
                value = row['Value']

                if key not in config:
                    config[key] = value if pd.isna(sub_key) else {}

                if not pd.isna(sub_key):
                    if sub_key.startswith('['):  # Für Listen wie CODER_SETTINGS
                        if not isinstance(config[key], list):
                            config[key] = []
                        index = int(sub_key.strip('[]'))
                        while len(config[key]) <= index:
                            config[key].append({})
                        if pd.isna(sub_sub_key):
                            config[key][index] = value
                        else:
                            config[key][index][sub_sub_key] = value
                    else:  # Für verschachtelte Dicts wie ATTRIBUTE_LABELS
                        if not isinstance(config[key], dict):
                            config[key] = {}
                        if pd.isna(sub_sub_key):
                            if key == 'BATCH_SIZE' or sub_key == 'BATCH_SIZE':
                                try:
                                    value = int(value)
                                    print(f"BATCH_SIZE aus Codebook geladen: {value}")
                                except (ValueError, TypeError):
                                    value = 5  # Standardwert
                                    print(f"Warnung: Ungültiger BATCH_SIZE Wert, verwende Standard: {value}")
                            config[key][sub_key] = value
                        else:
                            if sub_key not in config[key]:
                                config[key][sub_key] = {}
                            config[key][sub_key][sub_sub_key] = value

            # Prüfe auf ANALYSIS_MODE in der Konfiguration
            if 'ANALYSIS_MODE' in config:
                valid_modes = {'full', 'abductive', 'deductive'}
                if config['ANALYSIS_MODE'] not in valid_modes:
                    print(f"Warnung: Ungültiger ANALYSIS_MODE '{config['ANALYSIS_MODE']}' im Codebook. Verwende 'deductive'.")
                    config['ANALYSIS_MODE'] = 'deductive'
                else:
                    print(f"ANALYSIS_MODE aus Codebook geladen: {config['ANALYSIS_MODE']}")
            else:
                config['ANALYSIS_MODE'] = 'deductive'  # Standardwert
                print(f"ANALYSIS_MODE nicht im Codebook gefunden, verwende Standard: {config['ANALYSIS_MODE']}")

            # Prüfe auf REVIEW_MODE in der Konfiguration
            if 'REVIEW_MODE' in config:
                valid_modes = {'auto', 'manual', 'consensus', 'majority'}
                if config['REVIEW_MODE'] not in valid_modes:
                    print(f"Warnung: Ungültiger REVIEW_MODE '{config['REVIEW_MODE']}' im Codebook. Verwende 'auto'.")
                    config['REVIEW_MODE'] = 'consensus'
                else:
                    print(f"REVIEW_MODE aus Codebook geladen: {config['REVIEW_MODE']}")
            else:
                config['REVIEW_MODE'] = 'consensus'  # Standardwert
                print(f"REVIEW_MODE nicht im Codebook gefunden, verwende Standard: {config['REVIEW_MODE']}")


            # Stelle sicher, dass ATTRIBUTE_LABELS vorhanden ist
            if 'ATTRIBUTE_LABELS' not in config:
                config['ATTRIBUTE_LABELS'] = {'attribut1': 'Attribut1', 'attribut2': 'Attribut2', 'attribut3': 'Attribut3'}
            elif 'attribut3' not in config['ATTRIBUTE_LABELS']:
                # Füge attribut3 hinzu wenn noch nicht vorhanden
                config['ATTRIBUTE_LABELS']['attribut3'] = 'Attribut3'
            
            # Debug für attribut3
            if config['ATTRIBUTE_LABELS']['attribut3']:
                print(f"Drittes Attribut-Label geladen: {config['ATTRIBUTE_LABELS']['attribut3']}")

            self.config['CONFIG'] = self._sanitize_config(config)
            return True  # Explizite Rückgabe von True
        return False

    def _sanitize_config(self, config):
        """
        Bereinigt und validiert die Konfigurationswerte.
        Überschreibt Standardwerte mit Werten aus dem Codebook.
        
        Args:
            config: Dictionary mit rohen Konfigurationswerten
            
        Returns:
            dict: Bereinigtes Konfigurations-Dictionary
        """
        try:
            sanitized = {}
            script_dir = os.path.dirname(os.path.abspath(__file__))
            
            # Stelle sicher, dass OUTPUT_DIR immer gesetzt wird
            sanitized['OUTPUT_DIR'] = os.path.join(script_dir, 'output')
            os.makedirs(sanitized['OUTPUT_DIR'], exist_ok=True)
            print(f"Standard-Ausgabeverzeichnis gesichert: {sanitized['OUTPUT_DIR']}")
            
            for key, value in config.items():
                # Verzeichnispfade relativ zum aktuellen Arbeitsverzeichnis
                if key in ['DATA_DIR', 'OUTPUT_DIR']:
                    sanitized[key] = os.path.join(script_dir, str(value))
                    # Stelle sicher, dass Verzeichnis existiert
                    os.makedirs(sanitized[key], exist_ok=True)
                    print(f"Verzeichnis gesichert: {sanitized[key]}")
                
                # Numerische Werte für Chunking
                elif key in ['CHUNK_SIZE', 'CHUNK_OVERLAP']:
                    try:
                        # Konvertiere zu Integer und stelle sicher, dass die Werte positiv sind
                        sanitized[key] = max(1, int(value))
                        print(f"Übernehme {key} aus Codebook: {sanitized[key]}")
                    except (ValueError, TypeError):
                        # Wenn Konvertierung fehlschlägt, behalte Standardwert
                        default_value = CONFIG[key]
                        print(f"Warnung: Ungültiger Wert für {key}, verwende Standard: {default_value}")
                        sanitized[key] = default_value
                
                # Coder-Einstellungen mit Typkonvertierung
                elif key == 'CODER_SETTINGS':
                    sanitized[key] = [
                        {
                            'temperature': float(coder['temperature']) 
                                if isinstance(coder.get('temperature'), (int, float, str)) 
                                else 0.3,
                            'coder_id': str(coder.get('coder_id', f'auto_{i}'))
                        }
                        for i, coder in enumerate(value)
                    ]
                
                # Alle anderen Werte unverändert übernehmen
                else:
                    sanitized[key] = value

            # Verarbeitung des CODE_WITH_CONTEXT Parameters
            if 'CODE_WITH_CONTEXT' in config:
                # Konvertiere zu Boolean
                value = config['CODE_WITH_CONTEXT']
                if isinstance(value, str):
                    sanitized['CODE_WITH_CONTEXT'] = value.lower() in ('true', 'ja', 'yes', '1')
                else:
                    sanitized['CODE_WITH_CONTEXT'] = bool(value)
                print(f"Übernehme CODE_WITH_CONTEXT aus Codebook: {sanitized['CODE_WITH_CONTEXT']}")
            else:
                # Standardwert setzen
                sanitized['CODE_WITH_CONTEXT'] = True
                print(f"CODE_WITH_CONTEXT nicht in Codebook gefunden, verwende Standard: {sanitized['CODE_WITH_CONTEXT']}")

            # Verarbeitung des MULTIPLE_CODINGS Parameters
            if 'MULTIPLE_CODINGS' in config:
                # Konvertiere zu Boolean
                value = config['MULTIPLE_CODINGS']
                if isinstance(value, str):
                    sanitized['MULTIPLE_CODINGS'] = value.lower() in ('true', 'ja', 'yes', '1')
                else:
                    sanitized['MULTIPLE_CODINGS'] = bool(value)
                print(f"Übernehme MULTIPLE_CODINGS aus Codebook: {sanitized['MULTIPLE_CODINGS']}")
            else:
                # Standardwert setzen
                sanitized['MULTIPLE_CODINGS'] = True
                print(f"MULTIPLE_CODINGS nicht in Codebook gefunden, verwende Standard: {sanitized['MULTIPLE_CODINGS']}")
            
            # Verarbeitung des MULTIPLE_CODING_THRESHOLD Parameters
            if 'MULTIPLE_CODING_THRESHOLD' in config:
                try:
                    threshold = float(config['MULTIPLE_CODING_THRESHOLD'])
                    if 0.0 <= threshold <= 1.0:
                        sanitized['MULTIPLE_CODING_THRESHOLD'] = threshold
                        print(f"Übernehme MULTIPLE_CODING_THRESHOLD aus Codebook: {sanitized['MULTIPLE_CODING_THRESHOLD']}")
                    else:
                        print(f"Warnung: MULTIPLE_CODING_THRESHOLD muss zwischen 0.0 und 1.0 liegen, verwende Standard: 0.6")
                        sanitized['MULTIPLE_CODING_THRESHOLD'] = 0.6
                except (ValueError, TypeError):
                    print(f"Warnung: Ungültiger MULTIPLE_CODING_THRESHOLD Wert, verwende Standard: 0.6")
                    sanitized['MULTIPLE_CODING_THRESHOLD'] = 0.6
            else:
                sanitized['MULTIPLE_CODING_THRESHOLD'] = 0.6
                print(f"MULTIPLE_CODING_THRESHOLD nicht in Codebook gefunden, verwende Standard: {sanitized['MULTIPLE_CODING_THRESHOLD']}")

            # Verarbeitung des BATCH_SIZE Parameters
            if 'BATCH_SIZE' in config:
                try:
                    batch_size = int(config['BATCH_SIZE'])
                    if batch_size < 1:
                        print("Warnung: BATCH_SIZE muss mindestens 1 sein")
                        batch_size = 5
                    elif batch_size > 20:
                        print("Warnung: BATCH_SIZE > 20 könnte Performance-Probleme verursachen")
                    sanitized['BATCH_SIZE'] = batch_size
                    print(f"Finale BATCH_SIZE: {batch_size}")
                except (ValueError, TypeError):
                    print("Warnung: Ungültiger BATCH_SIZE Wert")
                    sanitized['BATCH_SIZE'] = 5
            else:
                print("BATCH_SIZE nicht in Codebook gefunden, verwende Standard: 5")
                sanitized['BATCH_SIZE'] = 5

            for key, value in config.items():
                if key == 'BATCH_SIZE':
                    continue

            # Stelle sicher, dass CHUNK_OVERLAP kleiner als CHUNK_SIZE ist
            if 'CHUNK_SIZE' in sanitized and 'CHUNK_OVERLAP' in sanitized:
                if sanitized['CHUNK_OVERLAP'] >= sanitized['CHUNK_SIZE']:
                    print(f"Warnung: CHUNK_OVERLAP ({sanitized['CHUNK_OVERLAP']}) muss kleiner sein als CHUNK_SIZE ({sanitized['CHUNK_SIZE']})")
                    sanitized['CHUNK_OVERLAP'] = sanitized['CHUNK_SIZE'] // 10  # 10% als Standardwert
                    
            return sanitized
            
        except Exception as e:
            print(f"Fehler bei der Konfigurationsbereinigung: {str(e)}")
            import traceback
            traceback.print_exc()
            # Verwende im Fehlerfall die Standard-Konfiguration
            return CONFIG
    

    def update_script_globals(self, globals_dict):
        """
        Aktualisiert die globalen Variablen mit den Werten aus der Config.
        """
        try:
            print("\nAktualisiere globale Variablen...")
            
            # Update DEDUKTIVE_KATEGORIEN
            if 'DEDUKTIVE_KATEGORIEN' in self.config:
                deduktive_kat = self.config['DEDUKTIVE_KATEGORIEN']
                if deduktive_kat and isinstance(deduktive_kat, dict):
                    globals_dict['DEDUKTIVE_KATEGORIEN'] = deduktive_kat
                    print(f"\nDEDUKTIVE_KATEGORIEN aktualisiert:")
                    for name in deduktive_kat.keys():
                        print(f"- {name}")
                else:
                    print("Warnung: Keine gültigen DEDUKTIVE_KATEGORIEN in Config")
            
            # WICHTIG: Update das CONFIG Dictionary komplett
            if 'CONFIG' in self.config:
                config_from_codebook = self.config['CONFIG']
                
                # Aktualisiere das globale CONFIG Dictionary
                if 'CONFIG' in globals_dict:
                    globals_dict['CONFIG'].update(config_from_codebook)
                    print(f"\nCONFIG Dictionary aktualisiert mit {len(config_from_codebook)} Einträgen")
                    
                    # Spezielle Debug-Ausgabe für MULTIPLE_CODINGS
                    if 'MULTIPLE_CODINGS' in config_from_codebook:
                        print(f"MULTIPLE_CODINGS aus Codebook: {config_from_codebook['MULTIPLE_CODINGS']}")
                        print(f"MULTIPLE_CODINGS in globalem CONFIG: {globals_dict['CONFIG'].get('MULTIPLE_CODINGS')}")
                    
                    # Spezielle Debug-Ausgabe für andere wichtige Settings
                    important_settings = ['ANALYSIS_MODE', 'REVIEW_MODE', 'CODE_WITH_CONTEXT', 'BATCH_SIZE']
                    for setting in important_settings:
                        if setting in config_from_codebook:
                            print(f"{setting} aus Codebook: {config_from_codebook[setting]}")
                else:
                    print("Warnung: CONFIG nicht im globalen Namespace gefunden")
            
            # Update andere Konfigurationswerte
            for key, value in self.config.items():
                if key not in ['DEDUKTIVE_KATEGORIEN', 'CONFIG'] and key in globals_dict:
                    if isinstance(value, dict) and isinstance(globals_dict[key], dict):
                        globals_dict[key].clear()
                        globals_dict[key].update(value)
                        print(f"Dict {key} aktualisiert")
                    else:
                        globals_dict[key] = value
                        print(f"Variable {key} aktualisiert")

            # Validiere Update
            if 'DEDUKTIVE_KATEGORIEN' in globals_dict:
                kat_count = len(globals_dict['DEDUKTIVE_KATEGORIEN'])
                print(f"\nFinale Validierung: {kat_count} Kategorien in globalem Namespace")
                if kat_count == 0:
                    print("Warnung: Keine Kategorien im globalen Namespace!")
            else:
                print("Fehler: DEDUKTIVE_KATEGORIEN nicht im globalen Namespace!")

            # Finale Validierung der MULTIPLE_CODINGS Einstellung
            if 'CONFIG' in globals_dict and 'MULTIPLE_CODINGS' in globals_dict['CONFIG']:
                final_multiple_codings = globals_dict['CONFIG']['MULTIPLE_CODINGS']
                print(f"\n🔍 FINALE MULTIPLE_CODINGS Einstellung: {final_multiple_codings}")
                if not final_multiple_codings:
                    print("✅ Mehrfachkodierung wurde DEAKTIVIERT")
                else:
                    print("ℹ️ Mehrfachkodierung bleibt AKTIVIERT")

        except Exception as e:
            print(f"Fehler beim Update der globalen Variablen: {str(e)}")
            import traceback
            traceback.print_exc()

    def _load_validation_config(self, wb):
        """Lädt die Validierungskonfiguration aus dem Codebook."""
        if 'CONFIG' not in wb.sheetnames:
            return {}
            
        validation_config = {
            'thresholds': {},
            'english_words': set(),
            'messages': {}
        }
        
        df = pd.read_excel(self.excel_path, sheet_name='CONFIG')
        validation_rows = df[df['Key'] == 'VALIDATION']
        
        for _, row in validation_rows.iterrows():
            sub_key = row['Sub-Key']
            sub_sub_key = row['Sub-Sub-Key']
            value = row['Value']
            
            if sub_key == 'ENGLISH_WORDS':
                if value == 'true':
                    validation_config['english_words'].add(sub_sub_key)
            elif sub_key == 'MESSAGES':
                validation_config['messages'][sub_sub_key] = value
            else:
                # Numerische Schwellenwerte
                try:
                    validation_config['thresholds'][sub_key] = float(value)
                except (ValueError, TypeError):
                    print(f"Warnung: Ungültiger Schwellenwert für {sub_key}: {value}")
        
        return validation_config

    def get_config(self):
        return self.config

    
"""
LLM Provider Module für QCA-AID
Unterstützt OpenAI und Mistral APIs mit einheitlicher Schnittstelle
"""

import os
import asyncio
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Union, Any
from dataclasses import dataclass
from dotenv import load_dotenv

# Lade Umgebungsvariablen
env_path = os.path.join(os.path.expanduser("~"), '.environ.env')
load_dotenv(env_path)

@dataclass
class LLMResponse:
    """Wrapper für LLM-Antworten um einheitliche Schnittstelle zu gewährleisten"""
    content: str
    model: str = ""
    usage: Dict = None
    
    def __init__(self, response):
        """
        Initialisiert LLMResponse basierend auf dem Provider-Response-Format
        
        Args:
            response: Rohe Antwort vom LLM Provider (OpenAI oder Mistral Format)
        """
        if hasattr(response, 'choices') and response.choices:
            # OpenAI Format
            self.content = response.choices[0].message.content
            self.model = getattr(response, 'model', '')
            self.usage = getattr(response, 'usage', None)
        elif hasattr(response, 'content'):
            # Mistral Format
            self.content = response.content
            self.model = getattr(response, 'model', '')
            self.usage = getattr(response, 'usage', None)
        else:
            # Fallback für unbekannte Formate
            self.content = str(response)
            self.model = "unknown"
            self.usage = None

class LLMProvider(ABC):
    """Abstrakte Basisklasse für LLM Provider"""
    
    def __init__(self):
        self.client = None
        self.model_name = None
        self.initialize_client()
    
    @abstractmethod
    def initialize_client(self):
        """Initialisiert den Client für den jeweiligen Provider"""
        pass
    
    @abstractmethod
    async def create_completion(self, 
                              model: str,
                              messages: List[Dict],
                              temperature: float = 0.7,
                              max_tokens: Optional[int] = None,
                              response_format: Optional[Dict] = None) -> Any:
        """Erstellt eine Chat Completion"""
        pass

class OpenAIProvider(LLMProvider):
    """OpenAI Provider Implementation"""
    
    def initialize_client(self):
        """Initialisiert den OpenAI Client"""
        try:
            from openai import AsyncOpenAI
            
            api_key = os.getenv('OPENAI_API_KEY')
            if not api_key:
                raise ValueError("OPENAI_API_KEY nicht in Umgebungsvariablen gefunden")
            
            self.client = AsyncOpenAI(api_key=api_key)
            print("OpenAI Client erfolgreich initialisiert")
            
        except ImportError:
            raise ImportError("OpenAI Bibliothek nicht installiert. Bitte installieren Sie: pip install openai")
        except Exception as e:
            raise Exception(f"Fehler bei OpenAI Client-Initialisierung: {str(e)}")
    
    async def create_completion(self, 
                              model: str,
                              messages: List[Dict],
                              temperature: float = 0.7,
                              max_tokens: Optional[int] = None,
                              response_format: Optional[Dict] = None) -> Any:
        """
        Erstellt eine OpenAI Chat Completion
        
        Args:
            model: Name des zu verwendenden Modells
            messages: Liste der Chat-Nachrichten
            temperature: Temperatur für die Antwortgenerierung
            max_tokens: Maximale Anzahl von Tokens (optional)
            response_format: Format der Antwort (optional, z.B. {"type": "json_object"})
            
        Returns:
            OpenAI ChatCompletion Response
        """
        try:
            # Erstelle Parameter-Dict
            params = {
                'model': model,
                'messages': messages,
                'temperature': temperature
            }
            
            # Füge optionale Parameter hinzu
            if max_tokens:
                params['max_tokens'] = max_tokens
            if response_format:
                params['response_format'] = response_format
            
            # API Call
            response = await self.client.chat.completions.create(**params)
            return response
            
        except Exception as e:
            print(f"Fehler bei OpenAI API Call: {str(e)}")
            raise

class MistralProvider(LLMProvider):
    """Mistral Provider Implementation"""
    
    def initialize_client(self):
        """Initialisiert den Mistral Client"""
        try:
            from mistralai import Mistral
            
            api_key = os.getenv('MISTRAL_API_KEY')
            if not api_key:
                raise ValueError("MISTRAL_API_KEY nicht in Umgebungsvariablen gefunden")
            
            self.client = Mistral(api_key=api_key)
            print("Mistral Client erfolgreich initialisiert")
            
        except ImportError:
            raise ImportError("Mistral Bibliothek nicht installiert. Bitte installieren Sie: pip install mistralai")
        except Exception as e:
            raise Exception(f"Fehler bei Mistral Client-Initialisierung: {str(e)}")
    
    async def create_completion(self, 
                              model: str,
                              messages: List[Dict],
                              temperature: float = 0.7,
                              max_tokens: Optional[int] = None,
                              response_format: Optional[Dict] = None) -> Any:
        """
        Erstellt eine Mistral Chat Completion
        
        Args:
            model: Name des zu verwendenden Modells
            messages: Liste der Chat-Nachrichten
            temperature: Temperatur für die Antwortgenerierung
            max_tokens: Maximale Anzahl von Tokens (optional)
            response_format: Format der Antwort (optional, wird bei Mistral ignoriert)
            
        Returns:
            Mistral ChatCompletion Response
        """
        try:
            # Erstelle Parameter-Dict
            params = {
                'model': model,
                'messages': messages,
                'temperature': temperature
            }
            
            # Füge optionale Parameter hinzu
            if max_tokens:
                params['max_tokens'] = max_tokens
            
            # Hinweis: Mistral unterstützt response_format möglicherweise nicht
            if response_format:
                print("Warnung: response_format wird von Mistral möglicherweise nicht unterstützt")
            
            # API Call (synchron, da Mistral möglicherweise kein async unterstützt)
            response = await self._make_async_call(params)
            return response
            
        except Exception as e:
            print(f"Fehler bei Mistral API Call: {str(e)}")
            raise
    
    async def _make_async_call(self, params):
        """Wrapper um synchrone Mistral Calls asynchron zu machen"""
        try:
            # Führe synchronen Call in Thread Pool aus
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None, 
                lambda: self.client.chat.complete(**params)
            )
            return response
        except Exception as e:
            print(f"Fehler bei async Mistral Call: {str(e)}")
            raise

class LLMProviderFactory:
    """Factory Klasse zur Erstellung von LLM Providern"""
    
    @staticmethod
    def create_provider(provider_name: str) -> LLMProvider:
        """
        Erstellt einen LLM Provider basierend auf dem Namen
        
        Args:
            provider_name: Name des Providers ('openai' oder 'mistral')
            
        Returns:
            LLMProvider: Initialisierter Provider
            
        Raises:
            ValueError: Wenn ein unbekannter Provider angefordert wird
        """
        provider_name = provider_name.lower().strip()
        
        print(f"Initialisiere LLM Provider: {provider_name}")
        
        try:
            if provider_name in ['openai', 'gpt']:
                return OpenAIProvider()
            elif provider_name in ['mistral', 'mistralai']:
                return MistralProvider()
            else:
                raise ValueError(f"Unbekannter LLM Provider: {provider_name}. "
                               f"Unterstützte Provider: 'openai', 'mistral'")
                
        except Exception as e:
            print(f"Fehler bei Provider-Erstellung: {str(e)}")
            raise

"""
    Export Utilities 
"""

import pandas as pd
from openpyxl.styles import Font, PatternFill, Alignment
from typing import Dict, List

def _sanitize_text_for_excel(text):
        """
        FIX: Erweiterte Textbereinigung für Excel-Export
        Entfernt Artefakte aus der Dokumentverarbeitung
        """
        if not text:
            return ""
        
        # Konvertiere zu String falls nicht bereits ein String
        if not isinstance(text, str):
            text = str(text)
        
        import re
        
        # FIX: Entferne Dateipfad-Artefakte
        # Entferne vollständige Dateipfade (file:///, C:/, etc.)
        text = re.sub(r'file:///[^\s\]]+', '', text)
        text = re.sub(r'[A-Za-z]:/[^\s\]]*\.txt', '', text)
        text = re.sub(r'/[^\s\]]*\.txt', '', text)
        
        
        # FIX: Entferne führende/schließende eckige Klammern ohne Inhalt oder mit Metadaten
        # Entferne ] am Textanfang
        text = re.sub(r'^\s*\]', '', text)
        # Entferne [ am Textende 
        text = re.sub(r'\[\s*$', '', text)
                
        # FIX: Entferne leere eckige Klammern und solche mit nur Leerzeichen
        text = re.sub(r'\[\s*\]', '', text)
        
        # FIX: Entferne Chunk-Trennzeichen und Metadaten-Marker
        text = re.sub(r'^\s*---+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*===+\s*$', '', text, flags=re.MULTILINE)
        
        # FIX: Entferne redundante eckige Klammern am Satzanfang
        # Wenn Sätze mit ] beginnen, entferne es
        text = re.sub(r'(?<=\s)\]([A-Z])', r'\1', text)
        text = re.sub(r'^\]([A-Z])', r'\1', text)
        
        # Standard-Bereinigung: Entferne problematische Steuerzeichen
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F\uFFFE\uFFFF]', '', text)
        
        # FIX: Behandle verschiedene Zeilenendetypen
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # FIX: Entferne übermäßig lange Zeilenumbrüche (mehr als 2 aufeinander)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # FIX: Normalisiere Leerzeichen - entferne mehrfache Leerzeichen
        text = re.sub(r' {2,}', ' ', text)
        
        # FIX: Entferne Leerzeichen vor Satzzeichen
        text = re.sub(r' +([.!?,:;])', r'\1', text)
        
        # FIX: Ersetze bekannte problematische Sonderzeichen
        problematic_chars = {
            '☺': ':)', '☻': ':)', '♥': '<3', '♦': 'diamond', '♣': 'club', '♠': 'spade',
            '†': '+', '‡': '++', '•': '*', '‰': 'promille', '™': '(TM)', '©': '(C)',
            '®': '(R)', '§': 'section', '¶': 'paragraph', '±': '+/-'
        }
        
        for char, replacement in problematic_chars.items():
            text = text.replace(char, replacement)
        
        # FIX: Normalisiere Unicode-Zeichen
        import unicodedata
        text = unicodedata.normalize('NFKD', text)
        
        # FIX: Entferne private use area characters
        text = re.sub(r'[\uE000-\uF8FF]', '', text)
        
        # FIX: Begrenze Textlänge für Excel-Zellen (Excel-Limit: 32.767 Zeichen)
        if len(text) > 32760:  # Sicherheitspuffer
            text = text[:32760] + "..."
        
        # FIX: Entferne führende/nachfolgende Leerzeichen und normalisiere
        text = text.strip()
        
        # FIX: Stelle sicher, dass der Text nicht leer ist nach der Bereinigung
        if not text.strip():
            return ""
            
        return text

def _generate_pastel_colors(num_colors):
    """
    Generiert eine Palette mit Pastellfarben.
    
    Args:
        num_colors (int): Anzahl der benötigten Farben
    
    Returns:
        List[str]: Liste von Hex-Farbcodes in Pastelltönen
    """
    import colorsys
    
    pastel_colors = []
    for i in range(num_colors):
        # Wähle Hue gleichmäßig über Farbkreis
        hue = i / num_colors
        # Konvertiere HSV zu RGB mit hoher Helligkeit und Sättigung
        rgb = colorsys.hsv_to_rgb(hue, 0.4, 0.95)
        # Konvertiere RGB zu Hex
        hex_color = 'FF{:02x}{:02x}{:02x}'.format(
            int(rgb[0] * 255), 
            int(rgb[1] * 255), 
            int(rgb[2] * 255)
        )
        pastel_colors.append(hex_color)
    
    return pastel_colors

    
def _format_confidence(confidence: dict) -> str:
    """Formatiert die Konfidenz-Werte für den Export"""
    try:
        if isinstance(confidence, dict):
            formatted_values = []
            # Verarbeite jeden Konfidenzwert einzeln
            for key, value in confidence.items():
                if isinstance(value, (int, float)):
                    formatted_values.append(f"{key}: {value:.2f}")
                elif isinstance(value, dict):
                    # Verarbeite verschachtelte Konfidenzwerte
                    nested_values = [f"{k}: {v:.2f}" for k, v in value.items() 
                                if isinstance(v, (int, float))]
                    if nested_values:
                        formatted_values.append(f"{key}: {', '.join(nested_values)}")
                elif isinstance(value, str):
                    formatted_values.append(f"{key}: {value}")
            
            return "\n".join(formatted_values)
        elif isinstance(confidence, (int, float)):
            return f"{float(confidence):.2f}"
        elif isinstance(confidence, str):
            return confidence
        else:
            return "0.00"
            
    except Exception as e:
        print(f"Fehler bei Konfidenz-Formatierung: {str(e)}")
        return "0.00"


# Abbruch der Analyse managen 
    
import keyboard
import threading
import time
import select
import sys
from typing import Optional, Callable, Any, TYPE_CHECKING

# Verwende TYPE_CHECKING um zirkuläre Imports zu vermeiden
if TYPE_CHECKING:
    from typing import TYPE_CHECKING
    # Hier würde normalerweise der Import stehen, aber wir verwenden Any

class EscapeHandler:
    """
    Handler für Escape-Taste um Kodierung sicher zu unterbrechen und Zwischenergebnisse zu speichern.
    """
    
    def __init__(self, analysis_manager: Any):  # Verwende Any statt 'IntegratedAnalysisManager'
        self.analysis_manager = analysis_manager
        self.escape_pressed = False
        self.user_wants_to_abort = False
        self.keyboard_thread = None
        self.monitoring = False
        self._keyboard_available = self._check_keyboard_availability()
        
    def _check_keyboard_availability(self) -> bool:
        """Prüft ob keyboard-Modul verfügbar ist"""
        try:
            import keyboard
            return True
        except ImportError:
            print("⚠️ 'keyboard' Modul nicht installiert. ESC-Handler nicht verfügbar.")
            print("   Installieren Sie mit: pip install keyboard")
            return False
        
    def start_monitoring(self):
        """Startet die Überwachung der Escape-Taste"""
        if not self._keyboard_available:
            return
            
        if self.monitoring:
            return
            
        self.monitoring = True
        self.escape_pressed = False
        self.user_wants_to_abort = False
        
        print("\n💡 Tipp: Drücken Sie ESC um die Kodierung sicher zu unterbrechen und Zwischenergebnisse zu speichern")
        
        # Starte Keyboard-Monitoring in separatem Thread
        self.keyboard_thread = threading.Thread(target=self._monitor_escape, daemon=True)
        self.keyboard_thread.start()
    
    def stop_monitoring(self):
        """Stoppt die Überwachung der Escape-Taste"""
        self.monitoring = False
        if self.keyboard_thread and self.keyboard_thread.is_alive():
            # Warte kurz auf das Ende des Threads
            self.keyboard_thread.join(timeout=1.0)
    
    def _monitor_escape(self):
        """Überwacht die Escape-Taste in separatem Thread"""
        if not self._keyboard_available:
            return
            
        try:
            import keyboard
            while self.monitoring:
                if keyboard.is_pressed('esc'):
                    if not self.escape_pressed:  # Verhindere mehrfache Verarbeitung
                        self.escape_pressed = True
                        self._handle_escape()
                    break
                time.sleep(0.1)  # Kleine Pause um CPU zu schonen
        except Exception as e:
            print(f"Fehler bei Escape-Überwachung: {str(e)}")
    
    def _handle_escape(self):
        """Behandelt das Drücken der Escape-Taste - VERBESSERT"""
        try:
            print("\n\n" + "="*60)
            print("🛑 ESCAPE-TASTE GEDRÜCKT - KODIERUNG UNTERBRECHEN?")
            print("="*60)
            
            # Zeige aktuellen Status
            current_status = self._get_current_status()
            print(f"\nAktueller Status:")
            print(f"- Verarbeitete Segmente: {current_status['processed_segments']}")
            print(f"- Erstellte Kodierungen: {current_status['total_codings']}")
            print(f"- Verstrichene Zeit: {current_status['elapsed_time']:.1f} Sekunden")
            
            if current_status['total_codings'] > 0:
                print(f"\n✅ {current_status['total_codings']} Kodierungen wurden bereits erstellt")
                print("   Diese können als Zwischenergebnisse exportiert werden.")
            else:
                print("\n⚠️  Noch keine Kodierungen vorhanden")
                print("   Ein Export würde leere Ergebnisse erzeugen.")
            
            print("\n" + "="*60)
            print("OPTIONEN:")
            print("j + ENTER = Kodierung beenden und Zwischenergebnisse exportieren")
            print("n + ENTER = Kodierung fortsetzen") 
            print("ESC       = Sofort beenden ohne Export")
            print("="*60)
            
            # Warte auf Benutzereingabe mit Timeout
            choice = self._get_user_choice_with_timeout(timeout=30)
            
            if choice == 'j':
                print("\n✅ Kodierung wird beendet - Zwischenergebnisse werden exportiert...")
                self.user_wants_to_abort = True
                self._trigger_safe_abort()
                
            elif choice == 'n':
                print("\n▶️  Kodierung wird fortgesetzt...")
                self.escape_pressed = False  # Reset für weitere ESC-Presses
                self.start_monitoring()  # Überwachung wieder starten
                
            elif choice == 'abort_immediately':
                print("\n🛑 Sofortiger Abbruch ohne Export...")
                self.user_wants_to_abort = True
                setattr(self.analysis_manager, '_immediate_abort', True)
                self._trigger_safe_abort()
                
            else:  # Timeout oder ungültige Eingabe
                print("\n⏰ Keine gültige Eingabe - Kodierung wird fortgesetzt...")
                self.escape_pressed = False
                self.start_monitoring()
                
        except Exception as e:
            print(f"Fehler bei Escape-Behandlung: {str(e)}")
            print("Kodierung wird fortgesetzt...")
            self.escape_pressed = False
            self.start_monitoring()
    
    def _get_user_choice_with_timeout(self, timeout: int = 30) -> str:
        """Holt Benutzereingabe mit Timeout - KORRIGIERT"""
        print(f"\nIhre Wahl (Timeout in {timeout} Sekunden): ", end="", flush=True)
        
        try:
            if sys.platform == "win32":
                # Windows-spezifische Implementierung mit verbesserter Eingabe-Behandlung
                try:
                    import msvcrt
                    start_time = time.time()
                    input_chars = []
                    
                    while time.time() - start_time < timeout:
                        if msvcrt.kbhit():
                            char = msvcrt.getch()
                            
                            # Handle verschiedene Eingabetypen
                            if char == b'\x1b':  # ESC
                                print("ESC (Sofortiger Abbruch)")
                                return 'abort_immediately'
                            elif char == b'\r':  # Enter
                                user_input = ''.join(input_chars).lower().strip()
                                print()  # Neue Zeile nach Enter
                                if user_input in ['j', 'n']:
                                    return user_input
                                else:
                                    print("Ungültige Eingabe. Bitte 'j' oder 'n' eingeben: ", end="", flush=True)
                                    input_chars = []
                                    continue
                            elif char == b'\x08':  # Backspace
                                if input_chars:
                                    input_chars.pop()
                                    print('\b \b', end="", flush=True)  # Löscht Zeichen visuell
                            else:
                                try:
                                    decoded_char = char.decode('utf-8')
                                    if decoded_char.isprintable():
                                        input_chars.append(decoded_char)
                                        print(decoded_char, end="", flush=True)
                                except UnicodeDecodeError:
                                    pass  # Ignoriere nicht-dekodierbare Zeichen
                        
                        time.sleep(0.05)  # Kürzere Pause für responsivere Eingabe
                    
                    print()  # Neue Zeile nach Timeout
                    return ''  # Timeout erreicht
                    
                except ImportError:
                    print("\nmsvcrt nicht verfügbar, verwende Fallback...")
                    return self._get_input_fallback(timeout)
                    
            else:
                # Unix/Linux-spezifische Implementierung
                return self._get_input_unix_with_timeout(timeout)
                
        except Exception as e:
            print(f"\nFehler bei Eingabe-Behandlung: {str(e)}")
            return self._get_input_fallback(timeout)

    def _get_input_unix_with_timeout(self, timeout: int) -> str:
        """Unix/Linux-spezifische Eingabe mit Timeout"""
        try:
            import select
            ready, _, _ = select.select([sys.stdin], [], [], timeout)
            if ready:
                choice = sys.stdin.readline().strip().lower()
                if choice in ['j', 'n']:
                    return choice
                elif choice == 'esc' or choice == '\x1b':
                    return 'abort_immediately'
                else:
                    print("Ungültige Eingabe.")
                    return ''
            else:
                print()  # Neue Zeile nach Timeout
                return ''  # Timeout
        except Exception as e:
            print(f"Fehler bei Unix-Eingabe: {str(e)}")
            return self._get_input_fallback(timeout)

    def _get_input_fallback(self, timeout: int) -> str:
        """Verbesserte Fallback Input-Methode"""
        try:
            print(f"\n(Einfache Eingabe-Methode - kein Timeout verfügbar)")
            print("Geben Sie 'j' für Beenden oder 'n' für Fortsetzen ein:")
            
            while True:
                try:
                    choice = input("Ihre Wahl [j/n]: ").strip().lower()
                    if choice in ['j', 'n']:
                        return choice
                    elif choice in ['esc', 'escape', 'abbruch']:
                        return 'abort_immediately'
                    else:
                        print("Ungültige Eingabe. Bitte 'j' oder 'n' eingeben.")
                except KeyboardInterrupt:
                    print("\nCtrl+C erkannt - sofortiger Abbruch")
                    return 'abort_immediately'
                except EOFError:
                    print("\nEingabe-Ende erkannt - Kodierung wird fortgesetzt")
                    return 'n'
                    
        except Exception as e:
            print(f"Fehler bei Fallback-Eingabe: {str(e)}")
            return 'n'  # Standard: Fortsetzen
    
    def _get_current_status(self) -> dict:
        """Holt den aktuellen Status der Analyse"""
        try:
            # Versuche get_progress_report zu verwenden
            if hasattr(self.analysis_manager, 'get_progress_report'):
                progress_report = self.analysis_manager.get_progress_report()
                return progress_report.get('progress', {
                    'processed_segments': 0,
                    'total_codings': 0,
                    'elapsed_time': 0
                })
            
            # Fallback: Direkte Attribut-Zugriffe
            processed_segments = 0
            total_codings = 0
            elapsed_time = 0
            
            if hasattr(self.analysis_manager, 'processed_segments'):
                processed_segments = len(getattr(self.analysis_manager, 'processed_segments', []))
            
            if hasattr(self.analysis_manager, 'coding_results'):
                total_codings = len(getattr(self.analysis_manager, 'coding_results', []))
            
            if hasattr(self.analysis_manager, 'start_time'):
                from datetime import datetime
                start_time = getattr(self.analysis_manager, 'start_time', None)
                if start_time:
                    elapsed_time = (datetime.now() - start_time).total_seconds()
            
            return {
                'processed_segments': processed_segments,
                'total_codings': total_codings,
                'elapsed_time': elapsed_time
            }
            
        except Exception as e:
            print(f"Fehler beim Ermitteln des Status: {str(e)}")
            return {
                'processed_segments': 0,
                'total_codings': 0,
                'elapsed_time': 0
            }
    
    def _trigger_safe_abort(self):
        """Löst einen sicheren Abbruch aus"""
        # Setze Flag für die Hauptschleife
        if hasattr(self.analysis_manager, '_should_abort'):
            self.analysis_manager._should_abort = True
        
        # Alternative: Setze ein neues Attribut
        setattr(self.analysis_manager, '_escape_abort_requested', True)
        
        print("🔄 Abbruch-Signal gesendet...")
    
    def should_abort(self) -> bool:
        """Prüft ob abgebrochen werden soll"""
        return self.user_wants_to_abort


# Hilfsfunktion für die Integration in bestehende Klassen
def add_escape_handler_to_manager(manager_instance) -> EscapeHandler:
    """
    Fügt einen EscapeHandler zu einem bestehenden Analysis Manager hinzu.
    
    Args:
        manager_instance: Instanz des IntegratedAnalysisManager
        
    Returns:
        EscapeHandler: Konfigurierter EscapeHandler
    """
    escape_handler = EscapeHandler(manager_instance)
    
    # Füge Escape-Handler als Attribut hinzu
    setattr(manager_instance, 'escape_handler', escape_handler)
    
    # Füge Abort-Flag hinzu falls nicht vorhanden
    if not hasattr(manager_instance, '_should_abort'):
        setattr(manager_instance, '_should_abort', False)
    
    # Füge Helper-Methode hinzu
    def check_escape_abort(self):
        """Prüft ob durch Escape abgebrochen werden soll"""
        return (getattr(self, '_should_abort', False) or 
                getattr(self, '_escape_abort_requested', False) or
                (hasattr(self, 'escape_handler') and self.escape_handler.should_abort()))
    
    # Binde die Methode an die Instanz
    import types
    manager_instance.check_escape_abort = types.MethodType(check_escape_abort, manager_instance)
    
    return escape_handler


# Decorator für automatische Escape-Handler Integration
def with_escape_handler(cls):
    """
    Decorator um automatisch einen Escape-Handler zu einer Klasse hinzuzufügen.
    
    Usage:
        @with_escape_handler
        class MyAnalysisManager:
            pass
    """
    original_init = cls.__init__
    
    def new_init(self, *args, **kwargs):
        original_init(self, *args, **kwargs)
        self.escape_handler = EscapeHandler(self)
        self._should_abort = False
        
        # Füge check_escape_abort Methode hinzu
        def check_escape_abort():
            return (getattr(self, '_should_abort', False) or 
                    getattr(self, '_escape_abort_requested', False) or
                    self.escape_handler.should_abort())
        
        self.check_escape_abort = check_escape_abort
    
    cls.__init__ = new_init
    return cls

# --- Klasse: DocumentReader ---
# Aufgabe: Laden und Vorbereiten des Analysematerials (Textdokumente, output)

class DocumentReader:
    def __init__(self, data_dir: str):
        self.data_dir = data_dir
        self.supported_formats = {'.docx', '.pdf', '.txt'}
        os.makedirs(data_dir, exist_ok=True)
        
        # Prüfe verfügbare Features
        if not DOCX_AVAILABLE:
            print("⚠️ python-docx nicht verfügbar - DOCX-Unterstützung deaktiviert")
            self.supported_formats.discard('.docx')
        
        if not PDF_AVAILABLE:
            print("⚠️ PyPDF2 nicht verfügbar - PDF-Unterstützung deaktiviert")
            self.supported_formats.discard('.pdf')
        
        print(f"\nDocumentReader initialisiert:")
        print(f"Verzeichnis: {os.path.abspath(data_dir)}")
        print(f"Unterstützte Formate: {', '.join(self.supported_formats)}")


    def clean_problematic_characters(self, text: str) -> str:
        """Verwendet die bereits vorhandene Funktion"""
        return _sanitize_text_for_excel(text)

    async def read_documents(self) -> Dict[str, str]:
        documents = {}
        try:
            all_files = os.listdir(self.data_dir)
            
            print("\nDateianalyse:")
            def is_supported_file(filename: str) -> bool:
                # Exclude backup and temporary files
                if any(ext in filename.lower() for ext in ['.bak', '.bkk', '.tmp', '~']):
                    return False
                # Get the file extension
                extension = os.path.splitext(filename)[1].lower()
                return extension in self.supported_formats

            supported_files = [f for f in all_files if is_supported_file(f)]
            
            print(f"\nGefundene Dateien:")
            for file in all_files:
                status = "✓" if is_supported_file(file) else "✗"
                print(f"{status} {file}")
            
            print(f"\nVerarbeite Dateien:")
            for filename in supported_files:
                try:
                    filepath = os.path.join(self.data_dir, filename)
                    extension = os.path.splitext(filename)[1].lower()
                    
                    print(f"\nLese: {filename}")
                    
                    if extension == '.docx':
                        content = self._read_docx(filepath)
                    elif extension == '.pdf':
                        content = self._read_pdf(filepath)
                    elif extension == '.txt':
                        content = self._read_txt(filepath)
                    else:
                        print(f"⚠ Nicht unterstütztes Format: {extension}")
                        continue
                    
                    if content and content.strip():
                        documents[filename] = content
                        print(f"✓ Erfolgreich eingelesen: {len(content)} Zeichen")
                    else:
                        print(f"⚠ Keine Textinhalte gefunden")
                
                except Exception as e:
                    print(f"✗ Fehler bei {filename}: {str(e)}")
                    print("Details:")
                    import traceback
                    traceback.print_exc()
                    continue

            print(f"\nVerarbeitungsstatistik:")
            print(f"- Dateien im Verzeichnis: {len(all_files)}")
            print(f"- Unterstützte Dateien: {len(supported_files)}")
            print(f"- Erfolgreich eingelesen: {len(documents)}")
            
            return documents

        except Exception as e:
            print(f"Fehler beim Einlesen der Dokumente: {str(e)}")
            import traceback
            traceback.print_exc()
            return documents

    def _read_txt(self, filepath: str) -> str:
        """Liest eine Textdatei ein."""
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

    def _read_docx(self, filepath: str) -> str:
        """
        Liest eine DOCX-Datei ein und extrahiert den Text mit ausführlicher Diagnose.
        Enthält zusätzliche Bereinigung für problematische Zeichen.
        """
        try:
            from docx import Document
            print(f"\nDetailierte Analyse von: {os.path.basename(filepath)}")
            
            # Öffne das Dokument mit zusätzlicher Fehlerbehandlung
            try:
                doc = Document(filepath)
            except Exception as e:
                print(f"  Fehler beim Öffnen der Datei: {str(e)}")
                print("  Versuche alternative Öffnungsmethode...")
                # Manchmal hilft es, die Datei zuerst zu kopieren
                import shutil
                temp_path = filepath + '.temp'
                shutil.copy2(filepath, temp_path)
                doc = Document(temp_path)
                os.remove(temp_path)

            # Sammle Dokumentinformationen
            paragraphs = []
            print("\nDokumentanalyse:")
            print(f"  Gefundene Paragraphen: {len(doc.paragraphs)}")
            
            # Verarbeite jeden Paragraphen einzeln
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    # Hier bereinigen wir Steuerzeichen und problematische Zeichen
                    clean_text = self.clean_problematic_characters(text)
                    paragraphs.append(clean_text)
                else:
                    print(f"  Paragraph {i+1}: Leer")

            # Wenn Paragraphen gefunden wurden
            if paragraphs:
                full_text = '\n'.join(paragraphs)
                print(f"\nErgebnis:")
                print(f"  ✓ {len(paragraphs)} Textparagraphen extrahiert")
                print(f"  ✓ Gesamtlänge: {len(full_text)} Zeichen")
                return full_text
            
            # Wenn keine Paragraphen gefunden wurden, suche in anderen Bereichen
            print("\nSuche nach alternativen Textinhalten:")
            
            # Prüfe Tabellen
            table_texts = []
            for i, table in enumerate(doc.tables):
                print(f"  Prüfe Tabelle {i+1}:")
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Auch hier bereinigen
                            clean_text = self.clean_problematic_characters(cell_text)
                            table_texts.append(clean_text)
                            print(f"    Zelleninhalt gefunden: {len(cell_text)} Zeichen")
            
            if table_texts:
                full_text = '\n'.join(table_texts)
                print(f"\nErgebnis:")
                print(f"  ✓ {len(table_texts)} Tabelleneinträge extrahiert")
                print(f"  ✓ Gesamtlänge: {len(full_text)} Zeichen")
                return full_text
                
            print("\n✗ Keine Textinhalte im Dokument gefunden")
            return ""
                
        except ImportError:
            print("\n✗ python-docx nicht installiert.")
            print("  Bitte installieren Sie das Paket mit:")
            print("  pip install python-docx")
            raise
        except Exception as e:
            print(f"\n✗ Unerwarteter Fehler beim DOCX-Lesen:")
            print(f"  {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    def _read_pdf(self, filepath: str) -> str:
        """
        Liest eine PDF-Datei ein und extrahiert den Text mit verbesserter Fehlerbehandlung.
        
        Args:
            filepath: Pfad zur PDF-Datei
                
        Returns:
            str: Extrahierter und bereinigter Text
        """
        try:
            import PyPDF2
            print(f"\nLese PDF: {os.path.basename(filepath)}")
            
            text_content = []
            with open(filepath, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    print(f"  Gefundene Seiten: {total_pages}")
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                # Bereinige den Text von problematischen Zeichen
                                cleaned_text = self.clean_problematic_characters(page_text)
                                text_content.append(cleaned_text)
                                print(f"  Seite {page_num}/{total_pages}: {len(cleaned_text)} Zeichen extrahiert")
                            else:
                                print(f"  Seite {page_num}/{total_pages}: Kein Text gefunden")
                                
                        except Exception as page_error:
                            print(f"  Fehler bei Seite {page_num}: {str(page_error)}")
                            # Versuche es mit alternativer Methode
                            try:
                                # Fallback: Extrahiere einzelne Textfragmente
                                print("  Versuche alternative Extraktionsmethode...")
                                if hasattr(page, 'extract_text'):
                                    fragments = []
                                    for obj in page.get_text_extraction_elements():
                                        if hasattr(obj, 'get_text'):
                                            fragments.append(obj.get_text())
                                    if fragments:
                                        fallback_text = ' '.join(fragments)
                                        cleaned_text = self.clean_problematic_characters(fallback_text)
                                        text_content.append(cleaned_text)
                                        print(f"  Alternative Methode erfolgreich: {len(cleaned_text)} Zeichen")
                            except:
                                print("  Alternative Methode fehlgeschlagen")
                                continue
                    
                except PyPDF2.errors.PdfReadError as pdf_error:
                    print(f"  PDF Lesefehler: {str(pdf_error)}")
                    print("  Versuche Fallback-Methode...")
                    
                    # Fallback-Methode, wenn direkte Lesemethode fehlschlägt
                    try:
                        from pdf2image import convert_from_path
                        from pytesseract import image_to_string
                        
                        print("  Verwende OCR-Fallback via pdf2image und pytesseract")
                        # Konvertiere PDF-Seiten zu Bildern
                        images = convert_from_path(filepath)
                        
                        for i, image in enumerate(images):
                            try:
                                # Extrahiere Text über OCR
                                ocr_text = image_to_string(image, lang='deu')
                                if ocr_text:
                                    # Bereinige den Text
                                    cleaned_text = self.clean_problematic_characters(ocr_text)
                                    text_content.append(cleaned_text)
                                    print(f"  OCR Seite {i+1}: {len(cleaned_text)} Zeichen extrahiert")
                                else:
                                    print(f"  OCR Seite {i+1}: Kein Text gefunden")
                            except Exception as ocr_error:
                                print(f"  OCR-Fehler bei Seite {i+1}: {str(ocr_error)}")
                                continue
                    
                    except ImportError:
                        print("  OCR-Fallback nicht verfügbar. Bitte installieren Sie pdf2image und pytesseract")
                    
            # Zusammenfassen des extrahierten Textes
            if text_content:
                full_text = '\n'.join(text_content)
                print(f"\nErgebnis:")
                print(f"  ✓ {len(text_content)} Textabschnitte extrahiert")
                print(f"  ✓ Gesamtlänge: {len(full_text)} Zeichen")
                return full_text
            else:
                print("\n✗ Kein Text aus PDF extrahiert")
                return ""
                
        except ImportError:
            print("PyPDF2 nicht installiert. Bitte installieren Sie: pip install PyPDF2")
            raise
        except Exception as e:
            print(f"Fehler beim PDF-Lesen: {str(e)}")
            import traceback
            traceback.print_exc()
            return ""  # Leerer String im Fehlerfall, damit der Rest funktioniert
    
    def _extract_metadata(self, filename: str) -> Tuple[str, str, str]:
        """
        Extrahiert Metadaten aus dem Dateinamen.
        Erwartet Format: attribut1_attribut2_attribut3.extension
        
        Args:
            filename (str): Name der Datei
            
        Returns:
            Tuple[str, str, str]: (attribut1, attribut2, attribut3)
        """
        try:
            name_without_ext = os.path.splitext(filename)[0]
            parts = name_without_ext.split('_')
            
            # Extrahiere bis zu drei Attribute, wenn verfügbar
            attribut1 = parts[0] if len(parts) >= 1 else ""
            attribut2 = parts[1] if len(parts) >= 2 else ""
            attribut3 = parts[2] if len(parts) >= 3 else ""
            
            return attribut1, attribut2, attribut3
        except Exception as e:
            print(f"Fehler beim Extrahieren der Metadaten aus {filename}: {str(e)}")
            return filename, "", ""


class ManualReviewGUI:
    """
    KORRIGIERT: GUI für kategorie-spezifisches manuelles Review
    """
    
    def perform_category_specific_review(self, segments_needing_review: List[Dict]) -> List[Dict]:
        """
        Führt manuelles Review für kategorie-spezifische Segmente durch
        
        Jedes Segment repräsentiert nur eine Hauptkategorie, daher
        sind die Entscheidungen fokussierter und methodisch korrekter
        """
        review_decisions = []
        
        for i, segment in enumerate(segments_needing_review, 1):
            print(f"\n🎯 Review {i}/{len(segments_needing_review)}: {segment['segment_id']}")
            print(f"   Kategorie: {segment['target_category']}")
            
            if segment['is_multiple_coding']:
                instance_info = segment['instance_info']
                print(f"   Teil {instance_info['instance_number']}/{instance_info['total_instances']} von {segment['original_segment_id']}")
                print(f"   Alle Kategorien dieses Segments: {', '.join(instance_info['all_categories'])}")
            
            # Zeige alle Kodierungen für diese spezifische Kategorie
            codings = segment['codings']
            
            print(f"\n   📋 {len(codings)} Kodierungen für Kategorie '{segment['target_category']}':")
            for j, coding in enumerate(codings, 1):
                coder = coding.get('coder_id', 'Unbekannt')
                subcats = coding.get('subcategories', [])
                confidence = coding.get('confidence', {})
                
                print(f"      {j}. {coder}: {segment['target_category']}")
                if subcats:
                    print(f"         Subkategorien: {', '.join(subcats)}")
                if isinstance(confidence, dict):
                    conf_val = confidence.get('total', 0.0)
                    print(f"         Konfidenz: {conf_val:.2f}")
            
            # Lade Textinhalt (von erster Kodierung)
            text_content = codings[0].get('text', 'Kein Text verfügbar')
            
            # Zeige GUI-Dialog für diese eine Kategorie
            decision = self._show_category_review_dialog(segment, text_content)
            
            if decision:
                review_decisions.append(decision)
        
        return review_decisions
    
    def _show_category_review_dialog(self, segment: Dict, text_content: str) -> Optional[Dict]:
        """
        Zeigt Review-Dialog für ein kategorie-spezifisches Segment
        
        Viel fokussierter als vorher, da nur eine Kategorie behandelt wird
        """
        # HIER würde die GUI-Implementierung kommen
        # Für Demo-Zwecke: Automatische Entscheidung basierend auf höchster Konfidenz
        
        codings = segment['codings']
        best_coding = max(codings, key=lambda x: self._extract_confidence_value(x))
        
        decision = best_coding.copy()
        decision.update({
            'segment_id': segment['segment_id'],
            'manual_review': True,
            'review_date': datetime.now().isoformat(),
            'review_justification': f"Automatisch gewählt: Höchste Konfidenz für Kategorie {segment['target_category']}",
            'original_segment_id': segment['original_segment_id'],
            'is_multiple_coding_instance': segment['is_multiple_coding'],
            'instance_info': segment['instance_info']
        })
        
        return decision
    
    def _extract_confidence_value(self, coding: Dict) -> float:
        """Extrahiert Konfidenzwert"""
        confidence = coding.get('confidence', {})
        if isinstance(confidence, dict):
            return confidence.get('total', 0.0)
        elif isinstance(confidence, (int, float)):
            return float(confidence)
        return 0.0


class ManualReviewComponent:
    """
    Komponente für die manuelle Überprüfung und Entscheidung bei Kodierungsunstimmigkeiten.
    Zeigt dem Benutzer Textstellen mit abweichenden Kodierungen und lässt ihn die finale Entscheidung treffen.
    """
    
    def __init__(self, output_dir: str):
        """
        Initialisiert die Manual Review Komponente.
        
        Args:
            output_dir (str): Verzeichnis für Export-Dokumente
        """
        self.output_dir = output_dir
        self.root = None
        self.review_results = []
        self.current_segment = None
        self.current_codings = None
        self.current_index = 0
        self.total_segments = 0
        self._is_processing = False
        
        # Import tkinter innerhalb der Methode, um Abhängigkeiten zu reduzieren
        self.tk = None
        self.ttk = None
        
    async def review_discrepancies(self, segment_codings: dict) -> list:
        """
        Führt einen manuellen Review-Prozess für Segmente mit abweichenden Kodierungen durch.
        
        Args:
            segment_codings: Dictionary mit Segment-ID als Schlüssel und Liste von Kodierungen als Werte
            
        Returns:
            list: Liste der finalen Kodierungsentscheidungen
        """
        try:
            # Importiere tkinter bei Bedarf
            import tkinter as tk
            from tkinter import ttk
            self.tk = tk
            self.ttk = ttk
            
            print("\n=== Manuelle Überprüfung von Kodierungsunstimmigkeiten ===")
            
            # Identifiziere Segmente mit abweichenden Kodierungen
            discrepant_segments = self._identify_discrepancies(segment_codings)
            
            if not discrepant_segments:
                print("Keine Unstimmigkeiten gefunden. Manueller Review nicht erforderlich.")
                return []
                
            self.total_segments = len(discrepant_segments)
            print(f"\nGefunden: {self.total_segments} Segmente mit Kodierungsabweichungen")
            
            # Starte das Tkinter-Fenster für den manuellen Review
            import asyncio
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, self._run_review_gui, discrepant_segments)
            
            print(f"\nManueller Review abgeschlossen: {len(self.review_results)} Entscheidungen getroffen")
            
            return self.review_results
            
        except Exception as e:
            print(f"Fehler beim manuellen Review: {str(e)}")
            import traceback
            traceback.print_exc()
            if self.root:
                try:
                    self.root.quit()
                    self.root.destroy()
                except:
                    pass
            return []
    
    def _identify_discrepancies(self, segment_codings: dict) -> list:
        """
        Identifiziert Segmente, bei denen verschiedene Kodierer zu unterschiedlichen Ergebnissen kommen.
        
        Args:
            segment_codings: Dictionary mit Segment-ID als Schlüssel und Liste von Kodierungen als Werte
            
        Returns:
            list: Liste von Tuples (segment_id, text, codings) mit Unstimmigkeiten
        """
        discrepancies = []
        
        for segment_id, codings in segment_codings.items():
            # Ignoriere Segmente mit nur einer Kodierung
            if len(codings) <= 1:
                continue
                
            # Prüfe auf Unstimmigkeiten in Hauptkategorien
            categories = set(coding.get('category', '') for coding in codings)
            
            # Prüfe auf menschliche Kodierer
            has_human_coder = any('human' in coding.get('coder_id', '') for coding in codings)
            
            # Wenn mehr als eine Kategorie ODER ein menschlicher Kodierer beteiligt
            if len(categories) > 1 or has_human_coder:
                # Hole den Text des Segments
                text = codings[0].get('text', '')
                if not text:
                    # Alternative Textquelle, falls 'text' nicht direkt verfügbar
                    text = codings[0].get('text_references', [''])[0] if codings[0].get('text_references') else ''
                
                discrepancies.append((segment_id, text, codings))
                
        print(f"Unstimmigkeiten identifiziert: {len(discrepancies)}/{len(segment_codings)} Segmente")
        return discrepancies
    
    async def review_discrepancies_direct(self, segment_codings: dict, skip_discrepancy_check: bool = False) -> list:
        """
        FIX: Neue Methode für ManualReviewComponent, die optional die Unstimmigkeits-Prüfung überspringt
        
        Args:
            segment_codings: Dictionary mit Segment-ID als Schlüssel und Liste von Kodierungen als Werte
            skip_discrepancy_check: Wenn True, behandle alle übergebenen Segmente als unstimmig
            
        Returns:
            list: Liste der finalen Kodierungsentscheidungen
        """
        try:
            # Importiere tkinter bei Bedarf
            import tkinter as tk
            from tkinter import ttk
            self.tk = tk
            self.ttk = ttk
            
            print("\n=== Manuelle Überprüfung von Kodierungsunstimmigkeiten ===")
            
            if skip_discrepancy_check:
                # FIX: Überspringe eigene Unstimmigkeits-Prüfung und verwende alle übergebenen Segmente
                print(f"🎯 Verwende alle {len(segment_codings)} übergebenen Segmente für Review (Prüfung übersprungen)")
                
                discrepant_segments = []
                for segment_id, codings in segment_codings.items():
                    if len(codings) > 1:  # Nur Segmente mit mehreren Kodierungen
                        # Hole den Text des Segments
                        text = codings[0].get('text', '')
                        if not text:
                            text = codings[0].get('text_references', [''])[0] if codings[0].get('text_references') else ''
                        
                        discrepant_segments.append((segment_id, text, codings))
                        
                print(f"📋 Direkte Übernahme: {len(discrepant_segments)} Segmente für Review")
            else:
                # Normale Unstimmigkeits-Identifikation
                discrepant_segments = self._identify_discrepancies(segment_codings)
            
            if not discrepant_segments:
                print("Keine Unstimmigkeiten gefunden. Manueller Review nicht erforderlich.")
                return []
                
            self.total_segments = len(discrepant_segments)
            print(f"\nGefunden: {self.total_segments} Segmente mit Kodierungsabweichungen")
            
            # Setze Review-Status zurück
            self.review_results = []
            
            # Starte das Tkinter-Fenster für den manuellen Review
            import asyncio
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, self._run_review_gui, discrepant_segments)
            
            print(f"\nManueller Review abgeschlossen: {len(self.review_results)} Entscheidungen getroffen")
            
            return self.review_results
            
        except Exception as e:
            print(f"Fehler beim manuellen Review: {str(e)}")
            import traceback
            traceback.print_exc()
            if self.root:
                try:
                    self.root.quit()
                    self.root.destroy()
                except:
                    pass
            return []
        
    def _run_review_gui(self, discrepant_segments: list):
        """
        Führt die grafische Benutzeroberfläche für den manuellen Review aus.
        
        Args:
            discrepant_segments: Liste von Segmenten mit Unstimmigkeiten
        """
        if self.root is not None:
            try:
                self.root.quit()
                self.root.destroy()
            except:
                pass
                
        self.root = self.tk.Tk()
        self.root.title("QCA-AID Manueller Review")
        self.root.geometry("1000x700")
        
        # Protokoll für das Schließen des Fensters
        self.root.protocol("WM_DELETE_WINDOW", self._on_closing)
        
        # Hauptframe
        main_frame = self.ttk.Frame(self.root)
        main_frame.pack(padx=10, pady=10, fill=self.tk.BOTH, expand=True)
        
        # Fortschrittsanzeige
        progress_frame = self.ttk.Frame(main_frame)
        progress_frame.pack(padx=5, pady=5, fill=self.tk.X)
        
        self.ttk.Label(progress_frame, text="Fortschritt:").pack(side=self.tk.LEFT, padx=5)
        progress_var = self.tk.StringVar()
        progress_var.set(f"Segment 1/{self.total_segments}")
        progress_label = self.ttk.Label(progress_frame, textvariable=progress_var)
        progress_label.pack(side=self.tk.LEFT, padx=5)
        
        # Text-Frame
        text_frame = self.ttk.LabelFrame(main_frame, text="Textsegment")
        text_frame.pack(padx=5, pady=5, fill=self.tk.BOTH, expand=True)
        
        text_widget = self.tk.Text(text_frame, height=10, wrap=self.tk.WORD)
        text_widget.pack(padx=5, pady=5, fill=self.tk.BOTH, expand=True)
        text_widget.config(state=self.tk.DISABLED)
        
        # Kodierungen-Frame
        codings_frame = self.ttk.LabelFrame(main_frame, text="Konkurrierende Kodierungen")
        codings_frame.pack(padx=5, pady=5, fill=self.tk.BOTH, expand=True)
        
        codings_canvas = self.tk.Canvas(codings_frame)
        scrollbar = self.ttk.Scrollbar(codings_frame, orient=self.tk.VERTICAL, command=codings_canvas.yview)
        
        codings_scrollable = self.ttk.Frame(codings_canvas)
        codings_scrollable.bind(
            "<Configure>",
            lambda e: codings_canvas.configure(
                scrollregion=codings_canvas.bbox("all")
            )
        )
        
        codings_canvas.create_window((0, 0), window=codings_scrollable, anchor="nw")
        codings_canvas.configure(yscrollcommand=scrollbar.set)
        
        codings_canvas.pack(side=self.tk.LEFT, fill=self.tk.BOTH, expand=True)
        scrollbar.pack(side=self.tk.RIGHT, fill=self.tk.Y)
        
        # Button-Frame
        button_frame = self.ttk.Frame(main_frame)
        button_frame.pack(padx=5, pady=10, fill=self.tk.X)
        
        self.ttk.Button(
            button_frame, 
            text="Vorheriges", 
            command=lambda: self._navigate(-1, text_widget, codings_scrollable, discrepant_segments, progress_var)
        ).pack(side=self.tk.LEFT, padx=5)
        
        self.ttk.Button(
            button_frame, 
            text="Nächstes", 
            command=lambda: self._navigate(1, text_widget, codings_scrollable, discrepant_segments, progress_var)
        ).pack(side=self.tk.LEFT, padx=5)
        
        self.ttk.Button(
            button_frame, 
            text="Abbrechen", 
            command=self._on_closing
        ).pack(side=self.tk.RIGHT, padx=5)
        
        # Begründung eingeben
        justification_frame = self.ttk.LabelFrame(main_frame, text="Begründung für Ihre Entscheidung")
        justification_frame.pack(padx=5, pady=5, fill=self.tk.X)
        
        justification_text = self.tk.Text(justification_frame, height=3, wrap=self.tk.WORD)
        justification_text.pack(padx=5, pady=5, fill=self.tk.X)
        
        # Initialisiere mit dem ersten Segment
        if discrepant_segments:
            self.current_index = 0
            self._update_display(text_widget, codings_scrollable, discrepant_segments, justification_text, progress_var)
        
        # Starte MainLoop
        self.root.mainloop()
    
    def _update_display(self, text_widget, codings_frame, discrepant_segments, justification_text, progress_var):
        """
        Aktualisiert die Anzeige für das aktuelle Segment.
        """
        # Aktualisiere Fortschrittsanzeige
        progress_var.set(f"Segment {self.current_index + 1}/{self.total_segments}")
        
        # Hole aktuelles Segment und Kodierungen
        segment_id, text, codings = discrepant_segments[self.current_index]
        self.current_segment = segment_id
        self.current_codings = codings
        
        # Setze Text
        text_widget.config(state=self.tk.NORMAL)
        text_widget.delete(1.0, self.tk.END)
        text_widget.insert(self.tk.END, text)
        text_widget.config(state=self.tk.DISABLED)
        
        # Begründungsfeld leeren
        justification_text.delete(1.0, self.tk.END)
        
        # Lösche alte Kodierungsoptionen
        for widget in codings_frame.winfo_children():
            widget.destroy()
            
        # Anzeige-Variable für die ausgewählte Kodierung
        selection_var = self.tk.StringVar()
        
        # Erstelle Radiobuttons für jede Kodierung
        for i, coding in enumerate(codings):
            coder_id = coding.get('coder_id', 'Unbekannt')
            category = coding.get('category', 'Keine Kategorie')
            subcategories = coding.get('subcategories', [])
            if isinstance(subcategories, tuple):
                subcategories = list(subcategories)
            confidence = 0.0
            
            # Extrahiere Konfidenzwert
            if isinstance(coding.get('confidence'), dict):
                confidence = coding['confidence'].get('total', 0.0)
            elif isinstance(coding.get('confidence'), (int, float)):
                confidence = float(coding['confidence'])
                
            # Formatiere die Subkategorien
            subcats_text = ', '.join(subcategories) if subcategories else 'Keine'
            
            # Erstelle Label-Text
            is_human = 'human' in coder_id
            coder_prefix = "[Mensch]" if is_human else "[Auto]"
            radio_text = f"{coder_prefix} {coder_id}: {category} ({confidence:.2f})\nSubkategorien: {subcats_text}"
            
            # Radiobutton mit Rahmen für bessere Sichtbarkeit
            coding_frame = self.ttk.Frame(codings_frame, relief=self.tk.GROOVE, borderwidth=2)
            coding_frame.pack(padx=5, pady=5, fill=self.tk.X)
            
            radio = self.ttk.Radiobutton(
                coding_frame,
                text=radio_text,
                variable=selection_var,
                value=str(i),
                command=lambda idx=i, j_text=justification_text: self._select_coding(idx, j_text)
            )
            radio.pack(padx=5, pady=5, anchor=self.tk.W)
            
            # Begründung anzeigen wenn vorhanden
            justification = coding.get('justification', '')
            if justification:
                just_label = self.ttk.Label(
                    coding_frame, 
                    text=f"Begründung: {justification[:150]}..." if len(justification) > 150 else f"Begründung: {justification}",
                    wraplength=500
                )
                just_label.pack(padx=5, pady=5, anchor=self.tk.W)
        
        # Eigene Kodierung als Option
        custom_frame = self.ttk.Frame(codings_frame, relief=self.tk.GROOVE, borderwidth=2)
        custom_frame.pack(padx=5, pady=5, fill=self.tk.X)
        
        custom_radio = self.ttk.Radiobutton(
            custom_frame,
            text="Eigene Entscheidung eingeben",
            variable=selection_var,
            value="custom",
            command=lambda: self._create_custom_coding(justification_text)
        )
        custom_radio.pack(padx=5, pady=5, anchor=self.tk.W)
        
        # Standardmäßig menschliche Kodierung auswählen, falls vorhanden
        for i, coding in enumerate(codings):
            if 'human' in coding.get('coder_id', ''):
                selection_var.set(str(i))
                self._select_coding(i, justification_text)
                break
    
    def _select_coding(self, coding_index, justification_text):
        """
        Ausgewählte Kodierung für das aktuelle Segment speichern.
        """
        self.selected_coding_index = coding_index
        
        # Hole die ausgewählte Kodierung
        selected_coding = self.current_codings[coding_index]
        
        # Fülle Begründung mit Vorschlag
        existing_just = selected_coding.get('justification', '')
        if existing_just:
            justification_text.delete(1.0, self.tk.END)
            justification_text.insert(self.tk.END, f"Übernommen von {selected_coding.get('coder_id', 'Kodierer')}: {existing_just}")
    
    def _create_custom_coding(self, justification_text):
        """
        Erstellt ein benutzerdefiniertes Kodierungsfenster.
        """
        custom_window = self.tk.Toplevel(self.root)
        custom_window.title("Eigene Kodierung")
        custom_window.geometry("600x500")
        
        input_frame = self.ttk.Frame(custom_window)
        input_frame.pack(padx=10, pady=10, fill=self.tk.BOTH, expand=True)
        
        # Hauptkategorie
        self.ttk.Label(input_frame, text="Hauptkategorie:").grid(row=0, column=0, padx=5, pady=5, sticky=self.tk.W)
        category_entry = self.ttk.Entry(input_frame, width=30)
        category_entry.grid(row=0, column=1, padx=5, pady=5, sticky=self.tk.W+self.tk.E)
        
        # Subkategorien
        self.ttk.Label(input_frame, text="Subkategorien (mit Komma getrennt):").grid(row=1, column=0, padx=5, pady=5, sticky=self.tk.W)
        subcats_entry = self.ttk.Entry(input_frame, width=30)
        subcats_entry.grid(row=1, column=1, padx=5, pady=5, sticky=self.tk.W+self.tk.E)
        
        # Begründung
        self.ttk.Label(input_frame, text="Begründung:").grid(row=2, column=0, padx=5, pady=5, sticky=self.tk.W)
        just_text = self.tk.Text(input_frame, height=5, width=30)
        just_text.grid(row=2, column=1, padx=5, pady=5, sticky=self.tk.W+self.tk.E)
        
        # Buttons
        button_frame = self.ttk.Frame(input_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)
        
        self.ttk.Button(
            button_frame, 
            text="Übernehmen",
            command=lambda: self._apply_custom_coding(
                category_entry.get(),
                subcats_entry.get(),
                just_text.get(1.0, self.tk.END),
                justification_text,
                custom_window
            )
        ).pack(side=self.tk.LEFT, padx=5)
        
        self.ttk.Button(
            button_frame,
            text="Abbrechen",
            command=custom_window.destroy
        ).pack(side=self.tk.LEFT, padx=5)
    
    def _apply_custom_coding(self, category, subcategories, justification, main_just_text, window):
        """
        Übernimmt die benutzerdefinierte Kodierung.
        """
        # Erstelle eine benutzerdefinierte Kodierung
        self.custom_coding = {
            'category': category,
            'subcategories': [s.strip() for s in subcategories.split(',') if s.strip()],
            'justification': justification.strip(),
            'coder_id': 'human_review',
            'confidence': {'total': 1.0, 'category': 1.0, 'subcategories': 1.0}
        }
        
        # Aktualisiere das Begründungsfeld im Hauptfenster
        main_just_text.delete(1.0, self.tk.END)
        main_just_text.insert(self.tk.END, f"Eigene Entscheidung: {justification.strip()}")
        
        # Schließe das Fenster
        window.destroy()
    
    def _navigate(self, direction, text_widget, codings_frame, discrepant_segments, progress_var):
        """
        Navigation zwischen den Segmenten und Speicherung der Entscheidung.
        """
        if self.current_segment is None or self.current_codings is None:
            return
            
        # Speichere aktuelle Entscheidung
        self._save_current_decision(text_widget)
        
        # Berechne neuen Index
        new_index = self.current_index + direction
        
        # Prüfe Grenzen
        if 0 <= new_index < len(discrepant_segments):
            self.current_index = new_index
            self._update_display(text_widget, codings_frame, discrepant_segments, text_widget, progress_var)
        elif new_index >= len(discrepant_segments):
            # Wenn wir am Ende angelangt sind, frage nach Abschluss
            if self.tk.messagebox.askyesno(
                "Review abschließen", 
                "Das war das letzte Segment. Möchten Sie den Review abschließen?"
            ):
                self.root.quit()
    
    def _save_current_decision(self, justification_text):
        """
        Speichert die aktuelle Entscheidung.
        """
        try:
            if hasattr(self, 'selected_coding_index'):
                # Normale Kodierungsentscheidung
                selected_coding = self.current_codings[self.selected_coding_index].copy()
                
                # Hole Begründung aus Textfeld
                justification = justification_text.get(1.0, self.tk.END).strip()
                
                # Aktualisiere die Kodierung
                selected_coding['segment_id'] = self.current_segment
                selected_coding['review_justification'] = justification
                selected_coding['manual_review'] = True
                selected_coding['review_date'] = datetime.now().isoformat()
                
                self.review_results.append(selected_coding)
                print(f"Entscheidung für Segment {self.current_segment} gespeichert: {selected_coding['category']}")
                
            elif hasattr(self, 'custom_coding'):
                # Benutzerdefinierte Kodierung
                custom = self.custom_coding.copy()
                custom['segment_id'] = self.current_segment
                custom['manual_review'] = True
                custom['review_date'] = datetime.now().isoformat()
                
                self.review_results.append(custom)
                print(f"Eigene Entscheidung für Segment {self.current_segment} gespeichert: {custom['category']}")
        
        except Exception as e:
            print(f"Fehler beim Speichern der Entscheidung: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def _on_closing(self):
        """Sicheres Schließen des Fensters mit vollständiger Ressourcenfreigabe"""
        try:
            if hasattr(self, 'root') and self.root:
                if self.tk.messagebox.askokcancel(
                    "Review beenden", 
                    "Möchten Sie den Review-Prozess wirklich beenden?\nGetroffene Entscheidungen werden gespeichert."
                ):
                    # Speichere aktuelle Entscheidung falls vorhanden
                    if self.current_segment is not None:
                        justification_text = None
                        for widget in self.root.winfo_children():
                            if isinstance(widget, self.tk.Text):
                                justification_text = widget
                                break
                        
                        if justification_text:
                            self._save_current_decision(justification_text)
                    
                    # Alle Tkinter-Variablen explizit löschen
                    for attr_name in dir(self):
                        attr = getattr(self, attr_name)
                        # Prüfen, ob es sich um eine Tkinter-Variable handelt
                        if hasattr(attr, '_tk'):
                            delattr(self, attr_name)
                    
                    # Fenster schließen
                    self.root.quit()
                    self.root.destroy()
                    self.root = None  # Wichtig: Referenz entfernen
        except:
            # Stelle sicher, dass Fenster auch bei Fehlern geschlossen wird
            if hasattr(self, 'root') and self.root:
                try:
                    self.root.quit()
                    self.root.destroy()
                    self.root = None
                except:
                    pass


def validate_category_specific_segments(category_segments: List[Dict]) -> bool:
    """
    Validiert die korrekte Erstellung kategorie-spezifischer Segmente
    """
    print("\n🔍 Validiere kategorie-spezifische Segmente...")
    
    validation_errors = []
    original_segments = defaultdict(list)
    
    for segment in category_segments:
        # Sammle Segmente nach ursprünglicher ID
        original_id = segment['original_segment_id']
        original_segments[original_id].append(segment)
        
        # Validiere Segment-Struktur
        if not segment.get('segment_id'):
            validation_errors.append(f"Segment ohne segment_id: {segment}")
        
        if not segment.get('target_category') and segment['target_category'] is not None:
            validation_errors.append(f"Segment ohne target_category: {segment['segment_id']}")
        
        if not segment.get('codings'):
            validation_errors.append(f"Segment ohne codings: {segment['segment_id']}")
    
    # Validiere Mehrfachkodierungs-Logik
    for original_id, segments in original_segments.items():
        if len(segments) > 1:
            # Mehrfachkodierung - prüfe Konsistenz
            instance_numbers = [s['instance_info']['instance_number'] for s in segments]
            total_instances = segments[0]['instance_info']['total_instances']
            
            # Prüfe aufeinanderfolgende Instanz-Nummern
            expected_numbers = list(range(1, total_instances + 1))
            if sorted(instance_numbers) != expected_numbers:
                validation_errors.append(f"Inkonsistente Instanz-Nummern für {original_id}: {instance_numbers}")
            
            # Prüfe eindeutige Kategorien
            categories = [s['target_category'] for s in segments]
            if len(set(categories)) != len(categories):
                validation_errors.append(f"Doppelte Kategorien für {original_id}: {categories}")
    
    if validation_errors:
        print("❌ Validierungsfehler gefunden:")
        for error in validation_errors[:5]:  # Zeige max 5 Fehler
            print(f"   • {error}")
        return False
    
    print("✅ Validierung erfolgreich")
    return True


def analyze_multiple_coding_impact(original_codings: List[Dict], category_segments: List[Dict]) -> Dict:
    """
    Analysiert den Einfluss der Mehrfachkodierungs-Korrektur
    """
    print("\n📊 Analysiere Mehrfachkodierungs-Einfluss...")
    
    # Zähle ursprüngliche Segmente
    original_segment_ids = set()
    for coding in original_codings:
        segment_id = coding.get('segment_id', '')
        if segment_id:
            original_segment_ids.add(segment_id)
    
    # Analysiere neue Segmente
    multiple_coding_segments = [s for s in category_segments if s['is_multiple_coding']]
    single_coding_segments = [s for s in category_segments if not s['is_multiple_coding']]
    
    # Gruppiere Mehrfachkodierungs-Segmente nach ursprünglicher ID
    multiple_groups = defaultdict(list)
    for segment in multiple_coding_segments:
        original_id = segment['original_segment_id']
        multiple_groups[original_id].append(segment)
    
    analysis = {
        'original_segments': len(original_segment_ids),
        'new_segments_total': len(category_segments),
        'single_coding_segments': len(single_coding_segments),
        'multiple_coding_groups': len(multiple_groups),
        'multiple_coding_segments': len(multiple_coding_segments),
        'expansion_factor': len(category_segments) / len(original_segment_ids) if original_segment_ids else 1,
        'category_distribution': defaultdict(int)
    }
    
    # Analysiere Kategorie-Verteilung
    for segment in category_segments:
        category = segment['target_category']
        if category:
            analysis['category_distribution'][category] += 1
    
    # Detailanalyse der Mehrfachkodierungen
    multiple_details = []
    for original_id, segments in multiple_groups.items():
        categories = [s['target_category'] for s in segments]
        multiple_details.append({
            'original_id': original_id,
            'categories': categories,
            'expansion': len(categories)
        })
    
    analysis['multiple_coding_details'] = multiple_details
    
    # Berichte
    print(f"   • Ursprüngliche Segmente: {analysis['original_segments']}")
    print(f"   • Neue Segmente gesamt: {analysis['new_segments_total']}")
    print(f"   • Einzelkodierung: {analysis['single_coding_segments']}")
    print(f"   • Mehrfachkodierungs-Gruppen: {analysis['multiple_coding_groups']}")
    print(f"   • Mehrfachkodierungs-Segmente: {analysis['multiple_coding_segments']}")
    print(f"   • Expansionsfaktor: {analysis['expansion_factor']:.2f}")
    
    if multiple_details:
        print(f"\n   📋 Top 3 Mehrfachkodierungs-Beispiele:")
        for detail in sorted(multiple_details, key=lambda x: x['expansion'], reverse=True)[:3]:
            print(f"      • {detail['original_id']}: {', '.join(detail['categories'])} ({detail['expansion']} Kategorien)")
    
    return analysis


def export_multiple_coding_report(analysis: Dict, output_dir: str):
    """
    Exportiert detaillierten Bericht über Mehrfachkodierungs-Behandlung
    """
    import json
    from pathlib import Path
    
    report_path = Path(output_dir) / "multiple_coding_analysis_report.json"
    
    report = {
        'timestamp': datetime.now().isoformat(),
        'summary': {
            'original_segments': analysis['original_segments'],
            'new_segments_total': analysis['new_segments_total'],
            'expansion_factor': analysis['expansion_factor'],
            'multiple_coding_groups': analysis['multiple_coding_groups']
        },
        'category_distribution': dict(analysis['category_distribution']),
        'multiple_coding_details': analysis['multiple_coding_details']
    }
    
    with open(report_path, 'w', encoding='utf-8') as f:
        json.dump(report, f, indent=2, ensure_ascii=False)
    
    print(f"📄 Mehrfachkodierungs-Bericht gespeichert: {report_path}")


# --- Klasse: PDFAnnotator ---
# Aufgabe: Annotierung des Original-PDFs mit Kategorie-basierten Highlights
import fitz  # PyMuPDF
import re
import os
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import json
from difflib import SequenceMatcher

class PDFAnnotator:
    """
    FIX: Annotiert Original-PDFs mit farbkodierten Highlights basierend auf QCA-AID Kodierungen
    
    Diese Klasse erstellt das gleiche PDF wie das Original, nur mit farbigen Highlights
    für kodierte Textabschnitte und vollständigen Annotationen mit Metadaten.
    """
    
    def __init__(self, results_exporter: 'ResultsExporter'):
        self.results_exporter = results_exporter
        self.category_colors = {}
        self.fuzzy_match_threshold = 0.85  # Ähnlichkeitsschwelle für Text-Matching
        
        # FIX: Farbpalette für Hauptkategorien (RGB-Werte für PyMuPDF)
        self.default_colors = {
            'Nicht kodiert': (0.8, 0.8, 0.8),  # Grau
            'Kategorie1': (1.0, 0.8, 0.8),     # Hellrot
            'Kategorie2': (0.8, 1.0, 0.8),     # Hellgrün  
            'Kategorie3': (0.8, 0.8, 1.0),     # Hellblau
            'Kategorie4': (1.0, 1.0, 0.8),     # Hellgelb
            'Kategorie5': (1.0, 0.8, 1.0),     # Hellmagenta
            'Kategorie6': (0.8, 1.0, 1.0),     # Hellcyan
        }
    
    def annotate_pdf_with_codings(self, 
                                 pdf_path: str, 
                                 codings: List[Dict], 
                                 chunks: Dict[str, List[str]], 
                                 output_path: str = None) -> str:
        """
        FIX: Hauptmethode zur PDF-Annotation
        
        Args:
            pdf_path: Pfad zum ursprünglichen PDF
            codings: Liste der Kodierungen aus QCA-AID
            chunks: Dictionary mit chunk_id -> text mapping
            output_path: Ausgabepfad (optional)
            
        Returns:
            str: Pfad zur annotierten PDF-Datei
        """
        print(f"\n🎨 Beginne PDF-Annotation: {Path(pdf_path).name}")
        
        # FIX: Öffne Original-PDF
        try:
            doc = fitz.open(pdf_path)
            print(f"   📄 PDF geladen: {len(doc)} Seiten")
        except Exception as e:
            print(f"❌ Fehler beim Öffnen der PDF: {e}")
            return None
        
        # FIX: Bereite Kodierungsdaten vor
        coding_map = self._prepare_coding_map(codings, chunks)
        print(f"   📋 {len(coding_map)} Textabschnitte zu annotieren")
        
        # FIX: Initialisiere Farbschema
        self._initialize_color_scheme(coding_map)
        
        # FIX: Annotiere jede Seite
        total_annotations = 0
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            annotations_on_page = self._annotate_page(page, coding_map, page_num + 1)
            total_annotations += annotations_on_page
            print(f"   📄 Seite {page_num + 1}: {annotations_on_page} Highlights hinzugefügt")
        
        # FIX: Erstelle Legende als erste Seite
        self._add_legend_page(doc)
        
        # FIX: Speichere annotierte PDF
        if not output_path:
            base_name = Path(pdf_path).stem
            output_path = f"{base_name}_annotiert.pdf"
        
        doc.save(output_path)
        doc.close()
        
        print(f"✅ PDF-Annotation abgeschlossen:")
        print(f"   📊 {total_annotations} Highlights erstellt")
        print(f"   💾 Gespeichert als: {output_path}")
        
        return output_path
    
    # FIX: Diese Korrekturen in PDFAnnotator Klasse in QCA_Utils.py vornehmen

    def _prepare_coding_map(self, codings: List[Dict], chunks: Dict[str, List[str]]) -> Dict[str, Dict]:
        """
        FIX: Korrigierte Version mit besserer Dateipfad-Behandlung und None-Kategorie-Handling
        """
        coding_map = {}
        
        print(f"\n   📋 Bereite {len(codings)} Kodierungen vor...")
        print(f"   📁 Verfügbare Dokumente: {len(chunks)}")
        
        for i, coding in enumerate(codings):
            segment_id = coding.get('segment_id', '')
            category = coding.get('category', 'Nicht kodiert')
            
            # FIX: Behandle None-Kategorien
            if category is None:
                category = 'Nicht kodiert'
            
            print(f"\n      🔍 Kodierung {i+1}: {segment_id} → {category}")
            
            try:
                if '_chunk_' not in segment_id:
                    print(f"          ❌ Ungültiges Segment-ID Format (kein '_chunk_')")
                    continue
                
                doc_name = segment_id.split('_chunk_')[0]
                chunk_part = segment_id.split('_chunk_')[1]
                
                if '-' in chunk_part:
                    chunk_id = int(chunk_part.split('-')[0])
                else:
                    chunk_id = int(chunk_part)
                
                print(f"          📂 Parsed: doc='{doc_name}', chunk_id={chunk_id}")
                
                if doc_name not in chunks:
                    print(f"          ❌ Dokument '{doc_name}' nicht in chunks gefunden")
                    continue
                
                doc_chunks = chunks[doc_name]
                
                if chunk_id >= len(doc_chunks):
                    print(f"          ❌ Chunk {chunk_id} nicht vorhanden (nur {len(doc_chunks)} Chunks)")
                    continue
                
                chunk_text = doc_chunks[chunk_id]
                
                if not chunk_text or len(str(chunk_text).strip()) < 10:
                    print(f"          ⚠️ Chunk-Text zu kurz oder leer")
                    continue
                
                text_content = str(chunk_text).strip()
                
                # FIX: Erweiterte Dateipfad-Erkennung und -Bereinigung
                if self._contains_file_path_artifacts(text_content):
                    # Versuche Text zu bereinigen statt komplett zu überspringen
                    cleaned_content = self._remove_file_path_artifacts(text_content)
                    if len(cleaned_content.strip()) >= 50:  # Mindestens 50 Zeichen nach Bereinigung
                        text_content = cleaned_content
                        print(f"          🔧 Dateipfad-Artefakte entfernt, verwende bereinigten Text")
                    else:
                        print(f"          ⚠️ Zu wenig Text nach Dateipfad-Bereinigung, überspringe")
                        continue
                
                clean_text = self._clean_text_for_matching(text_content)
                
                if len(clean_text) < 10:
                    print(f"          ⚠️ Bereinigter Text zu kurz: {len(clean_text)} Zeichen")
                    continue
                
                map_key = f"{segment_id}_{category}"
                
                coding_map[map_key] = {
                    'segment_id': segment_id,
                    'category': category,
                    'subcategories': coding.get('subcategories', []),
                    'justification': coding.get('justification', ''),
                    'confidence': coding.get('confidence', {}),
                    'original_text': text_content,
                    'clean_text': clean_text,
                    'doc_name': doc_name,
                    'chunk_id': chunk_id
                }
                
                print(f"          ✅ Bereit: {len(clean_text)} Zeichen")
                print(f"              Preview: '{clean_text[:80]}...'")
                
            except Exception as e:
                print(f"          ❌ Fehler: {e}")
                continue
        
        print(f"\n   📊 {len(coding_map)} Kodierungen erfolgreich vorbereitet")
        return coding_map

    def _contains_file_path_artifacts(self, text: str) -> bool:
        """
        FIX: Neue Methode zur Erkennung von Dateipfad-Artefakten
        """
        file_path_indicators = [
            'file:///',
            'OneDrive',
            'C:\\Users\\',
            '/Users/',
            '.txt',
            '.pdf',
            '.docx',
            'Projekte/Forschung'
        ]
        
        text_start = text[:200]  # Prüfe nur die ersten 200 Zeichen
        return any(indicator in text_start for indicator in file_path_indicators)

    def _remove_file_path_artifacts(self, text: str) -> str:
        """
        FIX: Neue Methode zur Entfernung von Dateipfad-Artefakten
        """
        import re
        
        # Entferne Zeilen die wie Dateipfade aussehen
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Überspringe Zeilen mit typischen Dateipfad-Mustern
            if (line.startswith('file:///') or 
                'OneDrive' in line or
                re.match(r'^[A-Z]:\\', line) or  # Windows Pfade
                line.startswith('/Users/') or   # Mac Pfade
                line.endswith(('.txt', '.pdf', '.docx')) or
                len(line) < 10):  # Sehr kurze Zeilen
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

    def _initialize_color_scheme(self, coding_map: Dict[str, Dict]) -> None:
        """
        FIX: Erweiterte Farbschema-Initialisierung mit None-Kategorie-Behandlung
        """
        categories = set()
        for coding_info in coding_map.values():
            category = coding_info['category']
            if category and category != 'None':  # FIX: Ignoriere None und leere Kategorien
                categories.add(category)
        
        # FIX: Nutze bestehende Farben vom ResultsExporter falls verfügbar
        if hasattr(self.results_exporter, 'category_colors') and self.results_exporter.category_colors:
            for category, hex_color in self.results_exporter.category_colors.items():
                if category in categories:  # Nur Kategorien die auch verwendet werden
                    self.category_colors[category] = self._hex_to_rgb(hex_color)
        
        # FIX: Ergänze fehlende Kategorien mit Standard-Farben
        available_colors = list(self.default_colors.values())
        color_index = 0
        
        for category in sorted(categories):
            if category not in self.category_colors:
                if color_index < len(available_colors):
                    self.category_colors[category] = available_colors[color_index]
                    color_index += 1
                else:
                    import random
                    self.category_colors[category] = (
                        0.7 + random.random() * 0.3,
                        0.7 + random.random() * 0.3,
                        0.7 + random.random() * 0.3
                    )
        
        print(f"   🎨 Farbschema initialisiert für {len(self.category_colors)} Kategorien")
        for cat, color in self.category_colors.items():
            print(f"      - {cat}: RGB{color}")

    def _group_codings_by_original_text(self, coding_map: Dict[str, Dict]) -> Dict[str, List[Dict]]:
        """
        FIX: Verbesserte Gruppierung mit None-Kategorie-Behandlung
        """
        text_groups = {}
        
        for map_key, coding_info in coding_map.items():
            category = coding_info['category']
            
            # FIX: Überspringe None-Kategorien bei der Gruppierung
            if not category or category == 'None':
                print(f"      ⚠️ Überspringe Kodierung mit leerer/None Kategorie: {map_key}")
                continue
            
            clean_text = coding_info.get('clean_text', coding_info.get('original_text', ''))
            
            if clean_text not in text_groups:
                text_groups[clean_text] = []
            
            text_groups[clean_text].append(coding_info)
        
        print(f"   📋 Text gruppiert in {len(text_groups)} Gruppen")
        return text_groups
    
    def _clean_text_for_matching(self, text: str) -> str:
        """
        FIX: Weniger aggressive Text-Bereinigung für besseres Matching
        """
        if not text:
            return ""
        
        # FIX: Behalte mehr Zeichen für besseres Matching
        clean = re.sub(r'\s+', ' ', text.strip())
        
        # FIX: Entferne nur wirklich problematische Zeichen
        clean = clean.replace('\uf0b7', '•')
        clean = clean.replace('\u2022', '•')
        
        # FIX: Behalte wichtige Satzzeichen und deutsche Zeichen
        clean = re.sub(r'[^\w\s.,!?;:()\-\"\'äöüßÄÖÜ€%]', '', clean)
        
        return clean
    
    def _initialize_color_scheme(self, coding_map: Dict[str, Dict]) -> None:
        """
        FIX: Initialisiert Farbschema basierend auf gefundenen Kategorien
        """
        # FIX: Sammle alle verwendeten Hauptkategorien
        categories = set()
        for coding_info in coding_map.values():
            categories.add(coding_info['category'])
        
        # FIX: Nutze bestehende Farben vom ResultsExporter falls verfügbar
        if hasattr(self.results_exporter, 'category_colors') and self.results_exporter.category_colors:
            for category, hex_color in self.results_exporter.category_colors.items():
                self.category_colors[category] = self._hex_to_rgb(hex_color)
        
        # FIX: Ergänze fehlende Kategorien mit Standard-Farben
        available_colors = list(self.default_colors.values())
        color_index = 0
        
        for category in sorted(categories):
            if category not in self.category_colors:
                if color_index < len(available_colors):
                    self.category_colors[category] = available_colors[color_index]
                    color_index += 1
                else:
                    # FIX: Generiere zufällige Pastellfarbe falls Standard-Farben aufgebraucht
                    import random
                    self.category_colors[category] = (
                        0.7 + random.random() * 0.3,  # Hell-Bereich 0.7-1.0
                        0.7 + random.random() * 0.3,
                        0.7 + random.random() * 0.3
                    )
        
        print(f"   🎨 Farbschema initialisiert für {len(self.category_colors)} Kategorien")
    
    def _hex_to_rgb(self, hex_color: str) -> Tuple[float, float, float]:
        """
        FIX: Konvertiert Hex-Farbe zu RGB (0-1 Bereich für PyMuPDF)
        """
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        
        try:
            r = int(hex_color[0:2], 16) / 255.0
            g = int(hex_color[2:4], 16) / 255.0
            b = int(hex_color[4:6], 16) / 255.0
            return (r, g, b)
        except:
            return (0.8, 0.8, 0.8)  # Fallback: Grau
    
    def _annotate_page(self, page, coding_map: Dict[str, Dict], page_num: int) -> int:
        """
        FIX: Überarbeitete Annotation mit präziser Höhen-Kalibrierung
        """
        annotations_added = 0
        page_text = page.get_text()
        
        print(f"\n   📄 Annotiere Seite {page_num} (Text-Länge: {len(page_text)} Zeichen)")
        
        text_groups = self._group_codings_by_original_text(coding_map)
        print(f"      📋 {len(text_groups)} Text-Gruppen zu verarbeiten")
        
        for i, (original_text, codings_group) in enumerate(text_groups.items(), 1):
            print(f"\n      🔍 Gruppe {i}/{len(text_groups)}: {len(codings_group)} Kodierungen")
            print(f"          Text-Länge: {len(original_text)} Zeichen")
            
            # FIX: Strategie 1 - Exakte Text-Grenzen durch Content-Matching
            y_start, y_end, x_left = self._find_exact_text_boundaries_by_content(page, original_text)
            
            if y_start is None:
                # FIX: Strategie 2 - Fallback mit Text-Position-Schätzung
                print(f"          ⚠️ Exakte Grenzen nicht gefunden, verwende Positions-Schätzung")
                matches = self._find_text_matches(page_text, original_text, original_text)
                
                if matches:
                    match_start, match_end = matches[0]
                    y_start, y_end, x_left = self._estimate_chunk_boundaries_from_text_position(
                        page, match_start, match_end, original_text
                    )
                else:
                    print(f"          ❌ Auch Fallback-Strategie fehlgeschlagen")
                    continue
            
            # FIX: Erstelle präzise, nicht-überlappende Sidebar-Balken
            sidebar_rects = self._create_non_overlapping_sidebar_rectangles(y_start, y_end, x_left, codings_group)
            
            # FIX: Erstelle Annotationen
            for rect_idx, (rect, coding_info) in enumerate(sidebar_rects):
                category = coding_info['category']
                color = self.category_colors.get(category, (0.8, 0.8, 0.8))
                
                # FIX: Sidebar-Balken mit hoher Deckkraft
                sidebar_annot = page.add_rect_annot(rect)
                sidebar_annot.set_colors({"stroke": color, "fill": color})
                sidebar_annot.set_opacity(0.9)
                
                # FIX: Annotation mit Details
                annotation_text = self._create_annotation_text(
                    coding_info, 
                    is_multiple=(len(codings_group) > 1), 
                    instance_nr=rect_idx+1
                )
                sidebar_annot.set_info(content=annotation_text)
                sidebar_annot.update()
                
                annotations_added += 1
                # print(f"            ✅ Balken erstellt: {category} (Höhe: {rect.height:.1f})")
        
        print(f"\n   📊 Seite {page_num}: {annotations_added} präzise Balken hinzugefügt")
        return annotations_added
    
    def _group_codings_by_original_text(self, coding_map: Dict[str, Dict]) -> Dict[str, List[Dict]]:
        """
        FIX: Gruppiert Kodierungen nach ursprünglichem Text für Mehrfachkodierungs-Behandlung
        """
        text_groups = {}
        
        for clean_text, coding_info in coding_map.items():
            # FIX: Extrahiere ursprüngliche Segment-ID ohne Mehrfachkodierungs-Suffix
            segment_id = coding_info['segment_id']
            original_segment_id = self._extract_original_segment_id(segment_id)
            
            # FIX: Gruppe nach ursprünglichem Text
            original_text = coding_info['original_text']
            
            if original_text not in text_groups:
                text_groups[original_text] = []
            
            text_groups[original_text].append(coding_info)
        
        return text_groups
    
    def _extract_original_segment_id(self, segment_id: str) -> str:
        """
        FIX: Extrahiert ursprüngliche Segment-ID (entfernt Mehrfachkodierungs-Suffixe)
        
        Beispiele:
        - "TEDFWI-1-01" → "TEDFWI-1"
        - "doc_chunk_5-02" → "doc_chunk_5"
        """
        if '-' in segment_id:
            parts = segment_id.rsplit('-', 1)
            if len(parts) == 2 and parts[1].isdigit() and len(parts[1]) <= 2:
                return parts[0]
        return segment_id
    
    def _create_multiple_coding_rectangles(self, base_rects: List, codings_group: List[Dict], page) -> List[Tuple]:
        """
        FIX: Erstellt Sidebar-Rechtecke für Mehrfachkodierung (vereinfacht)
        
        Bei mehr als einer Kodierung pro Segment: Sofort Sidebar-Markers verwenden
        """
        return self._create_sidebar_rectangles(base_rects, codings_group, page)
    
    def _create_sidebar_rectangles(self, base_rects: List, codings_group: List[Dict], page) -> List[Tuple]:
        """
        FIX: Korrigierte Sidebar-Rechtecke - schmale Streifen AM RAND der Text-Rechtecke
        """
        result_rects = []
        
        for i, coding_info in enumerate(codings_group):
            for base_rect in base_rects:
                # FIX: Schmale Streifen direkt am linken Rand des Text-Bereichs
                bar_width = 5    # Schmal aber sichtbar
                spacing = 1      # Minimaler Abstand
                
                # FIX: Positionierung INNERHALB des Text-Bereichs am linken Rand
                bar_rect = fitz.Rect(
                    base_rect.x0 + (bar_width + spacing) * i,     # Von links nach rechts innerhalb des Textes
                    base_rect.y0,                                  # Gleiche Höhe wie Text
                    base_rect.x0 + (bar_width + spacing) * i + bar_width,  # Breite des Balkens
                    base_rect.y1                                   # Gleiche Höhe wie Text
                )
                
                result_rects.append((bar_rect, coding_info))
                print(f"            📏 Sidebar {i+1}: {bar_rect} für {coding_info['category']}")
        
        return result_rects
    
    def _create_multiple_coding_highlight(self, page, rect, color, instance_nr: int, total_instances: int):
        """
        FIX: Erstellt Rechteck-Annotation für Sidebar-Mehrfachkodierung (vereinfacht)
        """
        # FIX: Verwende immer Rechteck-Annotation für Sidebar-Markers
        highlight = page.add_rect_annot(rect)
        
        # FIX: Setze Farbe mit hoher Deckkraft für bessere Sichtbarkeit
        highlight.set_colors({"stroke": color, "fill": color})
        highlight.set_opacity(0.8)  # 80% Deckkraft
        
        return highlight
    
    def _find_exact_text_boundaries_by_content(self, page, target_text: str) -> Tuple[Optional[float], Optional[float], Optional[float]]:
        """
        FIX: Findet exakte Text-Grenzen durch direkten Text-Abgleich
        
        Returns:
            Tuple[y_start, y_end, x_left] oder (None, None, None)
        """
        words = page.get_text("words")  # [(x0, y0, x1, y1, "word", block_no, line_no, word_no)]
        page_text = page.get_text()
        
        # FIX: Bereite Target-Text vor
        target_clean = self._clean_text_for_matching(target_text)
        target_words = [w.strip() for w in target_clean.split() if len(w.strip()) > 2]
        
        if len(target_words) < 5:
            print(f"          ❌ Zu wenige charakteristische Wörter: {len(target_words)}")
            return None, None, None
        
        print(f"          🔍 Suche {len(target_words)} Ziel-Wörter im PDF...")
        
        # FIX: Finde zusammenhängenden Textbereich
        # Strategie: Suche nach einer Sequenz von mindestens 5 aufeinanderfolgenden Wörtern
        best_match_start = None
        best_match_end = None
        best_match_score = 0
        
        # FIX: Sliding Window über PDF-Wörter
        for start_idx in range(len(words) - 5):
            # Nimm 10-Wort-Fenster für Vergleich
            window_size = min(10, len(words) - start_idx)
            pdf_window_words = []
            
            for i in range(start_idx, start_idx + window_size):
                pdf_word = words[i][4].lower().strip()
                if len(pdf_word) > 2:
                    pdf_window_words.append(pdf_word)
            
            if len(pdf_window_words) < 5:
                continue
            
            # FIX: Berechne Übereinstimmung mit Target-Wörtern
            matches = 0
            for target_word in target_words[:10]:  # Erste 10 Target-Wörter
                if any(target_word.lower() in pdf_word or pdf_word in target_word.lower() 
                       for pdf_word in pdf_window_words):
                    matches += 1
            
            match_score = matches / min(len(target_words), 10)
            
            # FIX: Wenn gute Übereinstimmung gefunden
            if match_score > 0.6 and match_score > best_match_score:
                best_match_score = match_score
                best_match_start = start_idx
                # Schätze Ende basierend auf Target-Länge
                estimated_length = min(len(target_words), 50)  # Maximal 50 Wörter pro Chunk
                best_match_end = min(start_idx + estimated_length, len(words) - 1)
                
                print(f"          ✅ Match gefunden: Score={match_score:.2f}, Wörter {start_idx}-{best_match_end}")
        
        if best_match_start is None:
            print(f"          ❌ Kein ausreichender Match gefunden (bester Score: {best_match_score:.2f})")
            return None, None, None
        
        # FIX: Bestimme exakte Koordinaten des gefundenen Bereichs
        start_word = words[best_match_start]
        end_word = words[best_match_end]
        
        y_start = start_word[1]  # y0 des ersten Wortes
        y_end = end_word[3]      # y1 des letzten Wortes
        x_left = min(start_word[0], end_word[0])  # Linkeste x-Koordinate
        
        # FIX: Erweitere um ein wenig Puffer für bessere Sichtbarkeit
        line_height = abs(y_end - y_start) / max(1, (best_match_end - best_match_start) // 8)  # Geschätzte Zeilenhöhe
        y_start = y_start - line_height * 0.1  # 10% Puffer oben
        y_end = y_end + line_height * 0.1      # 10% Puffer unten
        
        # print(f"          📐 Exakte Grenzen: Y={y_start:.1f}-{y_end:.1f} (Höhe: {y_end-y_start:.1f}), X={x_left:.1f}")
        
        # FIX: Validiere und korrigiere Koordinaten
        if y_start > y_end:
            print(f"          🔧 Korrigiere vertauschte Y-Koordinaten: {y_start:.1f} ↔ {y_end:.1f}")
            y_start, y_end = y_end, y_start
        
        # FIX: Mindesthöhe sicherstellen
        min_height = 20  # Mindestens 20 Pixel Höhe
        if (y_end - y_start) < min_height:
            print(f"          🔧 Erweitere zu geringe Höhe von {y_end - y_start:.1f} auf {min_height}")
            center_y = (y_start + y_end) / 2
            y_start = center_y - min_height / 2
            y_end = center_y + min_height / 2
        
        # FIX: Prüfe auf gültige Werte
        if any(coord is None or not (-10000 < coord < 10000) for coord in [y_start, y_end, x_left]):
            print(f"          ❌ Ungültige Koordinaten erkannt, verwende Fallback")
            return None, None, None
        
        # print(f"          ✅ Validierte Grenzen: Y={y_start:.1f}-{y_end:.1f} (Höhe: {y_end-y_start:.1f}), X={x_left:.1f}")
        
        return y_start, y_end, x_left
    
    def _find_exact_chunk_boundaries(self, page, target_text: str) -> Tuple[Optional[float], Optional[float], Optional[float], Optional[float]]:
        """
        FIX: Findet exakte Start- und End-Koordinaten eines Text-Chunks
        
        Returns:
            Tuple[y_start, y_end, x_left, x_right] oder (None, None, None, None) falls nicht gefunden
        """
        words = page.get_text("words")
        page_text = page.get_text()
        
        # FIX: Bereite Suchtext vor - erste und letzte charakteristische Wörter
        target_words = target_text.split()
        if len(target_words) < 3:
            return None, None, None, None
        
        # FIX: Charakteristische Wörter am Anfang und Ende
        start_keywords = []
        end_keywords = []
        
        # Sammle erste 5 längere Wörter (> 4 Zeichen) 
        for word in target_words[:10]:
            if len(word) > 4 and word.isalpha():
                start_keywords.append(word.lower())
                if len(start_keywords) >= 3:
                    break
        
        # Sammle letzte 5 längere Wörter (> 4 Zeichen)
        for word in reversed(target_words[-10:]):
            if len(word) > 4 and word.isalpha():
                end_keywords.append(word.lower())
                if len(end_keywords) >= 3:
                    break
        
        if not start_keywords or not end_keywords:
            return None, None, None, None
        
        print(f"          🎯 Suche Start-Wörter: {start_keywords}")
        print(f"          🎯 Suche End-Wörter: {end_keywords}")
        
        # FIX: Finde Start- und End-Positionen
        start_positions = []
        end_positions = []
        
        for word_info in words:
            word_text = word_info[4].lower()
            word_rect = fitz.Rect(word_info[:4])
            
            # Suche Start-Wörter
            if any(keyword in word_text or word_text in keyword for keyword in start_keywords):
                start_positions.append(word_rect)
                print(f"          ✅ Start-Wort gefunden: '{word_text}' an Y={word_rect.y0:.1f}")
            
            # Suche End-Wörter  
            if any(keyword in word_text or word_text in keyword for keyword in end_keywords):
                end_positions.append(word_rect)
                print(f"          ✅ End-Wort gefunden: '{word_text}' an Y={word_rect.y0:.1f}")
        
        if not start_positions or not end_positions:
            print(f"          ❌ Start oder End nicht gefunden (Start: {len(start_positions)}, End: {len(end_positions)})")
            return None, None, None, None
        
        # FIX: Bestimme Chunk-Grenzen
        # Nimm früheste Start-Position und späteste End-Position
        y_start = min(rect.y0 for rect in start_positions)
        y_end = max(rect.y1 for rect in end_positions)
        x_left = min(rect.x0 for rect in start_positions + end_positions)
        x_right = max(rect.x1 for rect in start_positions + end_positions)
        
        print(f"          📐 Chunk-Grenzen: Y={y_start:.1f}-{y_end:.1f}, X={x_left:.1f}-{x_right:.1f}")
        
        return y_start, y_end, x_left, x_right

    def _create_non_overlapping_sidebar_rectangles(self, y_start: float, y_end: float, x_left: float, codings_group: List[Dict]) -> List[Tuple]:
        """
        FIX: Erstellt nicht-überlappende Sidebar-Rechtecke
        """
        result_rects = []
        
        # FIX: Dynamische Balken-Größe basierend auf verfügbarer Höhe
        available_height = y_end - y_start
        num_codings = len(codings_group)
        
        # FIX: Berechne optimale Balken-Höhe
        if num_codings == 1:
            # Ein Balken: Nutze 80% der verfügbaren Höhe
            bar_height = available_height * 0.8
            bar_spacing = available_height * 0.1
        else:
            # Mehrere Balken: Teile Höhe gleichmäßig auf
            total_spacing = available_height * 0.2  # 20% für Abstände
            available_for_bars = available_height - total_spacing
            bar_height = available_for_bars / num_codings
            bar_spacing = total_spacing / (num_codings + 1)
        
        # FIX: Balken-Breite und Position
        bar_width = 12      # Sichtbare aber nicht störende Breite
        margin_from_text = 15  # Abstand vom Text
        
        # print(f"          📊 Balken-Layout: {num_codings} Balken, Höhe={bar_height:.1f}, Abstand={bar_spacing:.1f}")
        
        for i, coding_info in enumerate(codings_group):
            # FIX: Vertikale Position - von oben nach unten gestapelt
            bar_y_start = y_start + bar_spacing + i * (bar_height + bar_spacing)
            bar_y_end = bar_y_start + bar_height
            
            # FIX: Horizontale Position - links vom Text
            bar_x_start = x_left - margin_from_text - bar_width
            bar_x_end = x_left - margin_from_text
            
            # FIX: Stelle sicher, dass der Balken im sichtbaren Bereich ist
            if bar_x_start < 10:  # Mindestens 10px vom Seitenrand
                bar_x_start = 10
                bar_x_end = bar_x_start + bar_width
            
            bar_rect = fitz.Rect(bar_x_start, bar_y_start, bar_x_end, bar_y_end)
            result_rects.append((bar_rect, coding_info))
            
            # print(f"            📏 Balken {i+1}: {bar_rect} → {coding_info['category']}")
        
        return result_rects
    
    def _create_precise_sidebar_rectangles(self, y_start: float, y_end: float, x_left: float, codings_group: List[Dict]) -> List[Tuple]:
        """
        FIX: Erstellt präzise Sidebar-Rechtecke nur für den Chunk-Bereich
        """
        result_rects = []
        
        # FIX: Balken-Dimensionen
        bar_width = 8      # Sichtbar aber nicht störend
        spacing = 2        # Kleiner Abstand zwischen Balken
        margin = 10        # Abstand vom Text
        
        for i, coding_info in enumerate(codings_group):
            # FIX: Positionierung links vom Text-Bereich
            bar_x_start = x_left - margin - (bar_width + spacing) * (len(codings_group) - i)
            bar_x_end = bar_x_start + bar_width
            
            # FIX: Präzises Rechteck nur für diesen Chunk
            bar_rect = fitz.Rect(
                bar_x_start,    # Links vom Text
                y_start,        # Exakt Start des Chunks
                bar_x_end,      # Balken-Breite
                y_end           # Exakt Ende des Chunks
            )
            
            result_rects.append((bar_rect, coding_info))
            print(f"            📏 Präziser Sidebar {i+1}: {bar_rect} für {coding_info['category']}")
        
        return result_rects

    def _create_fallback_sidebar_rectangles(self, page, match_start: int, match_end: int, codings_group: List[Dict]) -> List[Tuple]:
        """
        FIX: Fallback für Sidebar-Rechtecke wenn präzise Erkennung fehlschlägt
        """
        result_rects = []
        page_rect = page.rect
        page_text = page.get_text()
        
        # FIX: Geschätzte Position basierend auf Text-Position
        text_ratio = match_start / max(1, len(page_text))
        
        # FIX: Realistische Schätzung für Chunk-Größe
        estimated_lines = min(10, max(3, (match_end - match_start) // 80))  # Etwa 80 Zeichen pro Zeile
        line_height = 12  # Typische Zeilenhöhe
        
        y_start = page_rect.y0 + 70 + (text_ratio * (page_rect.height - 140))
        y_end = y_start + (estimated_lines * line_height)
        x_left = page_rect.x0 + 50  # Typischer linker Textrand
        
        print(f"          📍 Fallback Chunk-Schätzung: Y={y_start:.1f}-{y_end:.1f}, Zeilen={estimated_lines}")
        
        # FIX: Erstelle Sidebar-Rechtecke
        bar_width = 8
        spacing = 2
        margin = 10
        
        for i, coding_info in enumerate(codings_group):
            bar_x_start = x_left - margin - (bar_width + spacing) * (len(codings_group) - i)
            bar_x_end = bar_x_start + bar_width
            
            bar_rect = fitz.Rect(bar_x_start, y_start, bar_x_end, y_end)
            result_rects.append((bar_rect, coding_info))
            print(f"            📏 Fallback Sidebar {i+1}: {bar_rect}")
        
        return result_rects
    
    
    def _estimate_chunk_boundaries_from_text_position(self, page, match_start: int, match_end: int, target_text: str) -> Tuple[float, float, float]:
        """
        FIX: Fallback-Methode mit besserer Schätzung
        """
        page_text = page.get_text()
        page_rect = page.rect
        
        # FIX: Genauere Positionsschätzung
        text_length = len(page_text)
        target_length = len(target_text)
        
        # Schätze Zeilen basierend auf Textlänge (etwa 80-100 Zeichen pro Zeile)
        estimated_lines = max(2, min(15, target_length // 85))
        line_height = 14  # Typische Zeilenhöhe in PDFs
        
        # Position im Dokument
        position_ratio = match_start / max(1, text_length)
        
        # FIX: Realistischer Y-Bereich
        text_area_top = page_rect.y0 + 60    # Typischer oberer Rand
        text_area_bottom = page_rect.y1 - 60  # Typischer unterer Rand
        text_area_height = text_area_bottom - text_area_top
        
        y_start = text_area_top + (position_ratio * text_area_height)
        y_end = y_start + (estimated_lines * line_height)
        
        # Stelle sicher, dass wir im Seitenbereich bleiben
        if y_end > text_area_bottom:
            y_end = text_area_bottom
            y_start = y_end - (estimated_lines * line_height)
        
        x_left = page_rect.x0 + 50  # Typischer linker Textrand
        
        print(f"          📍 Fallback-Schätzung: Y={y_start:.1f}-{y_end:.1f} (Zeilen: {estimated_lines}), X={x_left:.1f}")
        
        return y_start, y_end, x_left
    
    def _find_text_matches(self, page_text: str, clean_text: str, original_text: str) -> List[Tuple[int, int]]:
        """
        FIX: Vereinfachte Text-Suche nur für Fallback-Zwecke
        """
        matches = []
        
        # FIX: Nur einfache Strategien für Fallback
        # Strategie 1: Erste 100 Zeichen
        if len(clean_text) > 100:
            short_text = clean_text[:100]
            pos = page_text.find(short_text)
            if pos != -1:
                matches.append((pos, pos + len(clean_text)))
                print(f"      ✅ Fallback-Match gefunden an Position {pos}")
                return matches
        
        # Strategie 2: Erste 5 Wörter
        words = clean_text.split()[:5]
        if len(words) >= 3:
            search_phrase = ' '.join(words)
            pos = page_text.find(search_phrase)
            if pos != -1:
                estimated_end = pos + len(clean_text)
                matches.append((pos, estimated_end))
                print(f"      ✅ Wort-Fallback-Match gefunden an Position {pos}")
        
        return matches
    
    def _fuzzy_text_search(self, page_text: str, search_text: str, threshold: float = 0.7) -> List[Tuple[int, int]]:
        """
        FIX: Verbesserte Fuzzy-Suche mit anpassbarer Schwelle
        """
        from difflib import SequenceMatcher
        
        matches = []
        search_len = len(search_text)
        
        if search_len < 20:
            return matches
        
        # FIX: Suche in überlappenden Fenstern
        window_size = min(search_len, 200)  # Kleinere Fenster für besseres Matching
        step_size = window_size // 4        # Überlappung für bessere Abdeckung
        
        for i in range(0, len(page_text) - window_size + 1, step_size):
            candidate = page_text[i:i + window_size]
            similarity = SequenceMatcher(None, search_text[:window_size], candidate).ratio()
            
            if similarity >= threshold:
                matches.append((i, i + window_size))
                print(f"      ✅ Fuzzy-Match (Ähnlichkeit: {similarity:.2f}) an Position {i}")
                break  # Nehme nur den ersten guten Match
        
        return matches

    def _get_text_rectangles(self, page, match_start: int, match_end: int) -> List:
        """
        FIX: Korrigierte Text-Rechteck-Erkennung - nur der eigentliche Text-Bereich
        """
        rects = []
        
        try:
            # FIX: Nutze get_text("words") für präzise Koordinaten
            words = page.get_text("words")
            page_text = page.get_text()
            
            if match_start < 0 or match_end > len(page_text):
                return self._create_fallback_rect(page, match_start, match_end)
            
            target_text = page_text[match_start:match_end]
            
            # FIX: Erste und letzte Wörter des Targets finden
            target_words = target_text.split()
            if len(target_words) < 2:
                return self._create_fallback_rect(page, match_start, match_end)
            
            first_words = target_words[:3]  # Erste 3 Wörter
            last_words = target_words[-3:]  # Letzte 3 Wörter
            
            print(f"          🎯 Suche Bereich: '{' '.join(first_words)}' bis '{' '.join(last_words)}'")
            
            # FIX: Finde Start- und End-Koordinaten
            start_rects = []
            end_rects = []
            
            for word_info in words:
                word_text = word_info[4].lower()
                word_rect = fitz.Rect(word_info[:4])
                
                # Suche nach ersten Wörtern
                if any(first_word.lower() in word_text or word_text in first_word.lower() 
                       for first_word in first_words):
                    start_rects.append(word_rect)
                
                # Suche nach letzten Wörtern
                if any(last_word.lower() in word_text or word_text in last_word.lower() 
                       for last_word in last_words):
                    end_rects.append(word_rect)
            
            if start_rects and end_rects:
                # FIX: Bestimme Text-Bereich von erstem bis letztem gefundenen Wort
                min_x0 = min(rect.x0 for rect in start_rects + end_rects)
                min_y0 = min(rect.y0 for rect in start_rects + end_rects)
                max_x1 = max(rect.x1 for rect in start_rects + end_rects)
                max_y1 = max(rect.y1 for rect in start_rects + end_rects)
                
                # FIX: Begrenzter Text-Bereich (nicht die ganze Seite)
                text_rect = fitz.Rect(
                    min_x0,      # Linke Kante des Textes
                    min_y0,      # Obere Kante des Textes  
                    max_x1,      # Rechte Kante des Textes
                    max_y1       # Untere Kante des Textes
                )
                
                rects.append(text_rect)
                print(f"          ✅ Präziser Text-Bereich: {text_rect}")
            else:
                return self._create_fallback_rect(page, match_start, match_end)
            
        except Exception as e:
            print(f"          ❌ Fehler bei präziser Koordinaten-Suche: {e}")
            return self._create_fallback_rect(page, match_start, match_end)
        
        return rects
    
    def _create_fallback_rect(self, page, match_start: int, match_end: int) -> List:
        """
        FIX: Realistisches Fallback-Rechteck basierend auf Seitenlayout
        """
        page_rect = page.rect
        
        # FIX: Typisches PDF-Layout - Text nimmt mittleren Bereich ein
        text_margin_left = 50
        text_margin_right = 50  
        text_margin_top = 70
        
        # FIX: Geschätzte Position basierend auf Text-Position
        page_text = page.get_text()
        text_ratio = match_start / max(1, len(page_text))
        
        # FIX: Y-Position basierend auf Position im Text
        available_height = page_rect.height - 2 * text_margin_top
        y_position = page_rect.y0 + text_margin_top + (text_ratio * available_height)
        
        # FIX: Realistisches Text-Rechteck
        fallback_rect = fitz.Rect(
            page_rect.x0 + text_margin_left,           # Realistischer linker Rand
            y_position,                                # Geschätzte Y-Position
            page_rect.x1 - text_margin_right,          # Realistischer rechter Rand
            y_position + 40                            # Moderater Höhe (etwa 2-3 Textzeilen)
        )
        
        print(f"          📍 Fallback Text-Rechteck: {fallback_rect}")
        return [fallback_rect]
    
    def _create_annotation_text(self, coding_info: Dict, is_multiple: bool = False, instance_nr: int = 1) -> str:
        """
        FIX: Erstellt Annotations-Text mit allen Kodierungs-Informationen
        """
        lines = []
        
        # FIX: Mehrfachkodierungs-Header
        if is_multiple:
            lines.append(f"🔄 MEHRFACHKODIERUNG - Teil {instance_nr}")
            lines.append("")
        
        # FIX: Hauptkategorie
        lines.append(f"📋 Kategorie: {coding_info['category']}")
        
        # FIX: Subkategorien
        if coding_info['subcategories']:
            subcats = ', '.join(coding_info['subcategories'])
            lines.append(f"🔖 Subkategorien: {subcats}")
        
        # FIX: Konfidenz
        confidence = coding_info.get('confidence', {})
        if isinstance(confidence, dict) and 'total' in confidence:
            lines.append(f"📊 Konfidenz: {confidence['total']:.2f}")
        
        # FIX: Begründung (gekürzt)
        justification = coding_info.get('justification', '')
        if justification:
            short_justification = justification[:200] + "..." if len(justification) > 200 else justification
            lines.append(f"💭 Begründung: {short_justification}")
        
        # FIX: Segment-ID
        lines.append(f"🔢 Segment: {coding_info['segment_id']}")
        
        return '\n'.join(lines)
    
    def _add_legend_page(self, doc) -> None:
        """
        FIX: Fügt Legende als erste Seite hinzu
        """
        # FIX: Erstelle neue Seite am Anfang
        legend_page = doc.new_page(0, width=595, height=842)  # A4 Format
        
        # FIX: Titel
        title_rect = fitz.Rect(50, 50, 545, 80)
        legend_page.insert_text(title_rect.tl, "QCA-AID Kategorien-Legende", 
                               fontsize=20, color=(0, 0, 0))
        
        # FIX: Legende für jede Kategorie
        y_pos = 120
        for category, color in self.category_colors.items():
            # FIX: Farbfeld
            color_rect = fitz.Rect(50, y_pos, 80, y_pos + 20)
            legend_page.draw_rect(color_rect, color=color, fill=color)
            
            # FIX: Kategorie-Name
            text_rect = fitz.Rect(90, y_pos, 500, y_pos + 20)
            legend_page.insert_text(text_rect.tl, category, 
                                   fontsize=12, color=(0, 0, 0))
            
            y_pos += 30
        
        # FIX: Anweisungen
        instructions = [
            "",
            "Anweisungen:",
            "• Highlights zeigen kodierte Textabschnitte",
            "• Klicken Sie auf Highlights für Details",
            "• Annotationen enthalten Kategorie, Subkategorien und Begründung",
            "• Bei Mehrfachkodierung: Farbbalken links neben dem Text",
            "",
            f"Erstellt mit QCA-AID • {len(self.category_colors)} Kategorien identifiziert"
        ]
        
        y_pos += 30
        for instruction in instructions:
            legend_page.insert_text((50, y_pos), instruction, 
                                   fontsize=10, color=(0.3, 0.3, 0.3))
            y_pos += 15
        
        # print("   📖 Legende als erste Seite hinzugefügt")


class DocumentToPDFConverter:
    """
    FIX: Konvertiert TXT und DOCX Dateien zu PDF für einheitliche Annotation
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.temp_pdf_dir = os.path.join(output_dir, "temp_pdfs")
        os.makedirs(self.temp_pdf_dir, exist_ok=True)
        
        # FIX: Prüfe verfügbare Bibliotheken
        self.reportlab_available = self._check_reportlab()
        self.python_docx_available = self._check_python_docx()
    
    def _check_reportlab(self) -> bool:
        """Prüft ob ReportLab verfügbar ist"""
        try:
            import reportlab
            return True
        except ImportError:
            return False
    
    def _check_python_docx(self) -> bool:
        """Prüft ob python-docx verfügbar ist"""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def convert_txt_to_pdf(self, txt_path: str, output_path: str = None) -> str:
        """
        FIX: Konvertiert TXT-Datei zu PDF
        
        Args:
            txt_path: Pfad zur TXT-Datei
            output_path: Ausgabe-Pfad (optional)
            
        Returns:
            str: Pfad zur erstellten PDF
        """
        if not self.reportlab_available:
            print("   ❌ ReportLab nicht verfügbar für TXT→PDF Konvertierung")
            print("   💡 Installieren Sie mit: pip install reportlab")
            return None
        
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT
        
        if not output_path:
            base_name = os.path.splitext(os.path.basename(txt_path))[0]
            output_path = os.path.join(self.temp_pdf_dir, f"{base_name}.pdf")
        
        print(f"   📄 Konvertiere TXT zu PDF: {os.path.basename(txt_path)}")
        
        try:
            # FIX: Lese TXT-Datei
            with open(txt_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # FIX: Erstelle PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # FIX: Styles definieren
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                alignment=TA_LEFT,
                spaceAfter=6,
            )
            
            # FIX: Text in Paragraphen aufteilen
            story = []
            paragraphs = text_content.split('\n\n')  # Doppelte Zeilenumbrüche als Absätze
            
            for para_text in paragraphs:
                if para_text.strip():
                    # FIX: Bereinige Text für ReportLab
                    clean_text = para_text.replace('\n', ' ').strip()
                    # Escape spezielle Zeichen
                    clean_text = clean_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    para = Paragraph(clean_text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 6))
            
            # FIX: PDF erstellen
            doc.build(story)
            
            print(f"   ✅ TXT→PDF erfolgreich: {os.path.basename(output_path)}")
            return output_path
            
        except Exception as e:
            print(f"   ❌ Fehler bei TXT→PDF Konvertierung: {e}")
            return None
    
    def convert_docx_to_pdf(self, docx_path: str, output_path: str = None) -> str:
        """
        FIX: Konvertiert DOCX-Datei zu PDF
        
        Args:
            docx_path: Pfad zur DOCX-Datei
            output_path: Ausgabe-Pfad (optional)
            
        Returns:
            str: Pfad zur erstellten PDF
        """
        if not self.reportlab_available or not self.python_docx_available:
            print("   ❌ Benötigte Bibliotheken nicht verfügbar für DOCX→PDF")
            print("   💡 Installieren Sie mit: pip install reportlab python-docx")
            return None
        
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from docx import Document
        
        if not output_path:
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            output_path = os.path.join(self.temp_pdf_dir, f"{base_name}.pdf")
        
        print(f"   📄 Konvertiere DOCX zu PDF: {os.path.basename(docx_path)}")
        
        try:
            # FIX: Lese DOCX-Datei
            doc_docx = Document(docx_path)
            
            # FIX: Erstelle PDF
            doc_pdf = SimpleDocTemplate(output_path, pagesize=A4,
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)
            
            # FIX: Styles definieren
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                alignment=TA_LEFT,
                spaceAfter=6,
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                leading=16,
                alignment=TA_LEFT,
                spaceAfter=12,
                spaceBefore=12,
            )
            
            # FIX: DOCX-Inhalte extrahieren
            story = []
            
            for paragraph in doc_docx.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # FIX: Bereinige Text
                clean_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # FIX: Einfache Stil-Erkennung (könnte erweitert werden)
                if len(text) < 100 and (text.isupper() or text.startswith('#')):
                    # Vermutlich Überschrift
                    para = Paragraph(clean_text, heading_style)
                else:
                    # Normaler Text
                    para = Paragraph(clean_text, normal_style)
                
                story.append(para)
                story.append(Spacer(1, 6))
            
            # FIX: PDF erstellen
            doc_pdf.build(story)
            
            print(f"   ✅ DOCX→PDF erfolgreich: {os.path.basename(output_path)}")
            return output_path
            
        except Exception as e:
            print(f"   ❌ Fehler bei DOCX→PDF Konvertierung: {e}")
            return None
    
    def convert_document_to_pdf(self, file_path: str) -> str:
        """
        FIX: Universelle Konvertierung basierend auf Dateierweiterung
        
        Args:
            file_path: Pfad zur Eingabedatei
            
        Returns:
            str: Pfad zur PDF (Original-PDF oder konvertierte PDF)
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # FIX: Bereits PDF - einfach zurückgeben
            return file_path
        elif file_ext == '.txt':
            return self.convert_txt_to_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.convert_docx_to_pdf(file_path)
        else:
            print(f"   ⚠️ Nicht unterstütztes Format: {file_ext}")
            return None
    
    def cleanup_temp_pdfs(self):
        """
        FIX: Bereinigt temporäre PDF-Dateien
        """
        try:
            import shutil
            if os.path.exists(self.temp_pdf_dir):
                shutil.rmtree(self.temp_pdf_dir)
                print(f"   🧹 Temporäre PDFs bereinigt")
        except Exception as e:
            print(f"   ⚠️ Fehler bei Bereinigung: {e}")


# FIX: Erweiterte export_annotated_pdfs Methode in ResultsExporter Klasse

    def export_annotated_pdfs_all_formats(self, 
                                         codings: List[Dict], 
                                         chunks: Dict[str, List[str]], 
                                         data_dir: str) -> List[str]:
        """
        FIX: Erweiterte PDF-Annotation für alle Dateiformate (TXT, DOCX, PDF)
        
        Args:
            codings: Liste der finalen Kodierungen
            chunks: Dictionary mit chunk_id -> text mapping
            data_dir: Input-Verzeichnis mit Original-Dateien
            
        Returns:
            List[str]: Liste der Pfade zu erstellten annotierten PDFs
        """
        print(f"\n🎨 Beginne erweiterte PDF-Annotations-Export für alle Formate...")
        
        try:
            # FIX: Importiere PDF-Annotator
            from QCA_Utils import PDFAnnotator, DocumentToPDFConverter
        except ImportError:
            print("   ❌ PyMuPDF nicht verfügbar - PDF-Annotation übersprungen")
            return []
        
        # FIX: Initialisiere Konverter und Annotator
        pdf_converter = DocumentToPDFConverter(self.output_dir)
        pdf_annotator = PDFAnnotator(self)
        
        # FIX: Finde alle unterstützten Dateien
        supported_extensions = ['.pdf', '.txt', '.docx', '.doc']
        input_files = []
        
        try:
            for file in os.listdir(data_dir):
                file_path = os.path.join(data_dir, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                if os.path.isfile(file_path) and file_ext in supported_extensions:
                    input_files.append((file, file_path, file_ext))
                    
        except Exception as e:
            print(f"   ❌ Fehler beim Durchsuchen des Verzeichnisses: {e}")
            return []
        
        if not input_files:
            print("   ℹ️ Keine unterstützten Dateien im Input-Verzeichnis gefunden")
            return []
        
        print(f"   📁 {len(input_files)} Dateien gefunden:")
        for filename, _, ext in input_files:
            print(f"      • {filename} ({ext})")
        
        annotated_files = []
        
        # FIX: Verarbeite jede Datei
        for filename, file_path, file_ext in input_files:
            print(f"\n   📄 Verarbeite: {filename}")
            
            # FIX: Filtere Review-Kodierungen für diese Datei
            file_stem = os.path.splitext(filename)[0]
            file_codings = []
            
            for coding in codings:
                is_review_coding = (
                    coding.get('consensus_info') is not None or
                    coding.get('review_decision') is not None or
                    coding.get('selection_type') in ['consensus', 'majority', 'manual_priority'] or
                    len([c for c in codings if c.get('segment_id') == coding.get('segment_id')]) == 1
                )
                
                matches_file = (
                    file_stem in coding.get('document', '') or 
                    file_stem in coding.get('segment_id', '')
                )
                
                if is_review_coding and matches_file:
                    file_codings.append(coding)
            
            if not file_codings:
                print(f"      ⚠️ Keine Review-Kodierungen für {filename} gefunden")
                continue
            
            print(f"      📋 {len(file_codings)} Review-Kodierungen gefunden")
            
            # FIX: Konvertiere zu PDF falls nötig
            if file_ext == '.pdf':
                pdf_path = file_path
                print(f"      ✅ Bereits PDF")
            else:
                print(f"      🔄 Konvertiere {file_ext.upper()} zu PDF...")
                pdf_path = pdf_converter.convert_document_to_pdf(file_path)
                
                if not pdf_path:
                    print(f"      ❌ Konvertierung fehlgeschlagen")
                    continue
                
                print(f"      ✅ PDF erstellt: {os.path.basename(pdf_path)}")
            
            # FIX: Annotiere PDF
            try:
                output_filename = f"{file_stem}_QCA_annotiert.pdf"
                output_file = os.path.join(self.output_dir, output_filename)
                
                result_path = pdf_annotator.annotate_pdf_with_codings(
                    pdf_path,
                    file_codings,
                    chunks,
                    output_file
                )
                
                if result_path:
                    annotated_files.append(result_path)
                    print(f"      ✅ Annotiert: {os.path.basename(result_path)}")
                else:
                    print(f"      ❌ Annotation fehlgeschlagen")
                
            except Exception as e:
                print(f"      ❌ Fehler bei Annotation: {e}")
                continue
        
        # FIX: Bereinige temporäre Dateien
        pdf_converter.cleanup_temp_pdfs()
        
        print(f"\n✅ Erweiterte PDF-Annotation abgeschlossen: {len(annotated_files)} Dateien erstellt")
        return annotated_files
    
class DocumentToPDFConverter:
    """
    FIX: Konvertiert TXT und DOCX Dateien zu PDF für einheitliche Annotation
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.temp_pdf_dir = os.path.join(output_dir, "temp_pdfs")
        os.makedirs(self.temp_pdf_dir, exist_ok=True)
        
        # FIX: Prüfe verfügbare Bibliotheken
        self.reportlab_available = self._check_reportlab()
        self.python_docx_available = self._check_python_docx()
    
    def _check_reportlab(self) -> bool:
        """Prüft ob ReportLab verfügbar ist"""
        try:
            import reportlab
            return True
        except ImportError:
            return False
    
    def _check_python_docx(self) -> bool:
        """Prüft ob python-docx verfügbar ist"""
        try:
            import docx
            return True
        except ImportError:
            return False
    
    def convert_txt_to_pdf(self, txt_path: str, output_path: str = None) -> str:
        """
        FIX: Konvertiert TXT-Datei zu PDF
        
        Args:
            txt_path: Pfad zur TXT-Datei
            output_path: Ausgabe-Pfad (optional)
            
        Returns:
            str: Pfad zur erstellten PDF
        """
        if not self.reportlab_available:
            print("   ❌ ReportLab nicht verfügbar für TXT→PDF Konvertierung")
            print("   💡 Installieren Sie mit: pip install reportlab")
            return None
        
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT
        
        if not output_path:
            base_name = os.path.splitext(os.path.basename(txt_path))[0]
            output_path = os.path.join(self.temp_pdf_dir, f"{base_name}.pdf")
        
        print(f"   📄 Konvertiere TXT zu PDF: {os.path.basename(txt_path)}")
        
        try:
            # FIX: Lese TXT-Datei
            with open(txt_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # FIX: Erstelle PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # FIX: Styles definieren
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                alignment=TA_LEFT,
                spaceAfter=6,
            )
            
            # FIX: Text in Paragraphen aufteilen
            story = []
            paragraphs = text_content.split('\n\n')  # Doppelte Zeilenumbrüche als Absätze
            
            for para_text in paragraphs:
                if para_text.strip():
                    # FIX: Bereinige Text für ReportLab
                    clean_text = para_text.replace('\n', ' ').strip()
                    # Escape spezielle Zeichen
                    clean_text = clean_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    para = Paragraph(clean_text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 6))
            
            # FIX: PDF erstellen
            doc.build(story)
            
            print(f"   ✅ TXT→PDF erfolgreich: {os.path.basename(output_path)}")
            return output_path
            
        except Exception as e:
            print(f"   ❌ Fehler bei TXT→PDF Konvertierung: {e}")
            return None
    
    def convert_docx_to_pdf(self, docx_path: str, output_path: str = None) -> str:
        """
        FIX: Konvertiert DOCX-Datei zu PDF
        
        Args:
            docx_path: Pfad zur DOCX-Datei
            output_path: Ausgabe-Pfad (optional)
            
        Returns:
            str: Pfad zur erstellten PDF
        """
        if not self.reportlab_available or not self.python_docx_available:
            print("   ❌ Benötigte Bibliotheken nicht verfügbar für DOCX→PDF")
            print("   💡 Installieren Sie mit: pip install reportlab python-docx")
            return None
        
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from docx import Document
        
        if not output_path:
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            output_path = os.path.join(self.temp_pdf_dir, f"{base_name}.pdf")
        
        print(f"   📄 Konvertiere DOCX zu PDF: {os.path.basename(docx_path)}")
        
        try:
            # FIX: Lese DOCX-Datei
            doc_docx = Document(docx_path)
            
            # FIX: Erstelle PDF
            doc_pdf = SimpleDocTemplate(output_path, pagesize=A4,
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)
            
            # FIX: Styles definieren
            styles = getSampleStyleSheet()
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=12,
                alignment=TA_LEFT,
                spaceAfter=6,
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                leading=16,
                alignment=TA_LEFT,
                spaceAfter=12,
                spaceBefore=12,
            )
            
            # FIX: DOCX-Inhalte extrahieren
            story = []
            
            for paragraph in doc_docx.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # FIX: Bereinige Text
                clean_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # FIX: Einfache Stil-Erkennung (könnte erweitert werden)
                if len(text) < 100 and (text.isupper() or text.startswith('#')):
                    # Vermutlich Überschrift
                    para = Paragraph(clean_text, heading_style)
                else:
                    # Normaler Text
                    para = Paragraph(clean_text, normal_style)
                
                story.append(para)
                story.append(Spacer(1, 6))
            
            # FIX: PDF erstellen
            doc_pdf.build(story)
            
            print(f"   ✅ DOCX→PDF erfolgreich: {os.path.basename(output_path)}")
            return output_path
            
        except Exception as e:
            print(f"   ❌ Fehler bei DOCX→PDF Konvertierung: {e}")
            return None
    
    def convert_document_to_pdf(self, file_path: str) -> str:
        """
        FIX: Universelle Konvertierung basierend auf Dateierweiterung
        
        Args:
            file_path: Pfad zur Eingabedatei
            
        Returns:
            str: Pfad zur PDF (Original-PDF oder konvertierte PDF)
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # FIX: Bereits PDF - einfach zurückgeben
            return file_path
        elif file_ext == '.txt':
            return self.convert_txt_to_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.convert_docx_to_pdf(file_path)
        else:
            print(f"   ⚠️ Nicht unterstütztes Format: {file_ext}")
            return None
    
    def cleanup_temp_pdfs(self):
        """
        FIX: Bereinigt temporäre PDF-Dateien
        """
        try:
            import shutil
            if os.path.exists(self.temp_pdf_dir):
                shutil.rmtree(self.temp_pdf_dir)
                print(f"   🧹 Temporäre PDFs bereinigt")
        except Exception as e:
            print(f"   ⚠️ Fehler bei Bereinigung: {e}")


# FIX: Erweiterte export_annotated_pdfs Methode in ResultsExporter Klasse

    def export_annotated_pdfs_all_formats(self, 
                                         codings: List[Dict], 
                                         chunks: Dict[str, List[str]], 
                                         data_dir: str) -> List[str]:
        """
        FIX: Erweiterte PDF-Annotation für alle Dateiformate (TXT, DOCX, PDF)
        
        Args:
            codings: Liste der finalen Kodierungen
            chunks: Dictionary mit chunk_id -> text mapping
            data_dir: Input-Verzeichnis mit Original-Dateien
            
        Returns:
            List[str]: Liste der Pfade zu erstellten annotierten PDFs
        """
        print(f"\n🎨 Beginne erweiterte PDF-Annotations-Export für alle Formate...")
        
        try:
            # FIX: Importiere PDF-Annotator
            from QCA_Utils import PDFAnnotator, DocumentToPDFConverter
        except ImportError:
            print("   ❌ PyMuPDF nicht verfügbar - PDF-Annotation übersprungen")
            return []
        
        # FIX: Initialisiere Konverter und Annotator
        pdf_converter = DocumentToPDFConverter(self.output_dir)
        pdf_annotator = PDFAnnotator(self)
        
        # FIX: Finde alle unterstützten Dateien
        supported_extensions = ['.pdf', '.txt', '.docx', '.doc']
        input_files = []
        
        try:
            for file in os.listdir(data_dir):
                file_path = os.path.join(data_dir, file)
                file_ext = os.path.splitext(file)[1].lower()
                
                if os.path.isfile(file_path) and file_ext in supported_extensions:
                    input_files.append((file, file_path, file_ext))
                    
        except Exception as e:
            print(f"   ❌ Fehler beim Durchsuchen des Verzeichnisses: {e}")
            return []
        
        if not input_files:
            print("   ℹ️ Keine unterstützten Dateien im Input-Verzeichnis gefunden")
            return []
        
        print(f"   📁 {len(input_files)} Dateien gefunden:")
        for filename, _, ext in input_files:
            print(f"      • {filename} ({ext})")
        
        annotated_files = []
        
        # FIX: Verarbeite jede Datei
        for filename, file_path, file_ext in input_files:
            print(f"\n   📄 Verarbeite: {filename}")
            
            # FIX: Filtere Review-Kodierungen für diese Datei
            file_stem = os.path.splitext(filename)[0]
            file_codings = []
            
            for coding in codings:
                is_review_coding = (
                    coding.get('consensus_info') is not None or
                    coding.get('review_decision') is not None or
                    coding.get('selection_type') in ['consensus', 'majority', 'manual_priority'] or
                    len([c for c in codings if c.get('segment_id') == coding.get('segment_id')]) == 1
                )
                
                matches_file = (
                    file_stem in coding.get('document', '') or 
                    file_stem in coding.get('segment_id', '')
                )
                
                if is_review_coding and matches_file:
                    file_codings.append(coding)
            
            if not file_codings:
                print(f"      ⚠️ Keine Review-Kodierungen für {filename} gefunden")
                continue
            
            print(f"      📋 {len(file_codings)} Review-Kodierungen gefunden")
            
            # FIX: Konvertiere zu PDF falls nötig
            if file_ext == '.pdf':
                pdf_path = file_path
                print(f"      ✅ Bereits PDF")
            else:
                print(f"      🔄 Konvertiere {file_ext.upper()} zu PDF...")
                pdf_path = pdf_converter.convert_document_to_pdf(file_path)
                
                if not pdf_path:
                    print(f"      ❌ Konvertierung fehlgeschlagen")
                    continue
                
                print(f"      ✅ PDF erstellt: {os.path.basename(pdf_path)}")
            
            # FIX: Annotiere PDF
            try:
                output_filename = f"{file_stem}_QCA_annotiert.pdf"
                output_file = os.path.join(self.output_dir, output_filename)
                
                result_path = pdf_annotator.annotate_pdf_with_codings(
                    pdf_path,
                    file_codings,
                    chunks,
                    output_file
                )
                
                if result_path:
                    annotated_files.append(result_path)
                    print(f"      ✅ Annotiert: {os.path.basename(result_path)}")
                else:
                    print(f"      ❌ Annotation fehlgeschlagen")
                
            except Exception as e:
                print(f"      ❌ Fehler bei Annotation: {e}")
                continue
        
        # FIX: Bereinige temporäre Dateien
        pdf_converter.cleanup_temp_pdfs()
        
        print(f"\n✅ Erweiterte PDF-Annotation abgeschlossen: {len(annotated_files)} Dateien erstellt")
        return annotated_files